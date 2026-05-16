"""
DOC -> DOCX 文档转换工具

v37.7.6: 从 main.py 提取为共享模块，消除 WordBatchReplaceWorker 和 MainWindow 之间的重复实现。
"""

import os
import platform
import shutil
import subprocess
import tempfile

from privacyguard.utils.exceptions import ConversionError
from privacyguard.utils.security import validate_safe_path
from privacyguard.utils.temp_manager import TempFileManager


def resolve_soffice_cmd():
    """检测 LibreOffice 可执行文件路径（跨平台）。"""
    system = platform.system()
    if system == "Darwin":
        mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists(mac_path):
            return mac_path
    elif system == "Windows":
        program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
        program_files_x86 = os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)")
        candidates = [
            os.path.join(program_files, "LibreOffice", "program", "soffice.exe"),
            os.path.join(program_files_x86, "LibreOffice", "program", "soffice.exe"),
        ]
        for candidate in candidates:
            if os.path.exists(candidate):
                return candidate
    return shutil.which("soffice") or "soffice"


def convert_with_libreoffice(doc_path, temp_dir=None, max_retries=1, timeout=90):
    """使用 LibreOffice 将 .doc 转换为 .docx。

    Args:
        doc_path: 源 .doc 文件路径
        temp_dir: 临时目录（None 则自动创建）
        max_retries: 最大重试次数
        timeout: 单次转换超时秒数

    Returns:
        str: 转换后的 .docx 文件路径

    Raises:
        ConversionError: 转换失败时抛出
    """
    own_temp = temp_dir is None
    if own_temp:
        temp_dir = tempfile.mkdtemp(prefix="pg_doc_convert_")

    try:
        # 将源文件复制到临时目录，确保路径为纯英文
        temp_doc_path = os.path.join(temp_dir, "source.doc")
        shutil.copy2(doc_path, temp_doc_path)

        # 安全验证
        is_safe, error_msg = validate_safe_path(temp_doc_path, allowed_extensions=['.doc'])
        if not is_safe:
            raise ConversionError("文件路径不安全", error_msg)
        is_safe, error_msg = validate_safe_path(temp_dir)
        if not is_safe:
            raise ConversionError("临时目录不安全", error_msg)

        soffice_cmd = resolve_soffice_cmd()
        base_name = os.path.splitext(os.path.basename(doc_path))[0]

        for attempt in range(max_retries + 1):
            try:
                cmd = [soffice_cmd, "--headless", "--convert-to", "docx",
                       "--outdir", temp_dir, temp_doc_path]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)

                if result.returncode != 0:
                    stderr = result.stderr.strip() if result.stderr else "未知错误"
                    if attempt < max_retries:
                        continue
                    raise ConversionError("LibreOffice 转换失败", stderr)

                # 检查输出文件
                docx_path = os.path.join(temp_dir, "source.docx")
                if not os.path.exists(docx_path):
                    docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                if not os.path.exists(docx_path):
                    raise ConversionError("转换失败", "未生成 .docx 输出文件")

                return docx_path

            except subprocess.TimeoutExpired:
                if attempt < max_retries:
                    continue
                raise ConversionError("转换超时", f"LibreOffice 在 {timeout} 秒内未响应")

            except ConversionError:
                raise

            except (OSError, IOError, RuntimeError, ValueError) as e:
                raise ConversionError(f"LibreOffice 转换出错: {e}", "请尝试在 Word 中手动另存为 .docx 格式")

    finally:
        if own_temp:
            shutil.rmtree(temp_dir, ignore_errors=True)


def convert_with_antiword(doc_path, temp_dir=None):
    """使用 antiword 提取文本并创建 .docx。

    Args:
        doc_path: 源 .doc 文件路径
        temp_dir: 临时目录（None 则自动创建）

    Returns:
        str: 转换后的 .docx 文件路径

    Raises:
        ConversionError: 转换失败时抛出
    """
    from docx import Document as DocxDocument

    own_temp = temp_dir is None
    if own_temp:
        temp_dir = tempfile.mkdtemp(prefix="pg_doc_anti_")

    try:
        # 安全验证
        is_safe, error_msg = validate_safe_path(doc_path, allowed_extensions=['.doc'])
        if not is_safe:
            raise ConversionError("文件路径不安全", error_msg)

        result = subprocess.run(["antiword", doc_path], capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            stderr = result.stderr.strip() if result.stderr else "未知错误"
            raise ConversionError("antiword 转换失败", stderr)

        doc = DocxDocument()
        lines = result.stdout.splitlines()
        if not lines:
            doc.add_paragraph("")
        else:
            for line in lines:
                doc.add_paragraph(line)

        output_path = os.path.join(temp_dir, "source.docx")
        doc.save(output_path)
        return output_path

    finally:
        if own_temp:
            shutil.rmtree(temp_dir, ignore_errors=True)


def convert_doc_to_docx(doc_path, temp_dir=None):
    """将 .doc 转换为 .docx（自动尝试 LibreOffice -> antiword）。

    Args:
        doc_path: 源 .doc 文件路径
        temp_dir: 共享临时目录（如果提供，调用方负责清理）

    Returns:
        tuple: (docx_path, temp_dir) — temp_dir 用于调用方后续清理

    Raises:
        ConversionError: 所有转换方式均失败时抛出
    """
    # 如果调用方未提供临时目录，创建一个共享的
    if temp_dir is None:
        temp_dir = tempfile.mkdtemp(prefix="pg_doc_convert_")

    try:
        return convert_with_libreoffice(doc_path, temp_dir=temp_dir), temp_dir
    except (OSError, IOError, RuntimeError, ValueError, ConversionError) as libreoffice_error:
        try:
            return convert_with_antiword(doc_path, temp_dir=temp_dir), temp_dir
        except (OSError, IOError, RuntimeError, ValueError, ConversionError) as antiword_error:
            raise ConversionError(
                "DOC 转换失败",
                f"LibreOffice: {libreoffice_error}\nantiword: {antiword_error}"
            )
