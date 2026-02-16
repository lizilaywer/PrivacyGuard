#!/usr/bin/env python3
"""
Word 文档格式保持验证脚本

用于验证脱敏后文档的格式是否保持正确
"""

import os
import sys
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor


def create_test_document(output_path):
    """创建包含多种格式的测试文档"""
    doc = Document()

    # 添加标题
    doc.add_heading('格式保持测试文档', level=1)

    # 段落1: 粗体 + 红色
    para1 = doc.add_paragraph()
    run1 = para1.add_run('姓名：')
    run1.bold = True
    run2 = para1.add_run('张三')
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    para1.add_run('，电话：13812345678')

    # 段落2: 斜体 + 下划线
    para2 = doc.add_paragraph()
    run3 = para2.add_run('电子邮箱：')
    run3.italic = True
    run4 = para2.add_run('test@example.com')
    run4.underline = True

    # 段落3: 不同字号
    para3 = doc.add_paragraph()
    run5 = para3.add_run('身份证号：')
    run5.font.size = Pt(10)
    run6 = para3.add_run('110101199001011234')
    run6.font.size = Pt(14)
    run6.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    # 段落4: 混合格式
    para4 = doc.add_paragraph()
    run7 = para4.add_run('银行卡号：')
    run7.bold = True
    run8 = para4.add_run('6222021234567890123')
    run8.italic = True
    run8.underline = True

    # 添加表格
    doc.add_heading('表格测试', level=2)
    table = doc.add_table(rows=3, cols=3)

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '姓名'
    header_cells[1].text = '电话'
    header_cells[2].text = '身份证号'

    # 数据行
    table.rows[1].cells[0].text = '李四'
    table.rows[1].cells[1].text = '13987654321'
    table.rows[1].cells[2].text = '110101199001011235'

    table.rows[2].cells[0].text = '王五'
    table.rows[2].cells[1].text = '15011112222'
    table.rows[2].cells[2].text = '110101199002022234'

    # 添加普通段落
    doc.add_paragraph('这是一个普通段落，没有敏感信息。')

    # 添加项目符号
    doc.add_paragraph('重要提示：', style='List Bullet')
    doc.add_paragraph('请保护个人信息安全', style='List Bullet')

    doc.save(output_path)
    return doc


def analyze_document(doc_path):
    """分析文档的格式信息"""
    doc = Document(doc_path)

    print(f"\n=== 文档分析: {os.path.basename(doc_path)} ===")

    # 分析段落
    print(f"\n段落数量: {len(doc.paragraphs)}")

    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"\n段落 {idx}: \"{para.text[:50]}...\"")
            print(f"  样式: {para.style.name}")

            # 分析 runs
            if para.runs:
                print(f"  Runs: {len(para.runs)}")
                for run_idx, run in enumerate(para.runs):
                    if run.text:
                        attrs = []
                        if run.bold: attrs.append("粗体")
                        if run.italic: attrs.append("斜体")
                        if run.underline: attrs.append("下划线")
                        if run.font.size: attrs.append(f"字号:{run.font.size.pt}")
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            # RGBColor 需要特殊处理
                            try:
                                color_val = int(rgb)
                                attrs.append(f"颜色:#{color_val:06X}")
                            except:
                                attrs.append(f"颜色:{rgb}")
                        print(f"    Run {run_idx}: \"{run.text[:30]}\" [{', '.join(attrs) if attrs else '普通'}]")

    # 分析表格
    print(f"\n表格数量: {len(doc.tables)}")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}: {len(table.rows)}行 x {len(table.columns)}列")
        for row_idx, row in enumerate(table.rows):
            row_text = " | ".join([cell.text[:20] for cell in row.cells])
            print(f"  行 {row_idx}: {row_text}")


def verify_format_preservation(original_path, processed_path):
    """验证格式是否保持"""
    orig_doc = Document(original_path)
    proc_doc = Document(processed_path)

    print("\n\n=== 格式保持验证 ===")

    # 验证段落数量
    if len(orig_doc.paragraphs) != len(proc_doc.paragraphs):
        print(f"⚠ 段落数量不匹配: {len(orig_doc.paragraphs)} vs {len(proc_doc.paragraphs)}")
    else:
        print(f"✓ 段落数量匹配: {len(orig_doc.paragraphs)}")

    # 验证表格数量
    if len(orig_doc.tables) != len(proc_doc.tables):
        print(f"⚠ 表格数量不匹配: {len(orig_doc.tables)} vs {len(proc_doc.tables)}")
    else:
        print(f"✓ 表格数量匹配: {len(orig_doc.tables)}")

    # 验证关键格式
    print("\n格式详细对比:")
    for idx in range(min(len(orig_doc.paragraphs), len(proc_doc.paragraphs))):
        orig_para = orig_doc.paragraphs[idx]
        proc_para = proc_doc.paragraphs[idx]

        if orig_para.text.strip() or proc_para.text.strip():
            print(f"\n段落 {idx}:")
            print(f"  原文: \"{orig_para.text[:60]}\"")
            print(f"  脱敏: \"{proc_para.text[:60]}\"")

            # 检查 runs 数量（可能因脱敏而变化）
            print(f"  Runs: {len(orig_para.runs)} → {len(proc_para.runs)}")

            # 检查粗体
            orig_bold = any(r.bold for r in orig_para.runs if r.text)
            proc_bold = any(r.bold for r in proc_para.runs if r.text)
            if orig_bold == proc_bold:
                print(f"  ✓ 粗体: {orig_bold}")
            else:
                print(f"  ⚠ 粗体: {orig_bold} → {proc_bold}")

            # 检查斜体
            orig_italic = any(r.italic for r in orig_para.runs if r.text)
            proc_italic = any(r.italic for r in proc_para.runs if r.text)
            if orig_italic == proc_italic:
                print(f"  ✓ 斜体: {orig_italic}")
            else:
                print(f"  ⚠ 斜体: {orig_italic} → {proc_italic}")

    print("\n\n=== 验证完成 ===")


def main():
    """主函数"""
    import shutil

    # 创建测试文档
    test_file = tempfile.NamedTemporaryFile(suffix='_test.docx', delete=False)
    test_file.close()

    print("创建测试文档...")
    create_test_document(test_file.name)
    print(f"测试文档已创建: {test_file.name}")

    # 分析原文档
    analyze_document(test_file.name)

    # 复制文档用于处理
    processed_file = tempfile.NamedTemporaryFile(suffix='_processed.docx', delete=False)
    processed_file.close()

    # 这里可以调用实际的脱敏处理
    shutil.copy2(test_file.name, processed_file.name)
    print(f"\n处理后的文档: {processed_file.name}")
    print("(注: 此脚本仅创建测试文档，实际脱敏请运行主程序)")

    # 验证格式
    verify_format_preservation(test_file.name, processed_file.name)

    print(f"\n测试文档已保存至: {test_file.name}")
    print("您可以在主程序中打开此文档进行测试")


if __name__ == '__main__':
    main()
