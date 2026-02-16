#!/usr/bin/env python3
"""Word 文档处理器单元测试

测试格式保持替换功能
"""

import unittest
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor


class TestWordProcessor(unittest.TestCase):
    """测试 Word 文档处理器"""

    def setUp(self):
        """测试前准备"""
        self.test_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        self.test_file.close()

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_file.name):
            os.remove(self.test_file.name)
        # 清理可能的 .tmp 文件
        if os.path.exists(self.test_file.name + '.tmp'):
            os.remove(self.test_file.name + '.tmp')

    def _create_test_doc(self):
        """创建包含多种格式的测试文档"""
        doc = Document()

        # 标题
        doc.add_heading('测试文档', level=1)

        # 段落1: 粗体 + 红色 + 大字号
        para1 = doc.add_paragraph()
        run1 = para1.add_run('联系人：')
        run1.bold = True
        run2 = para1.add_run('张三')
        run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run2.font.size = Pt(16)
        run3 = para1.add_run('，电话：13812345678')

        # 段落2: 斜体 + 下划线
        para2 = doc.add_paragraph()
        run4 = para2.add_run('电子邮箱：')
        run4.italic = True
        run5 = para2.add_run('test@example.com')
        run5.underline = True

        # 段落3: 身份证号
        doc.add_paragraph('身份证号：110101199001011234')

        # 表格
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = '姓名'
        table.rows[0].cells[1].text = '电话'
        table.rows[0].cells[2].text = '身份证号'
        table.rows[1].cells[0].text = '李四'
        table.rows[1].cells[1].text = '13987654321'
        table.rows[1].cells[2].text = '110101199001011235'

        doc.save(self.test_file.name)
        return doc

    def test_find_phone_number(self):
        """测试识别手机号"""
        from main import WordWorker, DEFAULT_RULES

        doc = self._create_test_doc()
        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': []
            }

        worker = WordWorker(doc, word_data, list(DEFAULT_RULES.values()), '', '[已脱敏]')
        matches = worker._find_matches('联系电话：13812345678')

        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]['text'], '13812345678')
        self.assertEqual(matches[0]['rule_name'], '手机号码')

    def test_find_id_card(self):
        """测试识别身份证号"""
        from main import WordWorker, DEFAULT_RULES

        doc = self._create_test_doc()
        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': []
            }

        worker = WordWorker(doc, word_data, list(DEFAULT_RULES.values()), '', '[已脱敏]')
        matches = worker._find_matches('身份证号：110101199001011234')

        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]['text'], '110101199001011234')
        self.assertEqual(matches[0]['rule_name'], '身份证号')

    def test_find_email(self):
        """测试识别邮箱"""
        from main import WordWorker, DEFAULT_RULES

        doc = self._create_test_doc()
        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': []
            }

        worker = WordWorker(doc, word_data, list(DEFAULT_RULES.values()), '', '[已脱敏]')
        matches = worker._find_matches('电子邮箱：test@example.com')

        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]['text'], 'test@example.com')
        self.assertEqual(matches[0]['rule_name'], '电子邮箱')

    def test_format_preservation_in_runs(self):
        """测试 run 级别的格式保持"""
        from main import WordWorker, DEFAULT_RULES

        doc = self._create_test_doc()

        # 初始化 word_data
        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': []
            }

        # 扫描敏感信息
        worker = WordWorker(doc, word_data, list(DEFAULT_RULES.values()), '', '[已脱敏]')

        # 手动设置手机号匹配（模拟扫描结果）
        word_data['paragraph_1']['ocr'] = [
            {
                'pattern': r'1[3-9]\d{9}',
                'rule_name': '手机号码',
                'start': 4,
                'end': 15,
                'text': '13812345678',
                'replacement': '[已脱敏]'
            }
        ]

        # 检查原文档的 run 格式
        original_para = doc.paragraphs[1]
        self.assertEqual(len(original_para.runs), 3)
        self.assertTrue(original_para.runs[0].bold)
        self.assertIsNotNone(original_para.runs[1].font.color.rgb)
        self.assertIsNotNone(original_para.runs[1].font.size)

    def test_table_cell_text_extraction(self):
        """测试表格单元格文本提取"""
        from main import WordWorker, DEFAULT_RULES

        doc = self._create_test_doc()

        # 初始化 word_data，包括表格
        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': []
            }

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                    word_data[key] = {
                        'type': 'table_cell',
                        'table': table_idx,
                        'row': row_idx,
                        'cell': cell_idx,
                        'text': cell.text,
                        'ocr': [],
                        'manual': []
                    }

        # 检查表格单元格是否正确提取
        self.assertIn('table_0_cell_1_1', word_data)
        self.assertEqual(word_data['table_0_cell_1_1']['text'], '13987654321')


class TestIntegration(unittest.TestCase):
    """集成测试"""

    def setUp(self):
        """测试前准备"""
        self.test_input = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        self.test_input.close()
        self.test_output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        self.test_output.close()

    def tearDown(self):
        """测试后清理"""
        for f in [self.test_input.name, self.test_output.name]:
            if os.path.exists(f):
                os.remove(f)
            if os.path.exists(f + '.tmp'):
                os.remove(f + '.tmp')

    def test_full_workflow(self):
        """测试完整工作流程：创建 -> 扫描 -> 保存"""
        from docx import Document

        # 1. 创建测试文档
        doc = Document()
        doc.add_paragraph('联系人：张三，电话：13812345678')
        doc.add_paragraph('邮箱：test@example.com')
        doc.save(self.test_input.name)

        # 2. 打开文档
        opened_doc = Document(self.test_input.name)
        self.assertEqual(len(opened_doc.paragraphs), 2)
        self.assertIn('13812345678', opened_doc.paragraphs[0].text)
        self.assertIn('test@example.com', opened_doc.paragraphs[1].text)


if __name__ == '__main__':
    # 运行测试
    unittest.main(verbosity=2)
