import os
import tempfile
import unittest

from docx import Document

from main import WordBatchReplaceWorker


class TestWordBatchReplaceWorker(unittest.TestCase):

    def test_build_output_path_with_conflict(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_file = os.path.join(tmpdir, "sample.docx")
            with open(input_file, "w", encoding="utf-8") as f:
                f.write("x")

            worker = WordBatchReplaceWorker([], [], "[默认]")
            first = worker._build_output_path(input_file)
            # 制造重名冲突
            with open(first, "w", encoding="utf-8") as f:
                f.write("x")
            second = worker._build_output_path(input_file)

            self.assertTrue(first.endswith(".docx"))
            self.assertTrue(second.endswith(".docx"))
            self.assertNotEqual(first, second)

    def test_apply_rules_to_document(self):
        doc = Document()
        doc.add_paragraph("姓名：张三 电话：13812345678")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "身份证：110101199001011234"

        rules = [
            {"enabled": True, "mode": "exact", "find": "张三", "replace": "某某"},
            {"enabled": True, "mode": "regex", "find": r"\d{17}[\dXx]", "replace": "[证件号]"},
            {"enabled": True, "mode": "regex", "find": r"1[3-9]\d{9}", "replace": "[手机号]"},
        ]
        worker = WordBatchReplaceWorker([], rules, "[默认]")
        worker._apply_rules_to_document(doc)

        self.assertIn("某某", doc.paragraphs[0].text)
        self.assertIn("[手机号]", doc.paragraphs[0].text)
        self.assertIn("[证件号]", doc.tables[0].cell(0, 0).text)


if __name__ == "__main__":
    unittest.main()
