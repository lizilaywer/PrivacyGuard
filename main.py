import sys
import os
import fitz  # PyMuPDF
import re
import cv2
import numpy as np
import time
import shutil
import threading  # v36.5: 线程锁支持
import atexit  # v36.2: 用于确保临时文件清理
import tempfile  # v36.2: 临时文件管理
import traceback  # v37.0.5: 异常追踪
from pathlib import Path
from io import BytesIO
from PIL import Image
from bs4 import BeautifulSoup
from privacyguard.ocr.mixed_pdf import (
    collect_embedded_image_clip_rects,
    collect_image_block_ocr_hits,
)
from privacyguard.ocr.text_pdf import collect_text_pdf_hit_boxes
from privacyguard.utils.security import validate_safe_path, resource_path
from privacyguard.utils.exceptions import (
    PrivacyAppError,
    ConversionError,
    FileFormatError,
    SecurityError,
    MemoryLimitError,
    WorkerCancelledError,
)
from privacyguard.utils.temp_manager import TempFileManager
from privacyguard.workers.image_merge import ImageMergeWorker
from privacyguard.workers.word_worker import WordWorker as _ModularWordWorker
from privacyguard.workers.ocr_worker import OCRWorker as _ModularOCRWorker
from privacyguard.utils.doc_converter import convert_doc_to_docx as _shared_convert_doc_to_docx

# v37.0.5: 延迟导入 OCR 模块，便于错误处理
RapidOCR = None
OCR_INIT_ERROR = None

def init_ocr_engine():
    """v37.0.5: 安全初始化 OCR 引擎，捕获所有可能的错误"""
    global RapidOCR, OCR_INIT_ERROR
    if RapidOCR is not None:
        return True

    try:
        from rapidocr_onnxruntime import RapidOCR as _RapidOCR
        RapidOCR = _RapidOCR
        # 预热：创建一个测试实例验证 DLL 加载
        _ = _RapidOCR()
        print("[OCR] 引擎初始化成功")
        return True
    except ImportError as e:
        OCR_INIT_ERROR = f"OCR 模块未安装: {e}\n请运行: pip install rapidocr-onnxruntime"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        return False
    except OSError as e:
        OCR_INIT_ERROR = f"OCR DLL 加载失败: {e}\n可能缺少 Visual C++ 运行库"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        return False
    except Exception as e:
        OCR_INIT_ERROR = f"OCR 初始化失败: {type(e).__name__}: {e}"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        traceback.print_exc()
        return False

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QBoxLayout, QGridLayout,
                             QHBoxLayout, QPushButton, QLabel, QFileDialog,
                             QScrollArea, QMessageBox, QProgressBar, QFrame,
                             QDialog, QCheckBox, QGroupBox, QTextEdit, QSpinBox,
                             QRadioButton, QButtonGroup, QComboBox, QSizePolicy,
                             QTextBrowser, QLineEdit, QListWidget, QListWidgetItem,
                             QAbstractItemView, QSlider, QTableWidget, QMenu,
                             QTableWidgetItem, QHeaderView, QStyle)
from PyQt6.QtGui import QPixmap, QImage, QPainter, QPen, QColor, QWheelEvent, QCursor, QIcon, QDesktopServices
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QRectF, QPointF, QSettings, QMutex, QMutexLocker, QObject, pyqtSlot, QSize, QTimer, QUrl
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebChannel import QWebChannel
from PyQt6 import sip

# 导入主题系统
from theme import Theme

# v37.0.4: 简化配置系统 - 直接从 JSON 文件加载
import json


def read_app_version():
    """从统一版本文件读取基础版本号。"""
    version_file = Path(__file__).resolve().parent / "version.txt"
    try:
        return version_file.read_text(encoding="utf-8").strip()
    except OSError:
        return "37.7.6"

class SimpleConfig:
    """简化配置管理器 - 直接从 config.json 读取"""

    def __init__(self, config_path=None):
        self._config = {}
        if config_path is None:
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')
        self._config_path = config_path
        self._load_config()

    def _load_config(self):
        """加载配置文件"""
        try:
            if os.path.exists(self._config_path):
                with open(self._config_path, 'r', encoding='utf-8') as f:
                    self._config = json.load(f)
        except (OSError, IOError, json.JSONDecodeError) as e:
            print(f"[配置系统] 加载配置失败: {e}")

    def save(self):
        """将当前配置写回磁盘。"""
        try:
            with open(self._config_path, 'w', encoding='utf-8') as f:
                json.dump(self._config, f, indent=2, ensure_ascii=False)
            return True
        except (OSError, IOError, TypeError) as e:
            print(f"[配置系统] 保存配置失败: {e}")
            return False

    def get(self, key, default=None):
        """获取配置值（支持点分隔路径）"""
        keys = key.split('.')
        value = self._config
        for k in keys:
            if isinstance(value, dict) and k in value:
                value = value[k]
            else:
                return default
        return value

    def set(self, key, value, persist=True):
        """设置配置值（支持点分隔路径）

        Args:
            key: 配置键，支持点分隔路径
            value: 配置值
            persist: 是否立即保存到文件
        """
        keys = key.split('.')
        config = self._config
        # 遍历到倒数第二层，创建缺失的字典
        for k in keys[:-1]:
            if k not in config:
                config[k] = {}
            elif not isinstance(config[k], dict):
                config[k] = {}
            config = config[k]
        # 设置最终值
        config[keys[-1]] = value

        # 保存到文件
        if persist:
            self.save()

    def get_redaction_rules(self):
        """获取脱敏规则"""
        return self.get('redaction.default_rules', {})

# 初始化配置
config = SimpleConfig()

# === 核心防崩溃设置 ===
cv2.setNumThreads(0)
os.environ["OMP_NUM_THREADS"] = "1"

# === 软件配置 ===
# v37.0: 从配置读取，失败时使用硬编码后备
APP_NAME = config.get("app.name", "PrivacyGuard 脱敏卫士") if config else "PrivacyGuard 脱敏卫士"
APP_VERSION = read_app_version()
VERSION = f"{APP_VERSION} - Engineering Remediation"
PREVIEW_FONT_STACK = '"Segoe UI Variable", "Segoe UI", "Microsoft YaHei UI", "Microsoft YaHei", Arial, sans-serif'
WORD_PREVIEW_IMAGE_EXTENSION_MAP = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/tiff": ".tiff",
    "image/bmp": ".bmp",
    "image/webp": ".webp",
    "image/svg+xml": ".svg",
}
WORD_PREVIEW_BROKEN_IMAGE_DATA_URI = (
    "data:image/gif;base64,R0lGODlhAQABAIAAAAAAAP///ywAAAAAAQABAAACAUwAOw=="
)

# === 常量定义 ===
# v37.0: 从配置读取，失败时使用硬编码后备
MIN_RECT_WIDTH = config.get("ocr.min_rect_width", 5) if config else 5
PROGRESS_UPDATE_INTERVAL = config.get("ocr.progress_update_interval", 0.05) if config else 0.05
ZOOM_MIN = config.get("ocr.zoom_min", 0.5) if config else 0.5
ZOOM_MAX = config.get("ocr.zoom_max", 4.0) if config else 4.0
DEBUG_MODE = os.getenv('PRIVACYGUARD_DEBUG', 'False').lower() == 'true' if not config else config.get("advanced.debug_mode", False)

# === 默认规则库 ===
# v37.0: 从配置读取，支持新旧两种格式
if config:
    _rules_from_config = config.get_redaction_rules()
    DEFAULT_RULES = {}
    for name, rule in _rules_from_config.items():
        if isinstance(rule, dict):
            DEFAULT_RULES[name] = rule.get("pattern", "")
        else:
            DEFAULT_RULES[name] = str(rule)
else:
    DEFAULT_RULES = {
        "身份证号": r"(?<!\d)([1-9]\d{5}(19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]|\d{15})(?!\d)",
        "手机号码": r"(?<!\d)(1[3-9]\d{9})(?!\d)",
        "日期时间": r"\d{4}[年\-\.]\d{1,2}[月\-\.]\d{1,2}[日]?",
        "电子邮箱": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        "银行卡号": r"(?<!\d)([1-9]\d{12,18})(?!\d)",
        "印章": "__SEAL_DETECTION__"  # v37.5.0: 印章检测特殊标记
    }

WORD_RULE_SCHEMA_VERSION = 1


def normalize_word_replace_rules(rules, default_replacement_text="[已脱敏]"):
    """规范化多字段替换规则（会话级规则，不自动持久化）。"""
    normalized = []
    if not isinstance(rules, list):
        return normalized

    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
    mode_alias = {
        "exact": "exact",
        "regex": "regex",
        "精确": "exact",
        "正则": "regex"
    }

    for item in rules:
        if not isinstance(item, dict):
            continue
        enabled = bool(item.get("enabled", True))
        raw_mode = str(item.get("mode", "exact")).strip().lower()
        mode = mode_alias.get(raw_mode, "exact")

        find_text = str(item.get("find", "")).strip()
        if not find_text:
            continue

        replace_text = item.get("replace")
        if replace_text is None or str(replace_text) == "":
            replace_text = fallback_text
        else:
            replace_text = str(replace_text)

        normalized.append({
            "enabled": enabled,
            "mode": mode,
            "find": find_text,
            "replace": replace_text
        })

    return normalized


def resolve_word_preview_image_suffix(content_type):
    """根据 Mammoth 图片内容类型推导本地文件后缀。"""
    if not isinstance(content_type, str):
        return ".img"

    normalized = content_type.strip().lower()
    if normalized in WORD_PREVIEW_IMAGE_EXTENSION_MAP:
        return WORD_PREVIEW_IMAGE_EXTENSION_MAP[normalized]

    if "/" not in normalized:
        return ".img"

    subtype = normalized.split("/", 1)[1].split(";", 1)[0].strip()
    subtype = subtype.replace("+xml", "").replace("+zip", "")
    subtype = re.sub(r"[^a-z0-9]+", "", subtype)
    if not subtype:
        return ".img"
    return f".{subtype}"


def _range_overlaps(start, end, ranges):
    """判断区间是否与已有区间重叠。"""
    for s, e in ranges:
        if not (end <= s or start >= e):
            return True
    return False


def build_word_rule_matches(text, rules, default_replacement_text="[已脱敏]"):
    """根据规则查找文本匹配，执行策略：exact 优先于 regex，重叠先到先得。"""
    if not isinstance(text, str) or not text:
        return []

    normalized_rules = normalize_word_replace_rules(rules, default_replacement_text)
    selected = []
    occupied_ranges = []

    for target_mode in ("exact", "regex"):
        for rule_index, rule in enumerate(normalized_rules):
            if not rule.get("enabled", True) or rule.get("mode") != target_mode:
                continue

            pattern = re.escape(rule["find"]) if target_mode == "exact" else rule["find"]
            try:
                compiled = re.compile(pattern)
            except re.error:
                continue

            for matched in compiled.finditer(text):
                start = matched.start()
                end = matched.end()
                if start >= end:
                    continue
                if _range_overlaps(start, end, occupied_ranges):
                    continue

                selected.append({
                    "start": start,
                    "end": end,
                    "text": matched.group(0),
                    "replacement": rule["replace"],
                    "mode": target_mode,
                    "rule_index": rule_index,
                    "source": "rule"
                })
                occupied_ranges.append((start, end))

    selected.sort(key=lambda item: item["start"])
    return selected


def apply_rule_matches_to_text(text, matches):
    """将匹配区间应用到文本（倒序替换避免索引偏移）。"""
    if not isinstance(text, str) or not matches:
        return text

    output = text
    for match in sorted(matches, key=lambda item: item.get("start", 0), reverse=True):
        start = match.get("start")
        end = match.get("end")
        replacement = match.get("replacement", "")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < 0 or end > len(output) or start >= end:
            continue
        if replacement is None:
            replacement = ""
        if not isinstance(replacement, str):
            replacement = str(replacement)
        output = output[:start] + replacement + output[end:]
    return output


def apply_word_rules_to_text(text, rules, default_replacement_text="[已脱敏]"):
    """直接按规则替换文本并返回替换结果。"""
    matches = build_word_rule_matches(text, rules, default_replacement_text)
    return apply_rule_matches_to_text(text, matches)


def build_replaced_preview_segments(text, matches, default_replacement_text="[已脱敏]"):
    """根据匹配区间生成替换后文本分段（用于右侧预览高亮）。"""
    if not isinstance(text, str):
        return [{"type": "text", "value": ""}]
    if not matches:
        return [{"type": "text", "value": text}]

    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
    segments = []
    cursor = 0

    for match in sorted(matches, key=lambda item: item.get("start", 0)):
        start = match.get("start")
        end = match.get("end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < cursor or start < 0 or end > len(text) or start >= end:
            continue

        if start > cursor:
            segments.append({
                "type": "text",
                "value": text[cursor:start]
            })

        replacement = match.get("replacement", fallback_text)
        if replacement is None:
            replacement = fallback_text
        if not isinstance(replacement, str):
            replacement = str(replacement)

        segments.append({
            "type": "replacement",
            "value": replacement,
            "source": match.get("source", "rule"),
            "mode": match.get("mode", ""),
            "rule_name": match.get("rule_name", "")
        })
        cursor = end

    if cursor < len(text):
        segments.append({
            "type": "text",
            "value": text[cursor:]
        })

    if not segments:
        return [{"type": "text", "value": text}]
    return segments


def build_highlight_preview_segments(text, matches):
    """根据匹配区间生成原文高亮分段（用于左侧预览）。"""
    if not isinstance(text, str):
        return [{"type": "text", "value": ""}]
    if not matches:
        return [{"type": "text", "value": text}]

    segments = []
    cursor = 0
    for match in sorted(matches, key=lambda item: item.get("start", 0)):
        start = match.get("start")
        end = match.get("end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < cursor or start < 0 or end > len(text) or start >= end:
            continue

        if start > cursor:
            segments.append({"type": "text", "value": text[cursor:start]})

        segments.append({
            "type": "highlight",
            "value": text[start:end],
            "source": match.get("source", "manual"),
            "mode": match.get("mode", ""),
            "rule_name": match.get("rule_name", ""),
            "start": start,
            "end": end,
        })
        cursor = end

    if cursor < len(text):
        segments.append({"type": "text", "value": text[cursor:]})

    if not segments:
        return [{"type": "text", "value": text}]
    return segments


WORD_PREVIEW_BLOCK_SELECTOR = '[data-word-block="1"][data-key]'


def build_word_panel_update_script(block_updates):
    """构建仅更新正文块的 Word 预览增量刷新脚本。"""
    payload = json.dumps(block_updates or {}, ensure_ascii=False)
    return f"""
        (function() {{
            const updates = {payload};
            const elements = document.querySelectorAll('{WORD_PREVIEW_BLOCK_SELECTOR}');
            elements.forEach(function(el) {{
                const key = el.dataset.key;
                if (Object.prototype.hasOwnProperty.call(updates, key)) {{
                    const nextHtml = updates[key];
                    if (el.innerHTML !== nextHtml) {{
                        el.innerHTML = nextHtml;
                    }}
                }}
            }});
        }})();
    """


def should_reload_word_panel(source_changed, loaded_source_path, current_file_path, panel_ready):
    """判断 Word 预览面板是否需要重新加载完整文档。"""
    if source_changed:
        return True
    if not panel_ready:
        return True
    return loaded_source_path != current_file_path


def format_signed_percent(value):
    """将百分比格式化为适合界面展示的文案。"""
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = 0
    return f"{number:+d}%" if number else "0%"


def build_settings_nav_labels(enabled_rules, keyword_count, precision_is_default, ocr_adjust_value):
    """构建设置中心左侧导航标签。"""
    return [
        f"1 通用规则 · {max(0, int(enabled_rules))}项启用",
        f"2 自定义关键词 · {max(0, int(keyword_count))}条",
        f"3 扫描与微调 · {'默认' if precision_is_default else '已微调'}",
        f"4 OCR 检测框 · {format_signed_percent(ocr_adjust_value)}",
    ]


def build_settings_hero_tags(enabled_rules, keyword_count, enabled_word_rules, precision_is_default, ocr_adjust_value, scan_label):
    """构建设置页顶部的动态摘要标签。"""
    common_tag = (
        f"常用：规则 {max(0, int(enabled_rules))} 项 · "
        f"关键词 {max(0, int(keyword_count))} 条 · "
        f"Word {max(0, int(enabled_word_rules))} 条"
    )
    if precision_is_default:
        advanced_tag = f"高级：扫描推荐值 · OCR {format_signed_percent(ocr_adjust_value)}"
    else:
        normalized_scan_label = str(scan_label or "-").strip() or "-"
        advanced_tag = (
            f"高级：{normalized_scan_label} · "
            f"OCR {format_signed_percent(ocr_adjust_value)} · 已微调"
        )
    return common_tag, advanced_tag


def build_batch_result_rows(summary):
    """将批量替换 summary 转成结果表格行。"""
    if not isinstance(summary, dict):
        return []

    rows = []
    failed_items = summary.get("failed", []) if isinstance(summary.get("failed", []), list) else []
    success_items = summary.get("success", []) if isinstance(summary.get("success", []), list) else []

    for item in failed_items:
        if not isinstance(item, dict):
            continue
        input_path = str(item.get("input", "") or "")
        rows.append({
            "status": "失败",
            "status_key": "failed",
            "document": os.path.basename(input_path) if input_path else "未知文档",
            "detail": str(item.get("error", "") or "处理失败"),
            "action": "双击定位原文件",
            "open_path": input_path,
            "fallback_dir": os.path.dirname(input_path) if input_path else "",
        })

    for item in success_items:
        if not isinstance(item, dict):
            continue
        input_path = str(item.get("input", "") or "")
        output_path = str(item.get("output", "") or "")
        rows.append({
            "status": "成功",
            "status_key": "success",
            "document": os.path.basename(input_path) if input_path else "未知文档",
            "detail": os.path.basename(output_path) if output_path else "已生成输出文件",
            "action": "双击打开输出",
            "open_path": output_path,
            "fallback_dir": os.path.dirname(output_path) if output_path else "",
        })

    return rows


def summarize_batch_result_rows(rows):
    """汇总批量结果行数量。"""
    summary = {"total": 0, "success": 0, "failed": 0}
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        summary["total"] += 1
        status_key = row.get("status_key")
        if status_key == "success":
            summary["success"] += 1
        elif status_key == "failed":
            summary["failed"] += 1
    return summary


def build_batch_filter_labels(summary_counts, show_counts=False):
    """构建批量结果筛选按钮文案。"""
    counts = summary_counts if isinstance(summary_counts, dict) else {}
    total = max(0, int(counts.get("total", 0) or 0))
    success = max(0, int(counts.get("success", 0) or 0))
    failed = max(0, int(counts.get("failed", 0) or 0))
    if not show_counts:
        return {"all": "全部", "success": "成功", "failed": "失败"}
    return {
        "all": f"全部 {total}",
        "success": f"成功 {success}",
        "failed": f"失败 {failed}",
    }


def filter_batch_result_rows(rows, filter_mode):
    """按筛选模式过滤批量结果行。"""
    if filter_mode not in {"all", "success", "failed"}:
        filter_mode = "all"

    filtered = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        if filter_mode == "all":
            filtered.append(row)
        elif row.get("status_key") == filter_mode:
            filtered.append(row)
    return filtered


def build_batch_rule_summary_lines(rules, success_items, default_replacement_text="[已脱敏]"):
    """按规则生成批量替换摘要明细。"""
    normalized_rules = normalize_word_replace_rules(rules, default_replacement_text)
    if not normalized_rules:
        return []

    def _extract_rule_count(item, target_rule_index):
        if not isinstance(item, dict):
            return 0

        counts = item.get("rule_counts", [])
        if isinstance(counts, dict):
            try:
                return max(0, int(counts.get(str(target_rule_index), counts.get(target_rule_index, 0)) or 0))
            except (TypeError, ValueError):
                return 0

        if not isinstance(counts, list):
            return 0

        for entry in counts:
            if not isinstance(entry, dict):
                continue
            try:
                rule_index = int(entry.get("rule_index", -1))
                count = int(entry.get("count", 0) or 0)
            except (TypeError, ValueError):
                continue
            if rule_index == target_rule_index:
                return max(0, count)
        return 0

    lines = []
    for rule_index, rule in enumerate(normalized_rules, start=1):
        doc_parts = []
        replacement_text = rule.get("replace") or default_replacement_text

        for item in success_items or []:
            count = _extract_rule_count(item, rule_index - 1)
            if count <= 0:
                continue
            input_name = os.path.basename(str(item.get("input", "") or "")) or "未知文档"
            doc_parts.append(f"{input_name} 成功替换 {count} 条")

        if doc_parts:
            lines.append(
                f"{rule_index}、“{rule.get('find', '')}”替换为“{replacement_text}”，"
                + "，".join(doc_parts)
                + "；"
            )
        else:
            lines.append(
                f"{rule_index}、“{rule.get('find', '')}”替换为“{replacement_text}”，本轮未命中；"
            )

    return lines


def build_workbench_guidance(mode, batch_stage="rule_setup", has_results=False, compare_mode=False):
    """按当前模式生成顶部工作台的下一步引导标签。"""
    if mode == "pdf":
        first_step = "下一步：人工复核并导出" if has_results else "下一步：先点智能脱敏"
        return [
            first_step,
            "黑 / 白遮罩可立即切换",
            "支持手动画框补充脱敏",
        ]
    if mode == "word":
        compare_tip = "当前可隐藏对比预览" if compare_mode else "需要时可打开对比预览"
        first_step = "下一步：先检查替换规则" if not has_results else "下一步：复核替换结果"
        return [
            first_step,
            "原文预览与替换预览分开显示",
            compare_tip,
        ]
    if mode == "batch":
        if batch_stage == "running":
            return [
                "当前：正在批量替换文档",
                "可随时停止并保留已完成结果",
                "完成后可筛选成功 / 失败清单",
            ]
        if batch_stage in ("finished", "stopped"):
            return [
                "下一步：先看失败文档和原因",
                "可仅重试失败文档",
                "双击结果可打开输出或定位原文件",
            ]
        return [
            "下一步：确认规则后再开始执行",
            "这一步不会改动任何原文件",
            "建议至少启用一条 Word 替换规则",
        ]
    if mode == "image_merge":
        return [
            "下一步：确认图片顺序后开始合并",
            "支持多张图片自动合成为 PDF",
            "合并完成后会直接进入 PDF 脱敏",
        ]
    return [
        "支持拖拽导入，系统会自动分流",
        "PDF 走脱敏，Word 走替换",
        "多个 Word 会先进入批量规则确认",
        "多张图片可直接合并为 PDF",
    ]


def build_toolbar_mode_labels(mode, density_mode, has_results=False, enabled_word_rules=0):
    """构建工具栏在不同模式下的主动作文案。"""
    compact = density_mode != "wide"

    if mode == "pdf":
        if has_results:
            scan_text = "重脱" if compact else "重新脱敏"
            scan_tooltip = "重新执行 PDF 智能脱敏扫描"
        else:
            scan_text = "脱敏" if compact else "智能脱敏"
            scan_tooltip = "执行 PDF 智能脱敏扫描"
        save_text = "导出" if compact else "导出 PDF"
        save_tooltip = "导出当前 PDF 脱敏结果"
    elif mode == "word":
        if has_results:
            scan_text = "重替" if compact else "重新替换"
            scan_tooltip = "重新执行 Word 智能替换扫描"
        else:
            scan_text = "替换" if compact else "智能替换"
            scan_tooltip = "执行 Word 智能替换扫描"
        save_text = "导出" if compact else "导出 Word"
        save_tooltip = "导出当前 Word 替换结果"
    else:
        scan_text = "脱敏" if compact else "智能脱敏"
        save_text = "导出"
        scan_tooltip = "执行智能脱敏扫描"
        save_tooltip = "导出处理结果"

    if enabled_word_rules > 0:
        word_rules_text = f"规则 {enabled_word_rules}" if compact else f"替换规则 {enabled_word_rules}"
        word_rules_tooltip = f"打开 Word 替换规则（当前启用 {enabled_word_rules} 条）"
    else:
        word_rules_text = "规则" if compact else "替换规则"
        word_rules_tooltip = "打开 Word 替换规则"

    return {
        "open_text": "打开",
        "open_tooltip": "打开 PDF、Word 或图片文件",
        "scan_text": scan_text,
        "scan_tooltip": scan_tooltip,
        "save_text": save_text,
        "save_tooltip": save_tooltip,
        "word_rules_text": word_rules_text,
        "word_rules_tooltip": word_rules_tooltip,
    }


def _shift_density_mode(mode, order, step):
    """在既定密度序列里前后移动一档。"""
    if mode not in order:
        return mode
    index = order.index(mode)
    target = max(0, min(len(order) - 1, index + step))
    return order[target]


def resolve_workspace_density_mode(mode, width, height=0, scale=1.0):
    """解析主工作区工具栏密度档位，兼顾 Windows DPI 与窗口高度。"""
    width = max(int(width or 0), 1)
    height = max(int(height or 0), 0)
    scale = max(1.0, float(scale or 1.0))

    if mode == "pdf":
        wide_threshold = 1500
        compact_threshold = 1260
    elif mode == "word":
        wide_threshold = 1220
        compact_threshold = 980
    else:
        wide_threshold = 1080
        compact_threshold = 860

    if scale >= 1.5:
        wide_threshold += 90
        compact_threshold += 60
    elif scale >= 1.25:
        wide_threshold += 50
        compact_threshold += 30

    if height:
        if height >= 980:
            wide_threshold -= 50
            compact_threshold -= 30
        elif height <= 760:
            wide_threshold += 70
            compact_threshold += 40

    wide_threshold = max(compact_threshold + 80, wide_threshold)
    compact_threshold = max(720, min(compact_threshold, wide_threshold - 80))

    if width >= wide_threshold:
        return "wide"
    if width >= compact_threshold:
        return "compact"
    return "narrow"


def resolve_settings_density_mode(width, height=0, scale=1.0):
    """解析高级设置页密度档位，优先为 Windows 高 DPI 和不同窗口高度收口。"""
    width = max(int(width or 0), 1)
    height = max(int(height or 0), 0)
    scale = max(1.0, float(scale or 1.0))
    order = ["narrow", "compact", "roomy", "wide"]

    if width >= 1700:
        density_mode = "wide"
    elif width >= 1450:
        density_mode = "roomy"
    elif width >= 1260:
        density_mode = "compact"
    else:
        density_mode = "narrow"

    if scale >= 1.5:
        density_mode = _shift_density_mode(density_mode, order, -1)
        if width < 1360:
            density_mode = "narrow"
    elif scale >= 1.25 and density_mode == "wide" and width < 1760:
        density_mode = "roomy"

    if height:
        if height <= 820:
            density_mode = _shift_density_mode(density_mode, order, -1)
        elif height >= 980:
            if density_mode == "compact" and width >= 1380:
                density_mode = "roomy"
            elif density_mode == "roomy" and width >= 1600:
                density_mode = "wide"

    return density_mode


def merge_word_matches_with_priority(text, rules, default_replacement_text,
                                     manual_matches=None, ocr_matches=None):
    """合并规则替换、手动脱敏、OCR 脱敏区间，优先级：规则 > 手动 > OCR。"""
    manual_matches = manual_matches or []
    ocr_matches = ocr_matches or []
    text_len = len(text) if isinstance(text, str) else 0
    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"

    merged = []
    occupied_ranges = []

    def _append_candidates(candidates, source_name):
        for item in candidates:
            start = item.get("start")
            end = item.get("end")
            if not isinstance(start, int) or not isinstance(end, int):
                continue
            if start < 0 or end > text_len or start >= end:
                continue
            if _range_overlaps(start, end, occupied_ranges):
                continue

            replacement = item.get("replacement", fallback_text)
            if replacement is None:
                replacement = fallback_text
            if not isinstance(replacement, str):
                replacement = str(replacement)

            merged.append({
                "start": start,
                "end": end,
                "text": item.get("text", text[start:end] if isinstance(text, str) else ""),
                "replacement": replacement,
                "source": source_name,
                "mode": item.get("mode", "global"),
                "rule_name": item.get("rule_name", "")
            })
            occupied_ranges.append((start, end))

    _append_candidates(build_word_rule_matches(text, rules, fallback_text), "rule")
    _append_candidates(manual_matches, "manual")
    _append_candidates(ocr_matches, "ocr")
    merged.sort(key=lambda item: item["start"])
    return merged


def apply_range_to_runs(para, start, end, replacement):
    """在段落 run 列表上应用一次区间替换。"""
    if start >= end:
        return
    if not para.runs:
        return

    run_ranges = []
    cursor = 0
    for idx, run in enumerate(para.runs):
        text = run.text or ''
        run_start = cursor
        run_end = cursor + len(text)
        run_ranges.append((idx, run_start, run_end))
        cursor = run_end

    total_len = cursor
    if start < 0 or end > total_len:
        return

    start_run_idx = None
    start_offset = 0
    for idx, run_start, run_end in run_ranges:
        if start < run_end:
            start_run_idx = idx
            start_offset = start - run_start
            break
    if start_run_idx is None:
        start_run_idx = run_ranges[-1][0]
        start_offset = len(para.runs[start_run_idx].text or '')

    end_run_idx = None
    end_offset = 0
    for idx, run_start, run_end in run_ranges:
        if end <= run_end:
            end_run_idx = idx
            end_offset = end - run_start
            break
    if end_run_idx is None:
        end_run_idx = run_ranges[-1][0]
        end_offset = len(para.runs[end_run_idx].text or '')

    start_run = para.runs[start_run_idx]
    end_run = para.runs[end_run_idx]
    start_text = start_run.text or ''
    end_text = end_run.text or ''

    prefix = start_text[:start_offset]
    suffix = end_text[end_offset:]
    start_run.text = prefix + replacement + suffix

    if end_run_idx > start_run_idx:
        for idx in range(start_run_idx + 1, end_run_idx + 1):
            para.runs[idx].text = ''


def replace_matches_in_paragraph(para, matches, text_offset=0, fallback_replacement_text="[已脱敏]"):
    """按匹配区间替换段落文本，避免同词误替换和跨 run 漏替换。"""
    if not matches or not para.runs:
        return

    paragraph_text = ''.join(run.text for run in para.runs)
    if not paragraph_text:
        return

    text_len = len(paragraph_text)
    ranges = []
    seen = set()

    for match in matches:
        start = match.get('start')
        end = match.get('end')
        if not isinstance(start, int) or not isinstance(end, int):
            continue

        local_start = start - text_offset
        local_end = end - text_offset
        if local_start < 0 or local_end > text_len or local_start >= local_end:
            continue

        replacement = match.get('replacement', fallback_replacement_text)
        if replacement is None:
            replacement = fallback_replacement_text
        if not isinstance(replacement, str):
            replacement = str(replacement)

        key = (local_start, local_end, replacement)
        if key in seen:
            continue
        seen.add(key)
        ranges.append({
            'start': local_start,
            'end': local_end,
            'replacement': replacement
        })

    if not ranges:
        return

    ranges.sort(key=lambda item: (item['start'], -(item['end'] - item['start'])))
    filtered = []
    last_end = -1
    for item in ranges:
        if item['start'] < last_end:
            continue
        filtered.append(item)
        last_end = item['end']

    for item in reversed(filtered):
        apply_range_to_runs(para, item['start'], item['end'], item['replacement'])


# === 设置对话框 ===
class SettingsDialog(QDialog):
    """设置对话框 - v37.0: 支持配置持久化"""

    def __init__(self, parent=None, current_rules=None, use_enhance=False, custom_keywords="",
                 scan_level=2.0, offset_x=0, offset_w=0, replacement_text="[已脱敏]",
                 word_replace_rules=None,
                 config_manager=None):
        super().__init__(parent)
        self.config = config_manager

        # v37.4.1: 修复 Windows 深色模式下对话框显示问题
        # 设置窗口标志，确保对话框在深色系统主题下使用浅色样式
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        # v37.0: 从配置读取窗口尺寸，并结合当前屏幕做自适应约束
        if self.config:
            dialog_width = self.config.get("app.window.dialog_settings_width", 550)
            dialog_height = self.config.get("app.window.dialog_settings_height", 700)
        else:
            dialog_width, dialog_height = 550, 700

        self.setWindowTitle("高级设置")
        screen = QApplication.primaryScreen()
        if screen:
            available = screen.availableGeometry()
            max_w = min(max(760, int(available.width() * 0.92)), available.width())
            max_h = min(max(620, int(available.height() * 0.90)), available.height())
            preferred_w = min(max_w, 1180)
            preferred_h = min(max_h, 900)
            min_w = min(900, max_w)
            min_h = min(680, max_h)
            width = max(min(min(max(dialog_width, 960), preferred_w), max_w), min_w)
            height = max(min(min(max(dialog_height, 760), preferred_h), max_h), min_h)
            self.resize(width, height)
            self.setMinimumSize(min_w, min_h)
            self.setMaximumSize(max_w, max_h)
        else:
            self.resize(980, 760)
            self.setMinimumSize(820, 620)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.setSizeGripEnabled(True)
        self._settings_bound_window_handle = None

        # v37.4.1: 应用对话框主题样式
        self._apply_dialog_theme()

        self.selected_rules = []
        self.use_enhance = use_enhance
        self.custom_keywords = custom_keywords
        self.scan_level = scan_level
        self.offset_x = offset_x
        self.offset_w = offset_w
        self.replacement_text = replacement_text if isinstance(replacement_text, str) and replacement_text else "[已脱敏]"
        self.word_replace_rules = normalize_word_replace_rules(word_replace_rules or [], self.replacement_text)
        self.default_replacement_text = "[已脱敏]"
        self.recommended_rule_names = ["身份证号", "手机号码"]

        # v37.0: 从配置读取范围和标签
        if self.config:
            offset_config = self.config.get("redaction.offset", {})
            x_range = offset_config.get("x_range", [-20, 20])
            w_range = offset_config.get("w_range", [-20, 20])
            scan_config = self.config.get("redaction.scan", {})
            available_levels = scan_config.get("available_levels", [1.5, 2.0, 3.0])
            level_labels = scan_config.get("level_labels", {
                "1.5": "标准 (1.5x)",
                "2.0": "高精 (2.0x 推荐)",
                "3.0": "超精 (3.0x 最慢)"
            })
        else:
            x_range = [-20, 20]
            w_range = [-20, 20]
            available_levels = [1.5, 2.0, 3.0]
            level_labels = {
                "1.5": "标准 (1.5x)",
                "2.0": "高精 (2.0x 推荐)",
                "3.0": "超精 (3.0x 最慢)"
            }

        self.x_range = x_range
        self.w_range = w_range
        self.available_levels = available_levels
        self.level_labels = level_labels

        outer_layout = QVBoxLayout(self)
        outer_layout.setContentsMargins(18, 16, 18, 18)
        outer_layout.setSpacing(14)

        hero = QFrame(self)
        hero.setObjectName("settingsHero")
        hero_layout = QVBoxLayout(hero)
        hero_layout.setContentsMargins(18, 16, 18, 16)
        hero_layout.setSpacing(6)
        self.settings_hero_layout = hero_layout
        hero_title = QLabel("高级设置中心")
        hero_title.setObjectName("settingsTitle")
        self.settings_hero_title = hero_title
        hero_subtitle = QLabel("这里保存长期配置。PDF 的黑 / 白、单双页、缩放等即时操作，仍然放在主工作台里。")
        hero_subtitle.setObjectName("settingsSubtitle")
        hero_subtitle.setWordWrap(True)
        self.settings_hero_subtitle = hero_subtitle
        hero_layout.addWidget(hero_title)
        hero_layout.addWidget(hero_subtitle)
        hero_tag_layout = QHBoxLayout()
        hero_tag_layout.setSpacing(8)
        self.settings_hero_tag_layout = hero_tag_layout
        self.lbl_settings_common_tag = QLabel("常用设置：通用规则 / 关键词 / Word 替换")
        self.lbl_settings_common_tag.setObjectName("settingsHeroTag")
        self.lbl_settings_advanced_tag = QLabel("高级微调：扫描 / 覆盖 / OCR 检测框")
        self.lbl_settings_advanced_tag.setObjectName("settingsHeroTag")
        hero_tag_layout.addWidget(self.lbl_settings_common_tag)
        hero_tag_layout.addWidget(self.lbl_settings_advanced_tag)
        hero_tag_layout.addStretch()
        hero_layout.addLayout(hero_tag_layout)
        outer_layout.addWidget(hero)

        body_layout = QHBoxLayout()
        body_layout.setSpacing(18)
        self.settings_body_layout = body_layout
        outer_layout.addLayout(body_layout, stretch=1)

        sidebar = QFrame(self)
        sidebar.setObjectName("settingsSidebar")
        sidebar.setFixedWidth(236)
        self.settings_sidebar = sidebar
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(16, 16, 16, 16)
        sidebar_layout.setSpacing(12)
        self.settings_sidebar_layout = sidebar_layout

        nav_hint = QLabel("设置导航")
        nav_hint.setObjectName("settingsHint")
        self.settings_nav_hint = nav_hint
        sidebar_layout.addWidget(nav_hint)
        nav_subtitle = QLabel("点击左侧即可快速跳转到对应模块。")
        nav_subtitle.setObjectName("settingsSidebarSubtle")
        nav_subtitle.setWordWrap(True)
        self.settings_nav_subtitle = nav_subtitle
        sidebar_layout.addWidget(nav_subtitle)
        sidebar_layout.addWidget(self._create_settings_divider())

        self.settings_nav = QListWidget()
        self.settings_nav.setObjectName("settingsNav")
        self.settings_nav.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        self._settings_nav_base_titles = ["1 通用规则", "2 自定义关键词", "3 扫描与微调", "4 OCR 检测框"]
        for title in self._settings_nav_base_titles:
            self.settings_nav.addItem(title)
        sidebar_layout.addWidget(self.settings_nav, stretch=1)

        sidebar_meta_card = QFrame()
        sidebar_meta_card.setObjectName("settingsSidebarMetaCard")
        self.settings_sidebar_meta_card = sidebar_meta_card
        sidebar_meta_layout = QVBoxLayout(sidebar_meta_card)
        sidebar_meta_layout.setContentsMargins(12, 12, 12, 12)
        sidebar_meta_layout.setSpacing(10)
        sidebar_note = QLabel("建议：第一次使用只改规则和关键词；OCR 微调只在识别偏移时再调整。")
        sidebar_note.setWordWrap(True)
        sidebar_note.setObjectName("settingsSidebarNote")
        self.settings_sidebar_note = sidebar_note
        sidebar_meta_layout.addWidget(sidebar_note)
        sidebar_meta_layout.addWidget(self._create_settings_divider())
        self.settings_sidebar_status = QLabel("")
        self.settings_sidebar_status.setWordWrap(True)
        self.settings_sidebar_status.setObjectName("settingsSidebarStatus")
        sidebar_meta_layout.addWidget(self.settings_sidebar_status)
        sidebar_layout.addWidget(sidebar_meta_card)
        body_layout.addWidget(sidebar)

        content_scroll = QScrollArea(self)
        content_scroll.setWidgetResizable(True)
        content_scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.content_scroll = content_scroll

        content_widget = QWidget()
        self.content_widget = content_widget
        layout = QVBoxLayout(content_widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(14)

        overview_card = QFrame()
        overview_card.setObjectName("settingsOverview")
        self.settings_overview_card = overview_card
        overview_layout = QVBoxLayout(overview_card)
        overview_layout.setContentsMargins(18, 16, 18, 16)
        overview_layout.setSpacing(12)
        self.settings_overview_layout = overview_layout
        overview_title = QLabel("当前配置概览")
        overview_title.setObjectName("settingsOverviewTitle")
        self.settings_overview_title = overview_title
        overview_text = QLabel("先看这里，就能快速知道当前默认规则、关键词、Word 替换和 OCR 微调处于什么状态。")
        overview_text.setObjectName("settingsOverviewText")
        overview_text.setWordWrap(True)
        self.settings_overview_text = overview_text
        overview_layout.addWidget(overview_title)
        overview_layout.addWidget(overview_text)

        overview_metrics_layout = QGridLayout()
        overview_metrics_layout.setHorizontalSpacing(10)
        overview_metrics_layout.setVerticalSpacing(10)
        self.settings_overview_metrics_layout = overview_metrics_layout
        self.lbl_metric_rules = None
        self.lbl_metric_keywords = None
        self.lbl_metric_word_rules = None
        self.lbl_metric_ocr = None
        self.settings_metric_cards = []
        for key, title in [
            ("rules", "通用规则"),
            ("keywords", "自定义关键词"),
            ("word_rules", "Word 规则"),
            ("ocr", "OCR 调节"),
        ]:
            metric_card, metric_value = self._create_settings_metric_card(title)
            setattr(self, f"lbl_metric_{key}", metric_value)
            self.settings_metric_cards.append(metric_card)
            overview_metrics_layout.addWidget(metric_card, 0, len(self.settings_metric_cards) - 1)
        overview_layout.addLayout(overview_metrics_layout)

        overview_actions_layout = QGridLayout()
        overview_actions_layout.setHorizontalSpacing(8)
        overview_actions_layout.setVerticalSpacing(8)
        self.settings_overview_actions_layout = overview_actions_layout
        self._settings_quick_jump_titles = [
            ("去看通用规则", "通用规则"),
            ("去看关键词", "关键词"),
            ("去看扫描微调", "扫描微调"),
            ("去看 OCR", "OCR"),
        ]
        self.settings_quick_jump_buttons = []
        for row, (title, _short_title) in enumerate(self._settings_quick_jump_titles):
            jump_btn = QPushButton(title)
            jump_btn.setObjectName("settingsQuickJumpButton")
            jump_btn.clicked.connect(lambda _checked=False, target_row=row: self.settings_nav.setCurrentRow(target_row))
            jump_btn.setMinimumHeight(32)
            self.settings_quick_jump_buttons.append(jump_btn)
            overview_actions_layout.addWidget(jump_btn, 0, row)
        overview_layout.addLayout(overview_actions_layout)
        layout.addWidget(overview_card)

        # 1. 规则
        box_rules = QFrame()
        box_rules.setObjectName("settingsSectionCard")
        v_box = QVBoxLayout(box_rules)
        v_box.setContentsMargins(16, 16, 16, 16)
        v_box.setSpacing(12)
        rules_lead = QLabel("勾选后的规则会作为默认智能识别规则。第一次使用建议至少保留身份证号和手机号。")
        rules_lead.setObjectName("settingsSectionLead")
        rules_lead.setWordWrap(True)
        self.lbl_rules_summary = QLabel("")
        self.lbl_rules_summary.setObjectName("settingsSectionSummary")
        self.lbl_rules_summary.setWordWrap(True)
        v_box.addWidget(self._create_settings_section_header("1. 通用规则", rules_lead, self.lbl_rules_summary))
        rules_actions = QHBoxLayout()
        rules_actions.setSpacing(8)
        rules_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_rules_recommended = QPushButton("恢复推荐勾选")
        btn_rules_recommended.setObjectName("settingsInlineButton")
        btn_rules_recommended.clicked.connect(self._apply_recommended_rules)
        rules_actions.addWidget(btn_rules_recommended)
        btn_rules_all = QPushButton("全部勾选")
        btn_rules_all.setObjectName("settingsInlineButton")
        btn_rules_all.clicked.connect(self._select_all_rules)
        rules_actions.addWidget(btn_rules_all)
        btn_rules_clear = QPushButton("全部清空")
        btn_rules_clear.setObjectName("settingsInlineButton")
        btn_rules_clear.clicked.connect(self._clear_all_rules)
        rules_actions.addWidget(btn_rules_clear)
        rules_actions.addStretch()
        v_box.addLayout(rules_actions)
        rules_columns = QHBoxLayout()
        rules_columns.setSpacing(24)
        rules_left_col = QVBoxLayout()
        rules_left_col.setSpacing(6)
        rules_right_col = QVBoxLayout()
        rules_right_col.setSpacing(6)
        self.checks = {}
        rule_items = list(DEFAULT_RULES.items())
        split_index = (len(rule_items) + 1) // 2
        for index, (name, pattern) in enumerate(rule_items):
            cb = QCheckBox(name)
            if current_rules and pattern in current_rules: cb.setChecked(True)
            elif not current_rules and name in ["身份证号", "手机号码"]: cb.setChecked(True)
            self.checks[name] = cb
            cb.toggled.connect(self._refresh_rule_summary)
            if index < split_index:
                rules_left_col.addWidget(cb)
            else:
                rules_right_col.addWidget(cb)
        rules_left_col.addStretch()
        rules_right_col.addStretch()
        rules_columns.addLayout(rules_left_col, stretch=1)
        rules_columns.addLayout(rules_right_col, stretch=1)
        v_box.addLayout(rules_columns)
        layout.addWidget(box_rules)

        # 2. 关键词 + 统一替换文本 + Word 替换规则入口
        box_custom = QFrame()
        box_custom.setObjectName("settingsSectionCard")
        v_custom = QVBoxLayout(box_custom)
        v_custom.setContentsMargins(16, 16, 16, 16)
        v_custom.setSpacing(12)
        custom_lead = QLabel("这里适合录入姓名、单位、案号等固定内容；Word 替换规则适合更精细的定向替换。")
        custom_lead.setObjectName("settingsSectionLead")
        custom_lead.setWordWrap(True)
        self.lbl_custom_summary = QLabel("")
        self.lbl_custom_summary.setObjectName("settingsSectionSummary")
        self.lbl_custom_summary.setWordWrap(True)
        v_custom.addWidget(self._create_settings_section_header("2. 自定义关键词", custom_lead, self.lbl_custom_summary))
        custom_row = QGridLayout()
        custom_row.setHorizontalSpacing(28)
        custom_row.setVerticalSpacing(14)
        self.settings_custom_row = custom_row

        left_panel_widget = QFrame()
        left_panel_widget.setObjectName("settingsFieldCard")
        self.settings_left_field_card = left_panel_widget
        left_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        left_panel = QVBoxLayout(left_panel_widget)
        left_panel.setContentsMargins(16, 16, 16, 16)
        left_panel.setSpacing(10)
        left_title = QLabel("关键词列表")
        left_title.setObjectName("settingsFieldTitle")
        left_note = QLabel("按行输入固定敏感词。适合姓名、机构、案号等直接命中的内容。")
        left_note.setObjectName("settingsFieldNote")
        left_note.setWordWrap(True)
        left_panel.addWidget(left_title)
        left_panel.addWidget(left_note)
        left_panel.addWidget(self._create_settings_divider())
        custom_actions = QHBoxLayout()
        custom_actions.setSpacing(8)
        custom_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_clear_keywords = QPushButton("清空关键词")
        btn_clear_keywords.setObjectName("settingsInlineButton")
        btn_clear_keywords.clicked.connect(self._clear_custom_keywords)
        custom_actions.addWidget(btn_clear_keywords)
        custom_actions.addStretch()
        left_panel.addLayout(custom_actions)
        self.txt_custom = QTextEdit()
        self.txt_custom.setPlaceholderText("例如：法院 张三 (支持多行)")
        self.txt_custom.setPlainText(custom_keywords)
        self.txt_custom.setMinimumHeight(120)
        self.txt_custom.setMaximumHeight(190)
        self.txt_custom.textChanged.connect(self._refresh_custom_summary)
        left_panel.addWidget(self.txt_custom)
        custom_row.addWidget(left_panel_widget, 0, 0)

        right_panel_widget = QFrame()
        right_panel_widget.setObjectName("settingsFieldCard")
        right_panel_widget.setMinimumWidth(320)
        self.settings_right_field_card = right_panel_widget
        right_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        right_panel = QVBoxLayout(right_panel_widget)
        right_panel.setContentsMargins(16, 16, 16, 16)
        right_panel.setSpacing(10)
        right_title = QLabel("Word 替换规则")
        right_title.setObjectName("settingsFieldTitle")
        right_note = QLabel("这里适合维护更精细的 exact / regex 替换规则，批量 Word 会直接复用。")
        right_note.setObjectName("settingsFieldNote")
        right_note.setWordWrap(True)
        right_panel.addWidget(right_title)
        right_panel.addWidget(right_note)
        right_panel.addWidget(self._create_settings_divider())
        replacement_row = QHBoxLayout()
        replacement_row.setSpacing(8)
        replacement_row.addWidget(self._create_settings_form_label("统一替换文本:", 102))
        self.input_replacement_text = QLineEdit(self.replacement_text)
        self.input_replacement_text.setPlaceholderText("[已脱敏]")
        self.input_replacement_text.setMinimumWidth(240)
        self.input_replacement_text.textChanged.connect(self._refresh_custom_summary)
        replacement_row.addWidget(self.input_replacement_text)
        right_panel.addLayout(replacement_row)
        replacement_actions = QHBoxLayout()
        replacement_actions.setSpacing(8)
        replacement_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_replacement = QPushButton("恢复默认替换词")
        btn_reset_replacement.setObjectName("settingsInlineButton")
        btn_reset_replacement.clicked.connect(self._reset_replacement_text)
        replacement_actions.addWidget(btn_reset_replacement)
        replacement_actions.addStretch()
        right_panel.addLayout(replacement_actions)

        right_panel.addSpacing(8)
        right_panel.addWidget(self._create_settings_form_label("Word 替换规则:", 120))
        self.btn_edit_word_rules = QPushButton("打开替换规则设置")
        self.btn_edit_word_rules.clicked.connect(self._open_word_rules_editor)
        self.btn_edit_word_rules.setMinimumHeight(34)
        right_panel.addWidget(self.btn_edit_word_rules)

        self.lbl_word_rule_count = QLabel("")
        self.lbl_word_rule_count.setObjectName("settingsFieldNote")
        self.lbl_word_rule_count.setWordWrap(True)
        right_panel.addWidget(self.lbl_word_rule_count)
        word_rules_hint = QLabel("批量 Word 替换会直接复用这里的规则结构，不需要另外维护一套。")
        word_rules_hint.setWordWrap(True)
        word_rules_hint.setObjectName("settingsFieldNote")
        right_panel.addWidget(word_rules_hint)
        right_panel.addStretch()

        custom_row.addWidget(right_panel_widget, 0, 1)
        v_custom.addLayout(custom_row)
        layout.addWidget(box_custom)

        # 3. 精度与微调
        box_enhance = QFrame()
        box_enhance.setObjectName("settingsSectionCard")
        v_enhance = QVBoxLayout(box_enhance)
        v_enhance.setContentsMargins(16, 16, 16, 16)
        v_enhance.setSpacing(12)
        precision_lead = QLabel("只有当扫描偏移、覆盖范围不理想时再调整这里；默认设置更适合大多数文档。")
        precision_lead.setObjectName("settingsSectionLead")
        precision_lead.setWordWrap(True)
        self.lbl_precision_summary = QLabel("")
        self.lbl_precision_summary.setObjectName("settingsSectionSummary")
        self.lbl_precision_summary.setWordWrap(True)
        v_enhance.addWidget(self._create_settings_section_header("3. 精度与微调", precision_lead, self.lbl_precision_summary))
        precision_actions = QHBoxLayout()
        precision_actions.setSpacing(8)
        precision_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_precision = QPushButton("恢复推荐值")
        btn_reset_precision.setObjectName("settingsInlineButton")
        btn_reset_precision.clicked.connect(self._reset_precision_defaults)
        precision_actions.addWidget(btn_reset_precision)
        precision_actions.addStretch()
        v_enhance.addLayout(precision_actions)

        precision_cards = QGridLayout()
        precision_cards.setHorizontalSpacing(16)
        precision_cards.setVerticalSpacing(14)
        self.settings_precision_cards_layout = precision_cards

        scan_card = QFrame()
        scan_card.setObjectName("settingsFieldCard")
        self.settings_scan_card = scan_card
        scan_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        scan_card_layout = QVBoxLayout(scan_card)
        scan_card_layout.setContentsMargins(16, 16, 16, 16)
        scan_card_layout.setSpacing(10)
        scan_title = QLabel("扫描模式")
        scan_title.setObjectName("settingsFieldTitle")
        scan_note = QLabel("正常情况下保持默认即可。模式越高越细，但速度也会更慢。")
        scan_note.setObjectName("settingsFieldNote")
        scan_note.setWordWrap(True)
        scan_card_layout.addWidget(scan_title)
        scan_card_layout.addWidget(scan_note)
        scan_card_layout.addWidget(self._create_settings_divider())
        h_precision = QHBoxLayout()
        h_precision.addWidget(self._create_settings_form_label("当前模式:", 86))
        self.combo_precision = QComboBox()
        # v37.0: 从配置动态添加扫描级别选项
        for level in self.available_levels:
            label = self.level_labels.get(str(level), f"{level}x")
            self.combo_precision.addItem(label, level)
        idx = self.combo_precision.findData(scan_level)
        self.combo_precision.setCurrentIndex(idx if idx >=0 else 1)
        self.combo_precision.currentIndexChanged.connect(self._refresh_precision_summary)
        h_precision.addWidget(self.combo_precision)
        scan_card_layout.addLayout(h_precision)
        scan_card_layout.addStretch()
        precision_cards.addWidget(scan_card, 0, 0)

        calibrate_card = QFrame()
        calibrate_card.setObjectName("settingsFieldCard")
        self.settings_calibrate_card = calibrate_card
        calibrate_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        calibrate_card_layout = QVBoxLayout(calibrate_card)
        calibrate_card_layout.setContentsMargins(16, 16, 16, 16)
        calibrate_card_layout.setSpacing(10)
        calibrate_title = QLabel("覆盖微调")
        calibrate_title.setObjectName("settingsFieldTitle")
        calibrate_note = QLabel("只有在遮罩范围明显偏左、偏宽时再调，默认值更适合多数 PDF。")
        calibrate_note.setObjectName("settingsFieldNote")
        calibrate_note.setWordWrap(True)
        calibrate_card_layout.addWidget(calibrate_title)
        calibrate_card_layout.addWidget(calibrate_note)
        calibrate_card_layout.addWidget(self._create_settings_divider())
        h_calibrate = QHBoxLayout()
        v_cal_1 = QVBoxLayout()
        label_offset_x = QLabel("向左修正(px):")
        label_offset_x.setObjectName("settingsFieldLabel")
        v_cal_1.addWidget(label_offset_x)
        self.spin_offset_x = QSpinBox()
        self.spin_offset_x.setRange(x_range[0], x_range[1])
        self.spin_offset_x.setValue(offset_x)
        self.spin_offset_x.valueChanged.connect(self._refresh_precision_summary)
        v_cal_1.addWidget(self.spin_offset_x)

        v_cal_2 = QVBoxLayout()
        label_offset_w = QLabel("宽度收缩(px):")
        label_offset_w.setObjectName("settingsFieldLabel")
        v_cal_2.addWidget(label_offset_w)
        self.spin_offset_w = QSpinBox()
        self.spin_offset_w.setRange(w_range[0], w_range[1])
        self.spin_offset_w.setValue(offset_w)
        self.spin_offset_w.valueChanged.connect(self._refresh_precision_summary)
        v_cal_2.addWidget(self.spin_offset_w)

        h_calibrate.addLayout(v_cal_1)
        h_calibrate.addLayout(v_cal_2)
        calibrate_card_layout.addLayout(h_calibrate)
        calibrate_tip = QLabel("提示：对扫描区域 / 嵌入图片区域生效")
        calibrate_tip.setObjectName("settingsFieldNote")
        calibrate_tip.setWordWrap(True)
        calibrate_card_layout.addWidget(calibrate_tip)

        self.cb_enhance = QCheckBox("开启图像增强 (仅针对手写体)")
        self.cb_enhance.setChecked(use_enhance)
        self.cb_enhance.toggled.connect(self._refresh_precision_summary)
        calibrate_card_layout.addWidget(self.cb_enhance)
        calibrate_card_layout.addStretch()
        precision_cards.addWidget(calibrate_card, 0, 1)

        v_enhance.addLayout(precision_cards)
        layout.addWidget(box_enhance)

        # v37.4.0: 4. OCR 检测框调节（移除引擎选择，只保留 RapidOCR）
        box_ocr = QFrame()
        box_ocr.setObjectName("settingsSectionCard")
        v_ocr = QVBoxLayout(box_ocr)
        v_ocr.setContentsMargins(16, 16, 16, 16)
        v_ocr.setSpacing(12)
        ocr_lead = QLabel("检测框只在 OCR 结果偏大或偏小时再调。负值扩大框，正值收缩框，0 表示保持原样。")
        ocr_lead.setObjectName("settingsSectionLead")
        ocr_lead.setWordWrap(True)
        self.lbl_ocr_summary = QLabel("")
        self.lbl_ocr_summary.setObjectName("settingsSectionSummary")
        self.lbl_ocr_summary.setWordWrap(True)
        v_ocr.addWidget(self._create_settings_section_header("4. OCR 检测框调节", ocr_lead, self.lbl_ocr_summary))
        ocr_actions = QHBoxLayout()
        ocr_actions.setSpacing(8)
        ocr_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_ocr = QPushButton("恢复 0%")
        btn_reset_ocr.setObjectName("settingsInlineButton")
        btn_reset_ocr.clicked.connect(self._reset_ocr_adjustment)
        ocr_actions.addWidget(btn_reset_ocr)
        ocr_actions.addStretch()
        v_ocr.addLayout(ocr_actions)

        adjust_card = QFrame()
        adjust_card.setObjectName("settingsFieldCard")
        self.settings_adjust_card = adjust_card
        adjust_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        adjust_card_layout = QVBoxLayout(adjust_card)
        adjust_card_layout.setContentsMargins(16, 16, 16, 16)
        adjust_card_layout.setSpacing(10)
        adjust_title = QLabel("检测框调节")
        adjust_title.setObjectName("settingsFieldTitle")
        adjust_note = QLabel("当 OCR 框偏大或偏小时再调。负值扩大，正值收缩。")
        adjust_note.setObjectName("settingsFieldNote")
        adjust_note.setWordWrap(True)
        adjust_card_layout.addWidget(adjust_title)
        adjust_card_layout.addWidget(adjust_note)
        adjust_card_layout.addWidget(self._create_settings_divider())

        # v37.3.5: 检测框大小调节（支持负值扩大、正值收缩）
        h_adjust = QHBoxLayout()
        h_adjust.addWidget(self._create_settings_form_label("检测框调节:", 96))

        # 读取当前配置值（新配置名）
        adjust_ratio = config.get("ocr.box_adjust_ratio", 0.0) if config else 0.0

        self.slider_adjust = QSlider(Qt.Orientation.Horizontal)
        self.slider_adjust.setRange(-30, 50)  # -30% 到 +50%
        self.slider_adjust.setValue(int(adjust_ratio * 100))
        self.slider_adjust.valueChanged.connect(self._on_adjust_changed)
        h_adjust.addWidget(self.slider_adjust)

        self.label_adjust_value = QLabel(f"{int(adjust_ratio * 100)}%")
        self.label_adjust_value.setMinimumWidth(40)
        h_adjust.addWidget(self.label_adjust_value)

        adjust_card_layout.addLayout(h_adjust)

        adjust_info = QLabel("提示：负值扩大，正值收缩，0保持原样（默认0%）")
        adjust_info.setObjectName("settingsFieldNote")
        adjust_info.setWordWrap(True)
        adjust_card_layout.addWidget(adjust_info)

        # 说明标签
        info_text = (
            "\n引擎说明：\n"
            "• RapidOCR：默认 OCR 引擎，速度快，适合大文档批量处理\n"
        )

        info = QLabel(info_text)
        info.setObjectName("settingsFieldNote")
        info.setWordWrap(True)
        adjust_card_layout.addWidget(info)
        v_ocr.addWidget(adjust_card)
        layout.addWidget(box_ocr)

        layout.addStretch(1)
        content_scroll.setWidget(content_widget)
        body_layout.addWidget(content_scroll, stretch=1)

        footer = QFrame(self)
        footer.setObjectName("settingsFooter")
        footer_layout = QHBoxLayout(footer)
        footer_layout.setContentsMargins(18, 14, 18, 14)
        footer_layout.setSpacing(12)
        self.settings_footer_layout = footer_layout

        footer_note = QLabel("修改完成后点击“保存设置”，取消不会影响当前已生效配置。")
        footer_note.setObjectName("settingsFooterNote")
        footer_note.setWordWrap(True)
        self.settings_footer_note = footer_note
        footer_layout.addWidget(footer_note, stretch=1)

        footer_actions = QHBoxLayout()
        footer_actions.setSpacing(10)
        btn_cancel = QPushButton("取消")
        btn_cancel.setObjectName("settingsSecondaryButton")
        btn_cancel.clicked.connect(self.reject)
        btn_cancel.setMinimumHeight(40)
        btn_cancel.setMinimumWidth(116)
        self.btn_settings_cancel = btn_cancel
        footer_actions.addWidget(btn_cancel)

        btn_ok = QPushButton("保存设置")
        btn_ok.setObjectName("settingsPrimaryButton")
        btn_ok.clicked.connect(self.save_settings)
        btn_ok.setMinimumHeight(40)
        btn_ok.setMinimumWidth(164)
        self.btn_settings_save = btn_ok
        footer_actions.addWidget(btn_ok)
        footer_layout.addLayout(footer_actions)
        outer_layout.addWidget(footer)
        self._settings_sections = [box_rules, box_custom, box_enhance, box_ocr]
        self._settings_nav_syncing = False
        self.settings_nav.currentRowChanged.connect(self._scroll_to_settings_section)
        self.content_scroll.verticalScrollBar().valueChanged.connect(self._sync_settings_nav_from_scroll)
        self.settings_nav.setCurrentRow(0)
        self._refresh_word_rule_summary()
        self._refresh_rule_summary()
        self._refresh_precision_summary()
        self._refresh_ocr_summary()
        self._refresh_settings_layout_density()

    def _on_adjust_changed(self, value):
        """v37.3.5: 检测框调节滑块值变化回调"""
        self.label_adjust_value.setText(f"{value}%")
        self._refresh_ocr_summary()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._refresh_settings_layout_density()

    def showEvent(self, event):
        super().showEvent(event)
        self._bind_settings_window_handle_signals()
        QTimer.singleShot(0, self._refresh_settings_layout_density)

    def _get_settings_display_scale_factor(self):
        """返回设置窗口当前屏幕缩放，仅在 Windows 下生效。"""
        import platform

        if platform.system() != "Windows":
            return 1.0

        screen = None
        try:
            handle = self.windowHandle()
            if handle:
                screen = handle.screen()
        except Exception:
            screen = None

        if screen is None:
            screen = QApplication.primaryScreen()
        if screen is None:
            return 1.0

        try:
            scale = screen.logicalDotsPerInch() / 96.0
        except Exception:
            scale = 1.0

        return max(1.0, min(scale, 2.0))

    def _bind_settings_window_handle_signals(self):
        """绑定设置窗口跨屏信号，保证 DPI 切换后密度立即刷新。"""
        try:
            handle = self.windowHandle()
        except Exception:
            handle = None

        if handle is None or handle is self._settings_bound_window_handle:
            return

        try:
            handle.screenChanged.connect(self._on_settings_screen_changed)
        except Exception:
            pass
        self._settings_bound_window_handle = handle

    def _on_settings_screen_changed(self, _screen):
        """设置窗口切换屏幕后刷新密度。"""
        QTimer.singleShot(0, self._refresh_settings_layout_density)

    def _create_settings_metric_card(self, title):
        """创建设置页顶部概览指标卡。"""
        card = QFrame()
        card.setObjectName("settingsMetricCard")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(4)

        title_label = QLabel(title)
        title_label.setObjectName("settingsMetricLabel")
        value_label = QLabel("--")
        value_label.setObjectName("settingsMetricValue")
        value_label.setWordWrap(True)

        layout.addWidget(title_label)
        layout.addWidget(value_label)
        layout.addStretch()
        return card, value_label

    def _rebuild_settings_grid(self, layout, widgets, columns, stretches=None):
        """按列数重排设置页网格布局。"""
        if not layout:
            return
        while layout.count():
            layout.takeAt(0)

        max_slots = max(len(widgets), 4)
        for column in range(max_slots):
            layout.setColumnStretch(column, 0)
        for row in range(max_slots):
            layout.setRowStretch(row, 0)

        columns = max(1, columns)
        for index, widget in enumerate(widgets):
            row = index // columns
            column = index % columns
            layout.addWidget(widget, row, column)

        if stretches:
            for column, stretch in enumerate(stretches):
                layout.setColumnStretch(column, stretch)
        else:
            for column in range(columns):
                layout.setColumnStretch(column, 1)

    def _refresh_settings_layout_density(self):
        """按当前窗口宽度微调设置页整体比例。"""
        width = max(self.width(), 1)
        height = max(self.height(), 1)
        scale = self._get_settings_display_scale_factor()
        density_mode = resolve_settings_density_mode(width, height, scale)
        is_tall_window = height >= 980
        is_short_window = height <= 820
        is_very_wide_window = width >= 1620
        is_ultra_wide_window = width >= 1920
        is_cinema_wide_window = width >= 2140

        if density_mode == "wide":
            sidebar_width = 264
            body_spacing = 22
            hero_margins = (22, 18, 22, 18)
            overview_margins = (22, 18, 22, 18)
            footer_margins = (22, 14, 22, 14)
            hero_tag_spacing = 10
            metric_spacing = 12
            action_spacing = 10
            custom_spacing = 24
            precision_spacing = 18
            jump_min_width = 134
            quick_jump_compact = False
            overview_metric_columns = 4
            overview_action_columns = 4
            custom_columns = 2
            precision_columns = 2
            custom_stretches = (3, 2)
            precision_stretches = (1, 1)
            right_panel_max_width = 680
            right_panel_min_width = 430
            left_card_min_height = 348
            right_card_min_height = 348
            precision_card_min_height = 238
            adjust_card_min_height = 234
            cancel_min_width = 122
            save_min_width = 172
        elif density_mode == "roomy":
            sidebar_width = 248
            body_spacing = 18
            hero_margins = (18, 16, 18, 16)
            overview_margins = (18, 16, 18, 16)
            footer_margins = (18, 14, 18, 14)
            hero_tag_spacing = 8
            metric_spacing = 10
            action_spacing = 8
            custom_spacing = 22
            precision_spacing = 16
            jump_min_width = 124
            quick_jump_compact = False
            overview_metric_columns = 4
            overview_action_columns = 4
            custom_columns = 2
            precision_columns = 2
            custom_stretches = (3, 2)
            precision_stretches = (1, 1)
            right_panel_max_width = 620
            right_panel_min_width = 392
            left_card_min_height = 330
            right_card_min_height = 330
            precision_card_min_height = 224
            adjust_card_min_height = 220
            cancel_min_width = 116
            save_min_width = 164
        elif density_mode == "compact":
            sidebar_width = 232
            body_spacing = 16
            hero_margins = (16, 15, 16, 15)
            overview_margins = (16, 15, 16, 15)
            footer_margins = (16, 12, 16, 12)
            hero_tag_spacing = 8
            metric_spacing = 8
            action_spacing = 8
            custom_spacing = 18
            precision_spacing = 14
            jump_min_width = 114
            quick_jump_compact = False
            overview_metric_columns = 2
            overview_action_columns = 2
            custom_columns = 2
            precision_columns = 1
            custom_stretches = (7, 5)
            precision_stretches = (1,)
            right_panel_max_width = 540
            right_panel_min_width = 344
            left_card_min_height = 302
            right_card_min_height = 302
            precision_card_min_height = 208
            adjust_card_min_height = 204
            cancel_min_width = 112
            save_min_width = 156
        else:
            sidebar_width = 220
            body_spacing = 14
            hero_margins = (14, 14, 14, 14)
            overview_margins = (14, 14, 14, 14)
            footer_margins = (14, 12, 14, 12)
            hero_tag_spacing = 6
            metric_spacing = 8
            action_spacing = 6
            custom_spacing = 14
            precision_spacing = 12
            jump_min_width = 102
            quick_jump_compact = True
            overview_metric_columns = 2
            overview_action_columns = 2
            custom_columns = 1
            precision_columns = 1
            custom_stretches = (1,)
            precision_stretches = (1,)
            right_panel_max_width = 9999
            right_panel_min_width = 0
            left_card_min_height = 270
            right_card_min_height = 270
            precision_card_min_height = 194
            adjust_card_min_height = 194
            cancel_min_width = 108
            save_min_width = 148

        if scale >= 1.5:
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 6
            cancel_min_width += 8
            save_min_width += 10
            left_card_min_height += 14
            right_card_min_height += 14
            precision_card_min_height += 10
            adjust_card_min_height += 10
        elif scale >= 1.25:
            jump_min_width += 4
            cancel_min_width += 6
            save_min_width += 8
            left_card_min_height += 8
            right_card_min_height += 8
            precision_card_min_height += 6
            adjust_card_min_height += 6

        if is_tall_window:
            left_card_min_height += 18
            right_card_min_height += 18
            precision_card_min_height += 10
            adjust_card_min_height += 10
            hero_margins = (
                hero_margins[0] + 2,
                hero_margins[1] + 2,
                hero_margins[2] + 2,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 2,
                overview_margins[1] + 2,
                overview_margins[2] + 2,
                overview_margins[3] + 2,
            )
        elif is_short_window:
            body_spacing = max(12, body_spacing - 2)
            hero_tag_spacing = max(6, hero_tag_spacing - 1)
            metric_spacing = max(8, metric_spacing - 2)
            action_spacing = max(6, action_spacing - 2)
            custom_spacing = max(14, custom_spacing - 2)
            precision_spacing = max(12, precision_spacing - 2)
            left_card_min_height = max(270, left_card_min_height - 18)
            right_card_min_height = max(270, right_card_min_height - 18)
            precision_card_min_height = max(194, precision_card_min_height - 10)
            adjust_card_min_height = max(194, adjust_card_min_height - 10)
            footer_margins = (
                footer_margins[0],
                max(10, footer_margins[1] - 2),
                footer_margins[2],
                max(10, footer_margins[3] - 2),
            )

        if is_very_wide_window:
            sidebar_width += 10
            body_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            right_panel_max_width += 60 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 24 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10
            precision_card_min_height += 8
            adjust_card_min_height += 8
            hero_margins = (
                hero_margins[0] + 2,
                hero_margins[1] + 2,
                hero_margins[2] + 2,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 2,
                overview_margins[1] + 2,
                overview_margins[2] + 2,
                overview_margins[3] + 2,
            )

        if is_ultra_wide_window:
            sidebar_width += 8
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 6
            right_panel_max_width += 80 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 28 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10

        if is_cinema_wide_window:
            sidebar_width += 10
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 8
            right_panel_max_width += 120 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 36 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10
            precision_card_min_height += 8
            adjust_card_min_height += 8
            hero_margins = (
                hero_margins[0] + 4,
                hero_margins[1] + 2,
                hero_margins[2] + 4,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 4,
                overview_margins[1] + 2,
                overview_margins[2] + 4,
                overview_margins[3] + 2,
            )

        quick_jump_height = 34 if scale >= 1.25 else 32
        footer_button_height = 42 if scale >= 1.25 else 40

        if hasattr(self, "settings_sidebar"):
            self.settings_sidebar.setFixedWidth(sidebar_width)
        if hasattr(self, "settings_body_layout"):
            self.settings_body_layout.setSpacing(body_spacing)
        if hasattr(self, "settings_hero_layout"):
            self.settings_hero_layout.setContentsMargins(*hero_margins)
        if hasattr(self, "settings_hero_tag_layout"):
            self.settings_hero_tag_layout.setSpacing(hero_tag_spacing)
        if hasattr(self, "settings_overview_layout"):
            self.settings_overview_layout.setContentsMargins(*overview_margins)
        if hasattr(self, "settings_overview_metrics_layout"):
            self.settings_overview_metrics_layout.setHorizontalSpacing(metric_spacing)
            self.settings_overview_metrics_layout.setVerticalSpacing(metric_spacing)
            self._rebuild_settings_grid(
                self.settings_overview_metrics_layout,
                self.settings_metric_cards,
                overview_metric_columns,
            )
        if hasattr(self, "settings_overview_actions_layout"):
            self.settings_overview_actions_layout.setHorizontalSpacing(action_spacing)
            self.settings_overview_actions_layout.setVerticalSpacing(action_spacing)
            self._rebuild_settings_grid(
                self.settings_overview_actions_layout,
                self.settings_quick_jump_buttons,
                overview_action_columns,
            )
        if hasattr(self, "settings_footer_layout"):
            self.settings_footer_layout.setContentsMargins(*footer_margins)
        for index, button in enumerate(getattr(self, "settings_quick_jump_buttons", [])):
            button.setMinimumWidth(jump_min_width)
            button.setMinimumHeight(quick_jump_height)
            button.setMaximumHeight(quick_jump_height)
            if hasattr(self, "_settings_quick_jump_titles") and index < len(self._settings_quick_jump_titles):
                full_title, short_title = self._settings_quick_jump_titles[index]
                button.setText(short_title if quick_jump_compact else full_title)
        if hasattr(self, "btn_settings_cancel"):
            self.btn_settings_cancel.setMinimumWidth(cancel_min_width)
            self.btn_settings_cancel.setMinimumHeight(footer_button_height)
            self.btn_settings_cancel.setMaximumHeight(footer_button_height)
        if hasattr(self, "btn_settings_save"):
            self.btn_settings_save.setMinimumWidth(save_min_width)
            self.btn_settings_save.setMinimumHeight(footer_button_height)
            self.btn_settings_save.setMaximumHeight(footer_button_height)
        if hasattr(self, "settings_overview_text"):
            overview_text_width = 920 if density_mode == "wide" else 860 if density_mode == "roomy" else 760 if density_mode == "compact" else 680
            if is_very_wide_window:
                overview_text_width += 80
            if is_ultra_wide_window:
                overview_text_width += 80
            if is_cinema_wide_window:
                overview_text_width += 120
            self.settings_overview_text.setMaximumWidth(overview_text_width)
        if hasattr(self, "settings_hero_subtitle"):
            hero_subtitle_width = 1000 if density_mode == "wide" else 930 if density_mode == "roomy" else 840 if density_mode == "compact" else 720
            if is_very_wide_window:
                hero_subtitle_width += 100
            if is_ultra_wide_window:
                hero_subtitle_width += 100
            if is_cinema_wide_window:
                hero_subtitle_width += 140
            self.settings_hero_subtitle.setMaximumWidth(hero_subtitle_width)
        if hasattr(self, "settings_custom_row"):
            self.settings_custom_row.setHorizontalSpacing(custom_spacing)
            self.settings_custom_row.setVerticalSpacing(max(12, custom_spacing - 8))
            self._rebuild_settings_grid(
                self.settings_custom_row,
                [self.settings_left_field_card, self.settings_right_field_card],
                custom_columns,
                custom_stretches,
            )
        if hasattr(self, "settings_precision_cards_layout"):
            self.settings_precision_cards_layout.setHorizontalSpacing(precision_spacing)
            self.settings_precision_cards_layout.setVerticalSpacing(max(12, precision_spacing - 2))
            self._rebuild_settings_grid(
                self.settings_precision_cards_layout,
                [self.settings_scan_card, self.settings_calibrate_card],
                precision_columns,
                precision_stretches,
            )
        if hasattr(self, "settings_right_field_card"):
            self.settings_right_field_card.setMaximumWidth(right_panel_max_width)
            self.settings_right_field_card.setMinimumWidth(right_panel_min_width)
            self.settings_right_field_card.setMinimumHeight(right_card_min_height)
        if hasattr(self, "settings_left_field_card"):
            self.settings_left_field_card.setMinimumHeight(left_card_min_height)
        if hasattr(self, "settings_scan_card"):
            self.settings_scan_card.setMinimumHeight(precision_card_min_height)
        if hasattr(self, "settings_calibrate_card"):
            self.settings_calibrate_card.setMinimumHeight(precision_card_min_height)
        if hasattr(self, "settings_adjust_card"):
            self.settings_adjust_card.setMinimumHeight(adjust_card_min_height)
        self._apply_settings_density_styles(density_mode, scale, is_short_window)

    def _create_settings_section_title(self, title):
        """创建设置模块卡片内标题。"""
        title_label = QLabel(title)
        title_label.setObjectName("settingsSectionTitle")
        return title_label

    def _create_settings_section_header(self, title, lead_label, summary_label):
        """创建设置模块卡片内部统一头部。"""
        header = QFrame()
        header.setObjectName("settingsSectionHeader")
        layout = QVBoxLayout(header)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(10)
        layout.addWidget(self._create_settings_section_title(title))
        layout.addWidget(lead_label)
        layout.addWidget(summary_label)
        return header

    def _create_settings_action_hint(self, text):
        """创建设置分区内的轻量操作提示。"""
        hint = QLabel(text)
        hint.setObjectName("settingsActionHint")
        return hint

    def _create_settings_form_label(self, text, width=96):
        """创建设置子卡中的统一表单标签。"""
        label = QLabel(text)
        label.setObjectName("settingsFieldLabel")
        label.setFixedWidth(width)
        return label

    def _create_settings_divider(self):
        """创建设置子卡内的轻量分隔线。"""
        divider = QFrame()
        divider.setObjectName("settingsFieldDivider")
        divider.setMinimumHeight(1)
        divider.setMaximumHeight(1)
        return divider

    def _apply_settings_density_styles(self, density_mode, scale, is_short_window):
        """按当前设置页密度统一标题、说明、导航与按钮字级。"""
        theme = Theme.LIGHT
        if density_mode == "wide":
            hero_title_size = 21
            hero_subtitle_size = 13
            hero_tag_size = 12
            overview_title_size = 17
            overview_text_size = 12
            section_title_size = 16
            lead_size = 12
            summary_size = 12
            field_title_size = 13
            field_label_size = 12
            field_note_size = 12
            hint_size = 11
            nav_font_size = 12
            nav_padding_v = 12
            nav_padding_h = 14
            nav_margin_v = 3
            metric_label_size = 11
            metric_value_size = 17
            quick_jump_font_size = 12
            footer_note_size = 12
        elif density_mode == "roomy":
            hero_title_size = 20
            hero_subtitle_size = 12
            hero_tag_size = 11
            overview_title_size = 16
            overview_text_size = 11
            section_title_size = 15
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 11
            nav_font_size = 12
            nav_padding_v = 11
            nav_padding_h = 13
            nav_margin_v = 2
            metric_label_size = 11
            metric_value_size = 16
            quick_jump_font_size = 12
            footer_note_size = 11
        elif density_mode == "compact":
            hero_title_size = 19
            hero_subtitle_size = 12
            hero_tag_size = 11
            overview_title_size = 15
            overview_text_size = 11
            section_title_size = 15
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 10
            nav_font_size = 11
            nav_padding_v = 10
            nav_padding_h = 12
            nav_margin_v = 2
            metric_label_size = 11
            metric_value_size = 15
            quick_jump_font_size = 11
            footer_note_size = 11
        else:
            hero_title_size = 18
            hero_subtitle_size = 11
            hero_tag_size = 10
            overview_title_size = 15
            overview_text_size = 11
            section_title_size = 14
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 10
            nav_font_size = 11
            nav_padding_v = 10
            nav_padding_h = 11
            nav_margin_v = 2
            metric_label_size = 10
            metric_value_size = 15
            quick_jump_font_size = 11
            footer_note_size = 11

        if scale >= 1.5:
            hero_title_size += 1
            hero_subtitle_size += 1
            overview_title_size += 1
            section_title_size += 1
            lead_size += 1
            summary_size += 1
            field_title_size += 1
            field_label_size += 1
            field_note_size += 1
            nav_font_size += 1
            metric_value_size += 1
            quick_jump_font_size += 1
            footer_note_size += 1
            nav_padding_v += 1
            nav_padding_h += 1
        elif scale >= 1.25:
            hero_subtitle_size += 1
            overview_text_size += 1
            lead_size += 1
            summary_size += 1
            field_note_size += 1
            nav_padding_h += 1

        if is_short_window:
            hero_title_size = max(18, hero_title_size - 1)
            hero_subtitle_size = max(11, hero_subtitle_size - 1)
            section_title_size = max(14, section_title_size - 1)
            lead_size = max(11, lead_size - 1)
            summary_size = max(11, summary_size - 1)
            nav_padding_v = max(9, nav_padding_v - 1)
            nav_margin_v = 1

        if hasattr(self, "settings_hero_title"):
            self.settings_hero_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {hero_title_size}px; font-weight: 700;"
            )
        if hasattr(self, "settings_hero_subtitle"):
            self.settings_hero_subtitle.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hero_subtitle_size}px; line-height: 1.6;"
            )
        for label in [getattr(self, "lbl_settings_common_tag", None), getattr(self, "lbl_settings_advanced_tag", None)]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['primary']}; background-color: #E9F1FB; border: 1px solid {theme['border']}; "
                    f"border-radius: 10px; padding: 6px 10px; font-size: {hero_tag_size}px; font-weight: 700;"
                )
        if hasattr(self, "settings_overview_title"):
            self.settings_overview_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {overview_title_size}px; font-weight: 700;"
            )
        if hasattr(self, "settings_overview_text"):
            self.settings_overview_text.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {overview_text_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_nav_hint"):
            self.settings_nav_hint.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hint_size}px; font-weight: 700; letter-spacing: 0.05em;"
            )
        if hasattr(self, "settings_nav_subtitle"):
            self.settings_nav_subtitle.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_sidebar_note"):
            self.settings_sidebar_note.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_sidebar_status"):
            self.settings_sidebar_status.setStyleSheet(
                f"color: {theme['text_secondary']}; background-color: #F8FBFE; border: 1px solid {theme['border']}; "
                f"border-radius: 10px; padding: 8px 10px; font-size: {field_note_size}px; line-height: 1.6; font-weight: 600;"
            )
        if hasattr(self, "settings_footer_note"):
            self.settings_footer_note.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {footer_note_size}px; line-height: 1.6; padding-right: 8px;"
            )
        if hasattr(self, "settings_nav"):
            self.settings_nav.setStyleSheet(
                f"""
                QListWidget#settingsNav {{
                    background-color: transparent;
                    border: none;
                    outline: none;
                    padding: 0;
                }}
                QListWidget#settingsNav::item {{
                    background-color: transparent;
                    border: 1px solid transparent;
                    border-radius: 10px;
                    padding: {nav_padding_v}px {nav_padding_h}px;
                    margin: {nav_margin_v}px 0;
                    color: {theme["text"]};
                    font-size: {nav_font_size}px;
                    font-weight: 600;
                }}
                QListWidget#settingsNav::item:selected {{
                    background-color: {theme["hover"]};
                    border-color: {theme["border"]};
                    color: {theme["primary"]};
                }}
                QListWidget#settingsNav::item:hover {{
                    background-color: {theme["hover"]};
                }}
                """
            )
        for label in self.findChildren(QLabel, "settingsMetricLabel"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {metric_label_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsMetricValue"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {metric_value_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsSectionTitle"):
            label.setStyleSheet(
                f"color: {theme['primary']}; font-size: {section_title_size}px; font-weight: 700; background-color: transparent;"
            )
        for label in self.findChildren(QLabel, "settingsSectionLead"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {lead_size}px; line-height: 1.7;"
            )
        for label in self.findChildren(QLabel, "settingsSectionSummary"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; background-color: #F8FBFE; border: 1px solid {theme['border']}; "
                f"border-radius: 10px; padding: 8px 12px; font-size: {summary_size}px; font-weight: 600; line-height: 1.6;"
            )
        for label in self.findChildren(QLabel, "settingsFieldTitle"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {field_title_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsFieldLabel"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {field_label_size}px; font-weight: 600;"
            )
        for label in self.findChildren(QLabel, "settingsFieldNote"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        for label in self.findChildren(QLabel, "settingsActionHint"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hint_size}px; font-weight: 700; letter-spacing: 0.04em;"
            )
        for button in getattr(self, "settings_quick_jump_buttons", []):
            button.setStyleSheet(
                f"""
                QPushButton {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 8px 14px;
                    font-size: {quick_jump_font_size}px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background-color: {theme["hover"]};
                    border-color: {theme["primary"]};
                    color: {theme["primary"]};
                }}
                """
            )

    def _set_rule_checkbox_states(self, enabled_names):
        """批量更新规则勾选状态，避免多次触发联动刷新。"""
        enabled_names = set(enabled_names)
        for checkbox in self.checks.values():
            checkbox.blockSignals(True)
        try:
            for name, checkbox in self.checks.items():
                checkbox.setChecked(name in enabled_names)
        finally:
            for checkbox in self.checks.values():
                checkbox.blockSignals(False)
        self._refresh_rule_summary()

    def _apply_recommended_rules(self):
        """恢复首次使用推荐的通用规则勾选。"""
        self._set_rule_checkbox_states(self.recommended_rule_names)

    def _select_all_rules(self):
        """勾选全部通用规则。"""
        self._set_rule_checkbox_states(self.checks.keys())

    def _clear_all_rules(self):
        """清空所有通用规则勾选。"""
        self._set_rule_checkbox_states([])

    def _clear_custom_keywords(self):
        """清空自定义关键词文本框。"""
        self.txt_custom.clear()

    def _reset_replacement_text(self):
        """恢复默认统一替换文本。"""
        self.input_replacement_text.setText(self.default_replacement_text)

    def _reset_precision_defaults(self):
        """恢复推荐的扫描与微调设置。"""
        default_index = self.combo_precision.findData(2.0)
        if default_index >= 0:
            self.combo_precision.setCurrentIndex(default_index)
        self.spin_offset_x.setValue(0)
        self.spin_offset_w.setValue(0)
        self.cb_enhance.setChecked(False)
        self._refresh_precision_summary()

    def _reset_ocr_adjustment(self):
        """恢复 OCR 检测框默认调节值。"""
        self.slider_adjust.setValue(0)
        self._refresh_ocr_summary()

    def _scroll_to_settings_section(self, row):
        """左侧导航切换到对应设置区块。"""
        if getattr(self, "_settings_nav_syncing", False):
            return
        if row < 0 or row >= len(getattr(self, "_settings_sections", [])):
            return
        target = self._settings_sections[row]
        if not target or not hasattr(self, "content_scroll"):
            return
        self.content_scroll.verticalScrollBar().setValue(max(0, target.pos().y() - 8))

    def _sync_settings_nav_from_scroll(self, value):
        """滚动右侧内容时，同步高亮左侧设置导航。"""
        sections = getattr(self, "_settings_sections", [])
        if not sections or not hasattr(self, "settings_nav"):
            return

        active_row = 0
        threshold = value + 24
        for index, section in enumerate(sections):
            if section and section.pos().y() <= threshold:
                active_row = index
            else:
                break

        if self.settings_nav.currentRow() == active_row:
            return

        self._settings_nav_syncing = True
        self.settings_nav.blockSignals(True)
        self.settings_nav.setCurrentRow(active_row)
        self.settings_nav.blockSignals(False)
        self._settings_nav_syncing = False

    def _refresh_settings_overview(self):
        """刷新设置页顶部配置概览。"""
        if not all(hasattr(self, attr) for attr in [
            "lbl_metric_rules", "lbl_metric_keywords", "lbl_metric_word_rules", "lbl_metric_ocr"
        ]):
            return

        enabled_rules = len([name for name, cb in self.checks.items() if cb.isChecked()]) if hasattr(self, "checks") else 0
        total_rules = len(self.checks) if hasattr(self, "checks") else 0
        keywords = [line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()] if hasattr(self, "txt_custom") else []
        enabled_word_rules = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")]) if hasattr(self, "word_replace_rules") else 0
        total_word_rules = len(self.word_replace_rules) if hasattr(self, "word_replace_rules") else 0
        adjust_value = self.slider_adjust.value() if hasattr(self, "slider_adjust") else 0
        scan_label = self.combo_precision.currentText().strip() if hasattr(self, "combo_precision") else "-"
        precision_is_default = True
        if hasattr(self, "combo_precision") and hasattr(self, "spin_offset_x") and hasattr(self, "spin_offset_w") and hasattr(self, "cb_enhance"):
            precision_is_default = (
                self.combo_precision.currentData() == 2.0
                and self.spin_offset_x.value() == 0
                and self.spin_offset_w.value() == 0
                and not self.cb_enhance.isChecked()
            )

        self.lbl_metric_rules.setText(f"{enabled_rules} / {total_rules} 已启用")
        self.lbl_metric_keywords.setText(f"{len(keywords)} 条关键词")
        self.lbl_metric_word_rules.setText(f"{enabled_word_rules} / {total_word_rules} 条规则")
        self.lbl_metric_ocr.setText(f"{adjust_value}% · {scan_label}")
        if hasattr(self, "lbl_settings_common_tag") and hasattr(self, "lbl_settings_advanced_tag"):
            common_tag, advanced_tag = build_settings_hero_tags(
                enabled_rules,
                len(keywords),
                enabled_word_rules,
                precision_is_default,
                adjust_value,
                scan_label,
            )
            self.lbl_settings_common_tag.setText(common_tag)
            self.lbl_settings_common_tag.setToolTip(common_tag)
            self.lbl_settings_advanced_tag.setText(advanced_tag)
            self.lbl_settings_advanced_tag.setToolTip(advanced_tag)
        self._refresh_settings_sidebar()

    def _refresh_settings_sidebar(self):
        """刷新左侧导航和侧栏摘要。"""
        if not hasattr(self, "settings_nav"):
            return

        enabled_rules = len([name for name, cb in self.checks.items() if cb.isChecked()]) if hasattr(self, "checks") else 0
        keyword_count = len([line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()]) if hasattr(self, "txt_custom") else 0
        precision_is_default = True
        if hasattr(self, "combo_precision") and hasattr(self, "spin_offset_x") and hasattr(self, "spin_offset_w") and hasattr(self, "cb_enhance"):
            precision_is_default = (
                self.combo_precision.currentData() == 2.0
                and self.spin_offset_x.value() == 0
                and self.spin_offset_w.value() == 0
                and not self.cb_enhance.isChecked()
            )
        adjust_value = self.slider_adjust.value() if hasattr(self, "slider_adjust") else 0

        nav_labels = build_settings_nav_labels(enabled_rules, keyword_count, precision_is_default, adjust_value)
        for index, text in enumerate(nav_labels):
            item = self.settings_nav.item(index)
            if item and item.text() != text:
                item.setText(text)

        if hasattr(self, "settings_sidebar_status"):
            advanced_text = "扫描保持默认" if precision_is_default else "扫描参数已微调"
            self.settings_sidebar_status.setText(
                f"常用区：规则 {enabled_rules} 项、关键词 {keyword_count} 条。\n"
                f"高级区：{advanced_text} · OCR {format_signed_percent(adjust_value)}。"
            )

    def _refresh_word_rule_summary(self):
        enabled_count = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")])
        total_count = len(self.word_replace_rules)
        self.lbl_word_rule_count.setText(
            f"当前规则：{enabled_count} 条启用 / {total_count} 条总计\n"
            "点击“打开替换规则设置”可进入原有替换规则弹窗。"
        )
        self._refresh_custom_summary()

    def _refresh_rule_summary(self):
        enabled_names = [name for name, cb in self.checks.items() if cb.isChecked()]
        enabled_count = len(enabled_names)
        preview = "、".join(enabled_names[:3]) if enabled_names else "当前未启用任何通用规则"
        if enabled_count > 3:
            preview = f"{preview} 等 {enabled_count} 项"
        self.lbl_rules_summary.setText(f"当前启用：{preview}")
        self._refresh_settings_overview()

    def _refresh_custom_summary(self):
        keyword_lines = [line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()]
        keyword_count = len(keyword_lines)
        replacement_preview = self.input_replacement_text.text().strip() or "[已脱敏]"
        enabled_rule_count = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")])
        self.lbl_custom_summary.setText(
            f"自定义关键词 {keyword_count} 条 · 统一替换文本：{replacement_preview} · Word 规则：{enabled_rule_count} 条启用"
        )
        self._refresh_settings_overview()

    def _refresh_precision_summary(self):
        current_label = self.combo_precision.currentText().strip()
        enhance_text = "已开启图像增强" if self.cb_enhance.isChecked() else "图像增强关闭"
        self.lbl_precision_summary.setText(
            f"当前扫描模式：{current_label} · 向左修正 {self.spin_offset_x.value()} px · 宽度收缩 {self.spin_offset_w.value()} px · {enhance_text}"
        )
        self._refresh_settings_overview()

    def _refresh_ocr_summary(self):
        adjust_value = self.slider_adjust.value()
        if adjust_value < 0:
            trend = "扩大检测框"
        elif adjust_value > 0:
            trend = "收缩检测框"
        else:
            trend = "保持原始检测框"
        self.lbl_ocr_summary.setText(f"当前检测框调节：{adjust_value}% · {trend}")
        self._refresh_settings_overview()

    def _open_word_rules_editor(self):
        default_text = self.input_replacement_text.text().strip() or "[已脱敏]"
        dlg = WordReplaceRulesDialog(
            self,
            rules=self.word_replace_rules,
            default_replacement_text=default_text,
            title="Word 替换规则设置",
            apply_text="应用规则"
        )
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        self.word_replace_rules = dlg.rules
        self.input_replacement_text.setText(dlg.default_replacement_text)
        self._refresh_word_rule_summary()

    def save_settings(self):
        self.selected_rules = [DEFAULT_RULES[name] for name, cb in self.checks.items() if cb.isChecked()]
        # v37.5.0: 添加调试输出
        print(f"[Settings] 保存的规则: {self.selected_rules}")
        print(f"[Settings] 勾选的规则名称: {[name for name, cb in self.checks.items() if cb.isChecked()]}")
        self.use_enhance = self.cb_enhance.isChecked()
        self.custom_keywords = self.txt_custom.toPlainText().strip()
        self.scan_level = self.combo_precision.currentData()
        self.offset_x = self.spin_offset_x.value()
        self.offset_w = self.spin_offset_w.value()
        self.replacement_text = self.input_replacement_text.text().strip() or "[已脱敏]"
        self.word_replace_rules = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)

        # v37.4.0: 保存检测框调节比例
        self.box_adjust_ratio = self.slider_adjust.value() / 100.0

        # v37.0: 保存到配置文件
        if self.config:
            try:
                self.config.set("redaction.scan.default_level", self.scan_level, persist=False)
                self.config.set("redaction.offset.default_x", self.offset_x, persist=False)
                self.config.set("redaction.offset.default_w", self.offset_w, persist=False)
                # v37.4.0: 移除 OCR 引擎选择，只保留 RapidOCR
                # v37.3.5: 保存检测框调节比例（新配置名）
                self.config.set("redaction.custom_keywords", self.custom_keywords, persist=False)
                self.config.set("redaction.replacement_text", self.replacement_text, persist=False)
                self.config.set("ocr.box_adjust_ratio", self.box_adjust_ratio, persist=False)
                self.config.save()
            except Exception as e:
                print(f"[设置] 保存配置失败: {e}")

        self.accept()

    def _apply_dialog_theme(self):
        """应用对话框浅色主题样式（v37.4.1: 修复 Windows 深色模式显示问题）"""
        from theme import Theme
        theme = Theme.LIGHT

        self.setStyleSheet(f"""
            QDialog {{
                background-color: {theme["background"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QWidget {{
                background-color: {theme["background"]};
                color: {theme["text"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QFrame#settingsHero {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsOverview {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsMetricCard {{
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 12px;
            }}
            QLabel#settingsTitle {{
                color: {theme["text"]};
                font-size: 20px;
                font-weight: 700;
            }}
            QLabel#settingsOverviewTitle {{
                color: {theme["text"]};
                font-size: 16px;
                font-weight: 700;
            }}
            QLabel#settingsOverviewText {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsMetricLabel {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
            }}
            QLabel#settingsMetricValue {{
                color: {theme["text"]};
                font-size: 16px;
                font-weight: 700;
            }}
            QPushButton#settingsQuickJumpButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 14px;
                font-size: 12px;
                font-weight: 600;
            }}
            QPushButton#settingsQuickJumpButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
                color: {theme["primary"]};
            }}
            QLabel#settingsSubtitle {{
                color: {theme["text_secondary"]};
                font-size: 12px;
                line-height: 1.6;
            }}
            QLabel#settingsHeroTag {{
                color: {theme["primary"]};
                background-color: #E9F1FB;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 6px 10px;
                font-size: 11px;
                font-weight: 700;
            }}
            QFrame#settingsSidebar {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsSidebarMetaCard {{
                background-color: #FBFCFE;
                border: 1px solid {theme["border"]};
                border-radius: 14px;
            }}
            QFrame#settingsFooter {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 16px;
            }}
            QFrame#settingsFieldCard, QWidget#settingsInnerPanel {{
                background-color: #FBFCFE;
                border: 1px solid {theme["border"]};
                border-radius: 14px;
            }}
            QLabel#settingsHint {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
                letter-spacing: 0.05em;
            }}
            QLabel#settingsSidebarNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsSidebarSubtle {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsSidebarStatus {{
                color: {theme["text_secondary"]};
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 10px;
                font-size: 11px;
                line-height: 1.6;
                font-weight: 600;
            }}
            QLabel#settingsFooterNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
                padding-right: 8px;
            }}
            QLabel#settingsFieldTitle {{
                color: {theme["text"]};
                font-size: 12px;
                font-weight: 700;
            }}
            QLabel#settingsFieldLabel {{
                color: {theme["text"]};
                font-size: 11px;
                font-weight: 600;
            }}
            QLabel#settingsFieldNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QFrame#settingsFieldDivider {{
                background-color: {theme["border"]};
                border: none;
            }}
            QLabel#settingsSectionLead {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.7;
            }}
            QLabel#settingsSectionSummary {{
                color: {theme["text_secondary"]};
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: 600;
                line-height: 1.6;
            }}
            QListWidget#settingsNav {{
                background-color: transparent;
                border: none;
                outline: none;
                padding: 0;
            }}
            QListWidget#settingsNav::item {{
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 10px;
                padding: 10px 12px;
                margin: 2px 0;
                color: {theme["text"]};
                font-weight: 600;
            }}
            QListWidget#settingsNav::item:selected {{
                background-color: {theme["hover"]};
                border-color: {theme["border"]};
                color: {theme["primary"]};
            }}
            QListWidget#settingsNav::item:hover {{
                background-color: {theme["hover"]};
            }}
            QFrame#settingsSectionCard {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 16px;
            }}
            QFrame#settingsSectionHeader {{
                background-color: transparent;
                border: none;
            }}
            QLabel#settingsSectionTitle {{
                color: {theme["primary"]};
                font-size: 15px;
                font-weight: 700;
                background-color: transparent;
            }}
            QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QCheckBox {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
            }}
            QTextEdit {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 8px;
            }}
            QLineEdit {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px 8px;
                min-height: 22px;
            }}
            QComboBox {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px;
                min-width: 100px;
            }}
            QComboBox::drop-down {{
                border: none;
                width: 24px;
            }}
            QComboBox QAbstractItemView {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                selection-background-color: {theme["primary"]};
            }}
            QSpinBox {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px;
            }}
            QSlider {{
                background-color: transparent;
            }}
            QSlider::groove:horizontal {{
                border: none;
                height: 4px;
                background-color: {theme["border"]};
                border-radius: 2px;
            }}
            QSlider::handle:horizontal {{
                background-color: {theme["primary"]};
                border: none;
                width: 16px;
                height: 16px;
                margin: -6px 0;
                border-radius: 8px;
            }}
            QScrollArea {{
                background-color: transparent;
                border: none;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 10px;
                padding: 10px 20px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
            QPushButton:pressed {{
                background-color: {theme["pressed"]};
            }}
            QPushButton#settingsSecondaryButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                padding: 10px 18px;
            }}
            QPushButton#settingsSecondaryButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
            }}
            QLabel#settingsActionHint {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
            }}
            QPushButton#settingsInlineButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 9px;
                padding: 7px 12px;
                font-size: 11px;
                font-weight: 600;
            }}
            QPushButton#settingsInlineButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
                color: {theme["primary"]};
            }}
        """)


class WordReplaceRulesDialog(QDialog):
    """Word 多字段替换规则对话框（会话级规则，可导入/导出）。"""

    def __init__(self, parent=None, rules=None, default_replacement_text="[已脱敏]",
                 title="替换规则设置", apply_text="应用规则"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(780, 520)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

        self.rules = normalize_word_replace_rules(rules or [], default_replacement_text)
        self.default_replacement_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"

        main_layout = QVBoxLayout(self)

        header = QLabel("支持精确(exact)和正则(regex)两种模式。执行顺序：精确优先，其次正则；同模式按规则顺序。")
        header.setWordWrap(True)
        main_layout.addWidget(header)

        default_row = QHBoxLayout()
        default_row.addWidget(QLabel("默认替换文本:"))
        self.input_default_text = QLineEdit(self.default_replacement_text)
        self.input_default_text.setPlaceholderText("[已脱敏]")
        default_row.addWidget(self.input_default_text)
        main_layout.addLayout(default_row)

        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(["启用", "模式(exact/regex)", "查找文本", "替换文本"])
        self.table.verticalHeader().setVisible(False)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        main_layout.addWidget(self.table)

        row_btn_layout = QHBoxLayout()
        btn_add = QPushButton("新增规则")
        btn_del = QPushButton("删除规则")
        btn_up = QPushButton("上移")
        btn_down = QPushButton("下移")
        btn_import = QPushButton("导入JSON")
        btn_export = QPushButton("导出JSON")
        row_btn_layout.addWidget(btn_add)
        row_btn_layout.addWidget(btn_del)
        row_btn_layout.addWidget(btn_up)
        row_btn_layout.addWidget(btn_down)
        row_btn_layout.addStretch()
        row_btn_layout.addWidget(btn_import)
        row_btn_layout.addWidget(btn_export)
        main_layout.addLayout(row_btn_layout)

        btn_add.clicked.connect(self.add_rule_row)
        btn_del.clicked.connect(self.remove_selected_rule)
        btn_up.clicked.connect(lambda: self.move_selected_rule(-1))
        btn_down.clicked.connect(lambda: self.move_selected_rule(1))
        btn_import.clicked.connect(self.import_rules_json)
        btn_export.clicked.connect(self.export_rules_json)

        footer = QHBoxLayout()
        footer.addStretch()
        btn_cancel = QPushButton("取消")
        btn_apply = QPushButton(apply_text)
        footer.addWidget(btn_cancel)
        footer.addWidget(btn_apply)
        main_layout.addLayout(footer)

        btn_cancel.clicked.connect(self.reject)
        btn_apply.clicked.connect(self.apply_rules)

        for rule in self.rules:
            self.add_rule_row(rule)

        if self.table.rowCount() == 0:
            self.add_rule_row()

    def add_rule_row(self, rule=None):
        rule = rule or {"enabled": True, "mode": "exact", "find": "", "replace": ""}
        row = self.table.rowCount()
        self.table.insertRow(row)

        enabled_item = QTableWidgetItem()
        enabled_item.setFlags(
            Qt.ItemFlag.ItemIsEnabled |
            Qt.ItemFlag.ItemIsUserCheckable |
            Qt.ItemFlag.ItemIsSelectable
        )
        enabled_item.setCheckState(Qt.CheckState.Checked if rule.get("enabled", True) else Qt.CheckState.Unchecked)
        self.table.setItem(row, 0, enabled_item)

        mode = str(rule.get("mode", "exact")).strip().lower()
        if mode not in ("exact", "regex"):
            mode = "exact"
        self.table.setItem(row, 1, QTableWidgetItem(mode))
        self.table.setItem(row, 2, QTableWidgetItem(str(rule.get("find", ""))))
        self.table.setItem(row, 3, QTableWidgetItem(str(rule.get("replace", ""))))
        self.table.selectRow(row)

    def remove_selected_rule(self):
        row = self.table.currentRow()
        if row < 0:
            return
        self.table.removeRow(row)
        if self.table.rowCount() == 0:
            self.add_rule_row()

    def move_selected_rule(self, direction):
        current_row = self.table.currentRow()
        if current_row < 0:
            return
        target_row = current_row + direction
        if target_row < 0 or target_row >= self.table.rowCount():
            return

        current_data = []
        target_data = []
        for col in range(self.table.columnCount()):
            current_item = self.table.item(current_row, col)
            target_item = self.table.item(target_row, col)
            current_data.append(current_item.clone() if current_item else QTableWidgetItem())
            target_data.append(target_item.clone() if target_item else QTableWidgetItem())

        for col in range(self.table.columnCount()):
            self.table.setItem(current_row, col, target_data[col])
            self.table.setItem(target_row, col, current_data[col])

        self.table.selectRow(target_row)

    def _collect_rules_from_table(self, validate_regex=True):
        errors = []
        rules = []
        default_text = self.input_default_text.text().strip() or "[已脱敏]"

        for row in range(self.table.rowCount()):
            enabled_item = self.table.item(row, 0)
            mode_item = self.table.item(row, 1)
            find_item = self.table.item(row, 2)
            replace_item = self.table.item(row, 3)

            enabled = enabled_item is not None and enabled_item.checkState() == Qt.CheckState.Checked
            mode = mode_item.text().strip().lower() if mode_item else "exact"
            find_text = find_item.text().strip() if find_item else ""
            replace_text = replace_item.text() if replace_item else ""

            if mode not in ("exact", "regex"):
                errors.append(f"第 {row + 1} 行模式无效：{mode}（仅支持 exact/regex）")
                continue
            if enabled and not find_text:
                errors.append(f"第 {row + 1} 行查找文本不能为空")
                continue
            if enabled and mode == "regex" and validate_regex:
                try:
                    re.compile(find_text)
                except re.error as e:
                    errors.append(f"第 {row + 1} 行正则无效：{e}")
                    continue
            if not replace_text:
                replace_text = default_text

            rules.append({
                "enabled": enabled,
                "mode": mode,
                "find": find_text,
                "replace": replace_text
            })

        normalized = normalize_word_replace_rules(rules, default_text)
        return normalized, default_text, errors

    def apply_rules(self):
        rules, default_text, errors = self._collect_rules_from_table(validate_regex=True)
        if errors:
            QMessageBox.warning(self, "规则校验失败", "\n".join(errors))
            return

        self.rules = rules
        self.default_replacement_text = default_text
        self.accept()

    def import_rules_json(self):
        fname, _ = QFileDialog.getOpenFileName(self, "导入替换规则", "", "JSON 文件 (*.json)")
        if not fname:
            return

        try:
            with open(fname, "r", encoding="utf-8") as f:
                data = json.load(f)

            if not isinstance(data, dict):
                raise ValueError("JSON 根节点必须为对象")

            default_text = str(data.get("default_replacement_text", self.input_default_text.text().strip() or "[已脱敏]"))
            rules = data.get("rules", [])
            if not isinstance(rules, list):
                raise ValueError("rules 字段必须为数组")

            normalized = normalize_word_replace_rules(rules, default_text)

            self.input_default_text.setText(default_text)
            self.table.setRowCount(0)
            for rule in normalized:
                self.add_rule_row(rule)
            if self.table.rowCount() == 0:
                self.add_rule_row()

            QMessageBox.information(self, "导入成功", f"已导入 {len(normalized)} 条规则")
        except (OSError, IOError, ValueError, json.JSONDecodeError) as e:
            QMessageBox.critical(self, "导入失败", f"无法导入规则文件：\n{e}")

    def export_rules_json(self):
        rules, default_text, errors = self._collect_rules_from_table(validate_regex=True)
        if errors:
            QMessageBox.warning(self, "无法导出", "\n".join(errors))
            return

        fname, _ = QFileDialog.getSaveFileName(self, "导出替换规则", "word_replace_rules.json", "JSON 文件 (*.json)")
        if not fname:
            return

        if not fname.lower().endswith(".json"):
            fname = fname + ".json"

        payload = {
            "version": WORD_RULE_SCHEMA_VERSION,
            "default_replacement_text": default_text,
            "rules": rules
        }
        try:
            with open(fname, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "导出成功", f"规则已导出到：\n{fname}")
        except (OSError, IOError, ValueError) as e:
            QMessageBox.critical(self, "导出失败", f"无法导出规则文件：\n{e}")


# === 图片排序对话框 ===
class ImageListDialog(QDialog):
    """图片排序对话框 - 支持拖拽调整图片顺序"""
    def __init__(self, image_paths, parent=None):
        super().__init__(parent)

        # v37.4.1: 修复 Windows 深色模式下对话框显示问题
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        self.setWindowTitle("调整图片顺序")
        # v37.0: 从配置读取窗口尺寸
        if config:
            dialog_width = config.get("app.window.dialog_image_list_width", 600)
            dialog_height = config.get("app.window.dialog_image_list_height", 500)
        else:
            dialog_width, dialog_height = 600, 500
        self.resize(dialog_width, dialog_height)
        self.image_paths = image_paths

        layout = QVBoxLayout(self)

        # 说明标签
        info_label = QLabel("拖拽缩略图调整图片顺序，完成后点击「确认合并」")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)

        # 缩略图列表（支持拖拽）
        self.list_widget = QListWidget()
        self.list_widget.setIconSize(QSize(120, 120))
        self.list_widget.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)
        self.list_widget.setSelectionMode(QListWidget.SelectionMode.SingleSelection)

        # 添加缩略图
        for path in image_paths:
            item = QListWidgetItem(os.path.basename(path))
            item.setData(Qt.ItemDataRole.UserRole, path)
            # 生成缩略图
            try:
                pixmap = QPixmap(path).scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio)
                item.setIcon(QIcon(pixmap))
            except (OSError, IOError, ValueError) as e:
                # 如果缩略图生成失败，使用默认图标
                print(f"[ImageMergeDialog] 缩略图生成失败: {path}: {e}")
                pass
            self.list_widget.addItem(item)

        layout.addWidget(self.list_widget)

        # 按钮
        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("确认合并")
        btn_ok.clicked.connect(self.accept)
        btn_cancel = QPushButton("取消")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        # v37.4.1: 应用对话框主题样式
        self._apply_dialog_theme()

    def _apply_dialog_theme(self):
        """应用对话框浅色主题样式（v37.4.1: 修复 Windows 深色模式显示问题）"""
        from theme import Theme
        theme = Theme.LIGHT

        self.setStyleSheet(f"""
            QDialog {{
                background-color: {theme["background"]};
            }}
            QWidget {{
                background-color: {theme["background"]};
                color: {theme["text"]};
            }}
            QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QListWidget {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 8px;
            }}
            QListWidget::item {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                padding: 8px;
                border-radius: 4px;
            }}
            QListWidget::item:selected {{
                background-color: {theme["primary"]};
                color: white;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
        """)

    def get_ordered_paths(self):
        """获取排序后的图片路径"""
        paths = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            paths.append(item.data(Qt.ItemDataRole.UserRole))
        return paths

# === 配置常量 ===
# v37.0: 从配置读取，失败时使用硬编码后备
FEEDBACK_URL = config.get("app.feedback_url", "https://fcnwakmkeuz7.feishu.cn/share/base/form/shrcnEM1JEbdIKzdB400egj9lHe") if config else "https://fcnwakmkeuz7.feishu.cn/share/base/form/shrcnEM1JEbdIKzdB400egj9lHe"

# === 反馈对话框 ===
class FeedbackDialog(QDialog):
    """开发者信息与反馈对话框"""
    def __init__(self, parent=None):
        super().__init__(parent)

        # v37.4.1: 修复 Windows 深色模式下对话框显示问题
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        self.setWindowTitle("关于与反馈")
        # v37.0: 从配置读取窗口尺寸
        if config:
            dialog_width = config.get("app.window.dialog_feedback_width", 480)
            dialog_height = config.get("app.window.dialog_feedback_height", 600)
        else:
            dialog_width, dialog_height = 480, 600
        self.resize(dialog_width, dialog_height)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

        # 获取当前主题
        self.theme = Theme.LIGHT if not parent or not hasattr(parent, 'is_dark') or not parent.is_dark else Theme.DARK

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # === 标题区域 ===
        title_layout = QHBoxLayout()
        logo_label = QLabel()
        logo_label.setFixedSize(64, 64)
        logo_label.setStyleSheet(f"""
            background: {self.theme['primary']};
            border-radius: 12px;
            color: white;
            font-size: 24px;
            font-weight: bold;
        """)
        logo_label.setText("PG")
        logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_layout.addWidget(logo_label)

        title_info = QVBoxLayout()
        title_info.setSpacing(4)
        app_name = QLabel(APP_NAME)
        app_name.setStyleSheet(f"font-size: 20px; font-weight: bold; color: {self.theme['text']};")
        version_label = QLabel(f"版本 {VERSION}")
        version_label.setStyleSheet(f"font-size: 12px; color: {self.theme['text_secondary']};")
        title_info.addWidget(app_name)
        title_info.addWidget(version_label)
        title_layout.addLayout(title_info)
        title_layout.addStretch()
        layout.addLayout(title_layout)

        # === 分隔线 ===
        line1 = QFrame()
        line1.setFrameShape(QFrame.Shape.HLine)
        line1.setStyleSheet(f"background: {self.theme['border']}; max-height: 1px;")
        layout.addWidget(line1)

        # === 社交媒体账号 ===
        social_group = QGroupBox("关注开发者")
        social_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                color: {self.theme['text']};
                border: 1px solid {self.theme['border']};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 8px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 4px;
            }}
        """)
        social_layout = QVBoxLayout(social_group)

        # 微信公众号行
        wx_account = "池州汪律的Ai进化论"
        row1 = QHBoxLayout()
        wx_label = QLabel("微信公众号:")
        wx_label.setStyleSheet(f"color: {self.theme['text_secondary']};")
        row1.addWidget(wx_label)

        wx_account_label = QLabel(wx_account)
        wx_account_label.setStyleSheet(f"color: {self.theme['text']}; font-weight: 500;")
        wx_account_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        row1.addWidget(wx_account_label)
        row1.addStretch()

        wx_qr_btn = QPushButton("扫码关注")
        wx_qr_btn.setFixedSize(70, 24)
        wx_qr_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['primary']};
                border: none;
                border-radius: 4px;
                font-size: 11px;
                color: white;
            }}
            QPushButton:hover {{
                background: #0056CC;
            }}
        """)
        wx_qr_btn.clicked.connect(self._show_wx_qrcode)
        row1.addWidget(wx_qr_btn)
        social_layout.addLayout(row1)

        # 其他平台行
        row2 = QHBoxLayout()
        platforms_label = QLabel("抖音/小红书/B站（同号）:")
        platforms_label.setStyleSheet(f"color: {self.theme['text_secondary']};")
        row2.addWidget(platforms_label)

        other_account = "池州有个汪律师"
        other_label = QLabel(other_account)
        other_label.setStyleSheet(f"color: {self.theme['text']}; font-weight: 500;")
        other_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        row2.addWidget(other_label)
        row2.addStretch()

        copy_btn = QPushButton("复制")
        copy_btn.setFixedSize(50, 24)
        copy_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['hover']};
                border: none;
                border-radius: 4px;
                font-size: 11px;
                color: {self.theme['text']};
            }}
            QPushButton:hover {{
                background: {self.theme['pressed']};
            }}
        """)
        copy_btn.clicked.connect(lambda checked, a=other_account: self._copy_to_clipboard(a))
        row2.addWidget(copy_btn)
        social_layout.addLayout(row2)

        layout.addWidget(social_group)

        # === 开发者简介 ===
        dev_group = QGroupBox("开发者简介")
        dev_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                color: {self.theme['text']};
                border: 1px solid {self.theme['border']};
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 8px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 4px;
            }}
        """)
        dev_layout = QVBoxLayout(dev_group)
        dev_intro = QLabel(
            "<b>汪立</b><br><br>"
            "安徽始信律师事务所执业律师<br>"
            "全栈律师 | 前教师 | 退伍军人<br><br>"
            "📧 邮箱: <a href='mailto:491445490@qq.com'>491445490@qq.com</a>"
        )
        dev_intro.setWordWrap(True)
        dev_intro.setTextFormat(Qt.TextFormat.RichText)
        dev_intro.setOpenExternalLinks(True)
        dev_intro.setStyleSheet(f"""
            color: {self.theme['text']};
            line-height: 1.6;
            font-size: 13px;
        """)
        dev_layout.addWidget(dev_intro)
        layout.addWidget(dev_group)

        # === 操作按钮 ===
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        feedback_btn = QPushButton("📝 反馈建议")
        feedback_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['primary']};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: 500;
            }}
            QPushButton:hover {{
                background: {Theme.adjust_color(self.theme['primary'], -20)};
            }}
        """)
        feedback_btn.clicked.connect(self._open_feedback)
        btn_layout.addWidget(feedback_btn)

        # 新增：使用手册按钮
        manual_btn = QPushButton("📖 使用手册")
        manual_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['info'] if 'info' in self.theme else '#17a2b8'};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: 500;
            }}
            QPushButton:hover {{
                background: {Theme.adjust_color(self.theme['info'] if 'info' in self.theme else '#17a2b8', -20)};
            }}
        """)
        manual_btn.clicked.connect(self._open_manual)
        btn_layout.addWidget(manual_btn)

        donate_btn = QPushButton("☕ 打赏支持")
        donate_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['success']};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: 500;
            }}
            QPushButton:hover {{
                background: {Theme.adjust_color(self.theme['success'], -20)};
            }}
        """)
        donate_btn.clicked.connect(self._show_donate)
        btn_layout.addWidget(donate_btn)

        layout.addLayout(btn_layout)

        # === 分隔线 ===
        line2 = QFrame()
        line2.setFrameShape(QFrame.Shape.HLine)
        line2.setStyleSheet(f"background: {self.theme['border']}; max-height: 1px;")
        layout.addWidget(line2)

        # === 免责声明 ===
        disclaimer = QTextBrowser()
        disclaimer.setOpenExternalLinks(True)
        disclaimer.setMaximumHeight(160)
        disclaimer.setStyleSheet(f"""
            QTextBrowser {{
                background: {self.theme['surface']};
                border: 1px solid {self.theme['border']};
                border-radius: 8px;
                padding: 12px;
                color: {self.theme['text_secondary']};
                font-size: 11px;
            }}
        """)
        disclaimer.setHtml("""
            <p style="font-weight: bold; margin-bottom: 8px; color: #FF9500;">⚠️ 免责声明</p>
            <ol style="margin-left: 16px; line-height: 1.6;">
                <li>本软件免费仅供学习和个人使用，不构成任何法律建议。</li>
                <li>使用本软件进行文档脱敏处理后，用户需自行核实脱敏结果。</li>
                <li>开发者不对因使用本软件而产生的任何直接或间接损失承担责任。</li>
                <li>本软件不收集任何用户数据，所有处理均在本地完成。</li>
                <li>请勿将本软件用于任何违法用途，最终解释、修改权归汪立律师所有。</li>
            </ol>
        """)
        layout.addWidget(disclaimer)

        # === 关闭按钮 ===
        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['hover']};
                color: {self.theme['text']};
                border: none;
                border-radius: 8px;
                padding: 10px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                background: {self.theme['pressed']};
            }}
        """)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

    def _copy_to_clipboard(self, text):
        """复制文本到剪贴板"""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        # 显示复制成功提示
        if self.parent():
            QMessageBox.information(self.parent(), "复制成功", f"已复制: {text}")

    def _show_wx_qrcode(self):
        """显示微信公众号二维码对话框"""
        dialog = QDialog(self)
        dialog.setWindowTitle("关注微信公众号")
        dialog.setFixedSize(400, 480)
        dialog.setStyleSheet(f"""
            QDialog {{
                background: {self.theme['background']}
            }}
        """)

        layout = QVBoxLayout(dialog)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # 标题
        title = QLabel("扫码关注微信公众号")
        title.setStyleSheet(f"""
            color: {self.theme['text']};
            font-size: 16px;
            font-weight: bold;
            padding: 10px;
        """)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        # 公众号名称
        name_label = QLabel("池州汪律的Ai进化论")
        name_label.setStyleSheet(f"""
            color: {self.theme['primary']};
            font-size: 18px;
            font-weight: bold;
            padding: 5px;
        """)
        name_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(name_label)

        # 二维码图片
        qr_path = resource_path(os.path.join("assets", "wx_qrcode.png"))
        if os.path.exists(qr_path):
            qr_label = QLabel()
            pixmap = QPixmap(qr_path)
            if not pixmap.isNull():
                scaled_pixmap = pixmap.scaled(280, 280, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                qr_label.setPixmap(scaled_pixmap)
                qr_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                layout.addWidget(qr_label)
            else:
                qr_label.setText("二维码加载失败")
                qr_label.setStyleSheet(f"color: {self.theme['text']};")
                layout.addWidget(qr_label)
        else:
            qr_label = QLabel("请添加微信公众号二维码图片至\nassets/wx_qrcode.png")
            qr_label.setStyleSheet(f"color: {self.theme['text']};")
            qr_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(qr_label)

        # 提示文字
        hint = QLabel("微信扫一扫，关注公众号获取更多AI工具")
        hint.setStyleSheet(f"""
            color: {self.theme['text_secondary']};
            font-size: 12px;
            padding: 10px;
        """)
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # 关闭按钮
        close_btn = QPushButton("关闭")
        close_btn.setFixedSize(100, 32)
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['primary']};
                color: white;
                border: none;
                border-radius: 6px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                background: #0056CC;
            }}
        """)
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        dialog.exec()

    def _open_feedback(self):
        """打开反馈问卷链接"""
        import webbrowser
        webbrowser.open(FEEDBACK_URL)

    def _open_manual(self):
        """打开使用手册链接"""
        import webbrowser
        webbrowser.open("https://fcnwakmkeuz7.feishu.cn/docx/M9ojdaGUAoRVv7x3NCAcxkxenUe?from=from_copylink")

    def _show_donate(self):
        """显示打赏二维码对话框"""
        dialog = DonateDialog(self)
        dialog.exec()


# === 打赏对话框 ===
class DonateDialog(QDialog):
    """打赏二维码对话框"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("打赏支持")
        self.resize(360, 440)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

        # 获取当前主题
        self.theme = Theme.LIGHT if not parent or not hasattr(parent, 'is_dark') or not parent.is_dark else Theme.DARK

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # 感谢文字
        thanks_label = QLabel("感谢您的支持！")
        thanks_label.setStyleSheet(f"""
            font-size: 18px;
            font-weight: bold;
            color: {self.theme['text']};
        """)
        thanks_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(thanks_label)

        subtitle = QLabel("您的支持是我持续更新的动力 ❤️")
        subtitle.setStyleSheet(f"font-size: 13px; color: {self.theme['text_secondary']};")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(subtitle)

        # 二维码图片
        qr_label = QLabel()
        qr_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        qr_label.setFixedSize(260, 260)

        # 尝试加载二维码图片
        qr_path = resource_path(os.path.join("assets", "donate_qrcode.png"))
        if os.path.exists(qr_path):
            pixmap = QPixmap(qr_path)
            if not pixmap.isNull():
                qr_label.setPixmap(pixmap.scaled(240, 240, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
            else:
                qr_label.setText("二维码加载失败")
                qr_label.setStyleSheet(f"color: {self.theme['danger']}; font-size: 14px;")
        else:
            qr_label.setText("请添加二维码图片至\nassets/donate_qrcode.png")
            qr_label.setStyleSheet(f"""
                color: {self.theme['text_secondary']};
                font-size: 12px;
                background: {self.theme['hover']};
                border: 2px dashed {self.theme['border']};
                border-radius: 8px;
            """)

        layout.addWidget(qr_label, alignment=Qt.AlignmentFlag.AlignCenter)

        # 提示文字
        tip_label = QLabel("微信扫码打赏")
        tip_label.setStyleSheet(f"font-size: 13px; color: {self.theme['text_secondary']};")
        tip_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(tip_label)

        # 关闭按钮
        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['primary']};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background: {Theme.adjust_color(self.theme['primary'], -20)};
            }}
        """)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

class WordBatchReplaceWorker(QThread):
    """Word 批量替换线程（同一套规则应用全部文件）。"""

    progress_signal = pyqtSignal(int, int, str)      # processed, total, current_file
    file_done_signal = pyqtSignal(str, str)          # input_path, output_path
    file_error_signal = pyqtSignal(int, str, str)    # index, input_path, error_msg
    finished_signal = pyqtSignal(dict)               # summary

    def __init__(self, file_paths, rules, default_replacement_text):
        super().__init__()
        self.file_paths = list(file_paths or [])
        self.rules = normalize_word_replace_rules(rules or [], default_replacement_text)
        self.default_replacement_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
        self._error_decision = None
        self._decision_event = threading.Event()
        self._decision_lock = threading.Lock()
        self._temp_dirs = []

    def provide_error_decision(self, decision):
        """主线程调用：为当前错误设置决策（skip/stop）。"""
        with self._decision_lock:
            self._error_decision = decision
        self._decision_event.set()

    def _wait_for_error_decision(self):
        """等待主线程选择错误处理策略。"""
        self._decision_event.clear()
        while not self._decision_event.wait(0.1):
            if self.isInterruptionRequested():
                return "stop"
        with self._decision_lock:
            decision = self._error_decision or "skip"
            self._error_decision = None
        return decision

    def run(self):
        summary = {
            "total": len(self.file_paths),
            "processed": 0,
            "success": [],
            "failed": [],
            "stopped": False,
            "rules": list(self.rules),
        }

        try:
            total = len(self.file_paths)
            for idx, file_path in enumerate(self.file_paths):
                if self.isInterruptionRequested():
                    summary["stopped"] = True
                    break

                current_name = os.path.basename(file_path)
                self.progress_signal.emit(idx, total, current_name)

                try:
                    output_path, replace_stats = self._process_single_file(file_path)
                    summary["success"].append({
                        "input": file_path,
                        "output": output_path,
                        "total_replacements": max(0, int(replace_stats.get("total_replacements", 0) or 0))
                        if isinstance(replace_stats, dict) else 0,
                        "rule_counts": replace_stats.get("rule_counts", []) if isinstance(replace_stats, dict) else [],
                    })
                    self.file_done_signal.emit(file_path, output_path)
                except (OSError, IOError, RuntimeError, ValueError, ConversionError, PermissionError) as e:
                    error_msg = str(e)
                    summary["failed"].append({
                        "input": file_path,
                        "error": error_msg
                    })
                    self.file_error_signal.emit(idx, file_path, error_msg)
                    decision = self._wait_for_error_decision()
                    if decision == "stop":
                        summary["stopped"] = True
                        break
                finally:
                    self.progress_signal.emit(idx + 1, total, current_name)

            summary["processed"] = len(summary["success"]) + len(summary["failed"])
            self.finished_signal.emit(summary)
        finally:
            self._cleanup_temp_dirs()

    def _process_single_file(self, file_path):
        from docx import Document

        is_safe, error_msg = validate_safe_path(file_path, allowed_extensions=[".doc", ".docx"])
        if not is_safe:
            raise ConversionError("输入文件路径不安全", error_msg)

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            source_docx = file_path
        elif ext == ".doc":
            source_docx = self._convert_doc_to_docx(file_path)
        else:
            raise ConversionError("不支持的文件格式", f"{file_path}")

        doc = Document(source_docx)
        replace_stats = self._apply_rules_to_document(doc)
        output_path = self._build_output_path(file_path)
        doc.save(output_path)
        return output_path, replace_stats

    def _apply_rules_to_document(self, doc):
        rule_counts = {}
        total_replacements = 0

        def _collect_match_counts(matches):
            nonlocal total_replacements
            if not matches:
                return
            total_replacements += len(matches)
            for match in matches:
                try:
                    rule_index = int(match.get("rule_index", -1))
                except (TypeError, ValueError, AttributeError):
                    continue
                if rule_index < 0:
                    continue
                rule_counts[rule_index] = rule_counts.get(rule_index, 0) + 1

        for para in doc.paragraphs:
            text = ''.join(run.text for run in para.runs)
            if not text:
                continue
            matches = build_word_rule_matches(text, self.rules, self.default_replacement_text)
            if matches:
                _collect_match_counts(matches)
                replace_matches_in_paragraph(para, matches, text_offset=0,
                                             fallback_replacement_text=self.default_replacement_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    if not cell_text:
                        continue
                    cell_matches = build_word_rule_matches(cell_text, self.rules, self.default_replacement_text)
                    if not cell_matches:
                        continue
                    _collect_match_counts(cell_matches)

                    para_offset = 0
                    paragraphs = list(cell.paragraphs)
                    for idx, para in enumerate(paragraphs):
                        original_para_len = len(''.join(run.text for run in para.runs))
                        replace_matches_in_paragraph(
                            para,
                            cell_matches,
                            text_offset=para_offset,
                            fallback_replacement_text=self.default_replacement_text
                        )
                        para_offset += original_para_len
                        if idx < len(paragraphs) - 1:
                            para_offset += 1

        return {
            "total_replacements": total_replacements,
            "rule_counts": [
                {"rule_index": rule_index, "count": count}
                for rule_index, count in sorted(rule_counts.items())
            ],
        }

    def _build_output_path(self, file_path):
        base_path = os.path.splitext(file_path)[0]
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        candidate = f"{base_path}__replaced_{timestamp}.docx"
        if not os.path.exists(candidate):
            return candidate

        suffix = 1
        while True:
            candidate_with_suffix = f"{base_path}__replaced_{timestamp}_{suffix}.docx"
            if not os.path.exists(candidate_with_suffix):
                return candidate_with_suffix
            suffix += 1

    def _convert_doc_to_docx(self, doc_path):
        """v37.7.6: 委托给共享转换模块。"""
        docx_path, temp_dir = _shared_convert_doc_to_docx(doc_path)
        self._temp_dirs.append(temp_dir)
        return docx_path

    def _cleanup_temp_dirs(self):
        for temp_dir in list(self._temp_dirs):
            try:
                if os.path.isdir(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except (OSError, IOError):
                pass
        self._temp_dirs.clear()

# === 单页画布 (完全复制 v7.0 的 PDFCanvas 实现) ===
class SinglePageCanvas(QLabel):
    # 保留信号以兼容双页模式
    rect_added = pyqtSignal(int, QRectF)
    rect_removed = pyqtSignal(int, int, bool)
    zoom_request = pyqtSignal(float)
    page_change_request = pyqtSignal(int)  # 翻页请求信号：正值=下一页，负值=上一页

    def __init__(self, page_index=0, parent=None):
        super().__init__(parent)
        # 完全复制 v7.0 的初始化
        self.setMouseTracking(True)
        self.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
        self.setAutoFillBackground(True)

        self.page_index = page_index  # 新增：用于双页模式
        self.zoom_scale = 1.0
        self.rects_ocr = []
        self.rects_manual = []
        self.mask_color = QColor(0, 0, 0)

        self.drawing = False
        self.start_point = QPointF()
        self.current_rect = QRectF()

    def set_mask_color(self, color):
        """v7.0 方法"""
        self.mask_color = color
        self.update()

    def update_content(self, pixmap, scale, ocr_rects, manual_rects):
        """v7.0 方法 - 直接引用列表，不复制！"""
        self.setPixmap(pixmap)
        self.zoom_scale = scale
        self.rects_ocr = ocr_rects  # ← v7.0 直接引用，不复制
        self.rects_manual = manual_rects  # ← v7.0 直接引用，不复制
        self.update()

    # === v22.7: 完全回归 v7.0 - mousePressEvent 处理左右键 ===
    def mousePressEvent(self, event):
        """v7.0 风格：左键画框，右键删除"""
        if not self.pixmap():
            return

        # 左键画框
        if event.button() == Qt.MouseButton.LeftButton:
            self.drawing = True
            self.start_point = event.position()
            self.current_rect = QRectF(self.start_point, self.start_point)
            self.update()

        # 右键删除 (v7.0 风格)
        elif event.button() == Qt.MouseButton.RightButton:
            click_pos = event.position()
            if DEBUG_MODE:
                print(f"\n[DEBUG] === 右键点击 === 页面{self.page_index}, 位置({click_pos.x():.2f}, {click_pos.y():.2f})")
                print(f"[DEBUG] 手动框数量: {len(self.rects_manual)}, OCR框数量: {len(self.rects_ocr)}")

            deleted = False

            # 先删手动框（从后往前，优先删除上层）
            for i in range(len(self.rects_manual) - 1, -1, -1):
                screen_rect = self.pdf_to_screen(self.rects_manual[i])
                if DEBUG_MODE:
                    print(f"[DEBUG] 检查手动框[{i}]: {screen_rect}, contains={screen_rect.contains(click_pos)}")
                if screen_rect.contains(click_pos):
                    del self.rects_manual[i]
                    deleted = True
                    if DEBUG_MODE:
                        print(f"[DEBUG] ✓ 删除手动框[{i}]")
                    break

            # 再删 OCR 框
            if not deleted:
                for i in range(len(self.rects_ocr) - 1, -1, -1):
                    screen_rect = self.pdf_to_screen(self.rects_ocr[i])
                    if DEBUG_MODE:
                        print(f"[DEBUG] 检查OCR框[{i}]: {screen_rect}, contains={screen_rect.contains(click_pos)}")
                    if screen_rect.contains(click_pos):
                        del self.rects_ocr[i]
                        deleted = True
                        if DEBUG_MODE:
                            print(f"[DEBUG] ✓ 删除OCR框[{i}]")
                        break

            if deleted:
                self.update()
            else:
                if DEBUG_MODE:
                    print(f"[DEBUG] ✗ 未点击到任何矩形框")
            # 不需要 emit 信号！列表是共享引用，删除已经同步到主窗口

    def pdf_to_screen(self, rect):
        """v7.0 实现 - 带小的容错范围"""
        base_rect = QRectF(rect.x()*self.zoom_scale, rect.y()*self.zoom_scale,
                           rect.width()*self.zoom_scale, rect.height()*self.zoom_scale)
        # 扩展 2 像素容错范围，处理点击边界的情况
        return base_rect.adjusted(-2, -2, 2, 2)

    def paintEvent(self, event):
        """v7.0 实现"""
        super().paintEvent(event)
        if not self.pixmap(): return

        painter = QPainter(self)

        # 使用当前选中的颜色 (黑或白)
        painter.setBrush(self.mask_color)
        painter.setPen(Qt.PenStyle.NoPen)

        # 1. 绘制 AI 框
        for r in self.rects_ocr:
            sr = self.pdf_to_screen(r)
            painter.drawRect(sr)

        # 2. 绘制 手动框
        for r in self.rects_manual:
            sr = self.pdf_to_screen(r)
            painter.drawRect(sr)

        # 3. 绘制 拖拽框 (始终红色边框，提示用)
        if self.drawing and not self.current_rect.isEmpty():
            pen = QPen(QColor(255, 0, 0), 2)
            painter.setPen(pen)
            painter.setBrush(Qt.BrushStyle.NoBrush)
            painter.drawRect(self.current_rect)

    def mouseMoveEvent(self, event):
        """v7.0 实现"""
        if self.drawing and not self.start_point.isNull():
            current_pos = QPointF(event.position())
            self.current_rect = QRectF(self.start_point, current_pos).normalized()
            self.update()

    def mouseReleaseEvent(self, event):
        """v7.0 实现"""
        if event.button() == Qt.MouseButton.LeftButton and self.drawing:
            self.drawing = False
            if self.current_rect.width() > 5 and self.current_rect.height() > 5:
                real_rect = QRectF(
                    self.current_rect.x()/self.zoom_scale,
                    self.current_rect.y()/self.zoom_scale,
                    self.current_rect.width()/self.zoom_scale,
                    self.current_rect.height()/self.zoom_scale
                )
                # v7.0 直接添加到列表（列表是共享引用）
                self.rects_manual.append(real_rect)
            self.current_rect = QRectF()
            self.update()

    # v35.1: 增强滚轮事件 - 支持缩放和翻页
    # v37.3.6: 添加滚动阈值，防止 macOS 双指轻触误触发翻页
    SCROLL_THRESHOLD = 10  # 滚动阈值（像素），忽略小于此值的小幅滚动

    def wheelEvent(self, event: QWheelEvent):
        modifiers = QApplication.keyboardModifiers()
        delta = event.angleDelta().y()

        # v37.3.6: 忽略非常小的滚动量（macOS 双指轻触产生的噪声）
        if abs(delta) < self.SCROLL_THRESHOLD:
            # 小幅滚动只传递给父类处理正常滚动，不触发翻页
            super().wheelEvent(event)
            return

        # Ctrl/Cmd + 滚轮：缩放（保持原有功能）
        if modifiers == Qt.KeyboardModifier.ControlModifier or modifiers == Qt.KeyboardModifier.MetaModifier:
            event.accept()
            if delta > 0:
                self.zoom_request.emit(0.1)
            else:
                self.zoom_request.emit(-0.1)
        # Shift + 滚轮：快速翻页（一次2页）
        elif modifiers == Qt.KeyboardModifier.ShiftModifier:
            event.accept()
            if delta < 0:  # 向下滚动 = 下一页
                self.page_change_request.emit(2)
            else:  # 向上滚动 = 上一页
                self.page_change_request.emit(-2)
        # 普通滚轮：触发翻页信号（由 MainWindow 判断滚动条位置）
        else:
            # 发送翻页请求信号，由 MainWindow 处理边缘检测
            if delta < 0:
                self.page_change_request.emit(1)
            else:
                self.page_change_request.emit(-1)
            # 同时传递给父类以支持正常滚动
            super().wheelEvent(event)

# === OCR 线程 ===
# v37.7.6: 改为使用模块化 OCRWorker，自动注入 box_adjust_ratio
class OCRWorker(_ModularOCRWorker):
    """OCR 处理线程（兼容层：自动注入 config 中的 box_adjust_ratio）"""

    def __init__(self, pdf_path, rules, use_enhance, custom_keywords, scan_scale, off_x, off_w,
                 use_char_level_ocr: bool = False, seal_detection_enabled: bool = False):
        box_adjust_ratio = config.get("ocr.box_adjust_ratio", 0.0) if config else 0.0
        super().__init__(pdf_path, rules, use_enhance, custom_keywords, scan_scale, off_x, off_w,
                         use_char_level_ocr=use_char_level_ocr,
                         seal_detection_enabled=seal_detection_enabled,
                         box_adjust_ratio=box_adjust_ratio)

# === WebView Bridge：Python 与 JavaScript 通信 ===
class WebViewBridge(QObject):
    """Python 与 JavaScript 通信的桥梁"""

    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self._scroll_position = 0  # 保存的滚动位置
        self._pending_scroll_restore = False  # 是否有待恢复的滚动位置

    @pyqtSlot(str, int, int, str)
    def add_manual_redaction(self, key, start, end, selected_text):
        """添加精确手动脱敏（仅选中区域）

        Args:
            key: word_data 中的键（如 paragraph_0）
            start: 起始位置
            end: 结束位置
            selected_text: 被选中的文本
        """
        if key in self.main_window.word_data:
            # 检查重叠
            for existing in self.main_window.word_data[key]['manual']:
                if not (end <= existing['start'] or start >= existing['end']):
                    # 存在重叠，显示提示
                    QMessageBox.warning(self.main_window, "无法添加", "该区域与已有脱敏区域重叠")
                    return

            self.main_window.word_data[key]['manual'].append({
                'start': start,
                'end': end,
                'text': selected_text,
                'replacement': self.main_window.replacement_text,
                'mode': 'exact'  # 精确模式：只高亮选中区域
            })
            self.main_window.render_word_preview()
        else:
            # 添加调试日志
            print(f"警告: 键 '{key}' 不在 word_data 中")
            QMessageBox.warning(self.main_window, "添加失败", f"无法定位文本区域 (键: {key})")

    @pyqtSlot(str, str)
    def add_manual_redaction_global(self, key, selected_text):
        """添加全局手动脱敏（整篇相同文本）

        Args:
            key: 当前选中的 key（用于定位上下文）
            selected_text: 选中的文本
        """
        # 遍历所有 word_data，为每个包含该文本的位置添加脱敏
        for k, data in self.main_window.word_data.items():
            text = data['text']
            if not text:
                continue

            # 查找所有匹配位置
            import re
            pattern = re.escape(selected_text)
            for match in re.finditer(pattern, text):
                start = match.start()
                end = match.end()

                # 检查是否与已有脱敏重叠
                overlap = False
                for existing in data['manual']:
                    if not (end <= existing['start'] or start >= existing['end']):
                        overlap = True
                        break

                if not overlap:
                    data['manual'].append({
                        'start': start,
                        'end': end,
                        'text': selected_text,
                        'replacement': self.main_window.replacement_text,
                        'mode': 'global'  # 全局模式：高亮所有相同文本
                    })

        self.main_window.render_word_preview()

    @pyqtSlot(str, int, int)
    def remove_manual_redaction(self, key, start, end):
        """删除手动脱敏

        Args:
            key: word_data 中的键
            start: 起始位置
            end: 结束位置
        """
        if key in self.main_window.word_data:
            manual_list = self.main_window.word_data[key]['manual']
            # 查找要删除的项
            target_item = None
            for i, item in enumerate(manual_list):
                if item['start'] == start and item['end'] == end:
                    target_item = item
                    # 检查是否是全局模式
                    if item.get('mode') == 'global':
                        # 全局模式：删除所有相同文本的脱敏
                        text_to_remove = item['text']
                        self.remove_global_redaction(text_to_remove)
                    else:
                        # 精确模式：只删除当前项
                        manual_list.pop(i)
                        self.main_window.render_word_preview()
                    return

    def remove_global_redaction(self, text):
        """删除所有全局模式的脱敏（批量撤销）

        Args:
            text: 要删除的文本
        """
        count = 0
        for key, data in self.main_window.word_data.items():
            manual_list = data['manual']
            # 从后往前删除，避免索引问题
            for i in range(len(manual_list) - 1, -1, -1):
                item = manual_list[i]
                if item.get('mode') == 'global' and item['text'] == text:
                    manual_list.pop(i)
                    count += 1

        if count > 0:
            print(f"[撤销] 删除了 {count} 个全局脱敏: {text}")
            self.main_window.render_word_preview()

    def get_scroll_position(self):
        """获取当前保存的滚动位置"""
        return self._scroll_position

    def set_scroll_position(self, position):
        """设置要恢复的滚动位置"""
        self._scroll_position = position
        self._pending_scroll_restore = True

    def clear_pending_scroll_restore(self):
        """清除待恢复标志"""
        self._pending_scroll_restore = False

    def has_pending_scroll_restore(self):
        """检查是否有待恢复的滚动位置"""
        return self._pending_scroll_restore

    @pyqtSlot(str, float)
    def report_word_preview_scroll(self, panel_id, ratio):
        """接收 Word 双栏预览滚动比例并同步到另一侧。"""
        try:
            panel = str(panel_id or "").strip().lower()
            ratio_value = float(ratio)
        except (TypeError, ValueError):
            return
        self.main_window._sync_word_compare_scroll(panel, ratio_value)

# === Word 文档处理线程 ===
# v37.7.6: 改为使用模块化 WordWorker，补充 default_rules 参数
class WordWorker(_ModularWordWorker):
    """Word 文档智能脱敏线程（兼容层：自动注入 DEFAULT_RULES）"""

    def __init__(self, word_doc, word_data, rules, custom_keywords, replacement_text):
        super().__init__(word_doc, word_data, rules, custom_keywords,
                         replacement_text, default_rules=DEFAULT_RULES)

# === Word 预览交互式 JavaScript 代码常量 ===
# v36.4: 提取为常量，避免 _inject_interactive_html 函数过长
_INTERACTIVE_JS_CODE = r"""
<script>
    let pyBridge = null;
    let webChannelReady = false;

    document.addEventListener('DOMContentLoaded', function() {
        // 初始化 QWebChannel
        if (typeof qt !== 'undefined' && qt.webChannelTransport) {
            try {
                new QWebChannel(qt.webChannelTransport, function(channel) {
                    pyBridge = channel.objects.pyBridge;
                    if (pyBridge) {
                        webChannelReady = true;
                    }
                });
            } catch (e) {
                console.error('QWebChannel error:', e);
            }
        }
        setupContextMenu();
    });

    // v35.1: 备用右键位置（用于复杂 DOM 结构中 getSelection() 返回空值的情况）
    let lastContextMenuEvent = null;

    function setupContextMenu() {
        // v35.1: 使用捕获阶段监听（更可靠，在事件冒泡前捕获）
        document.addEventListener('contextmenu', function(e) {
            e.preventDefault();
            lastContextMenuEvent = {
                clientX: e.clientX,
                clientY: e.clientY,
                target: e.target
            };
            handleContextMenu(e);
        }, true);  // true = 捕获阶段

        // v35.1: mousedown 事件预先保存右键位置（备用方案）
        document.addEventListener('mousedown', function(e) {
            if (e.button === 2) {  // 右键
                lastContextMenuEvent = {
                    clientX: e.clientX,
                    clientY: e.clientY,
                    target: e.target
                };
            }
        }, true);

        document.addEventListener('click', function(e) {
            if (!e.target.closest('#redaction-menu')) {
                hideContextMenu();
            }
        });
    }

    // v35.1: 增强的右键菜单处理函数
    function handleContextMenu(e) {
        const target = e.target;
        let selection = window.getSelection();
        let selectedText = selection.toString().trim();

        // 查找点击的目标是否是手动脱敏标记或其内部元素
        let highlightElement = target.closest('.manual-highlight');

        // 点击手动脱敏标记
        if (highlightElement) {
            console.log('[ContextMenu] 点击了手动脱敏标记');
            showRemoveMenu(e.clientX, e.clientY, highlightElement);
            return;
        }

        // 选择了文本（主要路径）
        if (selectedText.length > 0) {
            // v36.5: 移除敏感信息日志，仅记录操作类型
            console.log('[ContextMenu] 选择了文本（已隐藏内容）');
            try {
                const range = selection.getRangeAt(0);
                showAddMenu(e.clientX, e.clientY, selection, selectedText);
            } catch (err) {
                console.warn('[ContextMenu] getRangeAt 失败，尝试备用方案:', err);
                // 备用方案：使用预先保存的位置
                if (lastContextMenuEvent) {
                    tryFallbackTextDetection(e, target);
                }
            }
            return;
        }

        // v35.1: 备用方案 - 尝试从点击位置获取文本
        console.log('[ContextMenu] getSelection() 为空，尝试备用检测');
        tryFallbackTextDetection(e, target);
    }

    // v35.1: 备用文本检测（当 window.getSelection() 失败时）
    function tryFallbackTextDetection(e, target) {
        // 方案1: 检查目标元素是否包含文本
        let textElement = target;

        // 向上查找包含 data-key 的文本块
        for (let i = 0; i < 10 && textElement; i++) {
            if (textElement.dataset && textElement.dataset.key) {
                console.log('[tryFallbackTextDetection] 找到 data-key 元素:', textElement.dataset.key);
                // 显示一个提示菜单（v36.1 安全修复：使用配置对象）
                const menu = createMenu([
                    { text: '请在文本上拖动选择后再右键', disabled: true }
                ]);
                positionMenu(menu, e.clientX, e.clientY);
                setTimeout(hideContextMenu, 2000);
                return;
            }
            textElement = textElement.parentNode;
        }

        // 未找到 data-key，显示提示
        console.warn('[tryFallbackTextDetection] 未找到有效的文本块');
    }

    function showRemoveMenu(x, y, element) {
        const key = element.dataset.key;
        const start = parseInt(element.dataset.start);
        const end = parseInt(element.dataset.end);

        console.log('[showRemoveMenu] key:', key, 'start:', start, 'end:', end);

        // v36.1 安全修复：使用配置对象替代 HTML 字符串
        const menu = createMenu([
            { text: '❌ 撤销脱敏', action: 'remove', key: key, start: start, end: end }
        ]);
        positionMenu(menu, x, y);
        attachMenuHandlers();
    }

    function showAddMenu(x, y, selection, selectedText) {
        const range = selection.getRangeAt(0);
        const textInfo = findTextPosition(selectedText, range);

        // 即使 textInfo 为 null 也不直接返回，使用全局模式
        let buttonConfigs = [];

        if (!textInfo || textInfo.mode === 'global' || textInfo.key === '__GLOBAL__') {
            // 降级到全局模式：只显示全文脱敏选项
            console.log('[showAddMenu] 使用全局降级模式');
            // v36.1 安全修复：使用配置对象替代 HTML 字符串
            buttonConfigs = [
                { text: '📄 全文脱敏此内容', action: 'add-global-only', textData: selectedText }
            ];
        } else {
            // 正常情况：提供精确和全局两种选项
            buttonConfigs = [
                { text: '🎯 选中区域添加脱敏', action: 'add-exact', key: textInfo.key, start: textInfo.start, end: textInfo.end, textData: selectedText },
                { text: '📄 整篇相同字节添加脱敏', action: 'add-global', key: textInfo.key, textData: selectedText }
            ];
        }

        const menu = createMenu(buttonConfigs);
        positionMenu(menu, x, y);
        attachMenuHandlers();
    }

    function attachMenuHandlers() {
        const menu = document.getElementById('redaction-menu');
        if (!menu) return;

        menu.querySelectorAll('button').forEach(btn => {
            btn.addEventListener('click', function(e) {
                e.preventDefault();
                e.stopPropagation();
                const action = this.dataset.action;
                const key = this.dataset.key;
                const start = parseInt(this.dataset.start);
                const end = parseInt(this.dataset.end);
                const text = this.dataset.text;

                if (action === 'remove') {
                    callRemove(key, start, end);
                } else if (action === 'add-exact') {
                    callAddExact(key, start, end, text);
                } else if (action === 'add-global') {
                    callAddGlobal(key, text);
                } else if (action === 'add-global-only') {
                    // 全局降级模式，不需要 key 参数
                    callAddGlobal(null, text);
                }
            });
        });
    }

    function callRemove(key, start, end) {
        console.log('[callRemove] 调用撤销: key=' + key + ' start=' + start + ' end=' + end);
        console.log('[callRemove] pyBridge=', pyBridge, 'webChannelReady=', webChannelReady);
        if (pyBridge && webChannelReady) {
            try {
                pyBridge.remove_manual_redaction(key, start, end);
                console.log('[callRemove] ✓ 撤销调用成功');
            } catch(e) {
                console.error('[callRemove] ✗ 撤销调用失败:', e);
            }
        } else {
            console.error('[callRemove] ✗ pyBridge 或 webChannel 未就绪');
        }
        hideContextMenu();
    }

    function callAddExact(key, start, end, text) {
        if (pyBridge && webChannelReady) {
            pyBridge.add_manual_redaction(key, start, end, text);
        }
        hideContextMenu();
        const selection = window.getSelection();
        selection.removeAllRanges();
    }

    function callAddGlobal(key, text) {
        // v36.5: 移除敏感信息日志
        console.log('[callAddGlobal] 调用全局脱敏（文本已隐藏）');
        if (pyBridge && webChannelReady) {
            // key 为 null 时表示纯全局模式
            pyBridge.add_manual_redaction_global(key || '', text);
        }
        hideContextMenu();
        const selection = window.getSelection();
        selection.removeAllRanges();
    }

    function findTextPosition(selectedText, range) {
        try {
            // v36.5: 移除敏感信息日志
            console.log('[findTextPosition] ========== 开始查找 ==========');
            console.log('[findTextPosition] 选中文本（已隐藏内容）');
            console.log('[findTextPosition] Range:', {
                startContainer: range.startContainer?.nodeName,
                startOffset: range.startOffset,
                endContainer: range.endContainer?.nodeName,
                endOffset: range.endOffset
                // 移除 text 字段，避免泄露敏感信息
            });

            let container = range.commonAncestorContainer;
            if (!container) {
                console.warn('[findTextPosition] ✗ 无 commonAncestorContainer，使用全局降级');
                return createGlobalFallbackResult(selectedText);
            }

            console.log('[findTextPosition] commonAncestorContainer:', container.nodeName, container.nodeType);

            // 如果是文本节点，获取其父元素
            if (container.nodeType === 3) {
                container = container.parentNode;
            }

            // 查找包含 data-key 的元素（向上遍历）
            let element = container;
            let maxIterations = 50;
            let iterations = 0;
            let foundKey = null;

            while (element && iterations < maxIterations) {
                iterations++;
                if (element.dataset && element.dataset.key) {
                    foundKey = element.dataset.key;
                    console.log('[findTextPosition] ✓ 找到 data-key:', foundKey);
                    break;
                }
                element = element.parentNode;
                if (element === document.body || element === document.documentElement) {
                    console.warn('[findTextPosition] 到达文档顶部，未找到 data-key，尝试文本内容定位');
                    break;  // 不返回 null，继续尝试其他方法
                }
            }

            // 方法 1: 精确计算（如果有 data-key）
            if (foundKey && element) {
                const key = foundKey;

                // === 方法 1a: 直接使用 Range 计算位置（最可靠）===
                try {
                    const textNodes = [];
                    const walker = document.createTreeWalker(
                        element,
                        NodeFilter.SHOW_TEXT,
                        {
                            acceptNode: function(node) {
                                // 接受所有非空文本节点
                                return NodeFilter.FILTER_ACCEPT;
                            }
                        }
                    );

                    let node;
                    while (node = walker.nextNode()) {
                        textNodes.push(node);
                    }

                    console.log('[findTextPosition] 找到', textNodes.length, '个文本节点');

                    // 打印所有文本节点信息
                    for (let i = 0; i < Math.min(textNodes.length, 10); i++) {
                        const tn = textNodes[i];
                        console.log(`[findTextPosition] 节点[${i}]:`, {
                            text: tn.textContent.substring(0, 30),
                            isStart: tn === range.startContainer,
                            isEnd: tn === range.endContainer
                        });
                    }

                    // 计算起始位置
                    let startOffset = 0;
                    let startFound = false;

                    // 特殊处理：如果 startContainer 是元素节点，找到它的第一个文本节点
                    let startContainer = range.startContainer;
                    if (startContainer.nodeType === 1) {  // 元素节点
                        console.log('[findTextPosition] startContainer 是元素节点，查找第一个文本子节点');
                        for (let i = 0; i < textNodes.length; i++) {
                            if (element.contains(textNodes[i]) || element === textNodes[i].parentNode) {
                                startContainer = textNodes[i];
                                break;
                            }
                        }
                    }

                    for (let i = 0; i < textNodes.length; i++) {
                        const tn = textNodes[i];
                        if (!startFound) {
                            if (tn === startContainer || (startContainer && tn.contains && tn.contains(startContainer))) {
                                startOffset += range.startOffset;
                                startFound = true;
                                console.log('[findTextPosition] ✓ 起始节点匹配, offset:', range.startOffset);
                            } else {
                                startOffset += tn.textContent.length;
                            }
                        }
                    }

                    // 计算结束位置
                    let endOffset = 0;
                    let endFound = false;

                    let endContainer = range.endContainer;
                    if (endContainer.nodeType === 1) {
                        console.log('[findTextPosition] endContainer 是元素节点，查找第一个文本子节点');
                        for (let i = 0; i < textNodes.length; i++) {
                            if (element.contains(textNodes[i]) || element === textNodes[i].parentNode) {
                                endContainer = textNodes[i];
                                break;
                            }
                        }
                    }

                    for (let i = 0; i < textNodes.length; i++) {
                        const tn = textNodes[i];
                        if (!endFound) {
                            if (tn === endContainer || (endContainer && tn.contains && tn.contains(endContainer))) {
                                endOffset += range.endOffset;
                                endFound = true;
                                console.log('[findTextPosition] ✓ 结束节点匹配, offset:', range.endOffset);
                            } else {
                                endOffset += tn.textContent.length;
                            }
                        }
                    }

                    console.log('[findTextPosition] Range 计算结果:', { startFound, endFound, startOffset, endOffset });

                    if (startFound && endFound && startOffset >= 0 && endOffset > startOffset) {
                        console.log('[findTextPosition] ✓✓✓ Range 计算成功 ✓✓✓');
                        return { key, start: startOffset, end: endOffset };
                    }
                } catch (e) {
                    console.error('[findTextPosition] Range 计算出错:', e);
                }

                // === 方法 1b: 文本匹配（后备）===
                console.log('[findTextPosition] 尝试文本匹配方法...');
                let originalText = '';
                if (element.dataset.originalText) {
                    // 安全解码：使用 DOMParser 替代 innerHTML，防止 XSS
                    const parser = new DOMParser();
                    const doc = parser.parseFromString(element.dataset.originalText, 'text/html');
                    originalText = doc.body.textContent || '';
                    console.log('[findTextPosition] 原始文本长度:', originalText.length);
                }

                if (!originalText) {
                    originalText = element.textContent || '';
                }

                // 尝试直接匹配
                let foundStart = originalText.indexOf(selectedText);
                if (foundStart !== -1) {
                    console.log('[findTextPosition] ✓✓✓ 直接匹配成功 ✓✓✓');
                    return { key, start: foundStart, end: foundStart + selectedText.length };
                }

                // 尝试归一化匹配
                const normalizedSelected = selectedText.replace(/\s+/g, ' ').trim();
                const normalizedOriginal = originalText.replace(/\s+/g, ' ');
                foundStart = normalizedOriginal.indexOf(normalizedSelected);
                if (foundStart !== -1) {
                    console.log('[findTextPosition] ✓✓✓ 归一化匹配成功 ✓✓✓');
                    return { key, start: foundStart, end: foundStart + selectedText.length };
                }
            }

            // 方法 2: 文本内容定位（遍历所有 data-key 元素）
            console.log('[findTextPosition] 尝试文本内容定位方法...');
            const textResult = findPositionByTextContent(selectedText);
            if (textResult) {
                console.log('[findTextPosition] ✓✓✓ 文本内容定位成功 ✓✓✓');
                return textResult;
            }

            // 方法 3: 全局脱敏降级
            console.log('[findTextPosition] 所有精确方法失败，使用全局降级模式');
            return createGlobalFallbackResult(selectedText);

        } catch (e) {
            console.error('[findTextPosition] 异常:', e);
            return createGlobalFallbackResult(selectedText);
        }
    }

    function createGlobalFallbackResult(selectedText) {
        console.log('[findTextPosition] 创建全局降级结果，文本长度:', selectedText.length);
        return {
            key: '__GLOBAL__',
            start: 0,
            end: selectedText.length,
            mode: 'global',
            text: selectedText
        };
    }

    function findPositionByTextContent(selectedText) {
        // 遍历所有带有 data-key 的元素，查找包含选中文本的元素
        const elements = document.querySelectorAll('[data-key]');
        const normalizedSelected = selectedText.replace(/\s+/g, ' ').trim();

        for (const el of elements) {
            const key = el.dataset.key;
            if (!key || key === '__GLOBAL__') continue;

            let originalText = '';
            if (el.dataset.originalText) {
                // 安全解码：使用 DOMParser 替代 innerHTML，防止 XSS
                const parser = new DOMParser();
                const doc = parser.parseFromString(el.dataset.originalText, 'text/html');
                originalText = doc.body.textContent || '';
            } else {
                originalText = el.textContent || '';
            }

            const normalizedOriginal = originalText.replace(/\s+/g, ' ');
            const foundStart = normalizedOriginal.indexOf(normalizedSelected);

            if (foundStart !== -1) {
                console.log('[findTextPosition] 文本内容定位找到匹配，key:', key);
                return { key, start: foundStart, end: foundStart + selectedText.length };
            }
        }

        return null;
    }

    // v36.1 安全修复：使用 DOM 方法替代 innerHTML，防止 XSS
    function createMenu(buttonConfigs) {
        hideContextMenu();
        const menu = document.createElement('div');
        menu.id = 'redaction-menu';
        menu.style.cssText = 'position:fixed; background:white; border:1px solid #ddd; border-radius:8px; box-shadow:0 4px 12px rgba(0,0,0,0.15); padding:8px 0; z-index:10000; min-width:150px;';

        const buttonStyle = 'display:block; width:100%; padding:8px 16px; border:none; background:none; text-align:left; cursor:pointer; font-size:14px; color:#333;';

        // 安全创建按钮元素，避免使用 innerHTML
        buttonConfigs.forEach(config => {
            const btn = document.createElement('button');
            btn.style.cssText = buttonStyle;
            btn.onmouseover = () => btn.style.backgroundColor = '#f5f5f5';
            btn.onmouseout = () => btn.style.backgroundColor = 'transparent';

            // 设置按钮文本（自动转义 HTML）
            btn.textContent = config.text || '';

            // 设置 data 属性
            if (config.action) btn.setAttribute('data-action', config.action);
            if (config.key !== undefined) btn.setAttribute('data-key', config.key);
            if (config.start !== undefined) btn.setAttribute('data-start', config.start);
            if (config.end !== undefined) btn.setAttribute('data-end', config.end);
            if (config.textData !== undefined) btn.setAttribute('data-text', config.textData);

            // 设置 disabled 状态
            if (config.disabled) {
                btn.disabled = true;
                btn.style.color = '#999';
                btn.style.cursor = 'default';
            }

            menu.appendChild(btn);
        });

        document.body.appendChild(menu);
        return menu;
    }

    function positionMenu(menu, x, y) {
        menu.style.display = 'block';
        menu.style.left = x + 'px';
        menu.style.top = y + 'px';
    }

    function hideContextMenu() {
        const menu = document.getElementById('redaction-menu');
        if (menu) menu.remove();
    }
</script>
"""

# === 主窗口 ===
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} v{VERSION} - Powered by li (汪立律师)")

        # v37.0: 从配置读取窗口尺寸，失败时使用硬编码后备
        # v37.2.0: 读取 OCR 引擎配置
        if config:
            min_width = config.get("app.window.min_width", 900)
            min_height = config.get("app.window.min_height", 600)
            default_width = config.get("app.window.default_width", 1300)
            default_height = config.get("app.window.default_height", 900)
            self.replacement_text = config.get("redaction.replacement_text", "[已脱敏]")
            self.scan_level = config.get("redaction.scan.default_level", 2.0)
            self.offset_x = config.get("redaction.offset.default_x", 0)
            self.offset_w = config.get("redaction.offset.default_w", 0)
            self.custom_keywords = config.get("redaction.custom_keywords", "")
            # v37.4.0: 移除 OCR 引擎配置，只使用 RapidOCR
        else:
            min_width, min_height = 900, 600
            default_width, default_height = 1300, 900
            self.replacement_text = "[已脱敏]"
            self.scan_level = 2.0
            self.offset_x = 0
            self.offset_w = 0
            self.custom_keywords = ""

        # 窗口尺寸设置：最小尺寸 + 默认尺寸
        self.setMinimumSize(min_width, min_height)
        self.resize(default_width, default_height)

        # 窗口状态保存
        self.settings = QSettings("PrivacyGuard", "App")
        self._restore_window_state()

        self.doc = None
        self.word_doc = None  # Word 文档对象
        self.file_path = ""
        self.current_page = 0
        self.zoom_level = 1.0
        self.page_data = {}
        self._ocr_processed_pages = set()  # OCR 实际处理过的页（用于准确完成状态提示）
        self.word_data = {}  # Word 文档数据结构
        self.word_replace_rules = []  # 会话级多字段替换规则
        self.word_compare_mode = False  # Word 预览是否开启原文/替换后对比
        self.word_compare_user_hidden = False  # 用户主动隐藏右侧对比预览
        self._word_data_lock = QMutex()  # v36.5: 保护 word_data 线程安全
        self.doc_type = None  # 'pdf', 'docx', 'doc'
        self.current_ui_mode = "idle"  # idle / pdf / word / batch / image_merge
        self.batch_stage = "idle"  # idle / rule_setup / running / finished / stopped
        self.batch_selected_files = []
        self.batch_total_files = 0
        self.batch_processed_files = 0
        self.batch_success_count = 0
        self.batch_failed_count = 0
        self.batch_current_file = ""
        self.batch_last_summary = None
        self.batch_result_filter_mode = "all"
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self.info_bar_message = ""
        self.toolbar_density_mode = "wide"
        self._bound_window_handle = None
        self._button_density_metrics = {}
        self.active_rules = [DEFAULT_RULES.get("身份证号", ""), DEFAULT_RULES.get("手机号码", "")]
        self.use_enhance = False
        self.current_color = QColor(0, 0, 0)
        self.dual_view = False

        # 预先创建 word_preview
        self.word_preview = None
        self.word_preview_replaced = None
        self.bridge = None
        self.word_web_channel = None
        self._word_scroll_sync_timer = QTimer(self)
        self._word_scroll_sync_timer.setInterval(160)
        self._word_scroll_sync_timer.timeout.connect(self._poll_word_compare_scroll_sync)
        self._word_scroll_sync_polling = False
        self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}
        self._word_scroll_sync_pending_target = None
        self._word_scroll_sync_pending_ratio = None
        self._word_scroll_sync_generation = 0
        self._word_preview_assets_dir = None
        self._word_preview_assets_base_url = QUrl()
        self._reset_word_preview_cache()

        # 线程管理和临时文件管理（v24 稳定性优化）
        self.active_worker = None
        self.batch_worker = None
        self.active_task_type = None  # 'scan', 'batch_replace'
        self.worker_lock = QMutex()
        self.temp_manager = TempFileManager()

        # 注册退出清理
        import atexit
        atexit.register(self._app_exit_cleanup)

        self.setup_ui()

        # v37.6.0: 启用拖拽支持
        self.setAcceptDrops(True)
        self._drag_active = False  # 拖拽状态标记
        self._drag_valid = False   # 拖拽文件是否有效

    def _detect_system_theme(self):
        """检测系统主题（v35.1 新增）

        Returns:
            str: 'light' 或 'dark'
        """
        import platform

        system = platform.system()

        try:
            if system == 'Darwin':  # macOS
                import subprocess
                result = subprocess.run(
                    ['defaults', 'read', '-g', 'AppleInterfaceStyle'],
                    capture_output=True, text=True
                )
                if result.returncode == 0 and 'Dark' in result.stdout:
                    return 'dark'
                return 'light'

            elif system == 'Windows':
                import winreg
                key = winreg.OpenKey(
                    winreg.HKEY_CURRENT_USER,
                    r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize"
                )
                value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
                winreg.CloseKey(key)
                return 'light' if value == 1 else 'dark'

            else:  # Linux
                import os
                gtk_theme = os.environ.get('GTK_THEME', '').lower()
                if 'dark' in gtk_theme:
                    return 'dark'
                return 'light'

        except (OSError, IOError, KeyError, ValueError, ImportError) as e:
            # 检测失败（注册表读取错误、环境变量异常等），默认浅色
            print(f"[MainWindow] 主题检测失败: {e}")
            return 'light'

    def _restore_window_state(self):
        """恢复窗口状态"""
        geometry = self.settings.value("window_geometry")
        if geometry:
            self.restoreGeometry(geometry)

    def closeEvent(self, event):
        """保存窗口状态并清理临时文件"""
        self._app_exit_cleanup()
        self.settings.setValue("window_geometry", self.saveGeometry())
        super().closeEvent(event)

    def showEvent(self, event):
        """首次显示和重新显示窗口时，补绑屏幕缩放监听。"""
        super().showEvent(event)
        self._bind_window_handle_signals()
        QTimer.singleShot(0, self._refresh_toolbar_responsiveness)

    def resizeEvent(self, event):
        """窗口大小改变时自动重新适应页面"""
        super().resizeEvent(event)

        if hasattr(self, "toolbar"):
            self._refresh_toolbar_responsiveness()

        # 只在 PDF 模式且文档已加载时处理
        if not self.doc or self.current_page is None:
            return

        # 自动重新适应（保持页面完整显示）
        self.fit_page()

    def _bind_window_handle_signals(self):
        """绑定窗口句柄相关信号，便于跨屏和 DPI 切换时收口工具栏。"""
        try:
            handle = self.windowHandle()
        except Exception:
            handle = None

        if handle is None or handle is self._bound_window_handle:
            return

        try:
            handle.screenChanged.connect(self._on_window_screen_changed)
        except Exception:
            pass
        self._bound_window_handle = handle

    def _on_window_screen_changed(self, _screen):
        """窗口切换到不同屏幕时，重新计算工具栏密度。"""
        QTimer.singleShot(0, self._refresh_toolbar_responsiveness)

    def _reset_word_preview_cache(self):
        """重置 Word 预览缓存与待应用的局部更新状态。"""
        if hasattr(self, "_word_scroll_sync_timer"):
            self._invalidate_word_scroll_sync()
        self._cleanup_word_preview_assets_dir()
        self._word_base_html = None
        self._word_html_source_path = None
        self._word_tagged_html = None
        self._word_preview_document_html = None
        self._word_replaced_document_html = None
        self._word_replaced_html = None
        self._pending_word_preview_blocks = None
        self._pending_word_replaced_blocks = None
        self._word_preview_ready = False
        self._word_replaced_ready = False
        self._word_preview_loaded_source_path = None
        self._word_replaced_loaded_source_path = None
        self._word_preview_target_source_path = None
        self._word_replaced_target_source_path = None
        self._word_preview_assets_base_url = QUrl()

    def _cleanup_word_preview_assets_dir(self):
        """清理当前 Word 预览图片临时目录。"""
        asset_dir = getattr(self, "_word_preview_assets_dir", None)
        if not asset_dir:
            return

        try:
            if os.path.isdir(asset_dir):
                shutil.rmtree(asset_dir, ignore_errors=True)
        except Exception as e:
            print(f"[清理] 删除 Word 预览图片目录失败: {e}")

        temp_manager = getattr(self, "temp_manager", None)
        instance_lock = getattr(temp_manager, "_instance_lock", None)
        temp_dirs = getattr(temp_manager, "temp_dirs", None)
        if temp_manager is not None and instance_lock is not None and isinstance(temp_dirs, list):
            try:
                with instance_lock:
                    if asset_dir in temp_dirs:
                        temp_dirs.remove(asset_dir)
            except Exception:
                pass

        self._word_preview_assets_dir = None

    def _create_word_preview_asset_dir(self):
        """创建当前 Word 预览使用的图片资源目录。"""
        self._cleanup_word_preview_assets_dir()
        if hasattr(self, "temp_manager") and self.temp_manager is not None:
            asset_dir = self.temp_manager.create_temp_dir()
        else:
            asset_dir = tempfile.mkdtemp(prefix="pg_word_preview_")
        self._word_preview_assets_dir = asset_dir
        self._word_preview_assets_base_url = QUrl.fromLocalFile(os.path.join(asset_dir, ""))
        return asset_dir

    def _build_word_html_from_docx(self, docx_path):
        """将 DOCX 转成 HTML，并把嵌入图片提取到临时目录。"""
        import mammoth

        asset_dir = self._create_word_preview_asset_dir()
        image_counter = {"value": 0}

        def convert_image_to_file(image):
            image_counter["value"] += 1
            suffix = resolve_word_preview_image_suffix(getattr(image, "content_type", ""))
            image_name = f"word_image_{image_counter['value']:04d}{suffix}"
            target_path = os.path.join(asset_dir, image_name)
            try:
                with image.open() as image_source, open(target_path, "wb") as output_file:
                    shutil.copyfileobj(image_source, output_file)
                return {
                    "src": image_name,
                    "loading": "lazy",
                    "decoding": "async",
                }
            except Exception as e:
                print(f"[Word预览] 提取嵌入图片失败: {e}")
                return {
                    "src": WORD_PREVIEW_BROKEN_IMAGE_DATA_URI,
                    "loading": "lazy",
                    "decoding": "async",
                    "data-pg-image-error": "1",
                }

        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(convert_image_to_file),
            )

        if getattr(result, "messages", None):
            for message in result.messages:
                print(f"[Word预览] Mammoth: {message}")
        return result.value

    # ============== v37.6.0: 拖拽打开文件功能 ==============

    def dragEnterEvent(self, event):
        """拖拽进入事件 - 验证文件类型并提供视觉反馈"""
        self._drag_active = True
        self._drag_valid = False

        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()

            # 验证所有文件
            valid_exts = ['.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
            all_valid = True
            file_count = 0

            for url in urls:
                if url.isLocalFile():
                    file_count += 1
                    path = url.toLocalFile().lower()
                    if not any(path.endswith(ext) for ext in valid_exts):
                        all_valid = False
                        break

            # 必须有文件且都有效
            if file_count > 0 and all_valid:
                self._drag_valid = True
                event.acceptProposedAction()
                self._update_drag_visual_feedback(True)
            else:
                # 有文件但格式不支持
                if file_count > 0:
                    event.ignore()
                    self._update_drag_visual_feedback(False)
                else:
                    event.ignore()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        """拖拽移动事件 - 持续反馈"""
        if self._drag_active and event.mimeData().hasUrls():
            # 检查鼠标位置是否在预览区域
            pos = event.position().toPoint() if hasattr(event, 'position') else event.pos()
            if self._is_in_preview_area(pos):
                event.acceptProposedAction()
            else:
                event.ignore()
        else:
            event.ignore()

    def dragLeaveEvent(self, event):
        """拖拽离开事件 - 清除视觉反馈"""
        self._drag_active = False
        self._drag_valid = False
        self._update_drag_visual_feedback(None)

    def dropEvent(self, event):
        """拖放事件 - 处理文件"""
        self._drag_active = False
        self._update_drag_visual_feedback(None)

        if not event.mimeData().hasUrls():
            event.ignore()
            return

        urls = event.mimeData().urls()
        file_paths = [url.toLocalFile() for url in urls if url.isLocalFile()]

        if not file_paths:
            QMessageBox.warning(self, "无效文件", "只支持本地文件拖拽")
            event.ignore()
            return

        # 调用处理逻辑
        self._handle_dropped_files(file_paths)
        event.acceptProposedAction()

    def _is_in_preview_area(self, pos):
        """检查坐标是否在预览区域内"""
        # pos 是 MainWindow 坐标，转换到 scroll 局部坐标后做命中测试
        map_pos = self.scroll.mapFrom(self, pos)
        return self.scroll.rect().contains(map_pos)

    def _update_drag_visual_feedback(self, valid):
        """更新拖拽视觉反馈

        Args:
            valid: True(有效), False(无效), None(清除)
        """
        if valid is True:
            # 有效文件 - 绿色边框提示
            self.scroll.setStyleSheet(f"""
                QScrollArea {{
                    background-color: {Theme.LIGHT["scroll_area"]};
                    border-radius: {Theme.BORDER_RADIUS}px;
                    border: 3px solid #34C759;
                }}
            """)
        elif valid is False:
            # 无效文件 - 红色边框提示
            self.scroll.setStyleSheet(f"""
                QScrollArea {{
                    background-color: {Theme.LIGHT["scroll_area"]};
                    border-radius: {Theme.BORDER_RADIUS}px;
                    border: 3px solid #FF3B30;
                }}
            """)
        else:
            # 清除 - 恢复默认
            self.scroll.setStyleSheet(self.scroll_style.format(Theme.LIGHT["scroll_area"]))

    def _handle_dropped_files(self, file_paths):
        """处理拖拽的文件

        Args:
            file_paths: 文件路径列表
        """
        # 清理之前的状态
        self._cleanup_before_open()
        self._cleanup_temp_file()

        if len(file_paths) == 1:
            # 单个文件
            fname = file_paths[0]
            doc_type = self.detect_file_type(fname)

            try:
                if doc_type == 'pdf':
                    self._open_pdf_file(fname)
                elif doc_type == 'docx':
                    self._open_word_docx(fname)
                elif doc_type == 'doc':
                    self._open_word_doc(fname)
                elif doc_type == 'image':
                    self._open_images_merge([fname])
                else:
                    QMessageBox.warning(
                        self, "不支持的格式",
                        f"文件: {os.path.basename(fname)}\n\n"
                        f"请选择 PDF、Word 文档或图片文件"
                    )
            except (IOError, OSError, ValueError, ConversionError) as e:
                QMessageBox.critical(self, "错误", f"打开文件失败: {str(e)}")
        else:
            # 多个文件：支持图片合并或 Word 批量替换
            are_all_images = all(self.detect_file_type(f) == 'image' for f in file_paths)
            are_all_words = all(self.detect_file_type(f) in ('docx', 'doc') for f in file_paths)
            if are_all_images:
                try:
                    self._open_images_merge(file_paths)
                except (IOError, OSError, ValueError, ConversionError) as e:
                    QMessageBox.critical(self, "错误", f"合并图片失败: {str(e)}")
            elif are_all_words:
                self.start_batch_replace(file_paths=file_paths)
            else:
                QMessageBox.warning(
                    self, "不支持的混合拖拽",
                    "同时拖拽多个文件时，仅支持两种场景：\n"
                    "1. 全部是图片（自动合并为PDF）\n"
                    "2. 全部是Word（自动启动批量替换）"
                )

    def _show_drag_tooltip(self, file_paths, valid):
        """显示拖拽文件提示信息（可选增强）

        Args:
            file_paths: 文件路径列表
            valid: 是否有效
        """
        if not file_paths:
            return

        if len(file_paths) == 1:
            fname = os.path.basename(file_paths[0])
            doc_type = self.detect_file_type(file_paths[0])
            type_names = {
                'pdf': 'PDF 文档',
                'docx': 'Word 文档',
                'doc': 'Word 文档(旧版)',
                'image': '图片文件',
                'unknown': '未知格式'
            }
            tooltip = f"{fname}\n类型: {type_names.get(doc_type, '未知')}"
        else:
            tooltip = f"共 {len(file_paths)} 个文件"
            if valid:
                tooltip += "\n将合并为 PDF"
            else:
                tooltip += "\n格式不支持"

        # 使用 QToolTip 显示
        from PyQt6.QtWidgets import QToolTip
        from PyQt6.QtGui import QCursor
        QToolTip.showText(QCursor.pos(), tooltip)

    # ============== 拖拽功能结束 ==============

    def _cleanup_before_open(self):
        """v37.0.6: 打开新文档前的完整资源清理

        解决问题：
        - 打开新文档时卡顿/未响应
        - 文件选择窗口内容不显示
        """
        self._invalidate_word_scroll_sync()
        # 1. 停止并等待活跃的 worker 线程
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.requestInterruption()
            # 等待线程结束（最多 2 秒）
            self.active_worker.wait(2000)
            self.active_worker = None

        # 2. 清理 QWebEngineView（Word 预览）
        # QWebEngineView 占用大量资源，需要正确清理
        if hasattr(self, 'word_preview') and self.word_preview:
            try:
                # 停止加载
                self.word_preview.stop()
                # 清空内容
                self.word_preview.setHtml('')
                # 隐藏
                self.word_preview.hide()
            except Exception as e:
                print(f"[清理] 清理 word_preview 时出错: {e}")
        if hasattr(self, 'word_preview_replaced') and self.word_preview_replaced:
            try:
                self.word_preview_replaced.stop()
                self.word_preview_replaced.setHtml('')
                self.word_preview_replaced.hide()
            except Exception as e:
                print(f"[清理] 清理 word_preview_replaced 时出错: {e}")

        # 3. 关闭 PDF 文档
        if self.doc:
            try:
                self.doc.close()
            except Exception as e:
                print(f"[清理] 关闭 PDF 文档时出错: {e}")
            self.doc = None

        # 4. 重置状态变量
        self.word_doc = None
        self.word_data = {}
        self._reset_word_preview_cache()
        self.page_data = {}
        self._ocr_processed_pages = set()
        self.current_page = None
        self.doc_type = None
        self.file_path = None
        self.word_compare_mode = False
        self.word_compare_user_hidden = False
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._reset_batch_session_state()
        self._clear_info_bar_message()

        # 5. v37.0.8: 重置 canvas 显示状态（不删除固定实例）
        # canvas_left 和 canvas_right 是固定实例，在 setup_ui() 中创建
        # 只需清除显示内容，不需要删除
        # v37.0.9: 修复属性名错误 - 使用正确的 rects_manual 和 rects_ocr
        if hasattr(self, 'canvas_left') and self.canvas_left:
            try:
                # 检查 C++ 对象是否仍然有效
                _ = self.canvas_left.size()  # 如果对象已删除，这里会抛出异常
                self.canvas_left.clear()  # 清除显示
                self.canvas_left.page_index = 0  # 重置页面索引
                self.canvas_left.rects_manual = []  # 清除手动脱敏区域（正确的属性名）
                self.canvas_left.rects_ocr = []  # 清除 OCR 区域（正确的属性名）
            except RuntimeError:
                print("[清理] canvas_left 的 C++ 对象已被删除，跳过清理")

        if hasattr(self, 'canvas_right') and self.canvas_right:
            try:
                _ = self.canvas_right.size()
                self.canvas_right.clear()
                self.canvas_right.page_index = 1
                self.canvas_right.rects_manual = []
                self.canvas_right.rects_ocr = []
            except RuntimeError:
                print("[清理] canvas_right 的 C++ 对象已被删除，跳过清理")

        if hasattr(self, 'word_compare_container') and self.word_compare_container:
            self.word_compare_container.hide()
        if hasattr(self, 'canvas_container') and self.canvas_container:
            self.canvas_container.show()

        # 6. 处理待处理的 Qt 事件，确保 UI 响应
        QApplication.processEvents()
        self._sync_ui_mode()

        print("[清理] 打开新文档前的资源清理完成")

    def _has_active_open_context(self):
        """判断当前是否已有活跃文档/任务，需要在打开新文件前先做清理。"""
        return bool(
            self.doc
            or self.word_doc
            or self.file_path
            or self.doc_type
            or self.page_data
            or self.word_data
            or self.image_merge_in_progress
            or (self.batch_stage != "idle" and self.batch_selected_files)
            or self.current_ui_mode != "idle"
        )

    def _is_canvas_valid(self, canvas):
        """v37.0.9: 检查 canvas 的 C++ 对象是否仍然有效"""
        if canvas is None:
            return False
        try:
            # 尝试访问一个简单的属性来验证对象是否有效
            _ = canvas.size()
            return True
        except RuntimeError:
            # C++ 对象已被删除
            return False

    def _safe_canvas_update(self, canvas, pixmap, scale, ocr_rects, manual_rects):
        """v37.0.9: 安全地更新 canvas 内容"""
        if not self._is_canvas_valid(canvas):
            print(f"[警告] canvas 无效，跳过更新")
            return False
        try:
            canvas.update_content(pixmap, scale, ocr_rects, manual_rects)
            return True
        except RuntimeError as e:
            print(f"[错误] 更新 canvas 时出错: {e}")
            return False

    def _safe_canvas_set_mask_color(self, canvas, color):
        """v37.0.9: 安全地设置 canvas 遮罩颜色"""
        if not self._is_canvas_valid(canvas):
            return False
        try:
            canvas.set_mask_color(color)
            return True
        except RuntimeError as e:
            print(f"[错误] 设置 canvas 颜色时出错: {e}")
            return False

    def _is_word_web_view_valid(self, web_view):
        """判断 Word 预览 WebView 是否仍然可用。"""
        if web_view is None:
            return False
        if sip.isdeleted(web_view):
            return False
        try:
            page = web_view.page()
            if page is None or sip.isdeleted(page):
                return False
            _ = web_view.isHidden()
            return True
        except RuntimeError:
            return False

    def _invalidate_word_scroll_sync(self):
        """停用并失效当前 Word 双栏滚动同步链，避免异步回调撞上已销毁对象。"""
        if hasattr(self, "_word_scroll_sync_timer"):
            self._word_scroll_sync_timer.stop()
        self._word_scroll_sync_polling = False
        self._word_scroll_sync_pending_target = None
        self._word_scroll_sync_pending_ratio = None
        self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}
        self._word_scroll_sync_generation += 1

    def _app_exit_cleanup(self):
        """应用退出时的清理（v24 稳定性优化）"""
        self._invalidate_word_scroll_sync()
        # 取消正在运行的线程
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.requestInterruption()
            # 等待线程结束（最多 2 秒）
            self.active_worker.wait(2000)

        # 清理临时文件
        if hasattr(self, 'temp_manager'):
            self.temp_manager.cleanup()

        # 清理旧版临时文件
        self._cleanup_temp_file()

    # v22.4: 移除 eventFilter，直接在 SinglePageCanvas.mousePressEvent 中处理

    def setup_ui(self):
        # 统一上下文条：文档上下文 + 临时任务提示
        self.default_info_bar_text = "📝 支持直接拖拽导入，系统会按文件类型自动进入 PDF 脱敏、Word 替换、批量 Word 或图片合并。"
        self.workbench_panel = QFrame()
        self.workbench_panel.setObjectName("workbenchPanel")
        workbench_layout = QVBoxLayout(self.workbench_panel)
        workbench_layout.setContentsMargins(16, 10, 16, 10)
        workbench_layout.setSpacing(4)
        self.workbench_layout = workbench_layout

        context_top_layout = QHBoxLayout()
        context_top_layout.setSpacing(14)
        context_top_layout.setContentsMargins(0, 0, 0, 0)
        self.context_top_layout = context_top_layout
        workbench_text = QVBoxLayout()
        workbench_text.setContentsMargins(0, 0, 0, 0)
        workbench_text.setSpacing(4)
        self.workbench_text_layout = workbench_text

        self.lbl_workbench_title = QLabel("欢迎使用 PrivacyGuard")
        self.lbl_workbench_title.setObjectName("workbenchTitle")
        self.lbl_workbench_subtitle = QLabel("拖拽或打开文件即可开始处理。")
        self.lbl_workbench_subtitle.setObjectName("workbenchSubtitle")
        self.lbl_workbench_subtitle.setWordWrap(True)
        workbench_text.addWidget(self.lbl_workbench_title)
        workbench_text.addWidget(self.lbl_workbench_subtitle)

        self.lbl_workbench_focus = QLabel("开始")
        self.lbl_workbench_focus.setObjectName("workbenchFocus")
        self.btn_workbench_feedback = self.create_btn("使用/反馈", self.show_feedback, style="secondary")
        self.btn_workbench_feedback.hide()

        context_top_layout.addLayout(workbench_text, stretch=1)
        context_top_layout.addWidget(self.btn_workbench_feedback, alignment=Qt.AlignmentFlag.AlignVCenter)
        context_top_layout.addWidget(self.lbl_workbench_focus, alignment=Qt.AlignmentFlag.AlignVCenter)
        workbench_layout.addLayout(context_top_layout)

        workbench_guidance_layout = QGridLayout()
        workbench_guidance_layout.setHorizontalSpacing(8)
        workbench_guidance_layout.setVerticalSpacing(8)
        workbench_guidance_layout.setContentsMargins(0, 0, 0, 0)
        workbench_guidance_layout.setColumnStretch(0, 1)
        workbench_guidance_layout.setColumnStretch(1, 1)
        self.workbench_guidance_layout = workbench_guidance_layout
        self.workbench_guidance_labels = []
        for index in range(4):
            guidance_label = QLabel("")
            guidance_label.setObjectName("workbenchHintTag")
            guidance_label.setWordWrap(True)
            guidance_label.hide()
            row = index // 2
            col = index % 2
            workbench_guidance_layout.addWidget(guidance_label, row, col)
            self.workbench_guidance_labels.append(guidance_label)
        workbench_layout.addLayout(workbench_guidance_layout)

        self.info_bar = QLabel(self.default_info_bar_text)
        self.info_bar.setObjectName("contextMessage")
        self.info_bar.setWordWrap(True)
        self.info_bar.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        workbench_layout.addWidget(self.info_bar)

        # 工具栏
        toolbar = QFrame()
        toolbar.setObjectName("toolbarRoot")
        toolbar.setFixedHeight(54)
        self.toolbar = toolbar  # 保存引用

        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 7, 16, 7)
        tb_layout.setSpacing(10)
        self.toolbar_layout = tb_layout

        self.toolbar_primary_group, self.toolbar_primary_layout = self._create_toolbar_group("toolbarGroupStrong")
        tb_layout.addWidget(self.toolbar_primary_group)
        self.btn_open = self.create_btn("打开", self.open_pdf)
        self.toolbar_primary_layout.addWidget(self.btn_open)
        self.btn_scan = self.create_btn("智能脱敏", self.start_ocr, enabled=False, style="success")
        self.toolbar_primary_layout.addWidget(self.btn_scan)

        self.toolbar_word_group, self.toolbar_word_layout = self._create_toolbar_group()
        tb_layout.addWidget(self.toolbar_word_group)
        self.btn_settings = self.create_btn("高级设置", self.open_settings, style="secondary")
        self.toolbar_word_layout.addWidget(self.btn_settings)
        self.btn_compare_toggle = self.create_btn("对比预览", self.toggle_word_compare_preview, style="secondary")
        self.toolbar_word_layout.addWidget(self.btn_compare_toggle)

        self.toolbar_pdf_group, self.toolbar_pdf_layout = self._create_toolbar_group()
        tb_layout.addWidget(self.toolbar_pdf_group)
        self.rb_black = self.create_btn("黑遮罩", self.update_canvas_color, style="toggle")
        self.rb_black.setObjectName("toolbarToggleButton")
        self.rb_black.setCheckable(True)
        self.rb_black.setChecked(True)
        self.rb_white = self.create_btn("白遮罩", self.update_canvas_color, style="toggle")
        self.rb_white.setObjectName("toolbarToggleButton")
        self.rb_white.setCheckable(True)
        self.bg_color = QButtonGroup(self)
        self.bg_color.setExclusive(True)
        self.bg_color.addButton(self.rb_black)
        self.bg_color.addButton(self.rb_white)
        self.toolbar_pdf_layout.addWidget(self.rb_black)
        self.toolbar_pdf_layout.addWidget(self.rb_white)

        self.cb_dual = self.create_btn("双页", self._toggle_dual_toolbar, style="toggle")
        self.cb_dual.setObjectName("toolbarToggleButton")
        self.cb_dual.setCheckable(True)
        self.toolbar_pdf_layout.addWidget(self.cb_dual)

        self.btn_fit = self.create_btn("适应", self.fit_page, style="secondary")
        self.toolbar_pdf_layout.addWidget(self.btn_fit)
        self.btn_fit.hide()

        tb_layout.addStretch()

        self.toolbar_zoom_group, self.toolbar_zoom_layout = self._create_toolbar_group()
        tb_layout.addWidget(self.toolbar_zoom_group)
        self.lbl_zoom = QLabel("100%")
        self.lbl_zoom.setObjectName("toolbarMeta")
        self.lbl_zoom.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.btn_zoom_out = self.create_btn("-", self.zoom_out, style="icon", tooltip="缩小预览")
        self.toolbar_zoom_layout.addWidget(self.btn_zoom_out)
        self.toolbar_zoom_layout.addWidget(self.lbl_zoom)
        self.btn_zoom_in = self.create_btn("+", self.zoom_in, style="icon", tooltip="放大预览")
        self.toolbar_zoom_layout.addWidget(self.btn_zoom_in)

        self.toolbar_nav_group, self.toolbar_nav_layout = self._create_toolbar_group()
        tb_layout.addWidget(self.toolbar_nav_group)
        self.btn_go_first = self.create_btn("", self.go_first, style="icon", tooltip="跳到第一页")
        self.toolbar_nav_layout.addWidget(self.btn_go_first)
        self.btn_prev_page = self.create_btn("", lambda: self.change_page(-1), style="icon", tooltip="上一页")
        self.toolbar_nav_layout.addWidget(self.btn_prev_page)
        self.lbl_page = QLabel("0 / 0")
        self.lbl_page.setObjectName("toolbarMeta")
        self.lbl_page.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.toolbar_nav_layout.addWidget(self.lbl_page)
        self.btn_next_page = self.create_btn("", lambda: self.change_page(1 if not self.dual_view else 2), style="icon", tooltip="下一页")
        self.toolbar_nav_layout.addWidget(self.btn_next_page)
        self.btn_go_last = self.create_btn("", self.go_last, style="icon", tooltip="跳到最后一页")
        self.toolbar_nav_layout.addWidget(self.btn_go_last)

        self.toolbar_utility_group, self.toolbar_utility_layout = self._create_toolbar_group("toolbarUtilityGroup")
        tb_layout.addWidget(self.toolbar_utility_group)
        self.btn_fit_utility = self.create_btn("适应页面", self.fit_page, style="secondary")
        self.btn_fit_utility.hide()
        self.toolbar_utility_layout.addWidget(self.btn_fit_utility)
        self.btn_feedback = self.create_btn("使用/反馈", self.show_feedback, style="secondary")
        self.toolbar_utility_layout.addWidget(self.btn_feedback)
        self.btn_more = self.create_btn("更多", self._show_toolbar_more_menu, style="secondary")
        self.btn_more.setObjectName("toolbarMoreButton")
        self.toolbar_more_menu = QMenu(self)
        self.btn_more.hide()
        self.toolbar_utility_layout.addWidget(self.btn_more)
        self.btn_save = self.create_btn("导出", self.save_pdf, enabled=False)
        self.toolbar_utility_layout.addWidget(self.btn_save)
        self._apply_native_toolbar_icons()

        main = QWidget()
        main.setObjectName("appRoot")
        self.setCentralWidget(main)
        layout = QVBoxLayout(main)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.workbench_panel)
        layout.addWidget(toolbar)

        # 滚动区域
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.scroll_style = (
            f"background-color: {{0}}; "
            f"border-radius: {Theme.BORDER_RADIUS}px; "
            f"border: none;"
        )

        # v22.9: 使用固定的 canvas_container，通过隐藏/显示实现单/双页切换
        self.canvas_left = SinglePageCanvas(0)
        self.canvas_right = SinglePageCanvas(1)

        # 容器始终作为 scroll 的 widget
        self.canvas_container = QWidget()
        self.pdf_workspace_outer_layout = QVBoxLayout(self.canvas_container)
        self.pdf_workspace_outer_layout.setContentsMargins(14, 10, 14, 16)
        self.pdf_workspace_outer_layout.setSpacing(0)
        self.pdf_workspace_shell = QFrame()
        self.pdf_workspace_shell.setObjectName("previewWorkspaceCard")
        self.pdf_workspace_shell.setMaximumWidth(1940)
        self.pdf_workspace_shell.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        pdf_workspace_shell_layout = QVBoxLayout(self.pdf_workspace_shell)
        pdf_workspace_shell_layout.setContentsMargins(8, 8, 8, 8)
        pdf_workspace_shell_layout.setSpacing(0)
        self.pdf_workspace_shell_layout = pdf_workspace_shell_layout
        self.pdf_stage_content = QWidget()
        self.pdf_stage_content.setObjectName("previewStage")
        self.canvas_layout = QHBoxLayout(self.pdf_stage_content)
        self.canvas_layout.setContentsMargins(10, 10, 10, 10)
        self.canvas_layout.setSpacing(12)
        self.canvas_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop)
        self.canvas_left.setObjectName("pdfPageCanvas")
        self.canvas_right.setObjectName("pdfPageCanvas")
        self.canvas_layout.addWidget(self.canvas_left)
        self.canvas_layout.addWidget(self.canvas_right)
        pdf_workspace_shell_layout.addWidget(self.pdf_stage_content)
        pdf_workspace_row_layout = QHBoxLayout()
        pdf_workspace_row_layout.setContentsMargins(0, 0, 0, 0)
        pdf_workspace_row_layout.setSpacing(0)
        pdf_workspace_row_layout.addStretch(1)
        pdf_workspace_row_layout.addWidget(self.pdf_workspace_shell, 18)
        pdf_workspace_row_layout.addStretch(1)
        self.pdf_workspace_row_layout = pdf_workspace_row_layout
        self.pdf_workspace_outer_layout.addLayout(pdf_workspace_row_layout, 1)

        # 信号连接
        self.canvas_left.rect_added.connect(self.on_rect_added)
        self.canvas_left.rect_removed.connect(self.on_rect_removed)
        self.canvas_left.zoom_request.connect(self.handle_zoom_request)
        self.canvas_left.page_change_request.connect(self.handle_page_change_request)

        self.canvas_right.rect_added.connect(self.on_rect_added)
        self.canvas_right.rect_removed.connect(self.on_rect_removed)
        self.canvas_right.zoom_request.connect(self.handle_zoom_request)
        self.canvas_right.page_change_request.connect(self.handle_page_change_request)

        # 设置 canvas 大小策略
        for canvas in [self.canvas_left, self.canvas_right]:
            canvas.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Preferred)

        # 预先创建 Word 预览视图（左：原文，右：替换后）
        self.word_preview = QWebEngineView()
        self.word_preview_replaced = QWebEngineView()
        self.word_preview_replaced.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        self.word_preview.loadFinished.connect(self._on_word_preview_load_finished)
        self.word_preview_replaced.loadFinished.connect(self._on_word_replaced_load_finished)

        self.idle_workspace_container = QWidget()
        idle_outer_layout = QVBoxLayout(self.idle_workspace_container)
        idle_outer_layout.setContentsMargins(30, 20, 30, 30)
        idle_outer_layout.setSpacing(0)
        self.idle_outer_layout = idle_outer_layout

        idle_card = QFrame()
        idle_card.setObjectName("workspaceCard")
        idle_card.setMaximumWidth(1320)
        idle_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.idle_card = idle_card
        idle_card_layout = QVBoxLayout(idle_card)
        idle_card_layout.setContentsMargins(28, 26, 28, 26)
        idle_card_layout.setSpacing(16)
        self.idle_card_layout = idle_card_layout

        idle_hero_panel = QFrame()
        idle_hero_panel.setObjectName("idleHeroPanel")
        idle_hero_panel.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        idle_hero_layout = QVBoxLayout(idle_hero_panel)
        idle_hero_layout.setContentsMargins(0, 0, 0, 0)
        idle_hero_layout.setSpacing(12)
        self.idle_hero_panel = idle_hero_panel
        self.idle_hero_layout = idle_hero_layout

        idle_title_row_layout = QHBoxLayout()
        idle_title_row_layout.setContentsMargins(0, 0, 0, 0)
        idle_title_row_layout.setSpacing(18)
        self.idle_title_row_layout = idle_title_row_layout

        idle_headline_layout = QVBoxLayout()
        idle_headline_layout.setContentsMargins(0, 0, 0, 0)
        idle_headline_layout.setSpacing(6)
        self.idle_headline_layout = idle_headline_layout

        idle_title = QLabel("选择开始方式")
        idle_title.setObjectName("workspaceTitle")
        idle_subtitle = QLabel("打开或拖拽文件，系统会自动进入对应模式。")
        idle_subtitle.setObjectName("workspaceSubtitle")
        idle_subtitle.setWordWrap(True)
        idle_headline_layout.addWidget(idle_title)
        idle_headline_layout.addWidget(idle_subtitle)

        idle_title_tools_layout = QVBoxLayout()
        idle_title_tools_layout.setContentsMargins(0, 0, 0, 0)
        idle_title_tools_layout.setSpacing(8)
        self.idle_title_tools_layout = idle_title_tools_layout

        idle_badge_row_layout = QHBoxLayout()
        idle_badge_row_layout.setContentsMargins(0, 0, 0, 0)
        idle_badge_row_layout.setSpacing(8)
        self.idle_badge_row_layout = idle_badge_row_layout
        self.lbl_idle_offline_badge = QLabel("本地离线")
        self.lbl_idle_offline_badge.setObjectName("idleHeroBadge")
        self.lbl_idle_auto_badge = QLabel("自动分流")
        self.lbl_idle_auto_badge.setObjectName("idleHeroBadge")
        idle_badge_row_layout.addStretch()
        idle_badge_row_layout.addWidget(self.lbl_idle_offline_badge, alignment=Qt.AlignmentFlag.AlignVCenter)
        idle_badge_row_layout.addWidget(self.lbl_idle_auto_badge, alignment=Qt.AlignmentFlag.AlignVCenter)

        idle_title_tools_layout.addLayout(idle_badge_row_layout)
        idle_title_tools_layout.addStretch(1)

        idle_title_row_layout.addLayout(idle_headline_layout, stretch=1)
        idle_title_row_layout.addLayout(idle_title_tools_layout)
        idle_hero_layout.addLayout(idle_title_row_layout)

        idle_action_panel = QFrame()
        idle_action_panel.setObjectName("idleActionPanel")
        idle_action_panel_layout = QVBoxLayout(idle_action_panel)
        idle_action_panel_layout.setContentsMargins(0, 0, 0, 0)
        idle_action_panel_layout.setSpacing(8)
        self.idle_action_panel = idle_action_panel
        self.idle_action_panel_layout = idle_action_panel_layout

        idle_action_buttons_layout = QGridLayout()
        idle_action_buttons_layout.setContentsMargins(0, 0, 0, 0)
        idle_action_buttons_layout.setHorizontalSpacing(16)
        idle_action_buttons_layout.setVerticalSpacing(12)
        idle_action_buttons_layout.setColumnStretch(0, 1)
        idle_action_buttons_layout.setColumnStretch(1, 1)
        self.idle_action_buttons_layout = idle_action_buttons_layout

        self.idle_start_card = QFrame()
        self.idle_start_card.setObjectName("idleStartCard")
        idle_start_layout = QVBoxLayout(self.idle_start_card)
        idle_start_layout.setContentsMargins(18, 16, 18, 16)
        idle_start_layout.setSpacing(8)
        self.idle_start_layout = idle_start_layout
        self.lbl_idle_start_title = QLabel("开始处理")
        self.lbl_idle_start_title.setObjectName("idleStartTitle")
        self.lbl_idle_start_text = QLabel("选择文件或直接拖拽到窗口，系统会自动进入对应模式。")
        self.lbl_idle_start_text.setObjectName("idleStartText")
        self.lbl_idle_start_text.setWordWrap(True)
        self.btn_idle_open = self.create_btn("选择文件", self.open_pdf)
        self.btn_idle_open.setObjectName("idlePrimaryActionButton")
        self.btn_idle_open.setProperty("btn_style", "idle_primary")
        self.btn_idle_open.setStyleSheet(self._get_button_style("idle_primary"))
        self.lbl_idle_drop_hint = QLabel("支持直接拖拽到窗口")
        self.lbl_idle_drop_hint.setObjectName("idleDropHint")
        self.lbl_idle_drop_hint.setWordWrap(True)
        self.idle_start_footer = QWidget()
        self.idle_start_footer_layout = QVBoxLayout(self.idle_start_footer)
        self.idle_start_footer_layout.setContentsMargins(0, 0, 0, 0)
        self.idle_start_footer_layout.setSpacing(8)
        self.idle_start_footer_layout.addWidget(self.lbl_idle_drop_hint)
        self.idle_start_footer_layout.addWidget(self.btn_idle_open, alignment=Qt.AlignmentFlag.AlignLeft)
        idle_start_layout.addWidget(self.lbl_idle_start_title)
        idle_start_layout.addWidget(self.lbl_idle_start_text)
        idle_start_layout.addStretch(1)
        idle_start_layout.addWidget(self.idle_start_footer)

        self.idle_support_card = QFrame()
        self.idle_support_card.setObjectName("idleSupportCard")
        idle_support_layout = QVBoxLayout(self.idle_support_card)
        idle_support_layout.setContentsMargins(18, 16, 18, 16)
        idle_support_layout.setSpacing(8)
        self.lbl_idle_support_title = QLabel("开发者与支持")
        self.lbl_idle_support_title.setObjectName("idleSupportTitle")
        self.lbl_idle_support_text = QLabel("汪立 · 安徽始信律师事务所执业律师")
        self.lbl_idle_support_text.setObjectName("idleSupportText")
        self.lbl_idle_support_text.setWordWrap(True)
        self.lbl_idle_support_meta = QLabel("全栈律师｜前教师｜退伍军人")
        self.lbl_idle_support_meta.setObjectName("idleSupportMeta")
        self.lbl_idle_support_email = QLabel("<a href='mailto:491445490@qq.com'>491445490@qq.com</a>")
        self.lbl_idle_support_email.setObjectName("idleSupportEmail")
        self.lbl_idle_support_email.setTextFormat(Qt.TextFormat.RichText)
        self.lbl_idle_support_email.setOpenExternalLinks(True)
        self.lbl_idle_support_note = QLabel("遇到问题、想看手册或支持更新，都可以直接从这里进入。")
        self.lbl_idle_support_note.setObjectName("idleSupportNote")
        self.lbl_idle_support_note.setWordWrap(True)
        idle_support_actions_layout = QGridLayout()
        idle_support_actions_layout.setContentsMargins(0, 0, 0, 0)
        idle_support_actions_layout.setHorizontalSpacing(10)
        idle_support_actions_layout.setVerticalSpacing(10)
        idle_support_actions_layout.setColumnStretch(0, 1)
        idle_support_actions_layout.setColumnStretch(1, 1)
        idle_support_actions_layout.setColumnStretch(2, 1)
        self.idle_support_actions_layout = idle_support_actions_layout
        self.btn_idle_feedback = self.create_btn("反馈建议", self._open_feedback, style="primary")
        self.btn_idle_feedback.setObjectName("idleFeedbackActionButton")
        self.btn_idle_feedback.setProperty("btn_style", "primary")
        self.btn_idle_feedback.setStyleSheet(self._get_button_style("primary"))
        self.btn_idle_manual = self.create_btn("使用手册", self._open_manual, style="secondary")
        self.btn_idle_manual.setObjectName("idleManualActionButton")
        self.btn_idle_manual.setProperty("btn_style", "secondary")
        self.btn_idle_manual.setStyleSheet(self._get_button_style("secondary"))
        self.btn_idle_donate = self.create_btn("打赏支持", self._show_donate, style="success")
        self.btn_idle_donate.setObjectName("idleDonateActionButton")
        self.btn_idle_donate.setProperty("btn_style", "success")
        self.btn_idle_donate.setStyleSheet(self._get_button_style("success"))
        idle_support_layout.addWidget(self.lbl_idle_support_title)
        idle_support_layout.addWidget(self.lbl_idle_support_text)
        idle_support_layout.addWidget(self.lbl_idle_support_meta)
        idle_support_layout.addWidget(self.lbl_idle_support_email)
        idle_support_layout.addWidget(self.lbl_idle_support_note)
        idle_support_layout.addStretch(1)
        idle_support_layout.addLayout(idle_support_actions_layout)
        idle_action_buttons_layout.addWidget(self.idle_start_card, 0, 0)
        idle_action_buttons_layout.addWidget(self.idle_support_card, 0, 1)
        idle_action_panel_layout.addLayout(idle_action_buttons_layout)
        idle_hero_layout.addWidget(idle_action_panel)
        idle_card_layout.addWidget(idle_hero_panel)

        idle_flow_panel = QFrame()
        idle_flow_panel.setObjectName("idleFlowPanel")
        idle_flow_panel.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        idle_flow_layout = QHBoxLayout(idle_flow_panel)
        idle_flow_layout.setContentsMargins(0, 2, 0, 0)
        idle_flow_layout.setSpacing(10)
        self.idle_flow_panel = idle_flow_panel
        self.idle_flow_layout = idle_flow_layout

        self.lbl_idle_tip = QLabel("推荐流程")
        self.lbl_idle_tip.setObjectName("idleFlowTitle")
        idle_flow_layout.addWidget(self.lbl_idle_tip)

        workflow_layout = QHBoxLayout()
        workflow_layout.setSpacing(8)
        self.workflow_layout = workflow_layout
        self.workflow_step_labels = []
        for step_text in ["1 导入", "2 规则", "3 处理", "4 复核", "5 导出"]:
            step_label = QLabel(step_text)
            step_label.setObjectName("workflowStep")
            self.workflow_step_labels.append(step_label)
            workflow_layout.addWidget(step_label)
        workflow_layout.addStretch()
        idle_flow_layout.addLayout(workflow_layout, stretch=1)
        idle_card_layout.addWidget(idle_flow_panel)
        idle_card_layout.addSpacing(2)

        idle_section_panel = QFrame()
        idle_section_panel.setObjectName("idleSectionPanel")
        idle_section_panel.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        idle_section_layout = QVBoxLayout(idle_section_panel)
        idle_section_layout.setContentsMargins(0, 0, 0, 0)
        idle_section_layout.setSpacing(10)
        self.idle_section_panel = idle_section_panel
        self.idle_section_layout = idle_section_layout

        idle_section_header_layout = QHBoxLayout()
        idle_section_header_layout.setContentsMargins(0, 0, 0, 0)
        idle_section_header_layout.setSpacing(8)
        self.idle_section_header_layout = idle_section_header_layout

        self.lbl_idle_section = QLabel("四大功能")
        self.lbl_idle_section.setObjectName("idleSectionLabel")
        self.lbl_idle_section_hint = QLabel("按文件类型自动进入")
        self.lbl_idle_section_hint.setObjectName("idleSectionHint")
        idle_section_header_layout.addWidget(self.lbl_idle_section)
        idle_section_header_layout.addStretch()
        idle_section_header_layout.addWidget(self.lbl_idle_section_hint)
        idle_section_layout.addLayout(idle_section_header_layout)

        route_specs = [
            ("PDF 脱敏", "单文档", "打开 PDF，智能脱敏或手动画框。", "pdf"),
            ("Word 替换", "单文档", "打开 Word，替换并对比预览。", "word"),
            ("批量 Word", "批量处理", "导入多份 Word，确认规则后批量执行。", "batch"),
            ("图片合并 PDF", "图片工具", "导入图片，排序后生成 PDF。", "image"),
        ]

        idle_routes_container = QWidget()
        idle_routes_container.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        idle_routes_layout = QGridLayout(idle_routes_container)
        idle_routes_layout.setContentsMargins(0, 0, 0, 0)
        idle_routes_layout.setHorizontalSpacing(16)
        idle_routes_layout.setVerticalSpacing(16)
        idle_routes_layout.setColumnStretch(0, 1)
        idle_routes_layout.setColumnStretch(1, 1)
        self.idle_routes_container = idle_routes_container
        self.idle_routes_layout = idle_routes_layout
        self.idle_route_cards = []
        for index, (title_text, meta_text, desc_text, accent_key) in enumerate(route_specs):
            route_card = QFrame()
            route_card.setObjectName("routeCard")
            route_card.setProperty("routeTone", accent_key)
            route_card.setMinimumHeight(84)
            route_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
            self.idle_route_cards.append(route_card)
            route_card_layout = QVBoxLayout(route_card)
            route_card_layout.setContentsMargins(14, 10, 14, 10)
            route_card_layout.setSpacing(4)
            route_accent = QFrame()
            route_accent.setObjectName("routeCardAccent")
            route_accent.setProperty("routeAccent", accent_key)
            route_accent.setFixedHeight(4)
            route_head_layout = QHBoxLayout()
            route_head_layout.setContentsMargins(0, 0, 0, 0)
            route_head_layout.setSpacing(8)
            route_title = QLabel(title_text)
            route_title.setObjectName("routeCardTitle")
            route_meta = QLabel(meta_text)
            route_meta.setObjectName("routeCardMeta")
            route_meta.setProperty("routeTone", accent_key)
            route_meta.setAlignment(Qt.AlignmentFlag.AlignCenter)
            route_desc = QLabel(desc_text)
            route_desc.setObjectName("routeCardText")
            route_desc.setWordWrap(True)
            route_desc.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
            route_card_layout.addWidget(route_accent)
            route_head_layout.addWidget(route_title)
            route_head_layout.addStretch()
            route_head_layout.addWidget(route_meta, alignment=Qt.AlignmentFlag.AlignTop)
            route_card_layout.addLayout(route_head_layout)
            route_card_layout.addWidget(route_desc)
            idle_routes_layout.addWidget(route_card, index // 2, index % 2)
        idle_section_layout.addWidget(idle_routes_container, 0)
        idle_card_layout.addWidget(idle_section_panel)

        idle_card_row_layout = QHBoxLayout()
        idle_card_row_layout.setContentsMargins(0, 0, 0, 0)
        idle_card_row_layout.setSpacing(0)
        idle_card_row_layout.addStretch(1)
        idle_card_row_layout.addWidget(idle_card, 12)
        idle_card_row_layout.addStretch(1)
        self.idle_card_row_layout = idle_card_row_layout
        idle_outer_layout.addLayout(idle_card_row_layout, 1)

        self.batch_workspace_container = QWidget()
        batch_outer_layout = QVBoxLayout(self.batch_workspace_container)
        batch_outer_layout.setContentsMargins(26, 24, 26, 28)
        batch_outer_layout.setSpacing(0)
        self.batch_outer_layout = batch_outer_layout

        batch_card = QFrame()
        batch_card.setObjectName("batchWorkspaceCard")
        batch_card.setMaximumWidth(1500)
        batch_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Maximum)
        self.batch_card = batch_card
        batch_card_layout = QVBoxLayout(batch_card)
        batch_card_layout.setContentsMargins(28, 24, 28, 24)
        batch_card_layout.setSpacing(14)
        self.batch_card_layout = batch_card_layout

        batch_header_layout = QHBoxLayout()
        batch_header_layout.setSpacing(12)
        self.batch_header_layout = batch_header_layout
        batch_header_text = QVBoxLayout()
        batch_header_text.setSpacing(4)
        self.batch_header_text_layout = batch_header_text
        self.lbl_batch_title = QLabel("批量 Word 工作台")
        self.lbl_batch_title.setObjectName("batchTitle")
        self.lbl_batch_subtitle = QLabel("先确认文档替换规则，再执行批量替换。处理中可以停止，已完成文件会保留。")
        self.lbl_batch_subtitle.setObjectName("batchSubtitle")
        self.lbl_batch_subtitle.setWordWrap(True)
        batch_header_text.addWidget(self.lbl_batch_title)
        batch_header_text.addWidget(self.lbl_batch_subtitle)
        self.lbl_batch_stage_badge = QLabel("等待开始")
        self.lbl_batch_stage_badge.setObjectName("batchStageBadge")
        batch_header_layout.addLayout(batch_header_text, stretch=1)
        batch_header_layout.addWidget(self.lbl_batch_stage_badge, alignment=Qt.AlignmentFlag.AlignTop)
        batch_card_layout.addLayout(batch_header_layout)

        self.lbl_batch_meta = QLabel("批量模式会优先进入规则确认，再进入执行。")
        self.lbl_batch_meta.setObjectName("batchMeta")
        self.lbl_batch_meta.setWordWrap(True)
        batch_card_layout.addWidget(self.lbl_batch_meta)

        batch_stage_layout = QGridLayout()
        batch_stage_layout.setContentsMargins(0, 0, 0, 0)
        batch_stage_layout.setHorizontalSpacing(12)
        batch_stage_layout.setVerticalSpacing(12)
        self.batch_stage_layout = batch_stage_layout
        self.batch_stage_cards = []
        for title_text, note_text in [
            ("1 规则确认", "先核对文档数量、统一替换文本和 Word 规则。"),
            ("2 执行替换", "系统逐个处理文档，支持跳过异常文件。"),
            ("3 查看结果", "处理结束后集中查看成功、失败和输出结果。"),
        ]:
            step_card = QFrame()
            step_card.setObjectName("batchStepCard")
            step_layout = QVBoxLayout(step_card)
            step_layout.setContentsMargins(14, 12, 14, 12)
            step_layout.setSpacing(4)
            step_title = QLabel(title_text)
            step_title.setObjectName("batchStepTitle")
            step_note = QLabel(note_text)
            step_note.setObjectName("batchStepNote")
            step_note.setWordWrap(True)
            step_layout.addWidget(step_title)
            step_layout.addWidget(step_note)
            self.batch_stage_cards.append((step_card, step_title, step_note))
        self._rebuild_batch_stage_layout("wide")
        batch_card_layout.addLayout(batch_stage_layout)

        batch_metrics_layout = QGridLayout()
        batch_metrics_layout.setContentsMargins(0, 0, 0, 0)
        batch_metrics_layout.setHorizontalSpacing(10)
        batch_metrics_layout.setVerticalSpacing(10)
        self.batch_metrics_layout = batch_metrics_layout
        self.batch_metric_cards = []
        self.lbl_batch_metric_files = None
        self.lbl_batch_metric_files_note = None
        self.lbl_batch_metric_rules = None
        self.lbl_batch_metric_rules_note = None
        self.lbl_batch_metric_progress = None
        self.lbl_batch_metric_progress_note = None
        self.lbl_batch_metric_result = None
        self.lbl_batch_metric_result_note = None
        for key, title_text in [
            ("files", "已选文档"),
            ("rules", "启用规则"),
            ("progress", "当前进度"),
            ("result", "执行结果"),
        ]:
            metric_card, metric_value, metric_note = self._create_batch_metric_card(title_text)
            setattr(self, f"lbl_batch_metric_{key}", metric_value)
            setattr(self, f"lbl_batch_metric_{key}_note", metric_note)
            self.batch_metric_cards.append(metric_card)
        self._rebuild_batch_metrics_layout("wide")
        batch_card_layout.addLayout(batch_metrics_layout)

        batch_actions_layout = QGridLayout()
        batch_actions_layout.setContentsMargins(0, 2, 0, 0)
        batch_actions_layout.setHorizontalSpacing(10)
        batch_actions_layout.setVerticalSpacing(10)
        self.batch_actions_layout = batch_actions_layout
        self.btn_batch_edit_rules = self.create_btn("重新设置规则", self._reopen_batch_rule_setup, style="secondary")
        self.btn_batch_pick_files = self.create_btn("重新选择文档", self._start_batch_replace_from_workspace, style="secondary")
        self.btn_batch_retry_failed = self.create_btn("仅重试失败文档", self._retry_failed_batch_files, style="secondary")
        self.btn_batch_open_output = self.create_btn("打开输出位置", self._open_batch_output_location, style="secondary")
        self.batch_action_buttons = [
            self.btn_batch_edit_rules,
            self.btn_batch_pick_files,
            self.btn_batch_retry_failed,
            self.btn_batch_open_output,
        ]
        self._rebuild_batch_action_layout("wide")
        batch_card_layout.addLayout(batch_actions_layout)

        self.lbl_batch_current_file = QLabel("当前文件：尚未开始")
        self.lbl_batch_current_file.setObjectName("batchCurrentFile")
        self.lbl_batch_current_file.setWordWrap(True)
        batch_card_layout.addWidget(self.lbl_batch_current_file)

        batch_detail_layout = QGridLayout()
        batch_detail_layout.setContentsMargins(0, 0, 0, 0)
        batch_detail_layout.setHorizontalSpacing(14)
        batch_detail_layout.setVerticalSpacing(14)
        self.batch_detail_layout = batch_detail_layout

        self.batch_summary_section = QFrame()
        self.batch_summary_section.setObjectName("batchDetailSection")
        self.batch_summary_section.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Maximum)
        batch_summary_section_layout = QVBoxLayout(self.batch_summary_section)
        batch_summary_section_layout.setContentsMargins(0, 0, 0, 0)
        batch_summary_section_layout.setSpacing(8)
        self.batch_summary_section_layout = batch_summary_section_layout
        self.lbl_batch_summary_hint = QLabel("本轮摘要")
        self.lbl_batch_summary_hint.setObjectName("batchSectionLabel")
        batch_summary_section_layout.addWidget(self.lbl_batch_summary_hint)

        self.batch_summary_browser = QTextBrowser()
        self.batch_summary_browser.setObjectName("batchSummaryBrowser")
        self.batch_summary_browser.setOpenExternalLinks(False)
        self.batch_summary_browser.setMinimumHeight(148)
        self.batch_summary_browser.setMaximumHeight(220)
        self.batch_summary_browser.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        batch_summary_section_layout.addWidget(self.batch_summary_browser)

        self.batch_result_section = QFrame()
        self.batch_result_section.setObjectName("batchDetailSection")
        self.batch_result_section.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        batch_result_section_layout = QVBoxLayout(self.batch_result_section)
        batch_result_section_layout.setContentsMargins(0, 0, 0, 0)
        batch_result_section_layout.setSpacing(8)
        self.batch_result_section_layout = batch_result_section_layout
        self.lbl_batch_result_hint = QLabel("结果清单")
        self.lbl_batch_result_hint.setObjectName("batchSectionLabel")

        batch_result_toolbar = QHBoxLayout()
        batch_result_toolbar.setSpacing(8)
        batch_result_toolbar.setContentsMargins(0, 0, 0, 0)
        batch_result_toolbar.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        self.batch_result_toolbar = batch_result_toolbar
        self.lbl_batch_result_meta = QLabel("结果计数：等待本轮结果")
        self.lbl_batch_result_meta.setObjectName("batchResultMeta")
        self.lbl_batch_result_meta.setAlignment(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft)
        batch_result_toolbar.addWidget(self.lbl_batch_result_meta, stretch=1)
        self.btn_batch_filter_all = QPushButton("全部")
        self.btn_batch_filter_success = QPushButton("仅成功")
        self.btn_batch_filter_failed = QPushButton("仅失败")
        for filter_mode, button in [
            ("all", self.btn_batch_filter_all),
            ("success", self.btn_batch_filter_success),
            ("failed", self.btn_batch_filter_failed),
        ]:
            button.setObjectName("batchFilterButton")
            button.setCursor(Qt.CursorShape.PointingHandCursor)
            button.setAutoDefault(False)
            button.setDefault(False)
            button.clicked.connect(lambda _checked=False, mode=filter_mode: self._set_batch_result_filter_mode(mode))
            batch_result_toolbar.addWidget(button)
        batch_result_header_layout = QHBoxLayout()
        batch_result_header_layout.setContentsMargins(0, 0, 0, 0)
        batch_result_header_layout.setSpacing(12)
        batch_result_header_layout.addWidget(self.lbl_batch_result_hint, stretch=1, alignment=Qt.AlignmentFlag.AlignVCenter)
        batch_result_header_layout.addLayout(batch_result_toolbar, stretch=2)
        self.batch_result_header_layout = batch_result_header_layout
        batch_result_section_layout.addLayout(batch_result_header_layout)

        self.batch_result_table = QTableWidget(0, 4)
        self.batch_result_table.setObjectName("batchResultTable")
        self.batch_result_table.setHorizontalHeaderLabels(["状态", "输入文档", "结果说明", "操作"])
        self.batch_result_table.verticalHeader().setVisible(False)
        self.batch_result_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.batch_result_table.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.batch_result_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.batch_result_table.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.batch_result_table.setAlternatingRowColors(False)
        self.batch_result_table.setShowGrid(False)
        self.batch_result_table.setWordWrap(True)
        self.batch_result_table.setMinimumHeight(196)
        self.batch_result_table.setMaximumHeight(280)
        batch_header = self.batch_result_table.horizontalHeader()
        batch_header.setHighlightSections(False)
        batch_header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        batch_header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        batch_header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        batch_header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        if self.batch_result_table.horizontalHeaderItem(0):
            self.batch_result_table.horizontalHeaderItem(0).setTextAlignment(int(Qt.AlignmentFlag.AlignCenter))
        if self.batch_result_table.horizontalHeaderItem(3):
            self.batch_result_table.horizontalHeaderItem(3).setTextAlignment(int(Qt.AlignmentFlag.AlignCenter))
        self.batch_result_table.cellDoubleClicked.connect(self._open_batch_result_row)
        batch_result_section_layout.addWidget(self.batch_result_table)

        self.batch_log_section = QFrame()
        self.batch_log_section.setObjectName("batchDetailSection")
        self.batch_log_section.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        batch_log_section_layout = QVBoxLayout(self.batch_log_section)
        batch_log_section_layout.setContentsMargins(0, 0, 0, 0)
        batch_log_section_layout.setSpacing(8)
        self.batch_log_section_layout = batch_log_section_layout
        self.lbl_batch_log_hint = QLabel("处理动态")
        self.lbl_batch_log_hint.setObjectName("batchSectionLabel")
        batch_log_section_layout.addWidget(self.lbl_batch_log_hint)

        self.batch_log_list = QListWidget()
        self.batch_log_list.setObjectName("batchLogList")
        self.batch_log_list.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection)
        self.batch_log_list.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.batch_log_list.setAlternatingRowColors(False)
        self.batch_log_list.setMinimumHeight(260)
        batch_log_section_layout.addWidget(self.batch_log_list)

        self._rebuild_batch_detail_layout("wide")
        batch_card_layout.addLayout(batch_detail_layout, stretch=1)

        batch_card_row_layout = QHBoxLayout()
        batch_card_row_layout.setContentsMargins(0, 0, 0, 0)
        batch_card_row_layout.setSpacing(0)
        batch_card_row_layout.addStretch(1)
        batch_card_row_layout.addWidget(batch_card, 10)
        batch_card_row_layout.addStretch(1)
        self.batch_card_row_layout = batch_card_row_layout
        batch_outer_layout.addLayout(batch_card_row_layout)
        batch_outer_layout.addStretch(1)

        self.merge_workspace_container = QWidget()
        merge_outer_layout = QVBoxLayout(self.merge_workspace_container)
        merge_outer_layout.setContentsMargins(26, 24, 26, 28)
        merge_outer_layout.setSpacing(0)
        self.merge_outer_layout = merge_outer_layout

        merge_card = QFrame()
        merge_card.setObjectName("mergeWorkspaceCard")
        merge_card.setMaximumWidth(1500)
        merge_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Maximum)
        self.merge_card = merge_card
        merge_card_layout = QVBoxLayout(merge_card)
        merge_card_layout.setContentsMargins(28, 24, 28, 24)
        merge_card_layout.setSpacing(14)
        self.merge_card_layout = merge_card_layout
        merge_header_layout = QHBoxLayout()
        merge_header_layout.setSpacing(12)
        self.merge_header_layout = merge_header_layout
        merge_header_text = QVBoxLayout()
        merge_header_text.setSpacing(4)
        self.merge_header_text_layout = merge_header_text
        self.lbl_merge_title = QLabel("图片正在合并为 PDF")
        self.lbl_merge_title.setObjectName("workspaceTitle")
        self.lbl_merge_subtitle = QLabel("系统会先按当前顺序生成 PDF，完成后自动进入 PDF 脱敏工作台。")
        self.lbl_merge_subtitle.setObjectName("workspaceSubtitle")
        self.lbl_merge_subtitle.setWordWrap(True)
        merge_header_text.addWidget(self.lbl_merge_title)
        merge_header_text.addWidget(self.lbl_merge_subtitle)
        self.lbl_merge_stage_badge = QLabel("等待开始")
        self.lbl_merge_stage_badge.setObjectName("batchStageBadge")
        merge_header_layout.addLayout(merge_header_text, stretch=1)
        merge_header_layout.addWidget(self.lbl_merge_stage_badge, alignment=Qt.AlignmentFlag.AlignTop)
        self.lbl_merge_meta = QLabel("当前还没有开始合并。")
        self.lbl_merge_meta.setObjectName("workspaceHint")
        self.lbl_merge_meta.setWordWrap(True)
        merge_card_layout.addLayout(merge_header_layout)
        merge_card_layout.addWidget(self.lbl_merge_meta)

        merge_stage_layout = QGridLayout()
        merge_stage_layout.setContentsMargins(0, 0, 0, 0)
        merge_stage_layout.setHorizontalSpacing(12)
        merge_stage_layout.setVerticalSpacing(12)
        self.merge_stage_layout = merge_stage_layout
        self.merge_stage_cards = []
        for title_text, note_text in [
            ("1 整理顺序", "按当前拖入顺序准备图片，确认后开始生成 PDF。"),
            ("2 合并 PDF", "系统将图片依次写入 PDF，并同步显示当前进度。"),
            ("3 进入工作台", "合并完成后自动打开生成的 PDF，继续进入脱敏工作台。"),
        ]:
            step_card = QFrame()
            step_card.setObjectName("batchStepCard")
            step_layout = QVBoxLayout(step_card)
            step_layout.setContentsMargins(14, 12, 14, 12)
            step_layout.setSpacing(4)
            step_title = QLabel(title_text)
            step_title.setObjectName("batchStepTitle")
            step_note = QLabel(note_text)
            step_note.setObjectName("batchStepNote")
            step_note.setWordWrap(True)
            step_layout.addWidget(step_title)
            step_layout.addWidget(step_note)
            self.merge_stage_cards.append((step_card, step_title, step_note))
        self._rebuild_merge_stage_layout("wide")
        merge_card_layout.addLayout(merge_stage_layout)

        merge_metrics_layout = QGridLayout()
        merge_metrics_layout.setContentsMargins(0, 0, 0, 0)
        merge_metrics_layout.setHorizontalSpacing(10)
        merge_metrics_layout.setVerticalSpacing(10)
        self.merge_metrics_layout = merge_metrics_layout
        self.merge_metric_cards = []
        self.lbl_merge_metric_images = None
        self.lbl_merge_metric_images_note = None
        self.lbl_merge_metric_status = None
        self.lbl_merge_metric_status_note = None
        self.lbl_merge_metric_next = None
        self.lbl_merge_metric_next_note = None
        for key, title_text in [
            ("images", "待合并图片"),
            ("status", "当前状态"),
            ("next", "后续动作"),
        ]:
            metric_card, metric_value, metric_note = self._create_batch_metric_card(title_text)
            setattr(self, f"lbl_merge_metric_{key}", metric_value)
            setattr(self, f"lbl_merge_metric_{key}_note", metric_note)
            self.merge_metric_cards.append(metric_card)
        self._rebuild_merge_metrics_layout("wide")
        merge_card_layout.addLayout(merge_metrics_layout)
        merge_card_row_layout = QHBoxLayout()
        merge_card_row_layout.setContentsMargins(0, 0, 0, 0)
        merge_card_row_layout.setSpacing(0)
        merge_card_row_layout.addStretch(1)
        merge_card_row_layout.addWidget(merge_card, 10)
        merge_card_row_layout.addStretch(1)
        self.merge_card_row_layout = merge_card_row_layout
        merge_outer_layout.addLayout(merge_card_row_layout)
        merge_outer_layout.addStretch(1)

        self.word_compare_container = QWidget()
        word_compare_outer_layout = QVBoxLayout(self.word_compare_container)
        word_compare_outer_layout.setContentsMargins(14, 10, 14, 16)
        word_compare_outer_layout.setSpacing(0)
        self.word_compare_outer_layout = word_compare_outer_layout

        self.word_workspace_shell = QFrame()
        self.word_workspace_shell.setObjectName("previewWorkspaceCard")
        self.word_workspace_shell.setMaximumWidth(1940)
        self.word_workspace_shell.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        word_workspace_shell_layout = QVBoxLayout(self.word_workspace_shell)
        word_workspace_shell_layout.setContentsMargins(8, 8, 8, 8)
        word_workspace_shell_layout.setSpacing(6)
        self.word_workspace_shell_layout = word_workspace_shell_layout

        self.word_compare_header = QFrame()
        self.word_compare_header.setObjectName("wordCompareHeader")
        self.word_compare_header.setFixedHeight(28)
        word_header_layout = QHBoxLayout(self.word_compare_header)
        word_header_layout.setContentsMargins(0, 0, 0, 0)
        word_header_layout.setSpacing(12)
        self.word_header_layout = word_header_layout

        self.lbl_word_original_header = QLabel("原文预览")
        self.lbl_word_original_header.setObjectName("wordCompareLabel")
        self.lbl_word_original_header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.word_header_divider = QFrame()
        self.word_header_divider.setFrameShape(QFrame.Shape.NoFrame)
        self.word_header_divider.setFixedWidth(12)
        self.lbl_word_replaced_header = QLabel("替换后预览")
        self.lbl_word_replaced_header.setObjectName("wordCompareLabel")
        self.lbl_word_replaced_header.setAlignment(Qt.AlignmentFlag.AlignCenter)

        word_header_layout.addWidget(self.lbl_word_original_header, stretch=1)
        word_header_layout.addWidget(self.word_header_divider)
        word_header_layout.addWidget(self.lbl_word_replaced_header, stretch=1)
        word_workspace_shell_layout.addWidget(self.word_compare_header)

        self.word_compare_content = QWidget()
        self.word_compare_content.setObjectName("previewStage")
        word_compare_layout = QHBoxLayout(self.word_compare_content)
        word_compare_layout.setContentsMargins(0, 0, 0, 0)
        word_compare_layout.setSpacing(8)
        self.word_compare_layout = word_compare_layout

        self.word_preview_original_panel = QWidget()
        self.word_preview_original_panel.setObjectName("wordPreviewShell")
        self.word_preview_original_panel.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        original_panel_layout = QVBoxLayout(self.word_preview_original_panel)
        original_panel_layout.setContentsMargins(1, 1, 1, 1)
        original_panel_layout.setSpacing(0)
        self.original_panel_layout = original_panel_layout
        original_panel_layout.addWidget(self.word_preview)

        self.word_preview_replaced_panel = QWidget()
        self.word_preview_replaced_panel.setObjectName("wordPreviewShell")
        self.word_preview_replaced_panel.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        replaced_panel_layout = QVBoxLayout(self.word_preview_replaced_panel)
        replaced_panel_layout.setContentsMargins(1, 1, 1, 1)
        replaced_panel_layout.setSpacing(0)
        self.replaced_panel_layout = replaced_panel_layout
        replaced_panel_layout.addWidget(self.word_preview_replaced)

        word_compare_layout.addWidget(self.word_preview_original_panel, stretch=1)
        word_compare_layout.addWidget(self.word_preview_replaced_panel, stretch=1)
        word_workspace_shell_layout.addWidget(self.word_compare_content, stretch=1)
        word_workspace_row_layout = QHBoxLayout()
        word_workspace_row_layout.setContentsMargins(0, 0, 0, 0)
        word_workspace_row_layout.setSpacing(0)
        word_workspace_row_layout.addStretch(1)
        word_workspace_row_layout.addWidget(self.word_workspace_shell, 18)
        word_workspace_row_layout.addStretch(1)
        self.word_workspace_row_layout = word_workspace_row_layout
        word_compare_outer_layout.addLayout(word_workspace_row_layout, 1)

        # 创建主容器，包含 canvas_container 和 word_compare_container
        self.main_container = QWidget()
        self.main_layout = QVBoxLayout(self.main_container)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        self.main_layout.addWidget(self.idle_workspace_container)
        self.main_layout.addWidget(self.batch_workspace_container)
        self.main_layout.addWidget(self.merge_workspace_container)
        self.main_layout.addWidget(self.canvas_container)
        self.main_layout.addWidget(self.word_compare_container)

        # 默认隐藏 Word 预览容器
        self.batch_workspace_container.hide()
        self.merge_workspace_container.hide()
        self.word_compare_container.hide()
        self.word_preview_replaced_panel.hide()
        self.word_preview.hide()
        self.word_preview_replaced.hide()

        # 设置 container 为固定的 widget
        self.scroll.setWidget(self.main_container)
        # 默认单页模式：隐藏右页
        self.canvas_right.hide()

        layout.addWidget(self.scroll)

        # 进度条和取消按钮区域（v36.3: 添加取消按钮）
        self.progress_shell = QWidget()
        progress_layout = QHBoxLayout(self.progress_shell)
        progress_layout.setContentsMargins(0, 0, 0, 0)
        progress_layout.setSpacing(8)

        self.progress = QProgressBar()
        self.progress.setFixedHeight(24)
        self.progress.setTextVisible(True)
        self.progress.setFormat("%p%")
        progress_layout.addWidget(self.progress, stretch=1)

        # 取消扫描按钮（初始隐藏）
        self.btn_cancel_scan = QPushButton("取消")
        self.btn_cancel_scan.setFixedSize(60, 24)
        self.btn_cancel_scan.setToolTip("停止扫描并保留已扫描结果")
        self.btn_cancel_scan.clicked.connect(self.cancel_ocr_scan)
        self.btn_cancel_scan.setVisible(False)  # 初始隐藏
        progress_layout.addWidget(self.btn_cancel_scan)

        layout.addWidget(self.progress_shell)

        # 应用浅色主题样式
        self._apply_light_theme()
        self._sync_ui_mode()

    def _apply_light_theme(self):
        """应用浅色主题样式（v35.1: Windows 强制浅色主题）"""
        import platform

        theme = Theme.LIGHT

        # 主窗口背景样式（确保整体色调统一）
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {theme["background"]};
                color: {theme["text"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QWidget#appRoot {{
                background-color: {theme["background"]};
                color: {theme["text"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QLabel {{
                background-color: transparent;
                color: {theme["text"]};
            }}
        """)

        # 统一上下文条
        if hasattr(self, "workbench_panel"):
            self.workbench_panel.setStyleSheet(
                f"""
                QFrame#workbenchPanel {{
                    background-color: {theme["surface"]};
                    border-bottom: 1px solid {theme["border"]};
                }}
                QLabel#workbenchTitle {{
                    color: {theme["text"]};
                    font-size: 16px;
                    font-weight: 700;
                    background-color: transparent;
                }}
                QLabel#workbenchSubtitle {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    line-height: 1.4;
                    background-color: transparent;
                }}
                QLabel#workbenchFocus {{
                    background-color: #E9F1FB;
                    color: {theme["primary"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 14px;
                    padding: 5px 12px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QLabel#workbenchHintTag {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 7px 12px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QLabel#contextMessage {{
                    background-color: {theme["info_bar"]};
                    color: {theme["text_secondary"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 7px 12px;
                    font-weight: 600;
                    font-size: {Theme.FONT_SIZE_SMALL}px;
                }}
                """
            )

        # 工具栏
        self.toolbar.setStyleSheet(f"""
            QFrame#toolbarRoot {{
                background-color: {theme["surface"]};
                border-bottom: 1px solid {theme["border"]};
            }}
            QFrame#toolbarGroup, QFrame#toolbarGroupStrong, QFrame#toolbarUtilityGroup {{
                background-color: transparent;
                border: none;
            }}
            QFrame#toolbarUtilityGroup {{
                background-color: transparent;
                border: none;
            }}
            QFrame#toolbarDivider {{
                background-color: {theme["border"]};
                border: none;
                margin-top: 6px;
                margin-bottom: 6px;
            }}
        """)

        if hasattr(self, "lbl_mode_badge"):
            self.lbl_mode_badge.setStyleSheet(
                f"""
                QLabel#modeBadge {{
                    background-color: {theme["hover"]};
                    color: {theme["primary"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 14px;
                    padding: 6px 12px;
                    font-size: {Theme.FONT_SIZE_SMALL}px;
                    font-weight: 700;
                }}
                """
            )
        if hasattr(self, "main_container"):
            self.main_container.setStyleSheet(
                f"""
                QFrame#workspaceCard {{
                    background-color: {theme["surface"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 18px;
                }}
                QFrame#batchWorkspaceCard, QFrame#mergeWorkspaceCard {{
                    background-color: {theme["surface"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 18px;
                }}
                QFrame#batchDetailSection {{
                    background-color: #FBFCFE;
                    border: 1px solid {theme["border"]};
                    border-radius: 16px;
                }}
                QFrame#previewWorkspaceCard {{
                    background-color: #FBFCFE;
                    border: none;
                    border-radius: 18px;
                }}
                QWidget#previewStage {{
                    background-color: transparent;
                    border: none;
                }}
                QWidget#wordPreviewShell {{
                    background-color: #FFFFFF;
                    border: 1px solid {theme["border"]};
                    border-radius: 15px;
                }}
                QLabel#pdfPageCanvas {{
                    background-color: transparent;
                    border: none;
                    border-radius: 0px;
                }}
                QLabel#workspaceTitle, QLabel#batchTitle {{
                    color: {theme["text"]};
                    font-size: 20px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#workspaceSubtitle, QLabel#batchSubtitle {{
                    color: {theme["text_secondary"]};
                    font-size: 12px;
                    line-height: 1.7;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#workspaceLead {{
                    color: {theme["text_secondary"]};
                    font-size: 12px;
                    font-weight: 600;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#workspaceHint, QLabel#batchMeta, QLabel#batchCurrentFile {{
                    color: {theme["text"]};
                    font-size: 12px;
                    line-height: 1.7;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSectionLabel {{
                    color: {theme["text"]};
                    font-size: 13px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSectionHint {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    font-weight: 600;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleHeroBadge {{
                    color: {theme["primary"]};
                    background-color: #E9F1FB;
                    border: 1px solid {theme["border"]};
                    border-radius: 9px;
                    padding: 4px 9px;
                    font-size: 10px;
                    font-weight: 700;
                }}
                QFrame#idleHeroPanel {{
                    background-color: transparent;
                    border: none;
                }}
                QFrame#idleSectionPanel {{
                    background-color: transparent;
                    border: none;
                }}
                QFrame#idleFlowPanel {{
                    background-color: transparent;
                    border: none;
                }}
                QLabel#idleFlowTitle {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QFrame#idleActionPanel {{
                    background-color: transparent;
                    border: none;
                }}
                QFrame#idleStartCard, QFrame#idleSupportCard {{
                    background-color: #F9FBFE;
                    border: 1px solid {theme["border"]};
                    border-radius: 18px;
                }}
                QLabel#idleStartTitle {{
                    color: {theme["text"]};
                    font-size: 13px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleStartText {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    line-height: 1.6;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleDropHint {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    font-weight: 600;
                    line-height: 1.6;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSupportTitle {{
                    color: {theme["text"]};
                    font-size: 13px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSupportText {{
                    color: {theme["text"]};
                    font-size: 12px;
                    font-weight: 600;
                    line-height: 1.5;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSupportMeta {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    font-weight: 600;
                    line-height: 1.5;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSupportEmail {{
                    color: {theme["primary"]};
                    font-size: 11px;
                    font-weight: 600;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#idleSupportNote {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    line-height: 1.6;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QFrame#batchStepCard, QFrame#batchMetricCard {{
                    background-color: #FCFDFF;
                    border: 1px solid {theme["border"]};
                    border-radius: 15px;
                }}
                QLabel#batchStepTitle {{
                    color: {theme["text"]};
                    font-size: 12px;
                    font-weight: 700;
                    background-color: transparent;
                }}
                QLabel#batchStepNote, QLabel#batchMetricNote {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    line-height: 1.6;
                    background-color: transparent;
                }}
                QLabel#batchMetricTitle {{
                    color: {theme["text_secondary"]};
                    font-size: 11px;
                    font-weight: 700;
                    background-color: transparent;
                }}
                QLabel#batchMetricValue {{
                    color: {theme["text"]};
                    font-size: 18px;
                    font-weight: 700;
                    background-color: transparent;
                }}
                QFrame#routeCard {{
                    background-color: #FCFDFF;
                    border: 1px solid {theme["border"]};
                    border-radius: 16px;
                }}
                QFrame#routeCard[routeTone="pdf"] {{
                    background-color: #F7FBFF;
                    border-color: #D7E8F8;
                }}
                QFrame#routeCard[routeTone="word"] {{
                    background-color: #F6FBF8;
                    border-color: #D9EEE3;
                }}
                QFrame#routeCard[routeTone="batch"] {{
                    background-color: #FFF9F2;
                    border-color: #F1E1C7;
                }}
                QFrame#routeCard[routeTone="image"] {{
                    background-color: #F8FAFC;
                    border-color: #E0E7EF;
                }}
                QFrame#routeCardAccent {{
                    border: none;
                    border-radius: 2px;
                    background-color: {theme["primary"]};
                    min-height: 4px;
                    max-height: 4px;
                }}
                QFrame#routeCardAccent[routeAccent="pdf"] {{
                    background-color: #1177CC;
                }}
                QFrame#routeCardAccent[routeAccent="word"] {{
                    background-color: #1FA971;
                }}
                QFrame#routeCardAccent[routeAccent="batch"] {{
                    background-color: #D9822B;
                }}
                QFrame#routeCardAccent[routeAccent="image"] {{
                    background-color: #5B6B7A;
                }}
                QLabel#routeCardTitle {{
                    color: {theme["text"]};
                    font-size: 13px;
                    font-weight: 700;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#routeCardMeta {{
                    color: {theme["text_secondary"]};
                    background-color: {theme["hover"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 9px;
                    padding: 3px 8px;
                    min-width: 56px;
                    font-size: 10px;
                    font-weight: 700;
                }}
                QLabel#routeCardMeta[routeTone="pdf"] {{
                    color: #1177CC;
                    background-color: #E9F3FD;
                    border-color: #D5E8FA;
                }}
                QLabel#routeCardMeta[routeTone="word"] {{
                    color: #1B9263;
                    background-color: #EAF6F0;
                    border-color: #D8EDE2;
                }}
                QLabel#routeCardMeta[routeTone="batch"] {{
                    color: #B56A1C;
                    background-color: #FFF2E3;
                    border-color: #F1DEC3;
                }}
                QLabel#routeCardMeta[routeTone="image"] {{
                    color: #5B6B7A;
                    background-color: #EEF3F8;
                    border-color: #DEE7F0;
                }}
                QLabel#routeCardText {{
                    color: {theme["text_secondary"]};
                    font-size: 10px;
                    line-height: 1.6;
                    background-color: transparent;
                    border: none;
                    padding: 0;
                }}
                QLabel#workspaceTag {{
                    background-color: {theme["hover"]};
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 6px 12px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QLabel#batchStageBadge {{
                    background-color: #FFF6EA;
                    color: {theme["warning"]};
                    border: 1px solid #F0DFC4;
                    border-radius: 14px;
                    padding: 6px 12px;
                    font-size: 12px;
                    font-weight: 700;
                }}
                QLabel#batchLogHint, QLabel#batchSectionLabel {{
                    color: {theme["text_secondary"]};
                    font-size: 12px;
                    font-weight: 700;
                    background-color: transparent;
                }}
                QLabel#batchResultMeta {{
                    color: {theme["text_secondary"]};
                    font-size: 12px;
                    font-weight: 600;
                    background-color: transparent;
                }}
                QTableWidget#batchResultTable {{
                    background-color: #FFFFFF;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 15px;
                    gridline-color: transparent;
                    padding: 6px;
                }}
                QTableWidget#batchResultTable::item {{
                    padding: 10px 12px;
                    border-bottom: 1px solid #E9EEF5;
                }}
                QTableWidget#batchResultTable QHeaderView::section {{
                    background-color: #F7FAFD;
                    color: {theme["text_secondary"]};
                    border: none;
                    border-bottom: 1px solid {theme["border"]};
                    padding: 8px 10px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QListWidget#batchLogList {{
                    background-color: #FFFFFF;
                    border: 1px solid {theme["border"]};
                    border-radius: 15px;
                    padding: 6px;
                    outline: none;
                }}
                QListWidget#batchLogList::item {{
                    padding: 10px 12px;
                    border-bottom: 1px solid #E9EEF5;
                    color: {theme["text"]};
                }}
                QListWidget#batchLogList::item:selected {{
                    background-color: transparent;
                    color: {theme["text"]};
                }}
                QTextBrowser#batchSummaryBrowser {{
                    background-color: #FFFFFF;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 15px;
                    padding: 12px;
                    font-size: 12px;
                    line-height: 1.7;
                }}
                QLabel#workflowStep {{
                    color: {theme["text_secondary"]};
                    background-color: {theme["hover"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 5px 10px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                """
            )

        # 滚动区域
        self.scroll.setStyleSheet(self.scroll_style.format(theme["scroll_area"]))

        # Word 双栏预览头部（紧凑样式）
        if hasattr(self, "word_compare_header"):
            self.word_compare_header.setStyleSheet(f"""
                QFrame#wordCompareHeader {{
                    background-color: transparent;
                    border: none;
                }}
            """)
        if hasattr(self, "lbl_word_original_header"):
            self.lbl_word_original_header.setStyleSheet(
                f"""
                QLabel#wordCompareLabel {{
                    color: {theme['text_secondary']};
                    background-color: #F7FAFD;
                    border: 1px solid {theme["border"]};
                    border-radius: 12px;
                    padding: 4px 12px;
                    font-size: 12px;
                    font-weight: 700;
                }}
                """
            )
        if hasattr(self, "lbl_word_replaced_header"):
            self.lbl_word_replaced_header.setStyleSheet(
                f"""
                QLabel#wordCompareLabel {{
                    color: {theme['text_secondary']};
                    background-color: #F7FAFD;
                    border: 1px solid {theme["border"]};
                    border-radius: 12px;
                    padding: 4px 12px;
                    font-size: 12px;
                    font-weight: 700;
                }}
                """
            )
        if hasattr(self, "word_header_divider"):
            self.word_header_divider.setStyleSheet("background-color: transparent; border: none;")

        # 标签文本颜色
        text_color = theme["text"]
        toolbar_meta_style = f"""
            QLabel#toolbarMeta {{
                color: {text_color};
                background-color: {theme["hover"]};
                border: 1px solid {theme["border"]};
                border-radius: 9px;
                padding: 4px 10px;
                font-size: 12px;
                font-weight: 700;
                min-width: 54px;
            }}
        """
        self.lbl_zoom.setStyleSheet(toolbar_meta_style)
        self.lbl_page.setStyleSheet(toolbar_meta_style)

        for toggle_btn in [self.rb_black, self.rb_white, self.cb_dual]:
            self._apply_button_variant(toggle_btn, "toggle")

        self._refresh_toolbar_more_button_style()

        # 进度条
        self.progress.setStyleSheet(f"""
            QProgressBar {{
                background-color: #E5E5EA;
                border: none;
                border-radius: 3px;
                color: #1D1D1F;
                font-weight: 600;
                font-size: 12px;
                text-align: center;
            }}
            QProgressBar::chunk {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #007AFF, stop:1 #34C759);
                border-radius: 3px;
            }}
        """)
        self._refresh_mode_badge()
        self._refresh_word_compare_toggle()
        self._refresh_workbench_context()

    def _refresh_toolbar_more_button_style(self):
        """刷新更多按钮样式，保持响应式和隐藏菜单箭头逻辑一致。"""
        if not hasattr(self, "btn_more"):
            return
        self.btn_more.setStyleSheet(
            self._get_button_style("secondary") +
            """
            QPushButton#toolbarMoreButton {
                padding-right: 16px;
            }
            QPushButton#toolbarMoreButton::menu-indicator {
                image: none;
                width: 0px;
                height: 0px;
            }
            """
        )

    def _apply_native_toolbar_icons(self):
        """为导航按钮应用 Qt 原生图标，提升 Windows 下的一致性。"""
        if not hasattr(self, "btn_go_first"):
            return

        style = self.style()
        icon_map = [
            (self.btn_go_first, QStyle.StandardPixmap.SP_MediaSkipBackward),
            (self.btn_prev_page, QStyle.StandardPixmap.SP_ArrowBack),
            (self.btn_next_page, QStyle.StandardPixmap.SP_ArrowForward),
            (self.btn_go_last, QStyle.StandardPixmap.SP_MediaSkipForward),
        ]
        for button, standard_icon in icon_map:
            button.setIcon(style.standardIcon(standard_icon))
            button.setText("")

    def _set_status_badge_style(self, label, fg, bg):
        """为轻量状态标签设置统一的胶囊样式。"""
        if not label:
            return
        badge_font_size = getattr(self, "_status_badge_font_size", 12)
        label.setStyleSheet(
            f"""
            QLabel {{
                background-color: {bg};
                color: {fg};
                border: 1px solid {Theme.LIGHT["border"]};
                border-radius: 14px;
                padding: 6px 12px;
                font-size: {badge_font_size}px;
                font-weight: 700;
            }}
            """
        )

    def _set_info_bar_message(self, message):
        """设置顶部临时任务提示；空值时按当前模式决定是否隐藏。"""
        self.info_bar_message = str(message).strip() if message else ""
        self._refresh_info_bar_visibility()

    def _clear_info_bar_message(self):
        """清空顶部临时任务提示。"""
        self.info_bar_message = ""
        self._refresh_info_bar_visibility()

    def _refresh_info_bar_visibility(self):
        """避免顶部提示条与工作台摘要重复。"""
        if not hasattr(self, "info_bar"):
            return

        if self.current_ui_mode == "idle":
            if self.info_bar_message:
                self.info_bar.setText(self.info_bar_message)
                self.info_bar.setVisible(True)
            else:
                self.info_bar.setVisible(False)
            return

        if self.info_bar_message:
            self.info_bar.setText(self.info_bar_message)
            self.info_bar.setVisible(True)
        else:
            self.info_bar.setVisible(False)

    def _rebuild_idle_action_layout(self, density_mode):
        """让首页主动作与支持入口在桌面端双列，窄窗口再收成单列。"""
        if not hasattr(self, "idle_action_buttons_layout"):
            return

        layout = self.idle_action_buttons_layout
        while layout.count():
            layout.takeAt(0)

        if density_mode == "narrow":
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 0)
            layout.addWidget(self.idle_start_card, 0, 0)
            layout.addWidget(self.idle_support_card, 1, 0)
        else:
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 1)
            layout.addWidget(self.idle_start_card, 0, 0)
            layout.addWidget(self.idle_support_card, 0, 1)

    def _rebuild_idle_support_actions_layout(self, density_mode):
        """重排首页支持区动作按钮，宽窗口横排，窄窗口改为 2+1。"""
        if not hasattr(self, "idle_support_actions_layout"):
            return

        layout = self.idle_support_actions_layout
        while layout.count():
            layout.takeAt(0)

        buttons = [
            getattr(self, "btn_idle_feedback", None),
            getattr(self, "btn_idle_manual", None),
            getattr(self, "btn_idle_donate", None),
        ]
        buttons = [button for button in buttons if button]
        if not buttons:
            return

        if density_mode == "narrow":
            for column in range(3):
                layout.setColumnStretch(column, 0)
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 1)
            layout.addWidget(buttons[0], 0, 0)
            if len(buttons) > 1:
                layout.addWidget(buttons[1], 0, 1)
            if len(buttons) > 2:
                layout.addWidget(buttons[2], 1, 0, 1, 2)
        else:
            for column in range(3):
                layout.setColumnStretch(column, 1)
            for index, button in enumerate(buttons):
                layout.addWidget(button, 0, index)

    def _rebuild_idle_route_layout(self, density_mode):
        """重排首页入口卡，保证桌面端严格 2x2 对齐。"""
        if not hasattr(self, "idle_routes_layout"):
            return

        layout = self.idle_routes_layout
        while layout.count():
            layout.takeAt(0)

        if density_mode == "narrow":
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 0)
            for row in range(max(1, len(getattr(self, "idle_route_cards", [])) + 1)):
                layout.setRowStretch(row, 0)
            for index, route_card in enumerate(getattr(self, "idle_route_cards", [])):
                layout.addWidget(route_card, index, 0)
        else:
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 1)
            layout.setRowStretch(0, 0)
            layout.setRowStretch(1, 0)
            for index, route_card in enumerate(getattr(self, "idle_route_cards", [])):
                layout.addWidget(route_card, index // 2, index % 2)

    def _rebuild_batch_action_layout(self, density_mode):
        """重排批量页动作区，宽窗口横排，中窗口两列，窄窗口单列。"""
        if not hasattr(self, "batch_actions_layout"):
            return

        layout = self.batch_actions_layout
        while layout.count():
            layout.takeAt(0)

        buttons = list(getattr(self, "batch_action_buttons", []))
        if not buttons:
            return

        for column in range(max(4, len(buttons))):
            layout.setColumnStretch(column, 0)
        for row in range(max(4, len(buttons))):
            layout.setRowStretch(row, 0)

        if density_mode == "narrow":
            layout.setColumnStretch(0, 1)
            for index, button in enumerate(buttons):
                layout.addWidget(button, index, 0)
        elif density_mode == "compact":
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 1)
            for index, button in enumerate(buttons):
                layout.addWidget(button, index // 2, index % 2)
        else:
            for column in range(len(buttons)):
                layout.setColumnStretch(column, 1)
            for index, button in enumerate(buttons):
                layout.addWidget(button, 0, index)

    def _rebuild_batch_stage_layout(self, density_mode):
        """重排批量页阶段卡，桌面端横排，中窗口两列，窄窗口单列。"""
        if not hasattr(self, "batch_stage_layout"):
            return

        layout = self.batch_stage_layout
        while layout.count():
            layout.takeAt(0)

        cards = list(getattr(self, "batch_stage_cards", []))
        if not cards:
            return

        if density_mode == "narrow":
            columns = 1
        elif density_mode == "compact":
            columns = 2
        else:
            columns = 3

        for column in range(max(columns, len(cards))):
            layout.setColumnStretch(column, 0)
        for row in range(len(cards)):
            layout.setRowStretch(row, 0)

        for index, (frame, _title_label, _note_label) in enumerate(cards):
            row = index // columns
            column = index % columns
            layout.addWidget(frame, row, column)

        for column in range(columns):
            layout.setColumnStretch(column, 1)

    def _rebuild_batch_metrics_layout(self, density_mode):
        """重排批量页指标卡，宽窗口四列，中窗口两列，窄窗口单列。"""
        if not hasattr(self, "batch_metrics_layout"):
            return

        layout = self.batch_metrics_layout
        while layout.count():
            layout.takeAt(0)

        cards = list(getattr(self, "batch_metric_cards", []))
        if not cards:
            return

        if density_mode == "narrow":
            columns = 1
        elif density_mode == "compact":
            columns = 2
        else:
            columns = 4

        for column in range(max(columns, len(cards))):
            layout.setColumnStretch(column, 0)
        for row in range(len(cards)):
            layout.setRowStretch(row, 0)

        for index, card in enumerate(cards):
            row = index // columns
            column = index % columns
            layout.addWidget(card, row, column)

        for column in range(columns):
            layout.setColumnStretch(column, 1)

    def _rebuild_merge_stage_layout(self, density_mode):
        """重排图片合并阶段卡，桌面端横排，中窗口两列，窄窗口单列。"""
        if not hasattr(self, "merge_stage_layout"):
            return

        layout = self.merge_stage_layout
        while layout.count():
            layout.takeAt(0)

        cards = list(getattr(self, "merge_stage_cards", []))
        if not cards:
            return

        if density_mode == "narrow":
            columns = 1
        elif density_mode == "compact":
            columns = 2
        else:
            columns = 3

        for column in range(max(columns, len(cards))):
            layout.setColumnStretch(column, 0)
        for row in range(len(cards)):
            layout.setRowStretch(row, 0)

        for index, (frame, _title_label, _note_label) in enumerate(cards):
            row = index // columns
            column = index % columns
            layout.addWidget(frame, row, column)

        for column in range(columns):
            layout.setColumnStretch(column, 1)

    def _rebuild_merge_metrics_layout(self, density_mode):
        """重排图片合并指标卡，宽窗口三列，中窗口两列，窄窗口单列。"""
        if not hasattr(self, "merge_metrics_layout"):
            return

        layout = self.merge_metrics_layout
        while layout.count():
            layout.takeAt(0)

        cards = list(getattr(self, "merge_metric_cards", []))
        if not cards:
            return

        if density_mode == "narrow":
            columns = 1
        elif density_mode == "compact":
            columns = 2
        else:
            columns = 3

        for column in range(max(columns, len(cards))):
            layout.setColumnStretch(column, 0)
        for row in range(len(cards)):
            layout.setRowStretch(row, 0)

        for index, card in enumerate(cards):
            row = index // columns
            column = index % columns
            layout.addWidget(card, row, column)

        for column in range(columns):
            layout.setColumnStretch(column, 1)

    def _rebuild_batch_detail_layout(self, density_mode):
        """重排批量页摘要/结果/日志区，宽窗口双区，窄窗口单列。"""
        if not hasattr(self, "batch_detail_layout"):
            return

        layout = self.batch_detail_layout
        while layout.count():
            layout.takeAt(0)

        summary_section = getattr(self, "batch_summary_section", None)
        result_section = getattr(self, "batch_result_section", None)
        log_section = getattr(self, "batch_log_section", None)
        if not summary_section or not result_section or not log_section:
            return

        for column in range(3):
            layout.setColumnStretch(column, 0)
        for row in range(3):
            layout.setRowStretch(row, 0)

        if density_mode == "narrow":
            layout.setHorizontalSpacing(0)
            layout.setVerticalSpacing(12)
            layout.setColumnStretch(0, 1)
            layout.setColumnStretch(1, 0)
            layout.addWidget(summary_section, 0, 0)
            layout.addWidget(result_section, 1, 0)
            layout.addWidget(log_section, 2, 0)
            layout.setRowStretch(0, 0)
            layout.setRowStretch(1, 1)
            layout.setRowStretch(2, 1)
        elif density_mode == "compact":
            layout.setHorizontalSpacing(12)
            layout.setVerticalSpacing(12)
            layout.setColumnStretch(0, 11)
            layout.setColumnStretch(1, 13)
            layout.addWidget(summary_section, 0, 0)
            layout.addWidget(result_section, 0, 1)
            layout.addWidget(log_section, 1, 0, 1, 2)
            layout.setRowStretch(0, 1)
            layout.setRowStretch(1, 1)
        else:
            layout.setHorizontalSpacing(14 if density_mode == "wide" else 12)
            layout.setVerticalSpacing(14 if density_mode == "wide" else 12)
            layout.setColumnStretch(0, 8)
            layout.setColumnStretch(1, 18)
            layout.addWidget(summary_section, 0, 0)
            layout.addWidget(log_section, 1, 0)
            layout.addWidget(result_section, 0, 1, 2, 1)
            layout.setRowStretch(0, 0)
            layout.setRowStretch(1, 1)

    def _get_display_scale_factor(self):
        """返回当前显示环境的缩放因子，主要用于 Windows DPI 收口。"""
        import platform

        if platform.system() != "Windows":
            return 1.0

        screen = None
        try:
            handle = self.windowHandle()
            if handle:
                screen = handle.screen()
        except Exception:
            screen = None

        if screen is None:
            screen = QApplication.primaryScreen()
        if screen is None:
            return 1.0

        try:
            scale = screen.logicalDotsPerInch() / 96.0
        except Exception:
            scale = 1.0

        return max(1.0, min(scale, 2.0))

    def _refresh_windows_density_metrics(self, density_mode):
        """按 Windows DPI 和当前工具栏密度收口高度、间距与命中区。"""
        if not hasattr(self, "toolbar") or not hasattr(self, "toolbar_layout"):
            return

        scale = self._get_display_scale_factor()
        theme = Theme.LIGHT

        if scale >= 1.5:
            toolbar_height = 68
            button_height = 40
            icon_size = 38
            side_margin = 12
            vertical_margin = 8
            spacing = 7
            workbench_h_margin = 16
            workbench_v_margin = 11
            title_font_size = 17
            subtitle_font_size = 12
            chip_font_size = 12
            meta_font_size = 12
            workflow_font_size = 12
            compare_header_height = 34
            compare_header_font_size = 12
            progress_height = 28
            cancel_width = 72
            batch_result_row_height = 48
        elif scale >= 1.25:
            toolbar_height = 64
            button_height = 38
            icon_size = 36
            side_margin = 14
            vertical_margin = 8
            spacing = 8
            workbench_h_margin = 17
            workbench_v_margin = 11
            title_font_size = 16
            subtitle_font_size = 12
            chip_font_size = 12
            meta_font_size = 12
            workflow_font_size = 11
            compare_header_height = 32
            compare_header_font_size = 12
            progress_height = 26
            cancel_width = 68
            batch_result_row_height = 44
        else:
            toolbar_height = 58
            button_height = 36
            icon_size = 32
            side_margin = 16
            vertical_margin = 7
            spacing = 10
            workbench_h_margin = 18
            workbench_v_margin = 12
            title_font_size = 16
            subtitle_font_size = 11
            chip_font_size = 11
            meta_font_size = 12
            workflow_font_size = 11
            compare_header_height = 30
            compare_header_font_size = 12
            progress_height = 24
            cancel_width = 60
            batch_result_row_height = 40

        current_mode = getattr(self, "current_ui_mode", "idle")

        logical_height = max(760, self.height())
        logical_width = max(1200, self.width())
        is_tall_workspace_window = logical_height >= 980
        is_short_workspace_window = logical_height <= 760
        is_very_wide_workspace_window = logical_width >= 1760
        is_ultra_wide_workspace_window = logical_width >= 2140
        is_cinema_wide_workspace_window = logical_width >= 2460
        is_workspace_mode = current_mode in {"pdf", "word", "batch", "image_merge"}

        if density_mode == "narrow":
            side_margin = max(10, side_margin - 2)
            spacing = max(6, spacing - 1)
            title_font_size = max(15, title_font_size - 1)
        elif density_mode == "compact":
            side_margin = max(11, side_margin - 1)
            spacing = max(7, spacing - 1)

        if scale >= 1.5:
            toolbar_height += 2
            button_height += 2
            progress_height += 2
            batch_result_row_height += 2
        elif scale >= 1.25:
            toolbar_height += 1
            button_height += 1

        if is_tall_workspace_window:
            toolbar_height += 2
            vertical_margin += 1
            spacing += 1
            title_font_size += 1
            subtitle_font_size += 1 if scale >= 1.25 else 0
        elif is_short_workspace_window:
            toolbar_height = max(56, toolbar_height - 2)
            button_height = max(34, button_height - 1)
            vertical_margin = max(6, vertical_margin - 1)
            spacing = max(6, spacing - 1)
            workbench_v_margin = max(10, workbench_v_margin - 1)

        is_pdf_like_mode = current_mode in {"pdf", "image_merge"}

        if current_mode == "idle":
            toolbar_height = max(42, toolbar_height - 10)
            vertical_margin = max(5, vertical_margin - 2)
            spacing = max(6, spacing - 1)
        elif is_pdf_like_mode:
            toolbar_height += 4
            vertical_margin += 1

        self._workflow_step_font_size = workflow_font_size
        self._status_badge_font_size = chip_font_size
        self._toolbar_meta_font_size = meta_font_size
        self._button_density_metrics = {
            "button_font_size": 14 if scale >= 1.5 else 13,
            "button_padding_v": 8 if scale >= 1.5 else 7,
            "button_padding_h": 15 if scale >= 1.5 else 14,
            "icon_font_size": 16 if scale >= 1.25 else 14,
            "icon_padding_v": 5 if scale >= 1.25 else 4,
            "icon_padding_h": 10 if scale >= 1.25 else 8,
            "icon_min": 34 if scale >= 1.5 else (32 if scale >= 1.25 else 28),
        }
        if density_mode == "narrow":
            self._button_density_metrics["button_padding_h"] = max(12, self._button_density_metrics["button_padding_h"] - 2)
        elif density_mode == "compact":
            self._button_density_metrics["button_padding_h"] = max(13, self._button_density_metrics["button_padding_h"] - 1)

        toolbar_bottom_margin = vertical_margin + (4 if is_pdf_like_mode else (2 if current_mode != "idle" else 1))
        toolbar_extra_height = 5 if is_pdf_like_mode else (2 if current_mode != "idle" else 0)
        self.toolbar.setFixedHeight(toolbar_height + toolbar_extra_height)
        self.toolbar_layout.setContentsMargins(side_margin, vertical_margin, side_margin, toolbar_bottom_margin)
        self.toolbar_layout.setSpacing(max(10, spacing + 1))

        group_spacing = max(4, spacing - 4)
        group_margin_h = 0
        group_margin_v = 2 if is_pdf_like_mode else 1
        utility_spacing = group_spacing
        for group_layout in [
            getattr(self, "toolbar_primary_layout", None),
            getattr(self, "toolbar_word_layout", None),
            getattr(self, "toolbar_pdf_layout", None),
            getattr(self, "toolbar_zoom_layout", None),
            getattr(self, "toolbar_nav_layout", None),
        ]:
            if group_layout:
                group_layout.setContentsMargins(group_margin_h, group_margin_v, group_margin_h, group_margin_v)
                group_layout.setSpacing(group_spacing)
        if hasattr(self, "toolbar_utility_layout"):
            self.toolbar_utility_layout.setContentsMargins(0, group_margin_v, 0, group_margin_v)
            self.toolbar_utility_layout.setSpacing(utility_spacing)

        if hasattr(self, "workbench_layout"):
            self.workbench_layout.setContentsMargins(workbench_h_margin, workbench_v_margin, workbench_h_margin, workbench_v_margin)
            self.workbench_layout.setSpacing(6 if density_mode == "wide" else 5)
        if hasattr(self, "context_top_layout"):
            self.context_top_layout.setSpacing(14 if density_mode == "wide" else 10)
        if hasattr(self, "workbench_text_layout"):
            self.workbench_text_layout.setSpacing(4 if density_mode == "wide" else 3)

        workspace_stage_margin = 14 if density_mode == "wide" else (10 if density_mode == "compact" else 8)
        workspace_stage_top_margin = 10 if density_mode == "wide" else (8 if density_mode == "compact" else 6)
        workspace_stage_bottom_margin = 16 if density_mode == "wide" else (14 if density_mode == "compact" else 10)
        preview_shell_padding = 8 if density_mode == "wide" else (6 if density_mode == "compact" else 5)
        preview_content_padding = 10 if density_mode == "wide" else (8 if density_mode == "compact" else 6)
        preview_content_spacing = 12 if density_mode == "wide" else (10 if density_mode == "compact" else 8)
        compare_header_gap = 10 if density_mode == "wide" else 8
        batch_card_padding_h = 28 if density_mode == "wide" else (24 if density_mode == "compact" else 20)
        batch_card_padding_v = 24 if density_mode == "wide" else (20 if density_mode == "compact" else 18)
        batch_detail_padding_h = 18 if density_mode == "wide" else (16 if density_mode == "compact" else 14)
        batch_detail_padding_v = 16 if density_mode == "wide" else (14 if density_mode == "compact" else 12)
        batch_section_gap = 14 if density_mode == "wide" else (12 if density_mode == "compact" else 10)
        batch_minor_gap = 10 if density_mode == "wide" else (8 if density_mode == "compact" else 6)
        if is_workspace_mode:
            workspace_stage_margin = max(4, workspace_stage_margin - 6)
            workspace_stage_bottom_margin = max(8, workspace_stage_bottom_margin - 2)
        if is_very_wide_workspace_window:
            workspace_stage_margin = max(2, workspace_stage_margin - 2)
            batch_card_padding_h = max(20, batch_card_padding_h - 2)
            batch_detail_padding_h = max(14, batch_detail_padding_h - 1)
        if is_ultra_wide_workspace_window:
            workspace_stage_margin = max(2, workspace_stage_margin - 2)
            preview_shell_padding = max(3, preview_shell_padding - 1)
            preview_content_padding = max(4, preview_content_padding - 1)
            batch_card_padding_h = max(18, batch_card_padding_h - 2)
            batch_section_gap = max(10, batch_section_gap - 1)
        if is_cinema_wide_workspace_window:
            workspace_stage_margin = max(1, workspace_stage_margin - 1)
            preview_shell_padding = max(3, preview_shell_padding - 1)
            preview_content_padding = max(4, preview_content_padding - 1)
            batch_card_padding_h = max(16, batch_card_padding_h - 2)
            batch_detail_padding_h = max(12, batch_detail_padding_h - 1)
            batch_section_gap = max(9, batch_section_gap - 1)
        if is_workspace_mode and is_very_wide_workspace_window:
            workbench_h_margin = max(14, workbench_h_margin - 2)
        if is_workspace_mode and is_ultra_wide_workspace_window:
            workbench_h_margin = max(12, workbench_h_margin - 2)
        if is_workspace_mode and is_cinema_wide_workspace_window:
            workbench_h_margin = max(10, workbench_h_margin - 2)
        if getattr(self, "current_ui_mode", "idle") in {"pdf", "word"}:
            workspace_stage_top_margin = max(4, workspace_stage_top_margin - 2)
            preview_shell_padding = max(4, preview_shell_padding - 1)
            preview_content_padding = max(5, preview_content_padding - 2)
        if getattr(self, "current_ui_mode", "idle") == "pdf":
            workspace_stage_top_margin = max(3, workspace_stage_top_margin - 1)
            preview_shell_padding = max(3, preview_shell_padding - 1)
            preview_content_padding = max(4, preview_content_padding - 1)
        if is_tall_workspace_window and getattr(self, "current_ui_mode", "idle") in {"pdf", "word"}:
            workspace_stage_top_margin = max(2, workspace_stage_top_margin - 1)
            workspace_stage_bottom_margin = max(6, workspace_stage_bottom_margin - 2)
            preview_shell_padding = max(3, preview_shell_padding - 1)
            preview_content_padding = max(4, preview_content_padding - 1)
            preview_content_spacing = max(8, preview_content_spacing - 1)
        if is_tall_workspace_window and getattr(self, "current_ui_mode", "idle") == "batch":
            workspace_stage_top_margin = max(4, workspace_stage_top_margin - 1)
            workspace_stage_bottom_margin = max(8, workspace_stage_bottom_margin - 2)
            batch_card_padding_v = max(18, batch_card_padding_v - 2)
            batch_section_gap = max(10, batch_section_gap - 2)
        preview_top_padding = max(4, preview_content_padding - 1)
        if hasattr(self, "pdf_workspace_outer_layout"):
            self.pdf_workspace_outer_layout.setContentsMargins(
                workspace_stage_margin,
                workspace_stage_top_margin,
                workspace_stage_margin,
                workspace_stage_bottom_margin,
            )
        if hasattr(self, "word_compare_outer_layout"):
            self.word_compare_outer_layout.setContentsMargins(
                workspace_stage_margin,
                workspace_stage_top_margin,
                workspace_stage_margin,
                workspace_stage_bottom_margin,
            )
        if hasattr(self, "batch_outer_layout"):
            self.batch_outer_layout.setContentsMargins(
                workspace_stage_margin,
                workspace_stage_top_margin,
                workspace_stage_margin,
                workspace_stage_bottom_margin,
            )
        if hasattr(self, "merge_outer_layout"):
            self.merge_outer_layout.setContentsMargins(
                workspace_stage_margin,
                workspace_stage_top_margin,
                workspace_stage_margin,
                workspace_stage_bottom_margin,
            )
        preview_shell_cap = 2560 if density_mode == "wide" else (2240 if density_mode == "compact" else 1680)
        preview_shell_min = 1040
        if getattr(self, "current_ui_mode", "idle") == "word" and getattr(self, "word_compare_mode", False):
            preview_shell_cap = 2580 if density_mode == "wide" else (2260 if density_mode == "compact" else 1760)
            preview_shell_min = 1420 if density_mode == "wide" else (1220 if density_mode == "compact" else 980)
        elif getattr(self, "current_ui_mode", "idle") == "word" and not getattr(self, "word_compare_mode", False):
            preview_shell_cap = 2580 if density_mode == "wide" else (2220 if density_mode == "compact" else 1700)
            preview_shell_min = 1360 if density_mode == "wide" else (1160 if density_mode == "compact" else 940)
        elif getattr(self, "current_ui_mode", "idle") == "pdf":
            preview_shell_cap = 2600 if density_mode == "wide" else (2280 if density_mode == "compact" else 1760)
        if is_ultra_wide_workspace_window:
            preview_shell_cap += 140
        elif is_very_wide_workspace_window:
            preview_shell_cap += 80
        if is_cinema_wide_workspace_window:
            preview_shell_cap += 220
        preview_shell_width = min(
            preview_shell_cap,
            max(preview_shell_min, logical_width - (workspace_stage_margin * 2) - 8),
        )
        workspace_card_width = min(
            2280 if density_mode == "wide" else (1980 if density_mode == "compact" else 1500),
            max(1040, logical_width - (workspace_stage_margin * 2) - 16),
        )
        if is_ultra_wide_workspace_window:
            workspace_card_width = min(workspace_card_width + 140, logical_width - (workspace_stage_margin * 2) - 8)
        elif is_very_wide_workspace_window:
            workspace_card_width = min(workspace_card_width + 80, logical_width - (workspace_stage_margin * 2) - 8)
        if is_cinema_wide_workspace_window:
            workspace_card_width = min(workspace_card_width + 220, logical_width - (workspace_stage_margin * 2) - 8)
        if hasattr(self, "pdf_workspace_shell"):
            self.pdf_workspace_shell.setMaximumWidth(preview_shell_width)
        if hasattr(self, "pdf_workspace_row_layout"):
            center_stretch = 28 if density_mode == "wide" else (22 if density_mode == "compact" else 16)
            if is_very_wide_workspace_window:
                center_stretch += 4
            if is_ultra_wide_workspace_window:
                center_stretch += 2
            if is_cinema_wide_workspace_window:
                center_stretch += 4
            side_stretch = 0 if is_ultra_wide_workspace_window else 1
            self.pdf_workspace_row_layout.setStretch(0, side_stretch)
            self.pdf_workspace_row_layout.setStretch(1, center_stretch)
            self.pdf_workspace_row_layout.setStretch(2, side_stretch)
        if hasattr(self, "pdf_workspace_shell_layout"):
            self.pdf_workspace_shell_layout.setContentsMargins(
                preview_shell_padding,
                preview_shell_padding,
                preview_shell_padding,
                preview_shell_padding,
            )
        if hasattr(self, "word_workspace_shell"):
            self.word_workspace_shell.setMaximumWidth(preview_shell_width)
        if hasattr(self, "word_workspace_row_layout"):
            center_stretch = 28 if density_mode == "wide" else (22 if density_mode == "compact" else 16)
            if is_very_wide_workspace_window:
                center_stretch += 4
            if is_ultra_wide_workspace_window:
                center_stretch += 2
            if is_cinema_wide_workspace_window:
                center_stretch += 4
            side_stretch = 0 if is_ultra_wide_workspace_window else 1
            self.word_workspace_row_layout.setStretch(0, side_stretch)
            self.word_workspace_row_layout.setStretch(1, center_stretch)
            self.word_workspace_row_layout.setStretch(2, side_stretch)
        if hasattr(self, "word_workspace_shell_layout"):
            self.word_workspace_shell_layout.setContentsMargins(
                preview_shell_padding,
                preview_shell_padding,
                preview_shell_padding,
                preview_shell_padding,
            )
            self.word_workspace_shell_layout.setSpacing(max(4, batch_minor_gap - 3))
        if hasattr(self, "batch_card"):
            self.batch_card.setMaximumWidth(workspace_card_width)
            self.batch_card.setSizePolicy(
                QSizePolicy.Policy.Expanding,
                QSizePolicy.Policy.Expanding if is_tall_workspace_window else QSizePolicy.Policy.Maximum,
            )
        if hasattr(self, "batch_card_row_layout"):
            center_stretch = 18 if density_mode == "wide" else (14 if density_mode == "compact" else 10)
            if is_very_wide_workspace_window:
                center_stretch += 3
            if is_ultra_wide_workspace_window:
                center_stretch += 1
            if is_cinema_wide_workspace_window:
                center_stretch += 3
            side_stretch = 0 if is_ultra_wide_workspace_window else 1
            self.batch_card_row_layout.setStretch(0, side_stretch)
            self.batch_card_row_layout.setStretch(1, center_stretch)
            self.batch_card_row_layout.setStretch(2, side_stretch)
        if hasattr(self, "batch_card_layout"):
            self.batch_card_layout.setContentsMargins(
                batch_card_padding_h,
                batch_card_padding_v,
                batch_card_padding_h,
                batch_card_padding_v,
            )
            self.batch_card_layout.setSpacing(batch_section_gap)
        batch_left_rail_width = 380 if density_mode == "wide" else (350 if density_mode == "compact" else 16777215)
        if is_cinema_wide_workspace_window and batch_left_rail_width < 16777215:
            batch_left_rail_width = max(340, batch_left_rail_width - 20)
        if hasattr(self, "batch_summary_section"):
            self.batch_summary_section.setMaximumWidth(batch_left_rail_width)
            self.batch_summary_section.setMinimumWidth(0 if density_mode == "narrow" else max(280, batch_left_rail_width - 70))
        if hasattr(self, "batch_log_section"):
            self.batch_log_section.setMaximumWidth(batch_left_rail_width)
            self.batch_log_section.setMinimumWidth(0 if density_mode == "narrow" else max(280, batch_left_rail_width - 70))
        if hasattr(self, "batch_result_section"):
            self.batch_result_section.setMaximumWidth(16777215)
            batch_result_min_width = 0 if density_mode == "narrow" else (840 if density_mode == "wide" else 700)
            if is_cinema_wide_workspace_window and batch_result_min_width:
                batch_result_min_width += 140
            self.batch_result_section.setMinimumWidth(batch_result_min_width)
        if hasattr(self, "merge_card"):
            self.merge_card.setMaximumWidth(workspace_card_width)
            self.merge_card.setSizePolicy(
                QSizePolicy.Policy.Expanding,
                QSizePolicy.Policy.Expanding if is_tall_workspace_window else QSizePolicy.Policy.Maximum,
            )
        if hasattr(self, "merge_card_layout"):
            self.merge_card_layout.setContentsMargins(
                batch_card_padding_h,
                batch_card_padding_v,
                batch_card_padding_h,
                batch_card_padding_v,
            )
            self.merge_card_layout.setSpacing(batch_section_gap)
        if hasattr(self, "merge_card_row_layout"):
            center_stretch = 14 if density_mode == "wide" else (12 if density_mode == "compact" else 8)
            if is_very_wide_workspace_window:
                center_stretch += 2
            if is_ultra_wide_workspace_window:
                center_stretch += 1
            if is_cinema_wide_workspace_window:
                center_stretch += 3
            side_stretch = 0 if is_ultra_wide_workspace_window else 1
            self.merge_card_row_layout.setStretch(0, side_stretch)
            self.merge_card_row_layout.setStretch(1, center_stretch)
            self.merge_card_row_layout.setStretch(2, side_stretch)
        if hasattr(self, "idle_outer_layout"):
            idle_side_margin = max(14, side_margin + (6 if is_very_wide_workspace_window else 8))
            if is_ultra_wide_workspace_window:
                idle_side_margin = max(12, idle_side_margin - 2)
            self.idle_outer_layout.setContentsMargins(idle_side_margin, 14, idle_side_margin, 26)
        if hasattr(self, "idle_card_layout"):
            idle_card_margin = 28 if density_mode == "wide" else (24 if density_mode == "compact" else 20)
            self.idle_card_layout.setContentsMargins(idle_card_margin, idle_card_margin - 4, idle_card_margin, idle_card_margin - 6)
            self.idle_card_layout.setSpacing(14 if density_mode == "wide" else 11)
        if hasattr(self, "idle_card"):
            idle_card_width = min(
                workspace_card_width,
                1980 if density_mode == "wide" else (1680 if density_mode == "compact" else 1180),
            )
            if is_cinema_wide_workspace_window:
                idle_card_width = min(idle_card_width + 180, workspace_card_width)
            self.idle_card.setMaximumWidth(idle_card_width)
        if hasattr(self, "idle_hero_panel"):
            self.idle_hero_panel.setMaximumWidth(16777215)
        if hasattr(self, "idle_hero_layout"):
            self.idle_hero_layout.setSpacing(12 if density_mode == "wide" else 10)
        if hasattr(self, "idle_flow_panel"):
            self.idle_flow_panel.setMaximumWidth(16777215)
        if hasattr(self, "idle_section_panel"):
            self.idle_section_panel.setMaximumWidth(16777215)
        if hasattr(self, "idle_section_layout"):
            self.idle_section_layout.setSpacing(10 if density_mode == "wide" else 8)
        if hasattr(self, "idle_action_panel_layout"):
            self.idle_action_panel_layout.setContentsMargins(0, 2, 0, 2)
            self.idle_action_panel_layout.setSpacing(10 if density_mode == "wide" else 8)
        if hasattr(self, "idle_title_row_layout"):
            self.idle_title_row_layout.setSpacing(18 if density_mode == "wide" else 12)
        if hasattr(self, "idle_title_tools_layout"):
            self.idle_title_tools_layout.setSpacing(8 if density_mode == "wide" else 6)
        if hasattr(self, "idle_badge_row_layout"):
            self.idle_badge_row_layout.setSpacing(8 if density_mode == "wide" else 6)
        if hasattr(self, "idle_utility_row_layout"):
            self.idle_utility_row_layout.setSpacing(8 if density_mode == "wide" else 6)
        if hasattr(self, "idle_action_buttons_layout"):
            self.idle_action_buttons_layout.setHorizontalSpacing(16 if density_mode == "wide" else 12)
            self.idle_action_buttons_layout.setVerticalSpacing(10 if density_mode == "narrow" else 0)
        if hasattr(self, "idle_flow_layout"):
            self.idle_flow_layout.setContentsMargins(0, 2, 0, 0)
            self.idle_flow_layout.setSpacing(12 if density_mode == "wide" else 9)
        if hasattr(self, "idle_routes_layout"):
            self.idle_routes_layout.setHorizontalSpacing(16 if density_mode == "wide" else 12)
            self.idle_routes_layout.setVerticalSpacing(12 if density_mode == "wide" else 10)
        self._rebuild_idle_action_layout(density_mode)
        self._rebuild_idle_route_layout(density_mode)
        self._rebuild_batch_stage_layout(density_mode)
        self._rebuild_batch_metrics_layout(density_mode)
        self._rebuild_batch_action_layout(density_mode)
        self._rebuild_batch_detail_layout(density_mode)
        self._rebuild_merge_stage_layout(density_mode)
        self._rebuild_merge_metrics_layout(density_mode)
        for route_card in getattr(self, "idle_route_cards", []):
            route_card_height = 126 if density_mode == "wide" else (114 if density_mode == "compact" else 102)
            route_card.setMinimumHeight(route_card_height)
            route_card.setMaximumHeight(route_card_height + 18)
        if hasattr(self, "canvas_layout"):
            self.canvas_layout.setContentsMargins(
                preview_content_padding,
                preview_top_padding,
                preview_content_padding,
                preview_content_padding,
            )
            self.canvas_layout.setSpacing(preview_content_spacing)
        if hasattr(self, "word_compare_layout"):
            self.word_compare_layout.setContentsMargins(0, 0, 0, 0)
            self.word_compare_layout.setSpacing(max(4, preview_content_spacing - 5))
        if hasattr(self, "word_compare_header"):
            self.word_compare_header.setFixedHeight(compare_header_height)
        if hasattr(self, "word_header_layout"):
            self.word_header_layout.setSpacing(compare_header_gap)
            self.word_header_layout.setContentsMargins(0, 0, 0, 1)
        if hasattr(self, "original_panel_layout"):
            self.original_panel_layout.setContentsMargins(0, 0, 0, 0)
        if hasattr(self, "replaced_panel_layout"):
            self.replaced_panel_layout.setContentsMargins(0, 0, 0, 0)
        if hasattr(self, "batch_header_layout"):
            self.batch_header_layout.setContentsMargins(0, 0, 0, 2)
            self.batch_header_layout.setSpacing(batch_section_gap - 2)
        if hasattr(self, "batch_header_text_layout"):
            self.batch_header_text_layout.setSpacing(max(3, batch_minor_gap - 4))
        if hasattr(self, "batch_stage_layout"):
            self.batch_stage_layout.setContentsMargins(0, 0, 0, 0)
            self.batch_stage_layout.setHorizontalSpacing(batch_minor_gap + 2)
            self.batch_stage_layout.setVerticalSpacing(batch_minor_gap + 2)
        if hasattr(self, "batch_metrics_layout"):
            self.batch_metrics_layout.setContentsMargins(0, 0, 0, 0)
            self.batch_metrics_layout.setHorizontalSpacing(batch_minor_gap)
            self.batch_metrics_layout.setVerticalSpacing(batch_minor_gap)
        if hasattr(self, "merge_header_layout"):
            self.merge_header_layout.setContentsMargins(0, 0, 0, 2)
            self.merge_header_layout.setSpacing(batch_section_gap - 2)
        if hasattr(self, "merge_header_text_layout"):
            self.merge_header_text_layout.setSpacing(max(3, batch_minor_gap - 4))
        if hasattr(self, "merge_stage_layout"):
            self.merge_stage_layout.setContentsMargins(0, 0, 0, 0)
            self.merge_stage_layout.setHorizontalSpacing(batch_minor_gap + 2)
            self.merge_stage_layout.setVerticalSpacing(batch_minor_gap + 2)
        if hasattr(self, "merge_metrics_layout"):
            self.merge_metrics_layout.setContentsMargins(0, 0, 0, 0)
            self.merge_metrics_layout.setHorizontalSpacing(batch_minor_gap)
            self.merge_metrics_layout.setVerticalSpacing(batch_minor_gap)
        if hasattr(self, "batch_actions_layout"):
            self.batch_actions_layout.setContentsMargins(0, 2, 0, 0)
            self.batch_actions_layout.setHorizontalSpacing(batch_minor_gap)
            self.batch_actions_layout.setVerticalSpacing(batch_minor_gap)
        if hasattr(self, "batch_summary_section_layout"):
            self.batch_summary_section_layout.setContentsMargins(
                batch_detail_padding_h,
                batch_detail_padding_v,
                batch_detail_padding_h,
                batch_detail_padding_v,
            )
            self.batch_summary_section_layout.setSpacing(batch_minor_gap)
        if hasattr(self, "batch_result_section_layout"):
            self.batch_result_section_layout.setContentsMargins(
                batch_detail_padding_h,
                batch_detail_padding_v,
                batch_detail_padding_h,
                batch_detail_padding_v,
            )
            self.batch_result_section_layout.setSpacing(batch_minor_gap)
        if hasattr(self, "batch_log_section_layout"):
            self.batch_log_section_layout.setContentsMargins(
                batch_detail_padding_h,
                batch_detail_padding_v,
                batch_detail_padding_h,
                batch_detail_padding_v,
            )
            self.batch_log_section_layout.setSpacing(batch_minor_gap)
        if hasattr(self, "batch_result_toolbar"):
            self.batch_result_toolbar.setSpacing(max(6, batch_minor_gap - 1))
            self.batch_result_toolbar.setContentsMargins(0, 1, 0, 2)
            self.batch_result_toolbar.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        if hasattr(self, "batch_result_header_layout"):
            self.batch_result_header_layout.setContentsMargins(0, 0, 0, 0)
            self.batch_result_header_layout.setSpacing(batch_minor_gap + 4)
        if hasattr(self, "batch_summary_browser"):
            self.batch_summary_browser.setMinimumHeight(118 if density_mode == "wide" else (110 if density_mode == "compact" else 104))
            summary_max_height = 172 if density_mode == "wide" else (164 if density_mode == "compact" else 154)
            if is_tall_workspace_window:
                summary_max_height += 16
            self.batch_summary_browser.setMaximumHeight(summary_max_height)
        if hasattr(self, "batch_result_table"):
            result_min_height = 300 if density_mode == "wide" else (250 if density_mode == "compact" else 196)
            if is_tall_workspace_window:
                result_min_height += 40 if density_mode == "wide" else 28
            self.batch_result_table.setMinimumHeight(result_min_height)
            self.batch_result_table.setMaximumHeight(16777215)
            self.batch_result_table.setColumnWidth(0, 88 if density_mode == "wide" else (82 if density_mode == "compact" else 76))
            self.batch_result_table.setColumnWidth(3, 104 if density_mode == "wide" else (96 if density_mode == "compact" else 88))
        if hasattr(self, "batch_log_list"):
            log_min_height = 184 if density_mode == "wide" else (172 if density_mode == "compact" else 156)
            if is_tall_workspace_window:
                log_min_height += 24 if density_mode == "wide" else 18
            self.batch_log_list.setMinimumHeight(log_min_height)
            self.batch_log_list.setMaximumHeight(16777215)
        merge_metric_min_height = 108 if density_mode == "wide" else (100 if density_mode == "compact" else 94)
        for merge_card in getattr(self, "merge_metric_cards", []):
            if merge_card:
                merge_card.setMinimumHeight(merge_metric_min_height)
                merge_card.setMaximumHeight(16777215)

        icon_button_size = button_height
        for icon_btn in [
            getattr(self, "btn_zoom_out", None),
            getattr(self, "btn_zoom_in", None),
            getattr(self, "btn_go_first", None),
            getattr(self, "btn_prev_page", None),
            getattr(self, "btn_next_page", None),
            getattr(self, "btn_go_last", None),
        ]:
            if icon_btn:
                icon_btn.setFixedSize(icon_button_size, icon_button_size)
                icon_btn.setIconSize(QSize(max(14, icon_button_size - 14), max(14, icon_button_size - 14)))

        control_height = button_height
        group_height = control_height + (7 if is_pdf_like_mode else 4)
        for button in [
            getattr(self, "btn_open", None),
            getattr(self, "btn_scan", None),
            getattr(self, "btn_idle_open", None),
            getattr(self, "btn_idle_feedback", None),
            getattr(self, "btn_idle_manual", None),
            getattr(self, "btn_idle_donate", None),
            getattr(self, "btn_compare_toggle", None),
            getattr(self, "rb_black", None),
            getattr(self, "rb_white", None),
            getattr(self, "cb_dual", None),
            getattr(self, "btn_fit", None),
            getattr(self, "btn_fit_utility", None),
            getattr(self, "btn_settings", None),
            getattr(self, "btn_feedback", None),
            getattr(self, "btn_workbench_feedback", None),
            getattr(self, "btn_more", None),
            getattr(self, "btn_save", None),
        ]:
            if button:
                button.setMinimumHeight(control_height)
                button.setMaximumHeight(control_height)

        for label in [
            getattr(self, "lbl_zoom", None),
            getattr(self, "lbl_page", None),
        ]:
            if label:
                label.setMinimumHeight(control_height)
                label.setMaximumHeight(control_height)

        for group in [
            getattr(self, "toolbar_primary_group", None),
            getattr(self, "toolbar_word_group", None),
            getattr(self, "toolbar_pdf_group", None),
            getattr(self, "toolbar_zoom_group", None),
            getattr(self, "toolbar_nav_group", None),
            getattr(self, "toolbar_utility_group", None),
        ]:
            if group:
                group.setMinimumHeight(group_height)
                group.setMaximumHeight(group_height)

        idle_action_width = 220 if density_mode == "wide" else (190 if density_mode == "compact" else 160)
        if hasattr(self, "btn_idle_open"):
            idle_primary_height = control_height + 4
            self.btn_idle_open.setMinimumWidth(idle_action_width)
            self.btn_idle_open.setMaximumWidth(16777215 if density_mode != "narrow" else idle_action_width)
            self.btn_idle_open.setMinimumHeight(idle_primary_height)
            self.btn_idle_open.setMaximumHeight(idle_primary_height)
        if hasattr(self, "idle_start_card"):
            start_min_height = idle_primary_height + (78 if density_mode == "wide" else (72 if density_mode == "compact" else 84))
            self.idle_start_card.setMinimumHeight(start_min_height)
            self.idle_start_card.setMaximumHeight(16777215)
        if hasattr(self, "idle_start_footer_layout"):
            self.idle_start_footer_layout.setSpacing(8 if density_mode == "wide" else 6)
        if hasattr(self, "btn_idle_feedback"):
            idle_secondary_height = control_height + 4
            idle_support_action_width = 120 if density_mode == "wide" else (108 if density_mode == "compact" else 100)
            self.btn_idle_feedback.setMinimumWidth(idle_support_action_width)
            self.btn_idle_feedback.setMaximumWidth(16777215 if density_mode != "narrow" else idle_action_width)
            self.btn_idle_feedback.setMinimumHeight(idle_secondary_height)
            self.btn_idle_feedback.setMaximumHeight(idle_secondary_height)
        for button in [getattr(self, "btn_idle_manual", None), getattr(self, "btn_idle_donate", None)]:
            if button:
                idle_secondary_height = control_height + 4
                idle_support_action_width = 120 if density_mode == "wide" else (108 if density_mode == "compact" else 100)
                button.setMinimumWidth(idle_support_action_width)
                button.setMaximumWidth(16777215 if density_mode != "narrow" else idle_action_width)
                button.setMinimumHeight(idle_secondary_height)
                button.setMaximumHeight(idle_secondary_height)
        if hasattr(self, "idle_support_card"):
            support_min_height = idle_primary_height + (92 if density_mode == "wide" else (88 if density_mode == "compact" else 108))
            self.idle_support_card.setMinimumHeight(support_min_height)
            self.idle_support_card.setMaximumHeight(16777215)

        for label in [getattr(self, "lbl_zoom", None), getattr(self, "lbl_page", None)]:
            if label:
                label.setMinimumHeight(control_height)
                label.setMaximumHeight(control_height)
                label.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
                label.setStyleSheet(
                    f"""
                    QLabel#toolbarMeta {{
                        color: {theme["text"]};
                        background-color: {theme["hover"]};
                        border: 1px solid {theme["border"]};
                        border-radius: 9px;
                        padding: 4px 10px;
                        font-size: {meta_font_size}px;
                        font-weight: 700;
                        min-width: 54px;
                    }}
                    """
                )
        if hasattr(self, "lbl_batch_result_meta"):
            self.lbl_batch_result_meta.setMinimumHeight(max(28, control_height - 4))
            self.lbl_batch_result_meta.setMaximumHeight(max(28, control_height - 4))
        if hasattr(self, "lbl_batch_result_hint"):
            self.lbl_batch_result_hint.setMinimumHeight(max(28, control_height - 4))
            self.lbl_batch_result_hint.setMaximumHeight(max(28, control_height - 4))
        batch_filter_height = max(30, control_height - 4)
        for button in [
            getattr(self, "btn_batch_filter_all", None),
            getattr(self, "btn_batch_filter_success", None),
            getattr(self, "btn_batch_filter_failed", None),
        ]:
            if button:
                button.setMinimumHeight(batch_filter_height)
                button.setMaximumHeight(batch_filter_height)
                button.setMinimumWidth(86 if density_mode == "wide" else (78 if density_mode == "compact" else 72))
        batch_filter_height = max(28, control_height - 4)
        batch_filter_width = 96 if density_mode == "wide" else (88 if density_mode == "compact" else 80)
        for button in [
            getattr(self, "btn_batch_filter_all", None),
            getattr(self, "btn_batch_filter_success", None),
            getattr(self, "btn_batch_filter_failed", None),
        ]:
            if button:
                button.setMinimumHeight(batch_filter_height)
                button.setMaximumHeight(batch_filter_height)
                button.setMinimumWidth(batch_filter_width)
        batch_action_width = 168 if density_mode == "wide" else (154 if density_mode == "compact" else 140)
        for button in getattr(self, "batch_action_buttons", []):
            if button:
                button.setMinimumWidth(batch_action_width)

        if hasattr(self, "lbl_workbench_title"):
            self.lbl_workbench_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {title_font_size}px; font-weight: 700; background-color: transparent;"
            )
        for label in [getattr(self, "lbl_batch_title", None), getattr(self, "lbl_merge_title", None)]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['text']}; font-size: {title_font_size + 3}px; font-weight: 700; background-color: transparent;"
                )
        for label in [
            getattr(self, "lbl_workbench_subtitle", None),
            getattr(self, "lbl_batch_subtitle", None),
            getattr(self, "lbl_merge_subtitle", None),
        ]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['text_secondary']}; font-size: {subtitle_font_size}px; line-height: 1.7; background-color: transparent;"
                )
        for label in [
            getattr(self, "lbl_batch_meta", None),
            getattr(self, "lbl_batch_current_file", None),
            getattr(self, "lbl_merge_meta", None),
            getattr(self, "lbl_idle_tip", None),
            getattr(self, "lbl_idle_drop_hint", None),
        ]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['text_secondary'] if label in (getattr(self, 'lbl_idle_tip', None), getattr(self, 'lbl_idle_drop_hint', None)) else theme['text']}; font-size: {subtitle_font_size}px; line-height: 1.7; background-color: transparent;"
                )
        if hasattr(self, "lbl_idle_section"):
            self.lbl_idle_section.setStyleSheet(
                f"color: {theme['text']}; font-size: {max(12, title_font_size - 1)}px; font-weight: 700; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_section_hint"):
            self.lbl_idle_section_hint.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {max(10, subtitle_font_size - 1)}px; font-weight: 600; background-color: transparent;"
            )
        for label in [getattr(self, "lbl_idle_offline_badge", None), getattr(self, "lbl_idle_auto_badge", None)]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['primary']}; background-color: #E9F1FB; border: 1px solid {theme['border']}; border-radius: 9px; padding: 4px 9px; font-size: {max(10, subtitle_font_size - 1)}px; font-weight: 700;"
                )
        if hasattr(self, "lbl_idle_section_hint"):
            self.lbl_idle_section_hint.setVisible(density_mode != "narrow")
        if hasattr(self, "lbl_idle_auto_badge"):
            self.lbl_idle_auto_badge.setVisible(density_mode == "wide")
        if hasattr(self, "lbl_idle_drop_hint"):
            self.lbl_idle_drop_hint.setText(
                "支持直接拖拽到窗口，系统会自动分流" if density_mode == "wide"
                else ("支持直接拖拽到窗口" if density_mode == "compact" else "支持拖拽到窗口")
            )
        if hasattr(self, "lbl_idle_start_title"):
            self.lbl_idle_start_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {max(12, title_font_size - 1)}px; font-weight: 700; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_start_text"):
            self.lbl_idle_start_text.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {max(10, subtitle_font_size - 1)}px; line-height: 1.6; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_support_text"):
            self.lbl_idle_support_text.setStyleSheet(
                f"color: {theme['text']}; font-size: {max(11, subtitle_font_size)}px; font-weight: 600; line-height: 1.5; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_support_meta"):
            self.lbl_idle_support_meta.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {max(10, subtitle_font_size - 1)}px; font-weight: 600; line-height: 1.5; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_support_email"):
            self.lbl_idle_support_email.setStyleSheet(
                f"color: {theme['primary']}; font-size: {max(10, subtitle_font_size - 1)}px; font-weight: 600; background-color: transparent;"
            )
        if hasattr(self, "lbl_idle_support_note"):
            self.lbl_idle_support_note.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {max(10, subtitle_font_size - 1)}px; line-height: 1.6; background-color: transparent;"
            )
        if hasattr(self, "btn_idle_settings"):
            self.btn_idle_settings.setMinimumHeight(control_height)
            self.btn_idle_settings.setMaximumHeight(control_height)
            self.btn_idle_settings.setMinimumWidth(112 if density_mode == "wide" else 96)
        if hasattr(self, "btn_idle_feedback"):
            self.btn_idle_feedback.setMinimumHeight(control_height)
            self.btn_idle_feedback.setMaximumHeight(control_height)
            self.btn_idle_feedback.setMinimumWidth(112 if density_mode == "wide" else 96)
        if hasattr(self, "btn_idle_manual"):
            self.btn_idle_manual.setMinimumHeight(control_height)
            self.btn_idle_manual.setMaximumHeight(control_height)
            self.btn_idle_manual.setMinimumWidth(112 if density_mode == "wide" else 96)
        if hasattr(self, "btn_idle_donate"):
            self.btn_idle_donate.setMinimumHeight(control_height)
            self.btn_idle_donate.setMaximumHeight(control_height)
            self.btn_idle_donate.setMinimumWidth(112 if density_mode == "wide" else 96)
        if hasattr(self, "btn_workbench_feedback"):
            self.btn_workbench_feedback.setMinimumHeight(control_height)
            self.btn_workbench_feedback.setMaximumHeight(control_height)
            self.btn_workbench_feedback.setMinimumWidth(112 if density_mode == "wide" else 96)
        if hasattr(self, "lbl_workbench_focus"):
            self.lbl_workbench_focus.setStyleSheet(
                f"""
                QLabel#workbenchFocus {{
                    background-color: #E9F1FB;
                    color: {theme["primary"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 14px;
                    padding: 6px 12px;
                    font-size: {chip_font_size}px;
                    font-weight: 700;
                }}
                """
            )
        for label in getattr(self, "workbench_guidance_labels", []):
            label.setStyleSheet(
                f"""
                QLabel#workbenchHintTag {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 7px 12px;
                    font-size: {chip_font_size}px;
                    font-weight: 700;
                }}
                """
            )
        if hasattr(self, "info_bar"):
            self.info_bar.setStyleSheet(
                f"""
                QLabel#contextMessage {{
                    background-color: {theme["info_bar"]};
                    color: {theme["text_secondary"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 8px 12px;
                    font-weight: 600;
                    font-size: {subtitle_font_size}px;
                }}
                """
            )
        for label in [getattr(self, "lbl_word_original_header", None), getattr(self, "lbl_word_replaced_header", None)]:
            if label:
                label.setStyleSheet(
                    f"""
                    QLabel#wordCompareLabel {{
                        color: {theme['text_secondary']};
                        background-color: {theme['hover']};
                        border: 1px solid {theme['border']};
                        border-radius: 10px;
                        padding: 3px 10px;
                        font-size: {max(11, compare_header_font_size - 1)}px;
                        font-weight: 600;
                    }}
                    """
                )
        for label in [
            getattr(self, "lbl_batch_log_hint", None),
            getattr(self, "lbl_batch_summary_hint", None),
            getattr(self, "lbl_batch_result_hint", None),
        ]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['text_secondary']}; font-size: {subtitle_font_size}px; font-weight: 700; background-color: transparent;"
                )
        if hasattr(self, "batch_summary_browser"):
            self.batch_summary_browser.setStyleSheet(
                f"""
                QTextBrowser#batchSummaryBrowser {{
                    background-color: #FCFDFF;
                    color: {theme["text"]};
                    border: 1px solid #E4EBF3;
                    border-radius: 15px;
                    padding: 12px;
                    font-size: {subtitle_font_size}px;
                    line-height: 1.7;
                }}
                """
            )
        if hasattr(self, "batch_result_table"):
            self.batch_result_table.setStyleSheet(
                f"""
                QTableWidget#batchResultTable {{
                    background-color: {theme["surface"]};
                    color: {theme["text"]};
                    border: 1px solid #E4EBF3;
                    border-radius: 15px;
                    gridline-color: transparent;
                    padding: 6px;
                    font-size: {subtitle_font_size}px;
                }}
                QTableWidget#batchResultTable::item {{
                    padding: 8px 10px;
                    border-bottom: 1px solid #E9EEF5;
                }}
                QTableWidget#batchResultTable QHeaderView::section {{
                    background-color: #F7FAFD;
                    color: {theme["text_secondary"]};
                    border: none;
                    border-bottom: 1px solid #E4EBF3;
                    padding: 8px 10px;
                    font-size: {max(11, subtitle_font_size - 1)}px;
                    font-weight: 700;
                }}
                """
            )
            self.batch_result_table.verticalHeader().setDefaultSectionSize(batch_result_row_height)
            self.batch_result_table.horizontalHeader().setFixedHeight(batch_result_row_height - 2)
        if hasattr(self, "lbl_batch_result_meta"):
            self.lbl_batch_result_meta.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {max(11, subtitle_font_size - 1)}px; font-weight: 600; background-color: transparent;"
            )
        if hasattr(self, "progress"):
            self.progress.setFixedHeight(progress_height)
        if hasattr(self, "btn_cancel_scan"):
            self.btn_cancel_scan.setFixedSize(cancel_width, progress_height)
            self.btn_cancel_scan.setStyleSheet(self._get_button_style("secondary"))

        self._refresh_batch_result_filter_buttons()
        self._refresh_button_density_styles()
        self._rebuild_idle_support_actions_layout(density_mode)

    def _refresh_button_density_styles(self):
        """在 DPI 或密度切换后，重新应用所有按钮的字体和内边距。"""
        for button in self.findChildren(QPushButton):
            style_type = button.property("btn_style")
            if not style_type:
                continue
            if button is getattr(self, "btn_more", None):
                continue
            button.setStyleSheet(self._get_button_style(style_type))

        self._refresh_toolbar_more_button_style()
        self._apply_native_toolbar_icons()

    def _refresh_toolbar_group_visibility(self):
        """按当前模式和响应式显隐，收口工具栏分组容器。"""
        group_map = [
            (getattr(self, "toolbar_primary_group", None), [getattr(self, "btn_open", None), getattr(self, "btn_scan", None)]),
            (getattr(self, "toolbar_word_group", None), [getattr(self, "btn_settings", None), getattr(self, "btn_compare_toggle", None)]),
            (getattr(self, "toolbar_pdf_group", None), [getattr(self, "rb_black", None), getattr(self, "rb_white", None), getattr(self, "cb_dual", None)]),
            (getattr(self, "toolbar_zoom_group", None), [getattr(self, "btn_zoom_out", None), getattr(self, "lbl_zoom", None), getattr(self, "btn_zoom_in", None)]),
            (getattr(self, "toolbar_nav_group", None), [getattr(self, "btn_go_first", None), getattr(self, "btn_prev_page", None), getattr(self, "lbl_page", None), getattr(self, "btn_next_page", None), getattr(self, "btn_go_last", None)]),
            (getattr(self, "toolbar_utility_group", None), [getattr(self, "btn_fit_utility", None), getattr(self, "btn_feedback", None), getattr(self, "btn_more", None), getattr(self, "btn_save", None)]),
        ]

        for group, widgets in group_map:
            if not group:
                continue
            group.setVisible(any(widget and not widget.isHidden() for widget in widgets))

    def _set_toolbar_widget_width(self, widget, min_width=0, extra=28):
        """根据当前文案为工具栏控件设置稳定宽度，避免文字被硬挤压。"""
        if not widget or not hasattr(widget, "fontMetrics"):
            return

        text = widget.text() if hasattr(widget, "text") else ""
        metrics = widget.fontMetrics()
        text_width = metrics.horizontalAdvance(text) + extra
        hint_width = 0
        try:
            widget.ensurePolished()
            hint_width = widget.sizeHint().width()
        except Exception:
            hint_width = 0

        if widget.objectName() == "toolbarMoreButton":
            hint_width += 12

        width = max(text_width, hint_width)
        width = max(width, min_width)
        widget.setMinimumWidth(width)
        widget.setMaximumWidth(width)

    def _set_toolbar_group_width(self, group, width=None):
        """为工具栏分组设置统一外宽；传 None 时恢复内容驱动宽度。"""
        if not group:
            return
        if width is None:
            group.setMinimumWidth(0)
            group.setMaximumWidth(16777215)
            return
        group.setMinimumWidth(width)
        group.setMaximumWidth(width)

    def _refresh_toolbar_overflow_menu(self, is_pdf, is_word, density_mode):
        """把低频动作收进更多菜单，避免窄窗口时挤爆工具栏。"""
        if not hasattr(self, "toolbar_more_menu"):
            return

        self.toolbar_more_menu.clear()
        has_item = False

        if is_word and not self.btn_compare_toggle.isVisible():
            compare_text = "打开对比预览"
            if self.word_doc and self._has_word_replacement_candidates():
                compare_text = "显示对比预览" if self.word_compare_user_hidden else "隐藏对比预览"
            compare_action = self.toolbar_more_menu.addAction(compare_text, self.toggle_word_compare_preview)
            compare_action.setEnabled(bool(self.word_doc) and self._has_word_replacement_candidates())
            has_item = True

        if is_pdf and not getattr(self, "btn_fit_utility", None).isVisible():
            self.toolbar_more_menu.addAction("适应页面", self.fit_page)
            has_item = True

        if is_pdf and density_mode == "narrow":
            if has_item:
                self.toolbar_more_menu.addSeparator()
            self.toolbar_more_menu.addAction("跳到第一页", self.go_first)
            self.toolbar_more_menu.addAction("跳到最后一页", self.go_last)
            has_item = True

        self.btn_more.setVisible(has_item)

    def _show_toolbar_more_menu(self):
        """在按钮下方手动弹出更多菜单，避免系统菜单箭头把按钮布局挤坏。"""
        if not hasattr(self, "toolbar_more_menu") or self.toolbar_more_menu.isEmpty():
            return
        source_btn = self.sender()
        if not isinstance(source_btn, QPushButton):
            source_btn = getattr(self, "btn_more", None)
        if not source_btn or not source_btn.isVisible():
            return

        popup_pos = source_btn.mapToGlobal(source_btn.rect().bottomLeft())
        self.toolbar_more_menu.popup(popup_pos)

    def _toggle_dual_toolbar(self, checked=False):
        """工具栏双页按钮入口，保持按钮状态和画布状态同步。"""
        checked = bool(checked)
        if hasattr(self, "cb_dual") and self.cb_dual.isChecked() != checked:
            self.cb_dual.setChecked(checked)
        self.toggle_dual_view(checked)
        self._refresh_toolbar_responsiveness()
        self._refresh_workbench_context()

    def _refresh_toolbar_responsiveness(self):
        """根据窗口宽度做工具栏响应式降级，保证缩放时仍可读。"""
        if not hasattr(self, "toolbar"):
            return

        width = self.width() or self.toolbar.width()
        # Qt 返回的是逻辑像素宽度，这里不再额外按屏幕缩放除一次，
        # 否则 Retina / Windows 高缩放下会被过早判成窄布局。
        effective_width = width
        scale = self._get_display_scale_factor()
        density_height = self.height() if scale > 1.0 else 0
        mode = self.current_ui_mode
        is_idle = mode == "idle"
        is_pdf = mode == "pdf"
        is_word = mode == "word"
        is_image_merge = mode == "image_merge"
        has_mode_results = self._has_pdf_redactions() if is_pdf else (self._has_word_replacement_candidates() if is_word else False)
        enabled_word_rules = self._count_enabled_word_rules() if is_word else 0

        density_mode = resolve_workspace_density_mode(mode, effective_width, density_height, scale)

        self.toolbar_density_mode = density_mode
        self._refresh_windows_density_metrics(density_mode)
        label_config = build_toolbar_mode_labels(
            mode,
            density_mode,
            has_results=has_mode_results,
            enabled_word_rules=enabled_word_rules,
        )

        self.btn_open.setText(label_config["open_text"])
        self.btn_scan.setText(label_config["scan_text"])
        self.btn_open.setToolTip(label_config["open_tooltip"])
        self.btn_scan.setToolTip(label_config["scan_tooltip"])

        black_text = "黑遮罩" if density_mode == "wide" else ("黑遮" if density_mode == "compact" else "黑")
        white_text = "白遮罩" if density_mode == "wide" else ("白遮" if density_mode == "compact" else "白")
        dual_text = "双页" if density_mode != "narrow" else "双"
        fit_text = "适应页面" if density_mode != "narrow" else "适应"
        self.rb_black.setText(black_text)
        self.rb_white.setText(white_text)
        self.cb_dual.setText(dual_text)
        self.btn_fit.setText(fit_text)
        self.btn_fit_utility.setText(fit_text)
        self.btn_settings.setText("高级设置")
        self.btn_feedback.setText("使用/反馈")
        self.btn_workbench_feedback.setText("使用/反馈")
        self.btn_save.setText(label_config["save_text"])
        self.btn_more.setText("更多")
        self.rb_black.setToolTip("使用黑色遮罩涂抹")
        self.rb_white.setToolTip("使用白色遮罩涂抹")
        self.cb_dual.setToolTip("切换单双页预览")
        self.btn_fit.setToolTip("按窗口大小适应页面")
        self.btn_fit_utility.setToolTip("按窗口大小适应页面")
        self.btn_settings.setToolTip("打开高级设置")
        self.btn_feedback.setToolTip("查看使用说明或提交反馈")
        self.btn_workbench_feedback.setToolTip("查看使用说明或提交反馈")
        self.btn_save.setToolTip(label_config["save_tooltip"])
        self.btn_more.setToolTip("显示收纳的操作")

        self.btn_settings.setVisible(not is_idle)
        self.btn_feedback.setVisible((not is_idle) and not (is_pdf or is_image_merge))
        self.btn_workbench_feedback.setVisible(is_pdf or is_image_merge)
        self.btn_fit.setVisible(False)
        self.btn_fit_utility.setVisible(is_pdf)
        self.btn_compare_toggle.setVisible(is_word and density_mode != "narrow")

        if is_pdf and density_mode == "narrow":
            self.btn_go_first.hide()
            self.btn_go_last.hide()
        else:
            self.btn_go_first.setVisible(is_pdf)
            self.btn_go_last.setVisible(is_pdf)

        self._refresh_toolbar_overflow_menu(is_pdf, is_word, density_mode)

        if density_mode == "wide":
            text_button_floor = 82
            utility_button_floor = 98
            compare_button_floor = 104
            more_button_floor = 76
        elif density_mode == "compact":
            text_button_floor = 74
            utility_button_floor = 90
            compare_button_floor = 94
            more_button_floor = 70
        else:
            text_button_floor = 68
            utility_button_floor = 82
            compare_button_floor = 84
            more_button_floor = 64

        if scale >= 1.5:
            text_button_floor += 8
            utility_button_floor += 10
            compare_button_floor += 10
            more_button_floor += 8
        elif scale >= 1.25:
            text_button_floor += 4
            utility_button_floor += 6
            compare_button_floor += 6
            more_button_floor += 4

        for button in [self.btn_open, self.btn_scan, self.btn_settings, self.btn_compare_toggle,
                       self.btn_fit_utility, self.btn_feedback, self.btn_more, self.btn_save]:
            if button.isVisible():
                if button in (self.btn_feedback, self.btn_fit_utility):
                    min_width = utility_button_floor
                    extra = 40
                elif button is self.btn_settings:
                    min_width = utility_button_floor
                    extra = 34
                elif button is self.btn_compare_toggle:
                    min_width = compare_button_floor
                    extra = 34
                elif button is self.btn_more:
                    min_width = more_button_floor
                    extra = 34
                else:
                    min_width = text_button_floor
                    extra = 30
                self._set_toolbar_widget_width(button, min_width=min_width, extra=extra)

        for toggle in [self.rb_black, self.rb_white, self.cb_dual]:
            if toggle.isVisible():
                self._set_toolbar_widget_width(toggle, min_width=52, extra=28)

        if self.lbl_zoom.isVisible():
            self._set_toolbar_widget_width(self.lbl_zoom, min_width=58, extra=20)
        if self.lbl_page.isVisible():
            self._set_toolbar_widget_width(self.lbl_page, min_width=58, extra=22)

        toolbar_meta_widths = []
        for label in [self.lbl_zoom, self.lbl_page]:
            if label.isVisible():
                toolbar_meta_widths.append(label.minimumWidth())
        if toolbar_meta_widths:
            shared_toolbar_meta_width = max(toolbar_meta_widths)
            for label in [self.lbl_zoom, self.lbl_page]:
                if label.isVisible():
                    label.setMinimumWidth(shared_toolbar_meta_width)
                    label.setMaximumWidth(shared_toolbar_meta_width)

        zoom_group = getattr(self, "toolbar_zoom_group", None)
        nav_group = getattr(self, "toolbar_nav_group", None)
        if is_pdf and zoom_group and nav_group:
            try:
                zoom_group.ensurePolished()
                nav_group.ensurePolished()
                shared_group_width = max(zoom_group.sizeHint().width(), nav_group.sizeHint().width())
            except Exception:
                shared_group_width = None
            if shared_group_width:
                self._set_toolbar_group_width(zoom_group, shared_group_width)
                self._set_toolbar_group_width(nav_group, shared_group_width)
        else:
            self._set_toolbar_group_width(zoom_group, None)
            self._set_toolbar_group_width(nav_group, None)

        self._refresh_toolbar_group_visibility()

    def _refresh_workbench_guidance(self, guidance_items):
        """刷新主工作台顶部的下一步引导标签。当前版本默认隐藏，保持上下文条简洁。"""
        if not hasattr(self, "workbench_guidance_labels"):
            return
        for label in self.workbench_guidance_labels:
            label.hide()

    def _refresh_workflow_steps(self, active_index):
        """刷新顶部流程步骤，让主路径始终可见。"""
        if not hasattr(self, "workflow_step_labels"):
            return
        step_font_size = getattr(self, "_workflow_step_font_size", 11)

        for index, label in enumerate(self.workflow_step_labels):
            if index < active_index:
                fg = Theme.LIGHT["success"]
                bg = "#EAF8F1"
            elif index == active_index:
                fg = Theme.LIGHT["primary"]
                bg = "#E9F1FB"
            else:
                fg = Theme.LIGHT["text_secondary"]
                bg = Theme.LIGHT["hover"]

            label.setStyleSheet(
                f"""
                QLabel#workflowStep {{
                    color: {fg};
                    background-color: {bg};
                    border: 1px solid {Theme.LIGHT["border"]};
                    border-radius: 10px;
                    padding: 5px 10px;
                    font-size: {step_font_size}px;
                    font-weight: 700;
                }}
                """
            )

    def _create_batch_metric_card(self, title):
        """创建批量 Word 工作台指标卡。"""
        card = QFrame()
        card.setObjectName("batchMetricCard")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(4)

        title_label = QLabel(title)
        title_label.setObjectName("batchMetricTitle")
        value_label = QLabel("--")
        value_label.setObjectName("batchMetricValue")
        value_label.setWordWrap(True)
        note_label = QLabel("")
        note_label.setObjectName("batchMetricNote")
        note_label.setWordWrap(True)

        layout.addWidget(title_label)
        layout.addWidget(value_label)
        layout.addWidget(note_label)
        layout.addStretch()
        return card, value_label, note_label

    def _set_batch_step_style(self, frame, title_label, note_label, state, accent_fg=None, accent_bg=None):
        """按当前批量阶段刷新流程卡片样式。"""
        color_map = {
            "pending": (Theme.LIGHT["text_secondary"], "#FBFCFE", Theme.LIGHT["border"]),
            "active": (Theme.LIGHT["primary"], "#E9F1FB", Theme.LIGHT["primary"]),
            "done": (Theme.LIGHT["success"], "#EAF8F1", Theme.LIGHT["success"]),
        }
        fg, bg, border = color_map.get(state, color_map["pending"])
        if accent_fg and accent_bg and state == "active":
            fg = accent_fg
            bg = accent_bg
            border = accent_fg

        frame.setStyleSheet(
            f"""
            QFrame#batchStepCard {{
                background-color: {bg};
                border: 1px solid {border};
                border-radius: 14px;
            }}
            """
        )
        title_label.setStyleSheet(
            f"""
            QLabel#batchStepTitle {{
                color: {fg};
                font-size: 12px;
                font-weight: 700;
                background-color: transparent;
            }}
            """
        )
        note_label.setStyleSheet(
            f"""
            QLabel#batchStepNote {{
                color: {Theme.LIGHT["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
                background-color: transparent;
            }}
            """
        )

    def _build_batch_summary_text(self, total_selected, replacement_preview):
        """构建批量 Word 工作台摘要文本。"""
        rule_count = self._count_enabled_word_rules()
        processed = min(self.batch_processed_files, total_selected) if total_selected else 0
        status_text = "已停止" if self.batch_stage == "stopped" else "已完成"

        if self.batch_stage == "rule_setup":
            lines = [
                "当前还在规则确认阶段，这一步不会改动任何原文件。",
                f"已载入文档：{total_selected} 个",
                f"当前启用 Word 规则：{rule_count} 条",
                f"统一替换文本：{replacement_preview}",
                "",
                "建议先确认：",
                "1. 文档数量是否正确",
                "2. 统一替换文本是否符合当前案卷习惯",
                "3. 是否至少启用了一条 Word 替换规则",
            ]
            return "\n".join(lines)

        if self.batch_stage == "running":
            current_file = self.batch_current_file or "正在准备下一个文档"
            lines = [
                "系统正在逐个处理 Word 文档。",
                f"当前进度：{processed}/{total_selected}",
                f"当前文件：{current_file}",
                f"已成功：{self.batch_success_count} 个",
                f"已失败：{self.batch_failed_count} 个",
                "",
                "如果遇到异常文件，可以在弹窗里选择跳过当前文件，或停止本轮任务。",
            ]
            return "\n".join(lines)

        if self.batch_stage in ("finished", "stopped"):
            lines = [
                f"本轮批量替换{status_text}。",
                f"总文档数：{total_selected}",
                f"成功：{self.batch_success_count} 个",
                f"失败：{self.batch_failed_count} 个",
                f"统一替换文本：{replacement_preview}",
            ]

            summary = self.batch_last_summary or {}
            success_items = summary.get("success", []) if isinstance(summary, dict) else []
            failed_items = summary.get("failed", []) if isinstance(summary, dict) else []
            summary_rules = summary.get("rules", self.word_replace_rules) if isinstance(summary, dict) else self.word_replace_rules

            rule_lines = build_batch_rule_summary_lines(summary_rules, success_items, replacement_preview)
            if rule_lines:
                lines.append("")
                lines.append("本次替换规则：")
                lines.extend(rule_lines)

            if success_items:
                lines.append("")
                lines.append("最近成功输出（最多 5 条）：")
                for item in success_items[:5]:
                    output_path = item.get("output", "")
                    lines.append(f"- {os.path.basename(output_path) if output_path else '已生成输出文件'}")

            if failed_items:
                lines.append("")
                lines.append("失败详情（最多 5 条）：")
                for item in failed_items[:5]:
                    input_name = os.path.basename(item.get("input", ""))
                    error_text = item.get("error", "")
                    lines.append(f"- {input_name}: {error_text}")

            return "\n".join(lines)

        lines = [
            "批量 Word 工作台适合把同一套替换规则一次应用到多个 Word 文档。",
            "系统会先进入规则确认，再开始批量执行，最后集中展示结果。",
            f"统一替换文本：{replacement_preview}",
        ]
        return "\n".join(lines)

    def _count_enabled_general_rules(self):
        """统计启用的通用规则数量。"""
        return len([rule for rule in self.active_rules if isinstance(rule, str) and rule])

    def _count_enabled_word_rules(self):
        """统计启用的 Word 替换规则数量。"""
        normalized = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)
        return len([item for item in normalized if item.get("enabled", True) and item.get("find")])

    def _has_pdf_redactions(self):
        """当前 PDF 是否已有智能/手动涂抹结果。"""
        return any(data.get("ocr") or data.get("manual") for data in self.page_data.values())

    def _has_word_redactions(self):
        """当前 Word 是否已有智能/手动脱敏结果。"""
        return any(data.get("ocr") or data.get("manual") for data in self.word_data.values())

    def _reset_batch_session_state(self):
        """重置批量 Word 工作台状态。"""
        self.batch_stage = "idle"
        self.batch_selected_files = []
        self.batch_total_files = 0
        self.batch_processed_files = 0
        self.batch_success_count = 0
        self.batch_failed_count = 0
        self.batch_current_file = ""
        self.batch_last_summary = None
        self.batch_result_filter_mode = "all"
        if hasattr(self, "batch_log_list"):
            self.batch_log_list.clear()

    def _append_batch_log(self, text, level="info"):
        """向批量工作台追加一条最近动态。"""
        if not hasattr(self, "batch_log_list") or not text:
            return

        color_map = {
            "info": Theme.LIGHT["text"],
            "success": Theme.LIGHT["success"],
            "warning": Theme.LIGHT["warning"],
            "error": Theme.LIGHT["danger"],
        }

        item = QListWidgetItem(text)
        item.setForeground(QColor(color_map.get(level, Theme.LIGHT["text"])))
        self.batch_log_list.insertItem(0, item)
        while self.batch_log_list.count() > 30:
            self.batch_log_list.takeItem(self.batch_log_list.count() - 1)

    def _reopen_batch_rule_setup(self):
        """用当前已选文档重新进入批量规则确认。"""
        file_paths = list(self.batch_selected_files) if self.batch_selected_files else None
        self.start_batch_replace(file_paths=file_paths)

    def _start_batch_replace_from_workspace(self):
        """从批量工作台重新选择文件。"""
        self.start_batch_replace()

    def _get_batch_failed_inputs(self):
        """提取本轮批量替换失败的输入文件。"""
        summary = self.batch_last_summary if isinstance(self.batch_last_summary, dict) else {}
        failed_items = summary.get("failed", []) if isinstance(summary, dict) else []
        return [item.get("input") for item in failed_items if isinstance(item, dict) and item.get("input")]

    def _get_batch_success_outputs(self):
        """提取本轮批量替换成功输出文件。"""
        summary = self.batch_last_summary if isinstance(self.batch_last_summary, dict) else {}
        success_items = summary.get("success", []) if isinstance(summary, dict) else []
        return [item.get("output") for item in success_items if isinstance(item, dict) and item.get("output")]

    def _retry_failed_batch_files(self):
        """仅重试本轮失败的批量 Word 文档。"""
        failed_files = self._get_batch_failed_inputs()
        if not failed_files:
            QMessageBox.information(self, "提示", "当前没有可重试的失败文档。")
            return
        self.start_batch_replace(file_paths=failed_files)

    def _open_batch_output_location(self):
        """打开本轮批量替换的输出位置。"""
        output_files = self._get_batch_success_outputs()
        if not output_files:
            QMessageBox.information(self, "提示", "当前还没有可打开的输出文件。")
            return

        first_output = output_files[0]
        target_dir = os.path.dirname(first_output) or os.path.dirname(os.path.abspath(first_output))
        if not target_dir or not os.path.isdir(target_dir):
            QMessageBox.warning(self, "提示", "输出目录不存在，可能已被移动或删除。")
            return

        if not QDesktopServices.openUrl(QUrl.fromLocalFile(target_dir)):
            QMessageBox.warning(self, "提示", "无法打开输出目录，请手动前往对应路径查看。")

    def _get_batch_filter_button_style(self, active=False):
        """返回批量结果筛选按钮样式。"""
        theme = Theme.LIGHT
        metrics = getattr(self, "_button_density_metrics", {}) or {}
        font_size = metrics.get("button_font_size", 13) - 1
        padding_v = max(5, metrics.get("button_padding_v", 7) - 2)
        padding_h = max(10, metrics.get("button_padding_h", 14) - 2)
        border_radius = 11

        if active:
            return f"""
                QPushButton {{
                    background-color: #EAF2FC;
                    color: {theme["primary"]};
                    border: 1px solid #B9D3F2;
                    border-radius: {border_radius}px;
                    padding: {padding_v}px {padding_h}px;
                    font-size: {font_size}px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background-color: #E3EEF9;
                    border-color: #AFCBEC;
                }}
                QPushButton:disabled {{
                    background-color: #F0F4F9;
                    color: {theme["text_secondary"]};
                    border-color: {theme["border"]};
                }}
            """

        return f"""
            QPushButton {{
                background-color: #FCFDFF;
                color: {theme["text"]};
                border: 1px solid #E2EAF3;
                border-radius: {border_radius}px;
                padding: {padding_v}px {padding_h}px;
                font-size: {font_size}px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: #F6FAFE;
                border-color: #C7D8EA;
            }}
            QPushButton:disabled {{
                background-color: #F5F7FA;
                color: {theme["text_secondary"]};
                border-color: {theme["border"]};
            }}
        """

    def _refresh_batch_result_filter_buttons(self):
        """刷新批量结果筛选按钮样式与可用性。"""
        all_rows = build_batch_result_rows(self.batch_last_summary) if self.batch_stage in ("finished", "stopped") else []
        counts = summarize_batch_result_rows(all_rows)
        labels = build_batch_filter_labels(counts, show_counts=self.batch_stage in ("finished", "stopped"))
        buttons = [
            ("all", getattr(self, "btn_batch_filter_all", None)),
            ("success", getattr(self, "btn_batch_filter_success", None)),
            ("failed", getattr(self, "btn_batch_filter_failed", None)),
        ]
        enable_filters = self.batch_stage in ("finished", "stopped") and bool(all_rows)
        for mode, button in buttons:
            if not button:
                continue
            button.setText(labels.get(mode, button.text()))
            button.setEnabled(enable_filters)
            button.setCursor(Qt.CursorShape.PointingHandCursor if enable_filters else Qt.CursorShape.ArrowCursor)
            button.setStyleSheet(self._get_batch_filter_button_style(active=self.batch_result_filter_mode == mode))

    def _set_batch_result_filter_mode(self, mode):
        """切换批量结果筛选模式。"""
        if mode not in {"all", "success", "failed"}:
            mode = "all"
        self.batch_result_filter_mode = mode
        self._refresh_batch_result_filter_buttons()
        self._populate_batch_result_table()

    def _populate_batch_result_table(self):
        """将本轮批量结果表格化展示。"""
        if not hasattr(self, "batch_result_table"):
            return

        table = self.batch_result_table
        all_rows = []
        if self.batch_stage in ("finished", "stopped"):
            all_rows = build_batch_result_rows(self.batch_last_summary)

        result_counts = summarize_batch_result_rows(all_rows)
        if hasattr(self, "lbl_batch_result_meta"):
            if self.batch_stage in ("finished", "stopped"):
                self.lbl_batch_result_meta.setText(
                    f"结果计数：共 {result_counts['total']} 条 · 成功 {result_counts['success']} · 失败 {result_counts['failed']}"
                )
            else:
                self.lbl_batch_result_meta.setText("结果计数：等待本轮结果")
        self._refresh_batch_result_filter_buttons()

        rows = filter_batch_result_rows(all_rows, self.batch_result_filter_mode)

        if not rows:
            placeholder_map = {
                "rule_setup": ("等待开始", "当前还在规则确认阶段", "开始执行后这里会列出每个文档的结果", "先确认规则"),
                "running": ("执行中", "系统正在批量处理文档", "本轮结束后这里会集中列出成功与失败明细", "处理中"),
                "finished": ("已完成", "当前没有可展示的结果行", "如果这轮没有生成结果，请检查日志和弹窗", "查看日志"),
                "stopped": ("已停止", "当前没有可展示的结果行", "可以重试失败文档，或者重新选择文件再执行", "查看日志"),
                "idle": ("待执行", "尚未进入批量模式", "选择多个 Word 文档后，这里会展示完整结果清单", "等待开始"),
            }
            if self.batch_stage in ("finished", "stopped") and all_rows:
                status, document, detail, action = (
                    "筛选为空",
                    "当前筛选条件下没有结果",
                    "你可以切回“全部”，或者改看成功 / 失败结果。",
                    "切换筛选",
                )
            else:
                status, document, detail, action = placeholder_map.get(self.batch_stage, placeholder_map["idle"])
            rows = [{
                "status": status,
                "status_key": "placeholder",
                "document": document,
                "detail": detail,
                "action": action,
                "open_path": "",
                "fallback_dir": "",
            }]

        table.clearContents()
        table.setRowCount(len(rows))

        status_colors = {
            "success": Theme.LIGHT["success"],
            "failed": Theme.LIGHT["danger"],
            "placeholder": Theme.LIGHT["text_secondary"],
        }
        status_backgrounds = {
            "success": "#EAF8F1",
            "failed": "#FDECEC",
            "placeholder": "#F4F7FB",
        }

        for row_index, row_data in enumerate(rows):
            status_item = QTableWidgetItem(row_data.get("status", ""))
            status_item.setTextAlignment(int(Qt.AlignmentFlag.AlignCenter))
            status_item.setForeground(QColor(status_colors.get(row_data.get("status_key"), Theme.LIGHT["text"])))
            status_item.setBackground(QColor(status_backgrounds.get(row_data.get("status_key"), "#FBFCFE")))
            status_font = status_item.font()
            status_font.setBold(True)
            status_item.setFont(status_font)
            status_item.setData(Qt.ItemDataRole.UserRole, row_data)

            document_item = QTableWidgetItem(row_data.get("document", ""))
            document_item.setToolTip(row_data.get("open_path", "") or row_data.get("document", ""))

            detail_item = QTableWidgetItem(row_data.get("detail", ""))
            detail_tooltip = row_data.get("detail", "")
            if row_data.get("open_path"):
                detail_tooltip = f"{detail_tooltip}\n{row_data.get('open_path')}"
            detail_item.setToolTip(detail_tooltip.strip())

            action_item = QTableWidgetItem(row_data.get("action", ""))
            action_item.setTextAlignment(int(Qt.AlignmentFlag.AlignCenter))

            if row_data.get("status_key") == "placeholder":
                placeholder_flags = Qt.ItemFlag.ItemIsEnabled
                for item in [status_item, document_item, detail_item, action_item]:
                    item.setFlags(placeholder_flags)

            table.setItem(row_index, 0, status_item)
            table.setItem(row_index, 1, document_item)
            table.setItem(row_index, 2, detail_item)
            table.setItem(row_index, 3, action_item)

        table.resizeRowsToContents()

    def _open_batch_result_row(self, row, _column):
        """双击批量结果行：成功打开输出，失败定位源文档。"""
        if not hasattr(self, "batch_result_table"):
            return

        row_item = self.batch_result_table.item(row, 0)
        row_data = row_item.data(Qt.ItemDataRole.UserRole) if row_item else None
        if not isinstance(row_data, dict) or row_data.get("status_key") == "placeholder":
            return

        open_path = row_data.get("open_path", "")
        fallback_dir = row_data.get("fallback_dir", "")
        document_name = row_data.get("document", "所选文档")

        if open_path and os.path.exists(open_path):
            if QDesktopServices.openUrl(QUrl.fromLocalFile(open_path)):
                return
        if fallback_dir and os.path.isdir(fallback_dir):
            if QDesktopServices.openUrl(QUrl.fromLocalFile(fallback_dir)):
                return

        QMessageBox.warning(self, "提示", f"无法打开 {document_name} 对应的结果路径，请检查文件是否仍然存在。")

    def _refresh_batch_workspace(self):
        """刷新批量 Word 工作台文案。"""
        if not hasattr(self, "lbl_batch_title"):
            return

        stage_map = {
            "idle": ("批量 Word 工作台", "等待开始", Theme.LIGHT["secondary"], Theme.LIGHT["hover"]),
            "rule_setup": ("批量 Word 文档替换规则模式", "规则确认中", Theme.LIGHT["warning"], "#FFF3E6"),
            "running": ("批量 Word 替换执行模式", "执行中", Theme.LIGHT["primary"], "#E9F1FB"),
            "finished": ("批量 Word 替换结果", "已完成", Theme.LIGHT["success"], "#EAF8F1"),
            "stopped": ("批量 Word 替换结果", "已停止", Theme.LIGHT["danger"], "#FDECEC"),
        }
        title, badge_text, badge_fg, badge_bg = stage_map.get(
            self.batch_stage,
            stage_map["idle"],
        )

        self.lbl_batch_title.setText(title)
        self._set_status_badge_style(self.lbl_batch_stage_badge, badge_fg, badge_bg)
        self.lbl_batch_stage_badge.setText(badge_text)

        total_selected = self.batch_total_files or len(self.batch_selected_files)
        rule_count = self._count_enabled_word_rules()
        replacement_preview = self.replacement_text if isinstance(self.replacement_text, str) and self.replacement_text else "[已脱敏]"
        processed = min(self.batch_processed_files, total_selected) if total_selected else 0
        failed_inputs = self._get_batch_failed_inputs()
        success_outputs = self._get_batch_success_outputs()
        summary_parts = []

        if total_selected:
            summary_parts.append(f"已选文档 {total_selected} 个")
        if rule_count:
            summary_parts.append(f"启用规则 {rule_count} 条")
        summary_parts.append(f"统一替换文本：{replacement_preview}")

        if self.batch_stage == "rule_setup":
            subtitle = "先确认文档替换规则，再开始批量替换。这个阶段不会改动原文件。"
        elif self.batch_stage == "running":
            subtitle = "系统正在逐个处理文档。遇到问题时可以跳过当前文件，或者直接停止任务。"
        elif self.batch_stage in ("finished", "stopped"):
            subtitle = "这一轮批量任务已经结束。你可以查看结果摘要，也可以直接重新选择文件再跑一轮。"
            summary_parts.append(f"成功 {self.batch_success_count} 个")
            summary_parts.append(f"失败 {self.batch_failed_count} 个")
        else:
            subtitle = "批量模式会先进入规则确认，再进入执行。适合多个 Word 文档使用同一套替换规则。"

        self.lbl_batch_subtitle.setText(subtitle)
        self.lbl_batch_meta.setText(" · ".join(summary_parts))

        stage_index_map = {
            "idle": 0,
            "rule_setup": 0,
            "running": 1,
            "finished": 2,
            "stopped": 2,
        }
        active_index = stage_index_map.get(self.batch_stage, 0)
        active_fg = Theme.LIGHT["danger"] if self.batch_stage == "stopped" else None
        active_bg = "#FDECEC" if self.batch_stage == "stopped" else None
        if hasattr(self, "batch_stage_cards"):
            for index, (frame, title_label, note_label) in enumerate(self.batch_stage_cards):
                if index < active_index:
                    state = "done"
                elif index == active_index:
                    state = "active"
                else:
                    state = "pending"
                self._set_batch_step_style(frame, title_label, note_label, state, accent_fg=active_fg, accent_bg=active_bg)

        if hasattr(self, "lbl_batch_metric_files"):
            self.lbl_batch_metric_files.setText(f"{total_selected}" if total_selected else "--")
            self.lbl_batch_metric_files_note.setText("本轮已载入的 Word 文档数量")
        if hasattr(self, "lbl_batch_metric_rules"):
            self.lbl_batch_metric_rules.setText(f"{rule_count}" if rule_count else "--")
            self.lbl_batch_metric_rules_note.setText(f"统一替换文本：{replacement_preview}")
        if hasattr(self, "lbl_batch_metric_progress"):
            progress_value = f"{processed}/{total_selected}" if total_selected else "--"
            progress_note = "规则确认完成后开始执行"
            if self.batch_stage == "running":
                progress_note = f"成功 {self.batch_success_count} · 失败 {self.batch_failed_count}"
            elif self.batch_stage in ("finished", "stopped"):
                progress_note = f"{badge_text} · 全部文件已结束本轮处理"
            self.lbl_batch_metric_progress.setText(progress_value)
            self.lbl_batch_metric_progress_note.setText(progress_note)
        if hasattr(self, "lbl_batch_metric_result"):
            if self.batch_stage in ("finished", "stopped", "running"):
                result_value = f"{self.batch_success_count} / {self.batch_failed_count}"
                result_note = "成功 / 失败"
            else:
                result_value = "待执行"
                result_note = "开始后这里会汇总结果"
            self.lbl_batch_metric_result.setText(result_value)
            self.lbl_batch_metric_result_note.setText(result_note)

        if self.batch_current_file:
            self.lbl_batch_current_file.setText(f"当前文件：{self.batch_current_file}")
        elif total_selected:
            self.lbl_batch_current_file.setText("当前文件：等待开始处理")
        else:
            self.lbl_batch_current_file.setText("当前文件：尚未选择批量文档")

        if hasattr(self, "batch_summary_browser"):
            self.batch_summary_browser.setPlainText(self._build_batch_summary_text(total_selected, replacement_preview))
        if hasattr(self, "lbl_batch_result_hint"):
            if self.batch_stage in ("finished", "stopped"):
                self.lbl_batch_result_hint.setText("结果清单（双击成功行可打开输出，双击失败行可定位原文件）")
            elif self.batch_stage == "running":
                self.lbl_batch_result_hint.setText("结果清单（处理中，完成后这里会列出每个文档）")
            else:
                self.lbl_batch_result_hint.setText("结果清单")
        self._populate_batch_result_table()

        self.btn_batch_edit_rules.setEnabled(bool(self.batch_selected_files) and self.active_task_type != "batch_replace")
        self.btn_batch_pick_files.setEnabled(self.active_task_type != "batch_replace")
        show_result_actions = self.batch_stage in ("finished", "stopped")
        self.btn_batch_retry_failed.setVisible(show_result_actions)
        self.btn_batch_open_output.setVisible(show_result_actions)
        self.btn_batch_retry_failed.setEnabled(bool(failed_inputs) and self.active_task_type != "batch_replace")
        self.btn_batch_open_output.setEnabled(bool(success_outputs) and self.active_task_type != "batch_replace")

    def _refresh_merge_workspace(self):
        """刷新图片合并模式提示。"""
        if not hasattr(self, "lbl_merge_meta"):
            return

        total_images = self.image_merge_total_images
        in_progress = self.image_merge_in_progress and total_images > 0

        if in_progress:
            self.lbl_merge_title.setText("图片正在合并为 PDF")
            self.lbl_merge_subtitle.setText("系统会按当前顺序生成 PDF，完成后自动进入 PDF 脱敏工作台。")
            self.lbl_merge_meta.setText(
                f"当前共 {total_images} 张图片。合并完成后会自动打开生成的 PDF，继续进入 PDF 脱敏流程。"
            )
            self._set_status_badge_style(self.lbl_merge_stage_badge, Theme.LIGHT["primary"], "#E9F1FB")
            self.lbl_merge_stage_badge.setText("合并中")
            stage_states = ["done", "active", "pending"]
            if hasattr(self, "lbl_merge_metric_images"):
                self.lbl_merge_metric_images.setText(str(total_images))
                self.lbl_merge_metric_images_note.setText("当前已载入待合并的图片数量")
            if hasattr(self, "lbl_merge_metric_status"):
                self.lbl_merge_metric_status.setText("进行中")
                self.lbl_merge_metric_status_note.setText("底部进度条会同步显示合并进度")
            if hasattr(self, "lbl_merge_metric_next"):
                self.lbl_merge_metric_next.setText("PDF 脱敏")
                self.lbl_merge_metric_next_note.setText("完成后自动进入 PDF 脱敏工作台")
        else:
            self.lbl_merge_title.setText("图片合并为 PDF")
            self.lbl_merge_subtitle.setText("系统会先整理图片顺序，再生成 PDF 并自动打开工作台。")
            self.lbl_merge_meta.setText("当前还没有开始合并。")
            self._set_status_badge_style(self.lbl_merge_stage_badge, Theme.LIGHT["warning"], "#FFF3E6")
            self.lbl_merge_stage_badge.setText("等待开始")
            stage_states = ["active", "pending", "pending"]
            if hasattr(self, "lbl_merge_metric_images"):
                self.lbl_merge_metric_images.setText(str(total_images) if total_images else "--")
                self.lbl_merge_metric_images_note.setText("拖入多张图片后会自动进入该工作台")
            if hasattr(self, "lbl_merge_metric_status"):
                self.lbl_merge_metric_status.setText("未开始")
                self.lbl_merge_metric_status_note.setText("当前尚未开始合并图片")
            if hasattr(self, "lbl_merge_metric_next"):
                self.lbl_merge_metric_next.setText("等待图片")
                self.lbl_merge_metric_next_note.setText("拖入图片后开始合并并进入 PDF 工作台")

        for index, state in enumerate(stage_states):
            if index < len(getattr(self, "merge_stage_cards", [])):
                frame, title_label, note_label = self.merge_stage_cards[index]
                self._set_batch_step_style(frame, title_label, note_label, state)

    def _refresh_workbench_context(self):
        """刷新主工作台标题、步骤和下一步提示。"""
        if not hasattr(self, "lbl_workbench_title"):
            return

        mode = self.current_ui_mode
        active_step = 0
        focus_text = "第 1 步"
        focus_fg = Theme.LIGHT["primary"]
        focus_bg = "#E9F1FB"
        guidance_items = build_workbench_guidance("idle")
        show_focus_badge = mode != "idle"
        show_workbench_subtitle = mode != "idle"

        if mode == "pdf" and self.doc:
            total_pages = len(self.doc)
            current_page = (self.current_page + 1) if self.current_page is not None else 0
            has_results = self._has_pdf_redactions()
            active_step = 3 if has_results else 2
            focus_text = "PDF 脱敏"
            show_focus_badge = False
            self.lbl_workbench_title.setText("PDF 脱敏工作台")
            self.lbl_workbench_subtitle.setText(
                f"{os.path.basename(self.file_path or '')} · {current_page} / {total_pages} 页 · 黑 / 白即时切换"
            )
            guidance_items = build_workbench_guidance("pdf", has_results=has_results)
        elif mode == "word" and self.word_doc:
            paragraph_count = len(self.word_doc.paragraphs)
            table_count = len(self.word_doc.tables)
            has_results = self._has_word_redactions() or self._count_enabled_word_rules() > 0
            active_step = 3 if has_results else 2
            focus_text = "Word 替换"
            focus_fg = Theme.LIGHT["success"]
            focus_bg = "#EAF8F1"
            show_focus_badge = False
            compare_status = "已开启" if self.word_compare_mode else "已隐藏"
            self.lbl_workbench_title.setText("Word 替换工作台")
            self.lbl_workbench_subtitle.setText(
                f"{os.path.basename(self.file_path or '')} · 段落 {paragraph_count} · 表格 {table_count} · 对比 {compare_status}"
            )
            guidance_items = build_workbench_guidance("word", has_results=has_results, compare_mode=self.word_compare_mode)
        elif mode == "batch":
            if self.batch_stage == "running":
                active_step = 2
                focus_text = "批量执行中"
                focus_fg = Theme.LIGHT["primary"]
                focus_bg = "#E9F1FB"
                summary_text = (
                    f"已选 {self.batch_total_files or len(self.batch_selected_files)} 个文档 · "
                    f"成功 {self.batch_success_count} · 失败 {self.batch_failed_count}"
                )
            elif self.batch_stage in ("finished", "stopped"):
                active_step = 4
                focus_text = "批量结果"
                focus_fg = Theme.LIGHT["success"] if self.batch_stage == "finished" else Theme.LIGHT["danger"]
                focus_bg = "#EAF8F1" if self.batch_stage == "finished" else "#FDECEC"
                summary_text = (
                    f"已选 {self.batch_total_files or len(self.batch_selected_files)} 个文档 · "
                    f"成功 {self.batch_success_count} · 失败 {self.batch_failed_count}"
                )
            else:
                active_step = 1
                focus_text = "规则确认"
                focus_fg = Theme.LIGHT["warning"]
                focus_bg = "#FFF3E6"
                summary_text = (
                    f"已选 {self.batch_total_files or len(self.batch_selected_files)} 个文档 · "
                    f"启用规则 {self._count_enabled_word_rules()} 条"
                )

            self.lbl_workbench_title.setText("批量 Word 工作台")
            self.lbl_workbench_subtitle.setText(summary_text)
            guidance_items = build_workbench_guidance("batch", batch_stage=self.batch_stage)
        elif mode == "image_merge":
            active_step = 2
            focus_text = "图片合并"
            self.lbl_workbench_title.setText("图片合并为 PDF")
            self.lbl_workbench_subtitle.setText(f"当前待合并图片：{self.image_merge_total_images} 张 · 完成后自动进入 PDF 脱敏模式")
            guidance_items = build_workbench_guidance("image_merge")
        else:
            self.lbl_workbench_title.setText("欢迎使用 PrivacyGuard")
            self.lbl_workbench_subtitle.setText("拖拽或打开文件即可开始处理。")

        self.lbl_workbench_focus.setText(focus_text)
        self._set_status_badge_style(self.lbl_workbench_focus, focus_fg, focus_bg)
        self.lbl_workbench_focus.setVisible(show_focus_badge)
        self.lbl_workbench_subtitle.setVisible(show_workbench_subtitle)
        self._refresh_workbench_guidance(guidance_items)
        self._refresh_workflow_steps(active_step)
        self._refresh_batch_workspace()
        self._refresh_merge_workspace()

    def _apply_button_variant(self, btn, style_type):
        """根据当前主题为已创建按钮切换样式。"""
        if not btn:
            return
        btn.setProperty("btn_style", style_type)
        btn.setStyleSheet(self._get_button_style(style_type))

    def _refresh_mode_badge(self):
        """刷新顶部模式标识，让用户一眼知道当前在处理什么。"""
        if not hasattr(self, "lbl_mode_badge"):
            return

        mode_map = {
            "idle": ("等待导入", Theme.LIGHT["secondary"], Theme.LIGHT["hover"]),
            "pdf": ("PDF 脱敏模式", Theme.LIGHT["primary"], "#E9F1FB"),
            "word": ("Word 替换模式", Theme.LIGHT["success"], "#EAF8F1"),
            "batch": ("批量 Word 替换", Theme.LIGHT["warning"], "#FFF3E6"),
            "image_merge": ("图片合并中", Theme.LIGHT["primary"], "#E9F1FB"),
        }
        text, fg, bg = mode_map.get(
            self.current_ui_mode,
            ("等待导入", Theme.LIGHT["secondary"], Theme.LIGHT["hover"]),
        )
        self.lbl_mode_badge.setText(text)
        self.lbl_mode_badge.setStyleSheet(
            f"""
            QLabel#modeBadge {{
                background-color: {bg};
                color: {fg};
                border: 1px solid {Theme.LIGHT["border"]};
                border-radius: 14px;
                padding: 6px 12px;
                font-size: {Theme.FONT_SIZE_SMALL}px;
                font-weight: 700;
            }}
            """
        )

    def _set_ui_mode(self, mode):
        """集中管理不同业务模式下的工具显隐，避免 PDF / Word 控件混杂。"""
        self.current_ui_mode = mode
        is_idle = mode == "idle"
        is_pdf = mode == "pdf"
        is_word = mode == "word"
        is_batch = mode == "batch"
        is_image_merge = mode == "image_merge"

        pdf_widgets = [
            self.rb_black, self.rb_white, self.cb_dual, self.btn_fit_utility,
            self.lbl_zoom, self.btn_zoom_out, self.btn_zoom_in,
            self.btn_go_first, self.btn_prev_page, self.lbl_page,
            self.btn_next_page, self.btn_go_last,
        ]
        for widget in pdf_widgets:
            widget.setVisible(is_pdf)
        self.btn_fit.setVisible(False)

        word_widgets = [self.btn_compare_toggle]
        for widget in word_widgets:
            widget.setVisible(is_word)

        show_scan = is_pdf or is_word
        show_save = is_pdf or is_word
        self.btn_open.setVisible(not is_idle)
        self.btn_scan.setVisible(show_scan)
        self.btn_save.setVisible(show_save)

        if is_batch:
            self.btn_scan.setVisible(False)
            self.btn_save.setVisible(False)

        if is_idle:
            self.btn_scan.setVisible(False)
            self.btn_save.setVisible(False)

        if hasattr(self, "idle_workspace_container"):
            self.idle_workspace_container.setVisible(is_idle)
        if hasattr(self, "batch_workspace_container"):
            self.batch_workspace_container.setVisible(is_batch)
        if hasattr(self, "merge_workspace_container"):
            self.merge_workspace_container.setVisible(is_image_merge)
        if hasattr(self, "canvas_container"):
            self.canvas_container.setVisible(is_pdf)
        if hasattr(self, "word_compare_container"):
            self.word_compare_container.setVisible(is_word)
        if hasattr(self, "workbench_panel"):
            self.workbench_panel.setVisible(not is_idle)
        if hasattr(self, "toolbar"):
            self.toolbar.setVisible(not is_idle)
        if hasattr(self, "progress_shell"):
            self.progress_shell.setVisible(not is_idle)

        self._refresh_mode_badge()
        self._refresh_word_compare_toggle()
        self._refresh_toolbar_responsiveness()
        self._refresh_workbench_context()
        self._refresh_info_bar_visibility()

    def _sync_ui_mode(self):
        """根据当前运行状态推导界面模式。"""
        if self.active_task_type == "batch_replace":
            self._set_ui_mode("batch")
        elif self.image_merge_in_progress:
            self._set_ui_mode("image_merge")
        elif self.doc:
            self._set_ui_mode("pdf")
        elif self.word_doc:
            self._set_ui_mode("word")
        elif self.batch_stage != "idle" and self.batch_selected_files:
            self._set_ui_mode("batch")
        else:
            self._set_ui_mode("idle")

    def _refresh_word_compare_toggle(self):
        """刷新 Word 对比预览按钮状态。"""
        if not hasattr(self, "btn_compare_toggle"):
            return

        if not self.word_doc:
            self.btn_compare_toggle.setEnabled(False)
            self.btn_compare_toggle.setText("对比预览")
            self.btn_compare_toggle.setToolTip("请先打开 Word 文档")
            self._apply_button_variant(self.btn_compare_toggle, "secondary")
            self._refresh_toolbar_responsiveness()
            return

        has_candidates = self._has_word_replacement_candidates()
        self.btn_compare_toggle.setEnabled(has_candidates)
        if not has_candidates:
            self.btn_compare_toggle.setText(
                "对比预览（暂无结果）" if self.toolbar_density_mode == "wide" else "暂无对比"
            )
            self.btn_compare_toggle.setToolTip("设置替换规则或执行智能替换后，可查看对比预览")
            self._apply_button_variant(self.btn_compare_toggle, "secondary")
            self._refresh_toolbar_responsiveness()
            return

        if self.word_compare_user_hidden:
            self.btn_compare_toggle.setText(
                "显示对比预览" if self.toolbar_density_mode == "wide" else "显示对比"
            )
            self.btn_compare_toggle.setToolTip("显示右侧替换后预览")
            self._apply_button_variant(self.btn_compare_toggle, "secondary")
        else:
            self.btn_compare_toggle.setText(
                "隐藏对比预览" if self.toolbar_density_mode == "wide" else "隐藏对比"
            )
            self.btn_compare_toggle.setToolTip("隐藏右侧替换后预览")
            self._apply_button_variant(self.btn_compare_toggle, "primary")
        self._refresh_toolbar_responsiveness()

    def toggle_word_compare_preview(self):
        """允许用户主动隐藏或恢复 Word 右侧对比预览。"""
        if not self.word_doc:
            QMessageBox.information(self, "提示", "请先打开 Word 文档。")
            return

        if not self._has_word_replacement_candidates():
            QMessageBox.information(
                self,
                "提示",
                "当前还没有可对比的替换结果。\n请先设置替换规则或执行智能脱敏。"
            )
            self._refresh_word_compare_toggle()
            return

        self.word_compare_user_hidden = not self.word_compare_user_hidden
        self.render_word_preview()
        self._refresh_word_compare_toggle()

    def create_btn(self, text, func, enabled=True, style="primary", width=None, tooltip=""):
        """创建现代化按钮

        Args:
            text: 按钮文本
            func: 点击回调
            enabled: 是否启用
            style: 样式类型 (primary, secondary, success, danger, icon)
            width: 固定宽度（可选）
            tooltip: 工具提示
        """
        btn = QPushButton(text)
        btn.clicked.connect(func)
        btn.setEnabled(enabled)
        if tooltip:
            btn.setToolTip(tooltip)
        if width:
            btn.setFixedWidth(width)

        # 设置游标
        btn.setCursor(Qt.CursorShape.PointingHandCursor)
        btn.setAutoDefault(False)
        btn.setDefault(False)
        btn.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)

        # 应用样式
        btn.setStyleSheet(self._get_button_style(style))

        # 保存样式类型以便主题切换
        btn.setProperty("btn_style", style)

        return btn

    def _create_toolbar_group(self, object_name="toolbarGroup"):
        """创建工具栏分组容器，降低一整排独立按钮的视觉噪音。"""
        group = QFrame()
        group.setObjectName(object_name)
        layout = QHBoxLayout(group)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(6)
        layout.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        return group, layout

    def _create_toolbar_divider(self):
        divider = QFrame()
        divider.setObjectName("toolbarDivider")
        divider.setFixedWidth(1)
        divider.setMinimumHeight(24)
        return divider

    def _get_button_style(self, style_type):
        """获取按钮样式（浅色主题）"""
        theme = Theme.LIGHT
        metrics = getattr(self, "_button_density_metrics", {}) or {}
        button_font_size = metrics.get("button_font_size", 13)
        button_padding_v = metrics.get("button_padding_v", 7)
        button_padding_h = metrics.get("button_padding_h", 14)
        icon_font_size = metrics.get("icon_font_size", 14)
        icon_padding_v = metrics.get("icon_padding_v", 4)
        icon_padding_h = metrics.get("icon_padding_h", 8)
        icon_min = metrics.get("icon_min", 28)

        styles = {
            "primary": f"""
                QPushButton {{
                    background-color: {theme["primary"]};
                    color: white;
                    border: none;
                    border-radius: {Theme.BUTTON_RADIUS}px;
                    padding: {button_padding_v}px {button_padding_h}px;
                    min-height: 0px;
                    font-weight: 600;
                    font-size: {button_font_size}px;
                }}
                QPushButton:hover {{
                    background-color: {Theme.adjust_color(theme["primary"], -15)};
                }}
                QPushButton:pressed {{
                    background-color: {Theme.adjust_color(theme["primary"], -25)};
                }}
                QPushButton:disabled {{
                    background-color: {theme["border"]};
                    color: {theme["text_secondary"]};
                }}
            """,
            "secondary": f"""
                QPushButton {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: {Theme.BUTTON_RADIUS}px;
                    padding: {button_padding_v}px {button_padding_h}px;
                    min-height: 0px;
                    font-size: {button_font_size}px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background-color: {theme["hover"]};
                    border-color: {theme["primary"]};
                }}
                QPushButton:disabled {{
                    background-color: #F5F7FA;
                    color: {theme["text_secondary"]};
                    border-color: {theme["border"]};
                }}
            """,
            "idle_primary": f"""
                QPushButton {{
                    background-color: {theme["primary"]};
                    color: white;
                    border: none;
                    border-radius: 14px;
                    padding: {button_padding_v + 2}px {button_padding_h + 6}px;
                    min-height: 0px;
                    font-size: {max(button_font_size, 14)}px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background-color: {Theme.adjust_color(theme["primary"], -12)};
                }}
                QPushButton:pressed {{
                    background-color: {Theme.adjust_color(theme["primary"], -22)};
                }}
                QPushButton:disabled {{
                    background-color: {theme["border"]};
                    color: {theme["text_secondary"]};
                }}
            """,
            "idle_secondary": f"""
                QPushButton {{
                    background-color: #F5F8FC;
                    color: {theme["text"]};
                    border: 1px solid #D7E2EE;
                    border-radius: 14px;
                    padding: {button_padding_v + 2}px {button_padding_h + 4}px;
                    min-height: 0px;
                    font-size: {max(button_font_size, 14)}px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background-color: #EEF4FA;
                    border-color: #C9D9EA;
                }}
                QPushButton:pressed {{
                    background-color: #E6EEF7;
                    border-color: #C1D2E4;
                }}
                QPushButton:disabled {{
                    background-color: #F5F7FA;
                    color: {theme["text_secondary"]};
                    border-color: {theme["border"]};
                }}
            """,
            "success": f"""
                QPushButton {{
                    background-color: {theme["success"]};
                    color: white;
                    border: none;
                    border-radius: {Theme.BUTTON_RADIUS}px;
                    padding: {button_padding_v}px {button_padding_h}px;
                    min-height: 0px;
                    font-weight: 600;
                    font-size: {button_font_size}px;
                }}
                QPushButton:hover {{
                    background-color: {Theme.adjust_color(theme["success"], -15)};
                }}
                QPushButton:pressed {{
                    background-color: {Theme.adjust_color(theme["success"], -25)};
                }}
                QPushButton:disabled {{
                    background-color: {theme["border"]};
                    color: {theme["text_secondary"]};
                }}
            """,
            "danger": f"""
                QPushButton {{
                    background-color: {theme["danger"]};
                    color: white;
                    border: none;
                    border-radius: {Theme.BUTTON_RADIUS}px;
                    padding: {button_padding_v + 1}px {button_padding_h + 2}px;
                    font-weight: 600;
                    font-size: {max(button_font_size, Theme.FONT_SIZE_NORMAL)}px;
                }}
                QPushButton:hover {{
                    background-color: {Theme.adjust_color(theme["danger"], -15)};
                }}
                QPushButton:pressed {{
                    background-color: {Theme.adjust_color(theme["danger"], -25)};
                }}
            """,
            "icon": f"""
                QPushButton {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 8px;
                    padding: {icon_padding_v}px {icon_padding_h}px;
                    min-width: {icon_min}px;
                    min-height: {icon_min}px;
                    font-size: {icon_font_size}px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background-color: {theme["hover"]};
                    border-color: {theme["primary"]};
                }}
                QPushButton:pressed {{
                    background-color: {theme["pressed"]};
                }}
                QPushButton:disabled {{
                    background-color: #F5F7FA;
                    color: {theme["text_secondary"]};
                    border-color: {theme["border"]};
                }}
            """,
            "toggle": f"""
                QPushButton {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 9px;
                    padding: {button_padding_v}px {button_padding_h}px;
                    min-height: 0px;
                    font-size: {button_font_size}px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background-color: {theme["hover"]};
                    border-color: {theme["primary"]};
                }}
                QPushButton:checked {{
                    background-color: #E9F1FB;
                    color: {theme["primary"]};
                    border-color: #B8D0EA;
                }}
                QPushButton:pressed {{
                    background-color: {theme["pressed"]};
                }}
                QPushButton:disabled {{
                    background-color: #F5F7FA;
                    color: {theme["text_secondary"]};
                    border-color: {theme["border"]};
                }}
            """,
        }

        return styles.get(style_type, styles["primary"])

    def handle_zoom_request(self, delta):
        new_zoom = self.zoom_level + delta
        if new_zoom < ZOOM_MIN: new_zoom = ZOOM_MIN
        if new_zoom > ZOOM_MAX: new_zoom = ZOOM_MAX
        self.zoom_level = new_zoom
        self.render_view()

    def toggle_dual_view(self, checked):
        """
        v22.9: 简化的单/双页切换 - 保持 canvas_container 为固定 widget
        只隐藏/显示内部的 canvas
        """
        self.dual_view = checked

        if checked:
            # 双页模式：显示两个 canvas
            self.canvas_right.show()
        else:
            # 单页模式：只显示 left canvas
            self.canvas_right.hide()

        self.render_view()

    def update_canvas_color(self):
        if self.rb_black.isChecked(): self.current_color = QColor(0,0,0)
        else: self.current_color = QColor(255,255,255)
        self.render_view()

    def open_settings(self):
        # v37.0: 传递配置管理器以支持配置持久化
        # v37.4.0: 移除 OCR 引擎选择，只保留 RapidOCR

        dlg = SettingsDialog(self, self.active_rules, self.use_enhance, self.custom_keywords,
                            self.scan_level, self.offset_x, self.offset_w, self.replacement_text,
                            self.word_replace_rules,
                            config_manager=config)
        if dlg.exec():
            self.active_rules = dlg.selected_rules
            self.use_enhance = dlg.use_enhance
            self.custom_keywords = dlg.custom_keywords
            self.scan_level = dlg.scan_level
            self.offset_x = dlg.offset_x
            self.offset_w = dlg.offset_w
            self.replacement_text = dlg.replacement_text
            self.word_replace_rules = dlg.word_replace_rules
            if self.word_doc:
                if not self._has_word_replacement_candidates():
                    self.word_compare_user_hidden = False
                self.render_word_preview()
            else:
                self._refresh_workbench_context()
            self._clear_info_bar_message()
            msg = self.create_message_box(self, QMessageBox.Icon.Information, "成功", "设置已保存")
            msg.exec()

    def show_feedback(self):
        """显示反馈与开发者信息对话框"""
        dlg = FeedbackDialog(self)
        dlg.exec()

    def _open_feedback(self):
        """从主窗口直接打开反馈问卷链接。"""
        import webbrowser
        webbrowser.open(FEEDBACK_URL)

    def _open_manual(self):
        """从主窗口直接打开使用手册链接。"""
        import webbrowser
        webbrowser.open("https://fcnwakmkeuz7.feishu.cn/docx/M9ojdaGUAoRVv7x3NCAcxkxenUe?from=from_copylink")

    def _show_donate(self):
        """从主窗口直接打开打赏支持对话框。"""
        dialog = DonateDialog(self)
        dialog.exec()

    @staticmethod
    def create_message_box(parent, icon, title, text, buttons=QMessageBox.StandardButton.Ok, default_button=QMessageBox.StandardButton.Ok):
        """创建带有浅色主题样式的消息框（v37.4.1: 修复 Windows 深色模式显示问题）

        Args:
            parent: 父窗口
            icon: QMessageBox.Icon 类型
            title: 标题
            text: 内容文本
            buttons: 按钮类型
            default_button: 默认按钮

        Returns:
            QMessageBox: 配置好的消息框实例
        """
        msg = QMessageBox(parent)
        msg.setIcon(icon)
        msg.setWindowTitle(title)
        msg.setText(text)
        msg.setStandardButtons(buttons)
        msg.setDefaultButton(default_button)

        # 设置窗口标志，防止 Windows 强制应用深色模式
        msg.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        # 应用浅色主题样式
        from theme import Theme
        theme = Theme.LIGHT
        msg.setStyleSheet(f"""
            QMessageBox {{
                background-color: {theme["background"]};
            }}
            QMessageBox QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
                font-weight: 600;
                min-width: 80px;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
        """)

        return msg

    def _on_replacement_changed(self, text):
        """替换文本变化时的处理"""
        self.replacement_text = text if text.strip() else "[已脱敏]"

    def open_word_replace_rules(self):
        """打开 Word 多字段替换规则设置。"""
        if not self.word_doc:
            QMessageBox.information(self, "提示", "请先打开 Word 文档。")
            return

        dlg = WordReplaceRulesDialog(
            self,
            rules=self.word_replace_rules,
            default_replacement_text=self.replacement_text,
            title="Word 替换规则设置",
            apply_text="应用规则"
        )
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        self.word_replace_rules = dlg.rules
        self.replacement_text = dlg.default_replacement_text

        if not self._has_word_replacement_candidates():
            self.word_compare_user_hidden = False
            self._set_word_compare_mode(False)
            self.render_word_preview()
            QMessageBox.information(self, "提示", "当前无可预览的替换结果，已恢复单栏预览。")
            return

        self.render_word_preview()
        self._refresh_word_compare_toggle()
        self._refresh_workbench_context()
        self._set_info_bar_message("🧩 替换后预览已同步展示：规则替换 + 智能脱敏 + 手动脱敏")

    def _has_enabled_word_replace_rules(self):
        normalized = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)
        return any(item.get("enabled", True) and item.get("find") for item in normalized)

    def _has_word_replacement_candidates(self):
        """是否存在可在右侧预览中展示的替换结果（规则/OCR/手动）。"""
        if self._has_enabled_word_replace_rules():
            return True
        for data in self.word_data.values():
            if not isinstance(data, dict):
                continue
            if data.get("manual") or data.get("ocr"):
                return True
        return False

    def _set_word_compare_mode(self, enabled):
        """切换 Word 单栏/双栏预览。"""
        self.word_compare_mode = bool(enabled) and bool(self.word_doc)
        self.canvas_container.hide()
        self.word_compare_container.show()

        if self.word_compare_mode:
            if hasattr(self, "word_compare_header"):
                self.word_compare_header.show()
            if hasattr(self, "lbl_word_original_header"):
                self.lbl_word_original_header.setText("原文预览")
            if hasattr(self, "lbl_word_replaced_header"):
                self.lbl_word_replaced_header.show()
            if hasattr(self, "word_header_divider"):
                self.word_header_divider.show()
            self.word_preview_replaced_panel.show()
            self.word_preview_replaced.show()
            if hasattr(self, "word_compare_layout"):
                self.word_compare_layout.setStretch(0, 1)
                self.word_compare_layout.setStretch(1, 1)
        else:
            if hasattr(self, "word_compare_header"):
                self.word_compare_header.hide()
            if hasattr(self, "lbl_word_original_header"):
                self.lbl_word_original_header.setText("文档预览")
            if hasattr(self, "lbl_word_replaced_header"):
                self.lbl_word_replaced_header.hide()
            if hasattr(self, "word_header_divider"):
                self.word_header_divider.hide()
            self.word_preview_replaced_panel.hide()
            self.word_preview_replaced.hide()
            if hasattr(self, "word_compare_layout"):
                self.word_compare_layout.setStretch(0, 1)
                self.word_compare_layout.setStretch(1, 0)

        self._configure_word_scroll_sync_panels()

    def _build_replaced_preview_fragment(self, source_text, merged_matches):
        """根据合并后的匹配区间构建右侧预览片段（统一高亮样式）。"""
        from html import escape as html_escape

        source_label_map = {
            "rule": "规则替换",
            "manual": "手动脱敏",
            "ocr": "智能脱敏"
        }
        segments = build_replaced_preview_segments(source_text, merged_matches, self.replacement_text)
        fragment_parts = []

        for segment in segments:
            value = segment.get("value", "")
            if value is None:
                value = ""
            if not isinstance(value, str):
                value = str(value)
            escaped_value = html_escape(value).replace("\n", "<br/>")

            if segment.get("type") != "replacement":
                fragment_parts.append(escaped_value)
                continue

            source = str(segment.get("source", "rule"))
            title = source_label_map.get(source, "文本替换")
            rule_name = str(segment.get("rule_name", "")).strip()
            if rule_name:
                title = f"{title} ({rule_name})"

            title_attr = html_escape(title)
            source_attr = html_escape(source)
            fragment_parts.append(
                f'<mark class="replace-preview-highlight" data-source="{source_attr}" title="{title_attr}">{escaped_value}</mark>'
            )

        return "".join(fragment_parts)

    def _build_word_replaced_preview_html(self, base_html):
        """构建替换后 HTML（右侧只读预览）。"""
        if not base_html:
            return ""

        text_blocks = {}
        for key, data in self.word_data.items():
            text = data.get("text", "")
            if text:
                text_blocks[key] = {"text": text, "escaped": text}

        try:
            html = self._add_data_key_attributes(base_html, text_blocks)
            soup = BeautifulSoup(html, "html.parser")
            for element in soup.find_all(attrs={"data-key": True}):
                key = element.get("data-key")
                if not key or key not in self.word_data:
                    continue

                source_text = self.word_data[key].get("text", "")
                merged_matches = merge_word_matches_with_priority(
                    source_text,
                    self.word_replace_rules,
                    self.replacement_text,
                    manual_matches=self.word_data[key].get("manual", []),
                    ocr_matches=self.word_data[key].get("ocr", [])
                )
                if not merged_matches:
                    continue

                replaced_fragment = self._build_replaced_preview_fragment(source_text, merged_matches)
                element.clear()
                fragment_soup = BeautifulSoup(replaced_fragment, "html.parser")
                children = list(fragment_soup.contents)
                if children:
                    for child in children:
                        element.append(child)
                else:
                    element.append(source_text)

            replaced_html = self._wrap_html_document(str(soup))
            style = f"""
            <style>
                body {{ font-family: {PREVIEW_FONT_STACK}; padding: 20px; line-height: 1.6; }}
                p:empty {{ display: none; margin: 0; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                td, th {{ border: 1px solid #ddd; padding: 8px; }}
                mark.replace-preview-highlight {{
                    background-color: #ffd666;
                    color: #2b1f00;
                    border-radius: 3px;
                    padding: 0 2px;
                    box-decoration-break: clone;
                    -webkit-box-decoration-break: clone;
                    box-shadow: inset 0 -1px 0 rgba(43, 31, 0, 0.20);
                }}
            </style>
            """
            if "<head>" in replaced_html:
                replaced_html = replaced_html.replace("<head>", "<head>" + style, 1)
            else:
                replaced_html = style + replaced_html
            return replaced_html
        except (TypeError, ValueError, AttributeError, RuntimeError) as e:
            print(f"[WordReplacePreview] 构建替换预览失败: {e}")
            return self._wrap_html_document(base_html)

    def start_batch_replace(self, file_paths=None):
        """启动 Word 批量替换（支持 .docx + .doc）。"""
        if self.active_worker is not None and self.active_worker.isRunning():
            QMessageBox.warning(self, "提示", "当前有任务正在执行，请稍候。")
            return

        files = list(file_paths or [])
        if not files:
            app = QApplication.instance()
            original_style = app.styleSheet()
            app.setStyleSheet(self._get_file_dialog_style())
            try:
                files, _ = QFileDialog.getOpenFileNames(
                    self,
                    "选择要批量替换的 Word 文件",
                    "",
                    "Word 文档 (*.docx *.doc)"
                )
            finally:
                app.setStyleSheet(original_style)

        if not files:
            return

        self.batch_selected_files = list(files)
        self.batch_total_files = len(files)
        self.batch_processed_files = 0
        self.batch_success_count = 0
        self.batch_failed_count = 0
        self.batch_current_file = ""
        self.batch_last_summary = None
        self.batch_stage = "rule_setup"
        if hasattr(self, "batch_log_list"):
            self.batch_log_list.clear()
        self._append_batch_log(f"已载入 {len(files)} 个 Word 文档，等待确认批量替换规则。", "info")
        self._clear_info_bar_message()
        self._set_ui_mode("batch")
        rules_dlg = WordReplaceRulesDialog(
            self,
            rules=self.word_replace_rules,
            default_replacement_text=self.replacement_text,
            title="Word 批量替换规则设置",
            apply_text="开始批量替换"
        )
        if rules_dlg.exec() != QDialog.DialogCode.Accepted:
            self._clear_info_bar_message()
            self._reset_batch_session_state()
            self._sync_ui_mode()
            return

        self.word_replace_rules = rules_dlg.rules
        self.replacement_text = rules_dlg.default_replacement_text
        normalized_rules = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)
        if not any(item.get("enabled", True) and item.get("find") for item in normalized_rules):
            QMessageBox.warning(self, "提示", "请至少启用一条 Word 替换规则后再开始批量替换。")
            self.batch_stage = "rule_setup"
            self._append_batch_log("当前未启用任何规则，批量替换尚未开始。", "warning")
            self._set_ui_mode("batch")
            return

        self.progress.setValue(0)
        self.btn_cancel_scan.setVisible(True)
        self.btn_cancel_scan.setEnabled(True)
        self._set_info_bar_message(f"📚 批量替换准备中... 共 {len(files)} 个文件")
        self.batch_stage = "running"
        self._append_batch_log(
            f"开始执行批量替换：共 {len(files)} 个文件，启用 {self._count_enabled_word_rules()} 条规则。",
            "info"
        )

        self.batch_worker = WordBatchReplaceWorker(files, normalized_rules, self.replacement_text)
        self.active_worker = self.batch_worker
        self.active_task_type = "batch_replace"
        self._sync_ui_mode()

        self.batch_worker.progress_signal.connect(self._on_batch_replace_progress)
        self.batch_worker.file_done_signal.connect(self._on_batch_replace_file_done)
        self.batch_worker.file_error_signal.connect(self._on_batch_replace_file_error)
        self.batch_worker.finished_signal.connect(self._on_batch_replace_finished)
        self.batch_worker.start()

    def _on_batch_replace_progress(self, processed, total, current_file):
        percent = int(processed / total * 100) if total > 0 else 0
        self.batch_processed_files = processed
        self.batch_total_files = total
        self.batch_current_file = os.path.basename(current_file) if current_file else ""
        self.progress.setValue(percent)
        self._set_info_bar_message(f"📚 批量替换进行中: {processed}/{total} - {current_file}")
        self._refresh_workbench_context()

    def _on_batch_replace_file_done(self, input_path, output_path):
        self.batch_success_count += 1
        self.batch_current_file = os.path.basename(input_path)
        self._append_batch_log(
            f"已完成：{os.path.basename(input_path)} -> {os.path.basename(output_path)}",
            "success"
        )
        self._refresh_workbench_context()
        print(f"[BatchReplace] 完成: {input_path} -> {output_path}")

    def _on_batch_replace_file_error(self, index, input_path, error_msg):
        if not self.batch_worker:
            return

        self.batch_failed_count += 1
        self.batch_current_file = os.path.basename(input_path)
        self._append_batch_log(
            f"处理失败：{os.path.basename(input_path)} - {error_msg}",
            "error"
        )
        self._refresh_workbench_context()

        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("批量替换出错")
        msg.setText(
            f"文件处理失败（{index + 1}）：\n{os.path.basename(input_path)}\n\n"
            f"错误：{error_msg}\n\n"
            "是否跳过该文件继续处理？"
        )
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Abort)
        msg.setDefaultButton(QMessageBox.StandardButton.Yes)
        yes_btn = msg.button(QMessageBox.StandardButton.Yes)
        abort_btn = msg.button(QMessageBox.StandardButton.Abort)
        if yes_btn:
            yes_btn.setText("跳过继续")
        if abort_btn:
            abort_btn.setText("停止任务")

        choice = msg.exec()
        decision = "skip" if choice == QMessageBox.StandardButton.Yes else "stop"
        if decision == "skip":
            self._append_batch_log(f"已跳过：{os.path.basename(input_path)}，继续处理后续文档。", "warning")
        else:
            self._append_batch_log("用户选择停止批量任务，正在安全结束当前流程。", "warning")
        self.batch_worker.provide_error_decision(decision)

    def _on_batch_replace_finished(self, summary):
        self.active_worker = None
        self.batch_worker = None
        self.active_task_type = None
        self.btn_cancel_scan.setVisible(False)
        self.btn_cancel_scan.setEnabled(True)
        self.progress.setValue(0)

        total = int(summary.get("total", 0))
        success = summary.get("success", [])
        failed = summary.get("failed", [])
        stopped = bool(summary.get("stopped", False))
        success_count = len(success)
        failed_count = len(failed)
        self.batch_total_files = total
        self.batch_processed_files = total
        self.batch_success_count = success_count
        self.batch_failed_count = failed_count
        self.batch_last_summary = summary
        self.batch_current_file = ""
        self.batch_stage = "stopped" if stopped else "finished"

        status_text = "已停止" if stopped else "已完成"
        self._set_info_bar_message(f"📚 批量替换{status_text}: 成功 {success_count} / 失败 {failed_count}")
        self._append_batch_log(
            f"批量替换{status_text}：成功 {success_count} 个，失败 {failed_count} 个。",
            "warning" if stopped else "success"
        )
        self._sync_ui_mode()

        lines = [
            f"批量替换{status_text}",
            f"总文件数: {total}",
            f"成功: {success_count}",
            f"失败: {failed_count}"
        ]

        if success:
            lines.append("")
            lines.append("成功输出（最多显示 10 条）:")
            for item in success[:10]:
                lines.append(f"- {item.get('output', '')}")
            if len(success) > 10:
                lines.append(f"... 其余 {len(success) - 10} 条已省略")

        if failed:
            lines.append("")
            lines.append("失败详情（最多显示 10 条）:")
            for item in failed[:10]:
                lines.append(f"- {os.path.basename(item.get('input', ''))}: {item.get('error', '')}")
            if len(failed) > 10:
                lines.append(f"... 其余 {len(failed) - 10} 条已省略")

        QMessageBox.information(self, "批量替换结果", "\n".join(lines))

    def detect_file_type(self, fname):
        """检测文件类型"""
        ext = os.path.splitext(fname)[1].lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.doc':
            return 'doc'
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
            return 'image'
        else:
            return 'unknown'

    def _get_file_dialog_style(self):
        """获取文件对话框样式（v36.2: 使用系统默认按钮样式确保跨平台可读性）"""
        return f"""
            QFileDialog {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                font-family: {Theme.FONT_FAMILY};
                font-size: 12px;
            }}
            QFileDialog QLabel {{ color: #1D1D1F; }}
            /* 按钮使用系统默认样式，确保跨平台可读性 */
            /* macOS/Windows 会自动应用适合的按钮颜色 */
            QFileDialog QListView, QFileDialog QTreeView {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
            }}
            QFileDialog QListView::item:selected, QFileDialog QTreeView::item:selected {{
                background-color: #007AFF;
                color: white;
            }}
            QFileDialog QComboBox {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
                padding: 4px;
            }}
            QFileDialog QLineEdit {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
                padding: 4px;
            }}
        """

    def open_pdf(self):
        """打开文件（支持图片多选和多 Word 批量替换）"""
        try:
            has_active_context = self._has_active_open_context()

            # 已有文档/任务时，保持原有清理策略，避免新旧上下文互相污染。
            if has_active_context:
                self._cleanup_before_open()
                self._cleanup_temp_file()

            # v37.0.6: 使用原生文件对话框，更稳定
            # 不使用 DontUseNativeDialog，让系统处理渲染
            fnames, _ = QFileDialog.getOpenFileNames(
                self, "选择文件", "",
                "支持的文件 (*.pdf *.doc *.docx *.jpg *.jpeg *.png *.bmp *.tiff)"
            )

            if not fnames:
                return

            # 首次从空首页打开文件时，推迟到真正选中文件后再做清理，
            # 避免文件对话框弹出前触发首页抖动/重排。
            if not has_active_context:
                self._cleanup_before_open()
                self._cleanup_temp_file()

            # 根据选择数量处理
            if len(fnames) == 1:
                # 单个文件，按原有逻辑处理
                fname = fnames[0]
                doc_type = self.detect_file_type(fname)
                if doc_type == 'pdf':
                    self._open_pdf_file(fname)
                elif doc_type == 'docx':
                    self._open_word_docx(fname)
                elif doc_type == 'doc':
                    self._open_word_doc(fname)
                elif doc_type == 'image':
                    self._open_images_merge([fname])
                else:
                    QMessageBox.warning(self, "不支持的格式", "请选择 PDF、Word 文档或图片文件")
            else:
                # 多个文件：支持图片合并或 Word 批量替换
                are_all_images = all(self.detect_file_type(f) == 'image' for f in fnames)
                are_all_words = all(self.detect_file_type(f) in ('doc', 'docx') for f in fnames)
                if are_all_images:
                    self._open_images_merge(fnames)
                elif are_all_words:
                    self.start_batch_replace(file_paths=fnames)
                else:
                    QMessageBox.warning(self, "不支持的混合选择",
                        "同时选择多个文件时，仅支持两种场景：\n"
                        "1. 全部是图片（自动合并为PDF）\n"
                        "2. 全部是Word（自动启动批量替换）")
        except (IOError, OSError, ValueError, ConversionError) as e:
            QMessageBox.critical(self, "错误", f"打开文件失败: {str(e)}")

    def _open_pdf_file(self, fname):
        """内部方法：打开 PDF 文件"""
        # 关闭已打开的PDF文档（防止资源泄露）
        if self.doc:
            self.doc.close()
            self.doc = None

        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._reset_batch_session_state()
        self._clear_info_bar_message()
        self.file_path = fname
        self.doc = fitz.open(fname)
        self.doc_type = 'pdf'
        total = len(self.doc)
        self.page_data = {i: {'ocr': [], 'manual': []} for i in range(total)}
        self.current_page = 0
        self.word_doc = None
        self.word_data = {}
        self.word_compare_mode = False
        self.word_compare_user_hidden = False
        self.btn_scan.setEnabled(True)
        self.btn_save.setEnabled(True)

        # 切换显示：显示 canvas，隐藏 Word 预览
        self.canvas_container.show()
        self.word_compare_container.hide()
        self.word_preview.hide()
        self.word_preview_replaced.hide()
        self._sync_ui_mode()

        self.fit_page()

    def _open_word_docx(self, fname):
        """打开 DOCX 文件"""
        try:
            from docx import Document

            # 打开新文件时失效缓存，避免复用旧 HTML
            self._reset_word_preview_cache()
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            self._reset_batch_session_state()
            self.word_compare_mode = False
            self.word_compare_user_hidden = False
            self.word_doc = Document(fname)
            self.file_path = fname
            self.doc_type = 'docx'
            self.doc = None  # 清空 PDF 文档对象
            self.page_data = {}

            # 初始化 word_data 结构
            self.word_data = {}
            for idx, para in enumerate(self.word_doc.paragraphs):
                self.word_data[f'paragraph_{idx}'] = {
                    'type': 'paragraph',
                    'index': idx,
                    'text': para.text,
                    'ocr': [],
                    'manual': []
                }

            # 扫描表格
            for table_idx, table in enumerate(self.word_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                        self.word_data[key] = {
                            'type': 'table_cell',
                            'table': table_idx,
                            'row': row_idx,
                            'cell': cell_idx,
                            'text': cell.text,
                            'ocr': [],
                            'manual': []
                        }

            # 启用按钮
            self.btn_scan.setEnabled(True)
            self.btn_save.setEnabled(True)

            # 根据规则与现有脱敏结果决定是否默认显示双栏预览
            self.word_compare_mode = self._has_word_replacement_candidates()
            self._sync_ui_mode()

            # 显示 HTML 预览
            self.render_word_preview()

            self._clear_info_bar_message()

        except (IOError, OSError, ValueError, KeyError) as e:
            QMessageBox.critical(self, "错误", f"打开 Word 文档失败: {str(e)}")

    def _open_word_doc(self, fname):
        """打开 .doc 文件（通过转换为 .docx）"""
        import shutil
        import subprocess
        import tempfile

        # 检查系统支持（v35.1: 使用增强的跨平台检测）
        support_info = self._check_doc_support()
        method = support_info['recommended']

        if not method:
            # 无可用工具，显示安装指南
            self._show_doc_install_guide()
            return

        # 显示格式限制提示
        tool_name = "LibreOffice" if method == 'libreoffice' else "antiword"
        reply = QMessageBox.question(
            self,
            "格式提示",
            f".doc 是旧版 Word 格式，将使用 {tool_name} 转换。\n\n"
            f"{'（antiword 只保留纯文本，会丢失格式）' if method == 'antiword' else '转换后可能丢失部分格式。'}\n\n"
            "建议先在 Word 中另存为 .docx 格式以获得最佳效果。\n\n"
            "是否继续转换？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.No:
            return

        # 转换并打开
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            docx_path = self._convert_doc_to_docx(fname, method)
            if docx_path:
                # 使用转换后的文件打开
                self.converted_temp_file = docx_path  # 保存临时文件路径以便后续清理
                self._open_word_docx(docx_path)
                self._set_info_bar_message("📝 已完成 .doc 转换，现已进入 Word 替换工作台。")
        except (IOError, OSError, ValueError, RuntimeError) as e:
            QMessageBox.critical(self, "转换失败", f"无法转换 .doc 文件:\n{str(e)}")
        finally:
            QApplication.restoreOverrideCursor()

    def _check_doc_support(self):
        """检查系统是否支持 .doc 格式（v35.1: 增强跨平台检测）"""
        import shutil
        import platform

        system = platform.system()
        result = {
            'libreoffice': False,
            'antiword': False,
            'recommended': None
        }

        # 检查 LibreOffice
        if system == 'Darwin':  # macOS
            # 检查应用程序目录
            libreoffice_app = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
            if os.path.exists(libreoffice_app):
                result['libreoffice'] = True
            # 也检查 PATH
            elif shutil.which('soffice'):
                result['libreoffice'] = True
        elif system == 'Windows':
            # Windows: 检查常见安装路径
            program_files = os.environ.get('ProgramFiles', 'C:\\Program Files')
            program_files_x86 = os.environ.get('ProgramFiles(x86)', 'C:\\Program Files (x86)')
            libreoffice_paths = [
                os.path.join(program_files, 'LibreOffice', 'program', 'soffice.exe'),
                os.path.join(program_files_x86, 'LibreOffice', 'program', 'soffice.exe'),
            ]
            for path in libreoffice_paths:
                if os.path.exists(path):
                    result['libreoffice'] = True
                    break
        else:  # Linux
            if shutil.which('soffice') or shutil.which('libreoffice'):
                result['libreoffice'] = True

        # 检查 antiword
        if shutil.which('antiword'):
            result['antiword'] = True

        # 确定推荐方案
        if result['libreoffice']:
            result['recommended'] = 'libreoffice'
        elif result['antiword']:
            result['recommended'] = 'antiword'

        return result

    def _show_doc_install_guide(self):
        """显示 .doc 转换工具安装指南（v35.1 新增）"""
        import platform
        system = platform.system()

        if system == 'Darwin':  # macOS
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>在终端执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">brew install --cask libreoffice</code>

<h4 style="margin-top:16px;">或使用轻量级方案 antiword</h4>
<p>在终端执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">brew install antiword</code>
<p style="color:#666;margin-top:8px;">注：antiword 只能提取纯文本，会丢失格式</p>
"""
        elif system == 'Windows':
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>请从官网下载安装：</p>
<a href="https://www.libreoffice.org/download/download/">https://www.libreoffice.org/download/</a>
<p style="margin-top:8px;">选择 Windows 版本下载并安装</p>

<h4 style="margin-top:16px;">安装后重启本软件即可</h4>
"""
        else:  # Linux
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>根据您的发行版执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">
# Debian/Ubuntu
sudo apt install libreoffice

# Fedora
sudo dnf install libreoffice

# Arch Linux
sudo pacman -S libreoffice
</code>

<h4 style="margin-top:16px;">或使用轻量级方案 antiword</h4>
<code style="background:#f5f5f5;padding:8px;display:block;">
# Debian/Ubuntu
sudo apt install antiword

# Fedora
sudo dnf install antiword
</code>
<p style="color:#666;margin-top:8px;">注：antiword 只能提取纯文本，会丢失格式</p>
"""

        msg = QMessageBox(self)
        msg.setWindowTitle("缺少 .doc 转换工具")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(guide)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setStyleSheet("""
            QMessageBox {
                min-width: 450px;
            }
            QLabel {
                min-width: 400px;
            }
        """)
        msg.exec()

    def _convert_doc_to_docx(self, doc_path, method='libreoffice'):
        """v37.7.6: 委托给共享转换模块。"""
        from privacyguard.utils.doc_converter import (
            convert_with_libreoffice, convert_with_antiword,
        )
        try:
            if method == 'libreoffice':
                docx_path, temp_dir = _shared_convert_doc_to_docx(doc_path)
                self.converted_temp_file = docx_path
                return docx_path
            elif method == 'antiword':
                temp_dir = self.temp_manager.create_temp_dir()
                return convert_with_antiword(doc_path, temp_dir=temp_dir)
            else:
                raise ValueError(f"不支持的转换方法: {method}")
        except ConversionError:
            raise
        except (OSError, IOError, RuntimeError, ValueError) as e:
            raise ConversionError(f"转换出错: {e}", "请尝试在 Word 中手动另存为 .docx 格式")

    def _open_images_merge(self, image_paths):
        """处理图片合并为PDF"""
        try:
            self.image_merge_in_progress = True
            self.image_merge_total_images = len(image_paths)
            self._set_ui_mode("image_merge")
            # 1. 让用户排序（如果是多张图片）
            if len(image_paths) > 1:
                dlg = ImageListDialog(image_paths, self)
                if dlg.exec() != QDialog.DialogCode.Accepted:
                    self.image_merge_in_progress = False
                    self.image_merge_total_images = 0
                    self._sync_ui_mode()
                    return
                image_paths = dlg.get_ordered_paths()
                self.image_merge_total_images = len(image_paths)

            # 2. 生成输出路径（在第一张图片所在目录）
            base_dir = os.path.dirname(image_paths[0])
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_filename = f"merged_{timestamp}.pdf"
            final_path = os.path.join(base_dir, output_filename)

            # 3. 创建临时文件
            temp_pdf = self.temp_manager.create_temp_file(suffix='.pdf')

            # 4. 启动Worker合并
            self.merge_worker = ImageMergeWorker(image_paths, temp_pdf)
            self.merge_worker.progress_signal.connect(self.progress.setValue)
            self.merge_worker.finished_signal.connect(lambda p: self._on_merge_finished(p, final_path))
            self.merge_worker.error_signal.connect(self._on_merge_error)
            self.merge_worker.start()

            self._set_info_bar_message(f"正在合并 {len(image_paths)} 张图片...")
            self._refresh_workbench_context()

        except (IOError, OSError, ValueError, RuntimeError) as e:
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            self._sync_ui_mode()
            QMessageBox.critical(self, "错误", f"启动图片合并失败: {str(e)}")

    def _on_merge_finished(self, temp_pdf, final_path):
        """合并完成回调"""
        try:
            # 将临时文件移动到最终位置
            shutil.move(temp_pdf, final_path)
            self._set_info_bar_message(f"✓ 合并完成: {final_path}")
            self.image_merge_in_progress = False

            # 自动打开生成的PDF
            self._open_pdf_file(final_path)

        except (IOError, OSError, ValueError) as e:
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            QMessageBox.critical(self, "错误", f"保存合并文件失败: {str(e)}")

    def _on_merge_error(self, error_msg):
        """合并错误回调"""
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._set_info_bar_message("✗ 合并失败")
        self._sync_ui_mode()
        QMessageBox.critical(self, "合并失败", error_msg)

    def _cleanup_temp_file(self):
        """清理转换产生的临时文件（v36.1 安全修复：跨平台安全清理）"""
        import tempfile

        if hasattr(self, 'converted_temp_file') and self.converted_temp_file:
            try:
                if os.path.exists(self.converted_temp_file):
                    os.remove(self.converted_temp_file)
                    print(f"[清理] 已删除临时文件: {self.converted_temp_file}")

                # 清理临时目录（安全版本）
                temp_dir = os.path.dirname(self.converted_temp_file)

                # 获取系统临时目录（跨平台：Windows 返回 C:\Users\...\AppData\Local\Temp，macOS/Linux 返回 /tmp 或 /var/tmp）
                system_temp_dir = tempfile.gettempdir()

                # 规范化路径进行比较（处理大小写敏感/不敏感、路径分隔符等）
                norm_temp_dir = os.path.normcase(os.path.abspath(temp_dir))
                norm_system_temp = os.path.normcase(os.path.abspath(system_temp_dir))

                # 安全检查：只有当目录位于系统临时目录下，且不是系统临时目录本身时才删除
                if (norm_temp_dir.startswith(norm_system_temp + os.sep) and
                    norm_temp_dir != norm_system_temp):
                    try:
                        # 检查目录是否为空（防止误删文件）
                        if os.path.isdir(temp_dir) and not os.listdir(temp_dir):
                            os.rmdir(temp_dir)
                            print(f"[清理] 已删除空临时目录: {temp_dir}")
                        else:
                            print(f"[清理] 临时目录非空或不存在，跳过删除: {temp_dir}")
                    except OSError as e:
                        print(f"[清理] 删除临时目录时出错（可能是非空）: {temp_dir} - {e}")
                else:
                    print(f"[清理] 跳过非系统临时目录: {temp_dir}")

            except Exception as e:
                print(f"[清理] 清理临时文件时出错: {e}")
            finally:
                self.converted_temp_file = None

    def clamp_zoom(self, zoom, allow_below_min=False):
        """
        将缩放值限制在有效范围内 (v33.2)

        Args:
            zoom: 缩放值
            allow_below_min: 是否允许低于 ZOOM_MIN (用于自适应模式)
        """
        if allow_below_min:
            # 自适应模式：允许更小的缩放比例以完整显示页面
            return min(ZOOM_MAX, zoom)
        else:
            # 手动模式：保持正常限制，防止过小
            return max(ZOOM_MIN, min(ZOOM_MAX, zoom))

    def fit_page(self):
        """完整适应页面 - 根据窗口和页面尺寸动态计算缩放比例"""
        if not self.doc or self.current_page is None:
            return

        # 获取画布可用尺寸（减去边距）
        canvas_width = self.scroll.width() - 40  # 40px 边距
        canvas_height = self.scroll.height() - 40

        # 获取当前页面的实际尺寸（点单位）
        page = self.doc[self.current_page]
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height

        # 分别计算宽度和高度的缩放比例
        zoom_w = canvas_width / page_width
        zoom_h = canvas_height / page_height

        # 取较小值确保页面完整显示在窗口中
        self.zoom_level = min(zoom_w, zoom_h)

        # 限制在最大最小范围内 (v33.2: 允许突破 ZOOM_MIN 以完整显示)
        self.zoom_level = self.clamp_zoom(self.zoom_level, allow_below_min=True)

        # 重新渲染
        self.render_view()

    def render_view(self):
        if not self.doc: return
        # v37.0.9: 添加 canvas 有效性检查
        if not self._is_canvas_valid(self.canvas_left):
            print("[警告] canvas_left 无效，跳过渲染")
            return
        self._render_single_page(self.canvas_left, self.current_page)
        if self.dual_view:
            if self.current_page + 1 < len(self.doc):
                if self._is_canvas_valid(self.canvas_right):
                    self._render_single_page(self.canvas_right, self.current_page + 1)
                    self.canvas_right.show()
            else:
                if self._is_canvas_valid(self.canvas_right):
                    self.canvas_right.hide()

        total = len(self.doc)
        display = f"{self.current_page + 1}"
        if self.dual_view and self.current_page + 1 < total:
            display += f"-{self.current_page + 2}"
        self.lbl_page.setText(f"{display} / {total}")
        self.lbl_zoom.setText(f"{int(self.zoom_level * 100)}%")
        self._refresh_toolbar_responsiveness()
        self._refresh_workbench_context()

    def _render_single_page(self, canvas, page_idx):
        """v7.0 风格渲染 - 直接传递列表引用
        v37.0.9: 添加异常处理防止 canvas 被删除后崩溃
        """
        # 检查 canvas 有效性
        if not self._is_canvas_valid(canvas):
            return

        try:
            page = self.doc[page_idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(self.zoom_level, self.zoom_level))
            img_fmt = QImage.Format.Format_RGB888
            qimg = QImage(pix.samples, pix.width, pix.height, pix.stride, img_fmt).copy()
            data = self.page_data[page_idx]
            # 使用安全的更新方法
            self._safe_canvas_update(canvas, QPixmap.fromImage(qimg), self.zoom_level, data['ocr'], data['manual'])
            self._safe_canvas_set_mask_color(canvas, self.current_color)
        except RuntimeError as e:
            print(f"[错误] 渲染页面 {page_idx} 时出错: {e}")
        except Exception as e:
            print(f"[错误] 渲染页面 {page_idx} 时发生意外错误: {e}")

    def on_rect_added(self, page_idx, rect):
        """由于使用共享列表引用，canvas 已直接修改列表，这里只需刷新视图"""
        # self.page_data[page_idx]['manual'].append(rect)  # 不需要，canvas 已经添加
        self.render_view()

    def on_rect_removed(self, page_idx, rect_idx, is_manual):
        """由于使用共享列表引用，canvas 已直接修改列表，这里只需刷新视图"""
        # if is_manual: del self.page_data[page_idx]['manual'][rect_idx]  # 不需要，canvas 已经删除
        # else: del self.page_data[page_idx]['ocr'][rect_idx]
        self.render_view()

    def change_page(self, delta):
        if not self.doc: return
        step = 2 if self.dual_view else 1
        new_page = self.current_page + (delta * step)
        if new_page < 0: new_page = 0
        if new_page >= len(self.doc): return
        self.current_page = new_page
        self.render_view()
        self.scroll.verticalScrollBar().setValue(0)

    def go_first(self):
        if not self.doc: return
        self.current_page = 0
        self.render_view()
        self.scroll.verticalScrollBar().setValue(0)

    def go_last(self):
        if not self.doc: return
        self.current_page = len(self.doc) - 1
        if self.dual_view and self.current_page % 2 != 0: self.current_page -= 1
        self.render_view()
        self.scroll.verticalScrollBar().setValue(0)

    def handle_page_change_request(self, delta):
        """处理滚轮翻页请求（v35.1 新增）

        Args:
            delta: 翻页数量，正值=向后翻页，负值=向前翻页
                   1/-1 = 普通滚轮（需检测边缘）
                   2/-2 = Shift+滚轮（快速翻页）
        """
        if not self.doc:
            return

        # 快速翻页（Shift+滚轮）：直接翻页，不检测边缘
        if abs(delta) >= 2:
            self.change_page(1 if delta > 0 else -1)
            return

        # 普通滚轮：检测滚动条位置，只在边缘时翻页
        scroll_bar = self.scroll.verticalScrollBar()
        at_top = scroll_bar.value() <= scroll_bar.minimum() + 10
        at_bottom = scroll_bar.value() >= scroll_bar.maximum() - 10

        # 向上滚动且在顶部 → 上一页
        if delta < 0 and at_top and self.current_page > 0:
            self.change_page(-1)
            # 翻页后滚动到底部，便于连续向上翻页
            QApplication.processEvents()
            scroll_bar.setValue(scroll_bar.maximum())
        # 向下滚动且在底部 → 下一页
        elif delta > 0 and at_bottom and self.current_page < len(self.doc) - 1:
            self.change_page(1)
            # 翻页后滚动到顶部
            scroll_bar.setValue(0)

    def keyPressEvent(self, event):
        """键盘快捷键处理（v35.1 新增）

        快捷键列表：
        - PageUp: 上一页
        - PageDown: 下一页
        - Home: 首页
        - End: 尾页
        - Space/Shift+Space: 翻页
        - Ctrl/Cmd + +/-: 缩放
        - Ctrl/Cmd + 0: 适应页面
        """
        key = event.key()
        modifiers = event.modifiers()

        # PageUp: 上一页
        if key == Qt.Key.Key_PageUp:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(-5)  # Shift+PageUp 快速翻页
            else:
                self.change_page(-1)
        # PageDown: 下一页
        elif key == Qt.Key.Key_PageDown:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(5)  # Shift+PageDown 快速翻页
            else:
                self.change_page(1)
        # Home: 首页
        elif key == Qt.Key.Key_Home:
            self.go_first()
        # End: 尾页
        elif key == Qt.Key.Key_End:
            self.go_last()
        # Space: 下一页（Shift+Space: 上一页）
        elif key == Qt.Key.Key_Space:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(-1)
            else:
                self.change_page(1)
        # Ctrl/Cmd + Plus: 放大
        elif key == Qt.Key.Key_Plus or key == Qt.Key.Key_Equal:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_in()
            else:
                super().keyPressEvent(event)
        # Ctrl/Cmd + Minus: 缩小
        elif key == Qt.Key.Key_Minus:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_out()
            else:
                super().keyPressEvent(event)
        # Ctrl/Cmd + 0: 重置缩放
        elif key == Qt.Key.Key_0:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_level = 1.0
                self.render_view()
            else:
                super().keyPressEvent(event)
        else:
            super().keyPressEvent(event)

    def zoom_in(self):
        self.handle_zoom_request(0.25)

    def zoom_out(self):
        self.handle_zoom_request(-0.25)

    def start_ocr(self):
        """智能扫描 - 支持 PDF 和 Word（v37.0.5: 增强错误处理）"""
        # 线程安全检查：防止重复启动
        if self.active_worker is not None:
            if self.active_worker.isRunning():
                QMessageBox.warning(self, "提示", "正在处理中，请稍候...")
                return

        self._set_info_bar_message("🔍 正在扫描敏感信息...")
        self.btn_scan.setEnabled(False)
        self.btn_cancel_scan.setVisible(True)  # 显示取消按钮（v36.3）
        self.btn_cancel_scan.setEnabled(True)
        self.active_task_type = "scan"

        # PDF 处理
        if self.doc:
            # v37.5.0: 检测是否启用印章检测
            seal_detection_enabled = "__SEAL_DETECTION__" in self.active_rules
            print(f"[OCR] active_rules: {self.active_rules}")
            print(f"[OCR] 印章检测启用: {seal_detection_enabled}")
            self._ocr_processed_pages = set()
            # v37.4.0: 只使用 RapidOCR，移除 use_char_level_ocr 参数
            self.worker = OCRWorker(self.file_path, self.active_rules, self.use_enhance, self.custom_keywords,
                                    self.scan_level, self.offset_x, self.offset_w,
                                    seal_detection_enabled=seal_detection_enabled)
            self.active_worker = self.worker  # 追踪线程
            self.worker.progress_signal.connect(self.progress.setValue)
            # v36.4: 使用线程安全的逐页结果信号
            self.worker.page_result_signal.connect(self._on_ocr_page_result)
            # v37.0.5: 连接错误信号
            self.worker.error_signal.connect(self._on_ocr_error)
            # 先连接原有的完成处理，再连接清理
            self.worker.finished_signal.connect(self._on_ocr_finished_safe)
            self.worker.finished_signal.connect(self._on_worker_finished)
            self.worker.start()
        # Word 处理
        elif self.word_doc:
            self.worker = WordWorker(self.word_doc, self.word_data, self.active_rules,
                                     self.custom_keywords, self.replacement_text)
            self.active_worker = self.worker  # 追踪线程
            self.worker.progress_signal.connect(self.progress.setValue)
            # 先连接原有的完成处理，再连接清理
            self.worker.finished_signal.connect(self.word_scan_finished)
            self.worker.finished_signal.connect(self._on_worker_finished)
            self.worker.start()

    def cancel_ocr_scan(self):
        """取消智能脱敏扫描（v36.3）"""
        if self.active_worker and self.active_worker.isRunning():
            if self.active_task_type == "batch_replace":
                title = "确认停止"
                text = "确定要停止批量替换吗？\n已完成的文件会被保留。"
            else:
                title = "确认取消"
                text = "确定要停止扫描吗？\n已扫描的进度将被保留。"
            reply = QMessageBox.question(
                self,
                title,
                text,
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                if self.active_task_type == "batch_replace":
                    self._set_info_bar_message("⏹️ 正在停止批量替换...")
                    self._append_batch_log("收到停止指令，正在安全结束批量任务。", "warning")
                else:
                    self._set_info_bar_message("⏹️ 正在停止扫描...")
                self.btn_cancel_scan.setEnabled(False)  # 防止重复点击
                self.active_worker.requestInterruption()  # 请求中断
                if self.active_task_type == "batch_replace" and hasattr(self, "batch_worker") and self.batch_worker:
                    self.batch_worker.provide_error_decision("stop")
                # Worker会在完成后通过finished_signal通知主线程

    def _on_worker_finished(self):
        """v37.0.6: 工作线程完成后的清理 + 延迟错误显示"""
        # v37.0.6: 等待线程完全终止，防止死锁
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.wait(3000)  # 最多等待 3 秒

        was_cancelled = self.active_worker and self.active_worker.isInterruptionRequested()
        self.active_worker = None
        self.active_task_type = None
        self.btn_scan.setEnabled(True)
        self.btn_cancel_scan.setVisible(False)  # 隐藏取消按钮
        self.btn_cancel_scan.setEnabled(True)
        self._sync_ui_mode()

        if was_cancelled:
            self._set_info_bar_message("⏹️ 扫描已取消，已保留部分结果")
        else:
            self._set_info_bar_message("✅ 扫描完成！")

        # v37.0.6: 线程清理完成后，延迟显示错误对话框（非阻塞）
        if hasattr(self, '_pending_error_msg') and self._pending_error_msg:
            error_msg = self._pending_error_msg
            self._pending_error_msg = None
            QTimer.singleShot(100, lambda: self._show_deferred_error(error_msg))

    def _show_deferred_error(self, error_msg: str):
        """v37.0.6: 安全显示错误对话框（在线程清理完成后调用）"""
        QMessageBox.critical(
            self,
            "OCR 错误",
            f"智能脱敏遇到问题：\n\n{error_msg}\n\n"
            "可能的解决方案：\n"
            "1. 重新安装 OCR 依赖: pip install rapidocr-onnxruntime\n"
            "2. 安装 Visual C++ 运行库（Windows）\n"
            "3. 如果是扫描版 PDF，请尝试使用文本版 PDF"
        )

    def ocr_finished(self, results):
        """v23.1: 添加去重逻辑，解决重复矩形导致的点击2次问题；v36.3: 支持部分结果"""
        total_pages = len(self.page_data)
        scanned_pages = len(results)

        for page, rects in results.items():
            # 去重：移除重复或高度重叠的矩形
            deduped = self._deduplicate_rects(rects)
            self.page_data[page]['ocr'] = deduped
            if len(rects) != len(deduped):
                if DEBUG_MODE:
                    print(f"[DEBUG] 页面{page}: 去重前{len(rects)}个矩形, 去重后{len(deduped)}个")
        self.render_view()
        self.progress.setValue(0)

        # 判断是部分结果还是完整结果（v36.3）
        if scanned_pages < total_pages:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已扫描 {scanned_pages}/{total_pages} 页，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", "智能扫描已完成！")

    # v37.0.6: 非阻塞 OCR 错误处理
    def _on_ocr_error(self, error_msg: str):
        """v37.0.6: 非阻塞处理 OCR 错误（存储错误，延迟到线程清理后显示）"""
        print(f"[OCR ERROR] {error_msg}")
        self._set_info_bar_message(f"❌ OCR 错误: {error_msg[:50]}...")
        # v37.0.6: 存储错误消息，延迟到线程清理完成后显示
        # 避免在模态对话框阻塞主线程时形成死锁
        self._pending_error_msg = error_msg

    # v36.4: 线程安全的 OCR 结果处理方法
    def _on_ocr_page_result(self, page_num: int, rects: list):
        """v36.4: 线程安全 - 接收单页 OCR 结果（在主线程执行）"""
        self._ocr_processed_pages.add(page_num)
        if page_num in self.page_data:
            # 去重：移除重复或高度重叠的矩形
            deduped = self._deduplicate_rects(rects)
            self.page_data[page_num]['ocr'] = deduped
            if DEBUG_MODE and len(rects) != len(deduped):
                print(f"[DEBUG] 页面{page_num}: 去重前{len(rects)}个矩形, 去重后{len(deduped)}个")

    def _on_ocr_finished_safe(self, _):
        """v36.4: 线程安全 - OCR 完成处理（在主线程执行）

        参数 _ 是空字典，保留以兼容信号签名
        """
        self.render_view()
        self.progress.setValue(0)

        # 统计已扫描的页面
        scanned_pages = len(self._ocr_processed_pages) if self._ocr_processed_pages else 0
        total_pages = len(self.page_data)
        was_cancelled = self.active_worker is not None and self.active_worker.isInterruptionRequested()

        # 判断是部分结果还是完整结果（以取消状态为准，避免“无命中=已取消”误判）
        if was_cancelled:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已扫描 {scanned_pages}/{total_pages} 页，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", "智能扫描已完成！")

    def _deduplicate_rects(self, rects):
        """移除重复或高度重叠的矩形"""
        if not rects:
            return rects

        # 按中心点坐标和尺寸排序，便于去重
        sorted_rects = sorted(rects, key=lambda r: (r.x(), r.y(), r.width(), r.height()))
        deduped = []

        for rect in sorted_rects:
            is_duplicate = False
            for existing in deduped:
                # 检查是否与已有矩形高度重叠（IoU > 0.7 或中心点距离 < 5 像素）
                if self._is_overlapping(rect, existing):
                    is_duplicate = True
                    break

            if not is_duplicate:
                deduped.append(rect)

        return deduped

    def _is_overlapping(self, rect1, rect2, threshold=0.7):
        """检查两个矩形是否高度重叠"""
        # 计算交集
        x_left = max(rect1.x(), rect2.x())
        y_top = max(rect1.y(), rect2.y())
        x_right = min(rect1.x() + rect1.width(), rect2.x() + rect2.width())
        y_bottom = min(rect1.y() + rect1.height(), rect2.y() + rect2.height())

        if x_right < x_left or y_bottom < y_top:
            return False

        # 计算交集面积
        intersection = (x_right - x_left) * (y_bottom - y_top)
        area1 = rect1.width() * rect1.height()
        area2 = rect2.width() * rect2.height()

        # IoU > threshold 视为重复
        union = area1 + area2 - intersection
        iou = intersection / union if union > 0 else 0

        # 或者中心点距离 < 5 像素也视为重复
        center1_x = rect1.x() + rect1.width() / 2
        center1_y = rect1.y() + rect1.height() / 2
        center2_x = rect2.x() + rect2.width() / 2
        center2_y = rect2.y() + rect2.height() / 2
        distance = ((center1_x - center2_x)**2 + (center1_y - center2_y)**2)**0.5

        return iou > threshold or distance < 5

    def word_scan_finished(self, results):
        """Word 文档扫描完成（v36.5: 线程安全）"""
        results_copy = dict(results)
        scan_meta = results_copy.pop('__scan_meta__', None)

        # v36.5: 使用锁保护 word_data 访问
        with QMutexLocker(self._word_data_lock):
            total_items = len(self.word_data)
            if scan_meta:
                processed_items = int(scan_meta.get('processed_items', total_items))
                total_from_meta = int(scan_meta.get('total_items', total_items))
                if total_from_meta > 0:
                    total_items = total_from_meta
            else:
                processed_items = total_items

            self.word_data = results_copy

            # 统计扫描结果
            total_matches = sum(len(data['ocr']) for data in self.word_data.values())

        self.render_word_preview()
        self.progress.setValue(0)

        was_cancelled = False
        if scan_meta is not None:
            was_cancelled = bool(scan_meta.get('cancelled', False))
        elif self.active_worker is not None:
            was_cancelled = self.active_worker.isInterruptionRequested()

        # 判断是部分结果还是完整结果（以取消状态为准）
        if was_cancelled:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已处理 {processed_items}/{total_items} 个段落/单元格，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", f"智能扫描已完成！\n共发现 {total_matches} 处敏感信息")

    def render_word_preview(self):
        """渲染 Word 文档预览（HTML）"""
        if not self.word_doc:
            return

        try:
            source_changed = self._word_base_html is None or self._word_html_source_path != self.file_path

            # 仅在源文件变化或无缓存时进行 docx -> html 转换
            if source_changed:
                self._word_base_html = self._build_word_html_from_docx(self.file_path)
                self._word_html_source_path = self.file_path
                self._build_word_preview_documents()

            has_candidates = self._has_word_replacement_candidates()
            if not has_candidates:
                self.word_compare_user_hidden = False
            compare_enabled = has_candidates and not self.word_compare_user_hidden
            self.word_compare_mode = compare_enabled
            self._set_word_compare_mode(compare_enabled)
            self._refresh_word_compare_toggle()
            self.word_preview.show()

            # v37.6.1: 禁用 Word 预览的拖拽接受，让事件传递到 MainWindow
            # 解决 Word 打开后无法拖拽打开新文件的问题
            self.word_preview.setAcceptDrops(False)
            self.word_preview_replaced.setAcceptDrops(False)

            # 设置 WebChannel（仅首次）
            if not hasattr(self, 'bridge') or self.bridge is None:
                channel = QWebChannel(self)
                self.bridge = WebViewBridge(self, self)
                channel.registerObject("pyBridge", self.bridge)
                self.word_web_channel = channel
            if self.word_web_channel is not None:
                self.word_preview.page().setWebChannel(self.word_web_channel)
                self.word_preview_replaced.page().setWebChannel(self.word_web_channel)

            self._pending_word_preview_blocks = self._build_word_original_panel_updates()
            base_url = self._word_preview_assets_base_url if self._word_preview_assets_base_url.isValid() else QUrl()
            if should_reload_word_panel(
                source_changed,
                self._word_preview_loaded_source_path,
                self.file_path,
                self._word_preview_ready
            ):
                self._word_preview_target_source_path = self.file_path
                self.word_preview.setHtml(self._word_preview_document_html or "", base_url)
            else:
                self._apply_word_panel_updates(self.word_preview, self._pending_word_preview_blocks)

            if compare_enabled:
                self._pending_word_replaced_blocks = self._build_word_replaced_panel_updates()
                self._word_replaced_html = self._word_replaced_document_html
                if should_reload_word_panel(
                    source_changed,
                    self._word_replaced_loaded_source_path,
                    self.file_path,
                    self._word_replaced_ready
                ):
                    self._word_replaced_target_source_path = self.file_path
                    self.word_preview_replaced.setHtml(self._word_replaced_document_html or "", base_url)
                else:
                    self._apply_word_panel_updates(self.word_preview_replaced, self._pending_word_replaced_blocks)
            else:
                self._word_replaced_html = None
                self._pending_word_replaced_blocks = None
                self._word_replaced_ready = False
                self._word_replaced_loaded_source_path = None
                self._word_replaced_target_source_path = None
                self.word_preview_replaced.setHtml("")

            self._configure_word_scroll_sync_panels()
            if compare_enabled:
                QTimer.singleShot(0, self._sync_word_compare_scroll_from_original)

            self._refresh_workbench_context()

        except (IOError, OSError, ValueError, RuntimeError) as e:
            QMessageBox.critical(self, "错误", f"渲染预览失败: {str(e)}")

    def _build_word_preview_documents(self):
        """构建只需首屏加载一次的 Word 预览文档。"""
        base_html = self._wrap_html_document(self._word_base_html or "")
        text_blocks = self._build_word_text_blocks()
        tagged_html = self._add_data_key_attributes(base_html, text_blocks)

        style = f"""
        <style>
            body {{ font-family: {PREVIEW_FONT_STACK}; padding: 20px; line-height: 1.6; }}
            p:empty {{ display: none; margin: 0; }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; vertical-align: top; }}
            mark.ocr-highlight {{ background-color: #ffeb3b; color: #000; display: inline; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight {{ background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight:hover {{ box-shadow: 0 0 0 1px #e03131, 0 0 4px rgba(225, 49, 49, 0.5); }}
            mark.replace-preview-highlight {{
                background-color: #ffd666;
                color: #2b1f00;
                border-radius: 3px;
                padding: 0 2px;
                box-decoration-break: clone;
                -webkit-box-decoration-break: clone;
                box-shadow: inset 0 -1px 0 rgba(43, 31, 0, 0.20);
            }}
        </style>
        """

        if "<head>" in tagged_html:
            tagged_html = tagged_html.replace("<head>", "<head>" + style, 1)
        else:
            tagged_html = style + tagged_html

        scroll_restore_script = self._get_word_preview_scroll_restore_script()
        document_html = self._inject_interactive_html(tagged_html, scroll_restore_script)

        self._word_tagged_html = tagged_html
        self._word_preview_document_html = document_html
        self._word_replaced_document_html = document_html
        self._word_preview_ready = False
        self._word_replaced_ready = False

    def _get_word_preview_scroll_restore_script(self):
        return '''<script>
    (function() {
        const STORAGE_KEY = 'word_preview_scroll_pos';
        let memoryScrollPos = 0;
        let syncTimeout = null;
        let syncFrame = null;
        let syncInterval = null;
        let lastNotifiedRatio = null;
        window.__wordPreviewPanelId = '';
        window.__wordScrollSyncEnabled = false;
        window.__wordScrollSyncApplying = false;

        function isLocalStorageAvailable() {
            try {
                const test = '__localStorage_test__';
                localStorage.setItem(test, test);
                localStorage.removeItem(test);
                return true;
            } catch (e) {
                return false;
            }
        }

        const useLocalStorage = isLocalStorageAvailable();

        function saveScroll() {
            const scrollY = window.pageYOffset || document.documentElement?.scrollTop || document.body?.scrollTop || 0;
            memoryScrollPos = scrollY;
            if (useLocalStorage) {
                try {
                    localStorage.setItem(STORAGE_KEY, scrollY.toString());
                } catch (e) {}
            }
        }

        function getMaxScrollY() {
            const docEl = document.documentElement;
            const body = document.body;
            const docHeight = Math.max(
                docEl ? docEl.scrollHeight : 0,
                body ? body.scrollHeight : 0
            );
            return Math.max(0, docHeight - window.innerHeight);
        }

        function getScrollRatio() {
            const scrollY = window.pageYOffset || document.documentElement?.scrollTop || document.body?.scrollTop || 0;
            const maxScrollY = getMaxScrollY();
            if (maxScrollY <= 0) {
                return 0;
            }
            return Math.max(0, Math.min(1, scrollY / maxScrollY));
        }

        function notifyScrollSync(force) {
            if (!window.__wordScrollSyncEnabled || window.__wordScrollSyncApplying) {
                return;
            }
            if (!pyBridge || !webChannelReady || typeof pyBridge.report_word_preview_scroll !== 'function') {
                return;
            }
            const ratio = getScrollRatio();
            if (!force && lastNotifiedRatio !== null && Math.abs(ratio - lastNotifiedRatio) < 0.002) {
                return;
            }
            lastNotifiedRatio = ratio;
            try {
                pyBridge.report_word_preview_scroll(window.__wordPreviewPanelId || '', ratio);
            } catch (e) {}
        }

        function refreshSyncInterval() {
            if (syncInterval) {
                clearInterval(syncInterval);
                syncInterval = null;
            }
            if (!window.__wordScrollSyncEnabled) {
                return;
            }
            syncInterval = setInterval(function() {
                notifyScrollSync(false);
            }, 80);
        }

        window.__setWordPreviewPanelId = function(panelId) {
            window.__wordPreviewPanelId = panelId || '';
        };

        window.__setWordPreviewScrollSyncEnabled = function(enabled) {
            window.__wordScrollSyncEnabled = !!enabled;
            refreshSyncInterval();
            if (window.__wordScrollSyncEnabled) {
                notifyScrollSync(true);
            }
        };

        window.__getWordPreviewScrollRatio = function() {
            return getScrollRatio();
        };

        window.__applyExternalWordScrollRatio = function(ratio) {
            const numericRatio = Number(ratio);
            if (!Number.isFinite(numericRatio)) {
                return;
            }
            const maxScrollY = getMaxScrollY();
            const targetRatio = Math.max(0, Math.min(1, numericRatio));
            const targetY = maxScrollY <= 0 ? 0 : Math.round(maxScrollY * targetRatio);
            lastNotifiedRatio = targetRatio;
            window.__wordScrollSyncApplying = true;
            window.scrollTo(0, targetY);
            setTimeout(function() {
                window.__wordScrollSyncApplying = false;
            }, 55);
        };

        function restoreScroll() {
            let savedPos = null;
            if (useLocalStorage) {
                try {
                    savedPos = localStorage.getItem(STORAGE_KEY);
                } catch (e) {}
            }
            if (!savedPos) {
                savedPos = memoryScrollPos.toString();
            }

            if (savedPos) {
                const targetY = parseInt(savedPos, 10);
                if (!isNaN(targetY) && targetY > 0) {
                    window.scrollTo(0, targetY);
                    setTimeout(function() {
                        window.scrollTo(0, targetY);
                    }, 10);
                }
            }
        }

        saveScroll();
        let scrollTimeout;
        window.addEventListener('scroll', function() {
            clearTimeout(scrollTimeout);
            scrollTimeout = setTimeout(saveScroll, 36);
            clearTimeout(syncTimeout);
            if (syncFrame !== null) {
                cancelAnimationFrame(syncFrame);
            }
            syncFrame = requestAnimationFrame(function() {
                syncFrame = null;
                notifyScrollSync(false);
            });
            syncTimeout = setTimeout(function() {
                notifyScrollSync(false);
            }, 24);
        });
        window.addEventListener('beforeunload', saveScroll);
        document.addEventListener('visibilitychange', function() {
            if (document.hidden) {
                saveScroll();
            } else if (window.__wordScrollSyncEnabled) {
                notifyScrollSync(true);
            }
        });
        if (document.readyState === 'loading') {
            document.addEventListener('DOMContentLoaded', restoreScroll);
        }
        window.addEventListener('load', restoreScroll);
        requestAnimationFrame(restoreScroll);
        refreshSyncInterval();
    })();
    </script>'''

    def _build_word_text_blocks(self):
        text_blocks = {}
        for key, data in self.word_data.items():
            text = data.get("text", "")
            if text:
                text_blocks[key] = {"text": text, "escaped": text}
        return text_blocks

    def _build_word_original_panel_updates(self):
        updates = {}
        for key, data in self.word_data.items():
            source_text = data.get("text", "")
            merged_matches = merge_word_matches_with_priority(
                source_text,
                [],
                self.replacement_text,
                manual_matches=data.get("manual", []),
                ocr_matches=data.get("ocr", [])
            )
            updates[key] = self._build_word_original_preview_fragment(key, source_text, merged_matches)
        return updates

    def _build_word_original_preview_fragment(self, key, source_text, merged_matches):
        from html import escape as html_escape

        segments = build_highlight_preview_segments(source_text, merged_matches)
        parts = []
        for segment in segments:
            value = segment.get("value", "")
            if value is None:
                value = ""
            if not isinstance(value, str):
                value = str(value)
            escaped_value = html_escape(value).replace("\n", "<br/>")

            if segment.get("type") != "highlight":
                parts.append(escaped_value)
                continue

            source = str(segment.get("source", "manual"))
            css_class = "manual-highlight" if source == "manual" else "ocr-highlight"
            attrs = [
                f'class="{css_class}"',
                f'data-key="{html_escape(str(key))}"',
                f'data-start="{int(segment.get("start", 0))}"',
                f'data-end="{int(segment.get("end", 0))}"',
            ]
            title = "手动脱敏" if source == "manual" else str(segment.get("rule_name", "")).strip() or "智能脱敏"
            if title:
                attrs.append(f'title="{html_escape(title)}"')
            parts.append(f"<mark {' '.join(attrs)}>{escaped_value}</mark>")

        return "".join(parts)

    def _build_word_replaced_panel_updates(self):
        updates = {}
        for key, data in self.word_data.items():
            source_text = data.get("text", "")
            merged_matches = merge_word_matches_with_priority(
                source_text,
                self.word_replace_rules,
                self.replacement_text,
                manual_matches=data.get("manual", []),
                ocr_matches=data.get("ocr", [])
            )
            updates[key] = self._build_replaced_preview_fragment(source_text, merged_matches)
        return updates

    def _apply_word_panel_updates(self, web_view, block_updates):
        if web_view is None or not block_updates:
            return

        script = build_word_panel_update_script(block_updates)
        web_view.page().runJavaScript(script)

    def _configure_word_scroll_sync_panel(self, web_view, panel_id):
        """配置单个 Word 预览面板的滚动同步能力。"""
        if not self._is_word_web_view_valid(web_view):
            return
        sync_enabled = bool(
            self.word_doc
            and self.word_compare_mode
            and not self.word_compare_user_hidden
            and not web_view.isHidden()
        )
        panel_js = json.dumps(str(panel_id))
        enabled_js = "true" if sync_enabled else "false"
        web_view.page().runJavaScript(
            f"""
            if (window.__setWordPreviewPanelId) window.__setWordPreviewPanelId({panel_js});
            if (window.__setWordPreviewScrollSyncEnabled) window.__setWordPreviewScrollSyncEnabled({enabled_js});
            """
        )

    def _configure_word_scroll_sync_panels(self):
        """同步刷新左右 Word 预览面板的滚动联动状态。"""
        if hasattr(self, "word_preview"):
            self._configure_word_scroll_sync_panel(self.word_preview, "original")
        if hasattr(self, "word_preview_replaced"):
            self._configure_word_scroll_sync_panel(self.word_preview_replaced, "replaced")
        self._refresh_word_scroll_sync_timer()

    def _refresh_word_scroll_sync_timer(self):
        """按当前 Word 双栏状态启停滚动同步轮询。"""
        should_run = bool(
            self.word_doc
            and self.word_compare_mode
            and not self.word_compare_user_hidden
            and self._word_preview_ready
            and self._word_replaced_ready
            and self._is_word_web_view_valid(self.word_preview)
            and self._is_word_web_view_valid(self.word_preview_replaced)
            and not self.word_preview.isHidden()
            and not self.word_preview_replaced.isHidden()
        )
        if should_run:
            if not self._word_scroll_sync_timer.isActive():
                self._word_scroll_sync_timer.start()
        else:
            self._word_scroll_sync_timer.stop()
            self._word_scroll_sync_polling = False
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None
            self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}

    def _apply_word_scroll_ratio_to_panel(self, panel_id, ratio):
        """把滚动比例应用到指定 Word 预览面板。"""
        try:
            ratio_value = max(0.0, min(1.0, float(ratio)))
        except (TypeError, ValueError):
            return

        target_view = self.word_preview_replaced if panel_id == "replaced" else self.word_preview
        if not self._is_word_web_view_valid(target_view) or target_view.isHidden():
            return

        self._word_scroll_sync_pending_target = panel_id
        self._word_scroll_sync_pending_ratio = ratio_value
        target_view.page().runJavaScript(
            f"if (window.__applyExternalWordScrollRatio) window.__applyExternalWordScrollRatio({ratio_value:.6f});"
        )

    def _sync_word_compare_scroll(self, source_panel, ratio):
        """双栏对比模式下，同步左右预览滚动位置。"""
        if not self.word_compare_mode or self.word_compare_user_hidden:
            return
        if source_panel not in {"original", "replaced"}:
            return
        try:
            ratio_value = max(0.0, min(1.0, float(ratio)))
        except (TypeError, ValueError):
            return
        if (
            self._word_scroll_sync_pending_target == source_panel
            and self._word_scroll_sync_pending_ratio is not None
            and abs(ratio_value - float(self._word_scroll_sync_pending_ratio)) <= 0.02
        ):
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None
            return
        if source_panel == "original":
            self._apply_word_scroll_ratio_to_panel("replaced", ratio_value)
        else:
            self._apply_word_scroll_ratio_to_panel("original", ratio_value)

    def _sync_word_compare_scroll_from_original(self):
        """用左侧原文预览的当前位置对齐右侧替换预览。"""
        if not self.word_compare_mode or self.word_compare_user_hidden:
            return
        if not self._word_preview_ready or not self._word_replaced_ready:
            return
        if not self._is_word_web_view_valid(self.word_preview):
            return
        if not self._is_word_web_view_valid(self.word_preview_replaced):
            return
        if self.word_preview.isHidden() or self.word_preview_replaced.isHidden():
            return

        generation = self._word_scroll_sync_generation
        self.word_preview.page().runJavaScript(
            "window.__getWordPreviewScrollRatio ? window.__getWordPreviewScrollRatio() : 0;",
            lambda ratio, sync_generation=generation: self._sync_word_compare_scroll_from_original_callback(ratio, sync_generation),
        )

    def _sync_word_compare_scroll_from_original_callback(self, ratio, generation):
        """仅在同步代次仍有效时，把左侧滚动位置应用到右侧。"""
        if generation != self._word_scroll_sync_generation:
            return
        self._apply_word_scroll_ratio_to_panel("replaced", ratio)

    def _poll_word_compare_scroll_sync(self):
        """轮询双栏 Word 预览滚动位置，作为联动兜底方案。"""
        if self._word_scroll_sync_polling:
            return
        if not self._word_scroll_sync_timer.isActive():
            return
        if not self.word_compare_mode or self.word_compare_user_hidden:
            return
        if not self._word_preview_ready or not self._word_replaced_ready:
            return
        if not self._is_word_web_view_valid(self.word_preview):
            return
        if not self._is_word_web_view_valid(self.word_preview_replaced):
            return
        if self.word_preview.isHidden() or self.word_preview_replaced.isHidden():
            return

        self._word_scroll_sync_polling = True
        generation = self._word_scroll_sync_generation
        self.word_preview.page().runJavaScript(
            "window.__getWordPreviewScrollRatio ? window.__getWordPreviewScrollRatio() : 0;",
            lambda ratio, sync_generation=generation: self._handle_word_scroll_sync_original_ratio(ratio, sync_generation),
        )

    def _handle_word_scroll_sync_original_ratio(self, original_ratio, generation):
        """获取左侧比例后继续读取右侧比例。"""
        if generation != self._word_scroll_sync_generation:
            self._word_scroll_sync_polling = False
            return
        try:
            normalized_original = max(0.0, min(1.0, float(original_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_original = 0.0

        if not self._is_word_web_view_valid(self.word_preview_replaced) or self.word_preview_replaced.isHidden():
            self._word_scroll_sync_polling = False
            return

        self.word_preview_replaced.page().runJavaScript(
            "window.__getWordPreviewScrollRatio ? window.__getWordPreviewScrollRatio() : 0;",
            lambda replaced_ratio, sync_generation=generation: self._handle_word_scroll_sync_ratio_pair(normalized_original, replaced_ratio, sync_generation),
        )

    def _handle_word_scroll_sync_ratio_pair(self, original_ratio, replaced_ratio, generation):
        """比较双栏滚动位置，并同步变化更明显的一侧。"""
        if generation != self._word_scroll_sync_generation:
            self._word_scroll_sync_polling = False
            return
        try:
            normalized_original = max(0.0, min(1.0, float(original_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_original = 0.0
        try:
            normalized_replaced = max(0.0, min(1.0, float(replaced_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_replaced = 0.0

        previous_original = self._word_scroll_sync_last_ratios.get("original")
        previous_replaced = self._word_scroll_sync_last_ratios.get("replaced")
        self._word_scroll_sync_last_ratios = {
            "original": normalized_original,
            "replaced": normalized_replaced,
        }

        if previous_original is None or previous_replaced is None:
            self._word_scroll_sync_polling = False
            return

        delta_original = abs(normalized_original - previous_original)
        delta_replaced = abs(normalized_replaced - previous_replaced)
        threshold = 0.012

        if (
            self._word_scroll_sync_pending_target == "replaced"
            and abs(normalized_replaced - float(self._word_scroll_sync_pending_ratio or 0.0)) <= 0.02
        ):
            delta_replaced = 0.0
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None
        elif (
            self._word_scroll_sync_pending_target == "original"
            and abs(normalized_original - float(self._word_scroll_sync_pending_ratio or 0.0)) <= 0.02
        ):
            delta_original = 0.0
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None

        if delta_original >= threshold and delta_original > delta_replaced:
            self._apply_word_scroll_ratio_to_panel("replaced", normalized_original)
        elif delta_replaced >= threshold and delta_replaced > delta_original:
            self._apply_word_scroll_ratio_to_panel("original", normalized_replaced)

        self._word_scroll_sync_polling = False

    def _on_word_preview_load_finished(self, ok):
        self._word_preview_ready = bool(ok)
        self._word_preview_loaded_source_path = self._word_preview_target_source_path if ok else None
        if ok and self._pending_word_preview_blocks:
            self._apply_word_panel_updates(self.word_preview, self._pending_word_preview_blocks)
        if ok:
            self._configure_word_scroll_sync_panel(self.word_preview, "original")
            self._refresh_word_scroll_sync_timer()
            QTimer.singleShot(0, self._sync_word_compare_scroll_from_original)

    def _on_word_replaced_load_finished(self, ok):
        self._word_replaced_ready = bool(ok)
        self._word_replaced_loaded_source_path = self._word_replaced_target_source_path if ok else None
        if ok and self._pending_word_replaced_blocks:
            self._apply_word_panel_updates(self.word_preview_replaced, self._pending_word_replaced_blocks)
        if ok:
            self._configure_word_scroll_sync_panel(self.word_preview_replaced, "replaced")
            self._refresh_word_scroll_sync_timer()
            QTimer.singleShot(0, self._sync_word_compare_scroll_from_original)

    def _add_data_key_attributes(self, html, text_blocks):
        """使用 BeautifulSoup 为文本块添加 data-key 属性

        Args:
            html: HTML 字符串
            text_blocks: 文本块字典 {key: {'text': 原始文本, 'escaped': 转义文本}}

        Returns:
            修改后的 HTML 字符串
        """
        try:
            soup = BeautifulSoup(html, 'html.parser')

            # 支持的标签列表（扩展）
            target_tags = ['p', 'td', 'th', 'li', 'span', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'a']

            for key, info in text_blocks.items():
                original_text = info['text']
                if not original_text or not original_text.strip():
                    continue

                # 归一化文本：将多个空白字符合并为单个空格
                normalized_original = ' '.join(original_text.split())

                for tag_name in target_tags:
                    for element in soup.find_all(tag_name):
                        # 跳过已经有 data-key 的元素
                        if element.get('data-key'):
                            continue

                        # 获取元素的文本内容并归一化
                        element_text = element.get_text()
                        normalized_element = ' '.join(element_text.split())

                        # 比较归一化后的文本
                        if normalized_original == normalized_element:
                            element['data-key'] = key
                            element['data-original-text'] = original_text
                            element['data-word-block'] = '1'
                            break

            return str(soup)

        except (ImportError, AttributeError, TypeError, ValueError) as e:
            print(f"[警告] BeautifulSoup 处理失败，使用正则表达式后备方案: {e}")
            return self._add_data_key_regex_fallback(html, text_blocks)

    def _add_data_key_regex_fallback(self, html, text_blocks):
        """使用正则表达式为文本块添加 data-key 属性（后备方案）

        Args:
            html: HTML 字符串
            text_blocks: 文本块字典

        Returns:
            修改后的 HTML 字符串
        """
        from html import escape as html_escape

        for key, info in text_blocks.items():
            original_text = info['text']
            escaped_text = html_escape(original_text)

            # 尝试各种 HTML 标签
            patterns = [
                (
                    f'<p([^>]*)>({re.escape(escaped_text)})</p>',
                    f'<p\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</p>'
                ),
                (
                    f'<td([^>]*)>({re.escape(escaped_text)})</td>',
                    f'<td\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</td>'
                ),
                (
                    f'<li([^>]*)>({re.escape(escaped_text)})</li>',
                    f'<li\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</li>'
                ),
                (
                    f'<span([^>]*)>({re.escape(escaped_text)})</span>',
                    f'<span\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</span>'
                ),
                (
                    f'<div([^>]*)>({re.escape(escaped_text)})</div>',
                    f'<div\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</div>'
                ),
            ]

            for pattern, replacement in patterns:
                if f'data-key="{key}"' not in html:
                    if re.search(pattern, html):
                        html = re.sub(pattern, replacement, html, count=1)
                        break

        return html

    def _highlight_exact_match(self, html, match):
        """使用 BeautifulSoup 精确高亮指定位置的文本

        Args:
            html: HTML 字符串
            match: 包含 key, start, end, text 的匹配信息

        Returns:
            修改后的 HTML 字符串
        """
        from bs4 import BeautifulSoup, NavigableString

        key = match['key']
        start = match['start']
        end = match['end']
        text = match['text']

        try:
            soup = BeautifulSoup(html, 'html.parser')

            # 找到对应的容器元素
            container = soup.find(attrs={'data-key': key})
            if not container:
                print(f"[警告] 未找到 data-key={key} 的元素")
                return html

            # 遍历所有文本节点，找到对应位置
            current_pos = 0
            for child in list(container.descendants):
                if isinstance(child, NavigableString) and not isinstance(child, str):
                    continue
                if isinstance(child, str):
                    node_text = str(child)
                    node_start = current_pos
                    node_end = current_pos + len(node_text)

                    # 检查目标范围是否在这个节点内
                    if start >= node_start and end <= node_end:
                        # 计算在当前节点内的相对位置
                        rel_start = start - node_start
                        rel_end = end - node_start

                        # 分割文本并插入 mark 标签
                        before = node_text[:rel_start]
                        highlighted = node_text[rel_start:rel_end]
                        after = node_text[rel_end:]

                        # 创建新的标记
                        mark_tag = soup.new_tag('mark')
                        mark_tag['class'] = 'manual-highlight'
                        mark_tag['data-key'] = key
                        mark_tag['data-start'] = start
                        mark_tag['data-end'] = end
                        mark_tag['style'] = 'background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone;'
                        mark_tag.string = highlighted

                        # 替换原文本节点
                        parent = child.parent
                        # 创建新的节点序列
                        new_nodes = []
                        if before:
                            new_nodes.append(NavigableString(before))
                        new_nodes.append(mark_tag)
                        if after:
                            new_nodes.append(NavigableString(after))

                        # 替换原节点
                        child.replace_with(*new_nodes)

                        return str(soup)

                    current_pos = node_end

            print(f"[警告] 精确模式未找到位置: key={key}, start={start}, end={end}")
            return html

        except (ImportError, AttributeError, TypeError, ValueError, IndexError) as e:
            print(f"[错误] 精确高亮失败: {e}")
            return html

    def _highlight_sensitive_info(self, html):
        """在 HTML 中高亮显示敏感信息（使用 JavaScript 进行高亮）"""
        import json
        from html import escape
        import re

        # 构建所有匹配数据的 JSON
        matches_data = []
        # 构建文本块数据（用于标记段落/单元格）
        text_blocks = {}

        for key, data in self.word_data.items():
            text = data['text']
            if not text:
                continue

            # 记录文本块信息
            text_blocks[key] = {
                'text': text,
                'escaped': escape(text)
            }

            # OCR 匹配
            for match in data['ocr']:
                matches_data.append({
                    'key': key,
                    'start': match['start'],
                    'end': match['end'],
                    'text': match['text'],
                    'type': 'ocr',
                    'rule_name': match.get('rule_name', 'OCR')
                })

            # 手动脱敏
            for match in data['manual']:
                matches_data.append({
                    'key': key,
                    'start': match['start'],
                    'end': match['end'],
                    'text': match['text'],
                    'type': 'manual',
                    'mode': match.get('mode', 'exact')  # 添加模式标识，默认为 exact
                })

        # === 新方法：直接在 HTML 字符串中进行高亮替换 ===
        from html import escape as html_escape

        # 按文本长度降序排序，先处理长的匹配
        matches_data.sort(key=lambda x: len(x['text']), reverse=True)

        # 分离精确模式和全局模式
        exact_matches = [m for m in matches_data if m.get('mode') == 'exact']
        global_matches = [m for m in matches_data if m.get('mode') != 'exact']

        # === 第一步：为所有文本块添加 data-key 属性（使用 BeautifulSoup 替代正则表达式）===
        html = self._add_data_key_attributes(html, text_blocks)

        # 处理全局模式匹配（OCR 和全局手动脱敏）
        # 去重：相同文本只处理一次
        processed_global_texts = set()
        for match in global_matches:
            text = match['text']
            if text in processed_global_texts:
                continue
            processed_global_texts.add(text)

            # 需要转义 HTML 特殊字符
            escaped_text = html_escape(text)
            is_ocr = match['type'] == 'ocr'

            # 构建替换标记
            if is_ocr:
                replacement = f'<mark class="ocr-highlight" data-key="{match["key"]}" data-start="{match["start"]}" data-end="{match["end"]}" title="{match.get("rule_name", "")}" style="background-color: #ffeb3b; color: #000; display: inline; box-decoration-break: clone; -webkit-box-decoration-break: clone;">{escaped_text}</mark>'
            else:
                replacement = f'<mark class="manual-highlight" data-key="{match["key"]}" data-start="{match["start"]}" data-end="{match["end"]}" style="background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone;">{escaped_text}</mark>'

            # 使用正则表达式替换所有匹配
            # 使用正向预查避免匹配已经在标记中的文本
            pattern = re.compile(re.escape(escaped_text) + r'(?![^<]*>)')
            html = pattern.sub(replacement, html)

        # 处理精确模式匹配（只高亮特定位置的文本）
        # 使用 BeautifulSoup 进行精确位置定位
        for match in exact_matches:
            html = self._highlight_exact_match(html, match)

        # 为了兼容性，仍然生成 matches_json（但不再用于高亮）
        matches_json = json.dumps(matches_data, ensure_ascii=False)
        text_blocks_json = json.dumps(text_blocks, ensure_ascii=False)

        # 添加简化版的 JavaScript 高亮脚本（仅用于调试）
        highlight_script = '''
        <script>
        (function() {
            console.log('[Highlight] HTML 预高亮已完成');
            // 高亮已在 Python 端完成，这里不需要再做
        })();
        </script>
        '''

        # 添加样式
        style = f"""
        <style>
            body {{ font-family: {PREVIEW_FONT_STACK}; padding: 20px; line-height: 1.6; }}
            p:empty {{ display: none; margin: 0; }}
            mark.ocr-highlight {{ background-color: #ffeb3b; color: #000; display: inline; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight {{ background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight:hover {{ box-shadow: 0 0 0 1px #e03131, 0 0 4px rgba(225, 49, 49, 0.5); }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; }}
        </style>
        """

        # 插入样式和脚本
        if '<head>' in html:
            html = html.replace('<head>', '<head>' + style, 1)
        else:
            html = style + html

        # 在 </body> 前插入脚本，或者直接追加到末尾
        if '</body>' in html:
            html = html.replace('</body>', highlight_script + '</body>', 1)
        else:
            html = html + highlight_script

        return html

    def _inject_interactive_html(self, html, scroll_restore=''):
        """注入 JavaScript 交互逻辑用于右键菜单和脱敏操作

        v36.4: 重构为使用模块级常量 _INTERACTIVE_JS_CODE，简化函数逻辑

        Args:
            html: 要注入的 HTML
            scroll_restore: 滚动恢复脚本（可选）
        """
        # 包装 HTML 为完整文档（如果不是的话）
        html = self._wrap_html_document(html)

        # 注入脚本
        qwebchannel_js = '<script src="qrc:///qtwebchannel/qwebchannel.js"></script>'

        if '</head>' in html:
            html = html.replace('</head>', qwebchannel_js + _INTERACTIVE_JS_CODE + scroll_restore + '</head>')
        else:
            html = qwebchannel_js + _INTERACTIVE_JS_CODE + scroll_restore + html

        return html

    def _wrap_html_document(self, html):
        """将 HTML 包装成完整文档（如果不是完整文档的话）

        Args:
            html: 输入的 HTML 字符串

        Returns:
            完整的 HTML 文档字符串
        """
        is_full_document = '<html' in html.lower() or '<!doctype' in html.lower()

        if is_full_document:
            return html

        return f'''<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
    body {{ margin: 0; padding: 20px; font-family: {PREVIEW_FONT_STACK}; line-height: 1.6; }}
    img {{ max-width: 100%; height: auto; }}
    p {{ margin: 0 0 10px 0; }}
</style>
</head>
<body>
{html}
</body>
</html>'''

    def save_pdf(self):
        """保存脱敏后的文档 - 支持 PDF 和 Word - v37.3: 安全加固，脱敏区域永久化"""
        # v36: 应用文件对话框样式
        # v37.3: PDF 脱敏安全加固 - 脱敏区域永久嵌入，不可编辑
        app = QApplication.instance()
        original_style = app.styleSheet()

        # PDF 保存
        if self.doc:
            app.setStyleSheet(self._get_file_dialog_style())
            try:
                fname, _ = QFileDialog.getSaveFileName(self, "保存 PDF", "", "PDF Files (*.pdf)")
            finally:
                app.setStyleSheet(original_style)

            if fname:
                doc_save = None
                try:
                    QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
                    doc_save = fitz.open(self.file_path)
                    fill_col = (0, 0, 0) if self.current_color.name() == "#000000" else (1, 1, 1)

                    for i in range(len(doc_save)):
                        page = doc_save[i]

                        # v37.3.1: 修复内部编辑功能 - 使用副本避免修改原始数据
                        # 从 page_data 中获取脱敏区域列表
                        ocr_list = self.page_data[i].get('ocr', [])
                        manual_list = self.page_data[i].get('manual', [])

                        # 1. 添加脱敏注释
                        # v37.3.1: 重建 QRectF 确保不修改原始对象
                        for r in ocr_list + manual_list:
                            # 从 QRectF 提取坐标并重建，避免引用问题
                            x, y, w, h = r.x(), r.y(), r.width(), r.height()
                            rect = fitz.Rect(x, y, x + w, y + h)
                            annot = page.add_redact_annot(rect)
                            annot.set_colors(stroke=fill_col, fill=fill_col)
                            annot.update()

                        # v37.3: 安全加固 - 修改图像像素，彻底销毁原始内容
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

                        # v37.3: 安全加固 - 删除所有注释对象，防止被 PDF 编辑器修改
                        for annot in page.annots():
                            page.delete_annot(annot)

                    # v37.3: 安全加固 - 使用垃圾回收和压缩彻底删除未引用对象
                    doc_save.save(
                        fname,
                        garbage=4,        # 最大垃圾回收级别
                        deflate=True,     # 压缩内容流
                        clean=True        # 清理未引用对象
                    )

                    QApplication.restoreOverrideCursor()
                    QMessageBox.information(self, "成功", f"文件已安全保存至：\n{fname}")
                except (IOError, OSError, ValueError, RuntimeError) as e:
                    QApplication.restoreOverrideCursor()
                    QMessageBox.critical(self, "失败", str(e))
                finally:
                    if doc_save:
                        doc_save.close()
        # Word 保存
        elif self.word_doc:
            app.setStyleSheet(self._get_file_dialog_style())
            try:
                fname, _ = QFileDialog.getSaveFileName(self, "保存 Word", "", "Word 文档 (*.docx)")
            finally:
                app.setStyleSheet(original_style)

            if fname:
                self._save_word(fname)

    def _save_word(self, fname):
        """保存 Word 文档 - v24 改进版：详细错误处理 + 使用 TempFileManager + 合并 OCR 和 Manual 脱敏"""
        try:
            import shutil
            from docx import Document

            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

            # 使用 TempFileManager 管理临时文件
            temp_file = self.temp_manager.create_temp_file()
            shutil.copy2(self.file_path, temp_file)

            # 打开副本进行修改
            new_doc = Document(temp_file)

            # 遍历段落进行 run 级别的文本替换
            for para_idx, para in enumerate(new_doc.paragraphs):
                key = f'paragraph_{para_idx}'
                if key in self.word_data:
                    data = self.word_data[key]
                    source_text = data.get("text", "")
                    merged_matches = merge_word_matches_with_priority(
                        source_text,
                        self.word_replace_rules,
                        self.replacement_text,
                        manual_matches=data.get("manual", []),
                        ocr_matches=data.get("ocr", [])
                    )
                    if merged_matches:
                        replace_matches_in_paragraph(
                            para,
                            merged_matches,
                            text_offset=0,
                            fallback_replacement_text=self.replacement_text
                        )

            # 遍历表格进行 run 级别的文本替换
            for table_idx, table in enumerate(new_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                        if key in self.word_data:
                            data = self.word_data[key]
                            source_text = data.get("text", "")
                            merged_matches = merge_word_matches_with_priority(
                                source_text,
                                self.word_replace_rules,
                                self.replacement_text,
                                manual_matches=data.get("manual", []),
                                ocr_matches=data.get("ocr", [])
                            )

                            if merged_matches:
                                # 处理单元格内的所有段落（按 cell.text 的偏移映射）
                                para_offset = 0
                                paragraphs = list(cell.paragraphs)
                                for idx, para in enumerate(paragraphs):
                                    original_para_len = len(''.join(run.text for run in para.runs))
                                    replace_matches_in_paragraph(
                                        para,
                                        merged_matches,
                                        text_offset=para_offset,
                                        fallback_replacement_text=self.replacement_text
                                    )
                                    para_offset += original_para_len
                                    if idx < len(paragraphs) - 1:
                                        # python-docx 的 cell.text 使用换行拼接段落
                                        para_offset += 1

            # 保存文档
            new_doc.save(fname)

            QApplication.restoreOverrideCursor()
            QMessageBox.information(self, "成功", f"文件已保存至：\n{fname}")

        except PermissionError:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self, "保存失败",
                f"没有写入权限：\n{fname}\n\n"
                "建议：\n"
                "1. 检查文件是否被其他程序打开\n"
                "2. 检查文件夹权限\n"
                "3. 尝试保存到其他位置"
            )

        except (OSError, IOError, RuntimeError, ValueError, KeyError, AttributeError) as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self, "保存失败",
                f"保存 Word 文档时出错：\n{str(e)}\n\n"
                "请尝试：\n"
                "1. 重启应用\n"
                "2. 在 Word 中手动打开文件\n"
                "3. 导出错误日志以供分析"
            )

    def _replace_in_paragraph(self, para, matches, text_offset=0):
        """兼容入口：调用通用段落替换实现。"""
        replace_matches_in_paragraph(
            para,
            matches,
            text_offset=text_offset,
            fallback_replacement_text=self.replacement_text
        )

    def _apply_range_to_runs(self, para, start, end, replacement):
        """兼容入口：调用通用 run 区间替换实现。"""
        apply_range_to_runs(para, start, end, replacement)

    def _copy_run_format(self, target_run, source_run):
        """复制 run 的所有格式属性

        Args:
            target_run: 目标 run
            source_run: 源 run
        """
        # 字体属性
        if source_run.bold is not None:
            target_run.bold = source_run.bold
        if source_run.italic is not None:
            target_run.italic = source_run.italic
        if source_run.underline is not None:
            target_run.underline = source_run.underline
        if source_run.strike is not None:
            target_run.strike = source_run.strike

        # 字体名称和大小
        if source_run.font.name:
            try:
                target_run.font.name = source_run.font.name
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制字体名称失败: {e}")
        if source_run.font.size:
            target_run.font.size = source_run.font.size

        # 颜色
        if source_run.font.color and source_run.font.color.rgb:
            try:
                target_run.font.color.rgb = source_run.font.color.rgb
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制字体颜色失败: {e}")

        # 高亮
        if source_run.font.highlight_color:
            try:
                target_run.font.highlight_color = source_run.font.highlight_color
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制高亮颜色失败: {e}")

        # 下标/上标
        if source_run.font.subscript:
            target_run.font.subscript = True
        if source_run.font.superscript:
            target_run.font.superscript = True

if __name__ == "__main__":
    # v37.0.5: 全局异常钩子，防止未捕获异常导致崩溃
    def exception_hook(exc_type, exc_value, exc_traceback):
        """全局异常处理器"""
        error_msg = ''.join(traceback.format_exception(exc_type, exc_value, exc_traceback))
        print(f"[FATAL ERROR] 未捕获的异常:\n{error_msg}")

        # 尝试显示错误对话框
        try:
            if QApplication.instance():
                QMessageBox.critical(
                    None,
                    "程序错误",
                    f"程序遇到未预期的错误：\n\n{exc_type.__name__}: {exc_value}\n\n"
                    "请将此错误信息反馈给开发者。"
                )
        except Exception:
            pass

        # 调用默认异常处理器
        sys.__excepthook__(exc_type, exc_value, exc_traceback)

    sys.excepthook = exception_hook

    # v37.0.5: 线程异常钩子
    def thread_exception_hook(args):
        """线程异常处理器"""
        error_msg = ''.join(traceback.format_exception(args.exc_type, args.exc_value, args.exc_traceback))
        print(f"[THREAD ERROR] 线程异常:\n{error_msg}")

    threading.excepthook = thread_exception_hook

    # v37.0.5: 启动时预加载 OCR 引擎（可选，用于早期检测问题）
    if os.getenv('PRIVACYGUARD_PRELOAD_OCR', '').lower() == 'true':
        print("[INFO] 预加载 OCR 引擎...")
        init_ocr_engine()

    app = QApplication(sys.argv)
    
    # 设置应用图标
    icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', 'logo', 'export', '256', 'logo_default_256.png')
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))
        print(f"[INFO] 应用图标已加载: {icon_path}")
    else:
        print(f"[WARN] 应用图标未找到: {icon_path}")
    
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
